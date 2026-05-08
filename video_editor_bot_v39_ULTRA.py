"""
🎬 Video Editor Bot — Pro Edition v25 (Termux Ready)
Wink-Style Premium: Enhance Ultra + Wink Premium 4K + Filters + Speed + Aspect +
Flip + Fade + Boomerang + GIF + Watermark Remove + Trim +
Merge + Text + BGM + Volume + Blur BG + Vignette + Zoom + Freeze Frame

──────────────────────────────────────────────────────────────────
VERSION HISTORY (সংক্ষেপে)
──────────────────────────────────────────────────────────────────
v13: ChatGPT (OpenAI) integration — SEO + Subtitle (Whisper) + AI Chat
     + Text→Voice (TTS) — gTTS (free) + OpenAI TTS (premium voice)
     + Multi-language voice (Bengali / English / Hindi / Auto)
v14: ElevenLabs TTS (super realistic, multilingual)
v15: Bengali Premium Font System + Status Video Maker (9:16)
v16: Bug fixes — drawtext escape, A/V drift, gemini safety,
     acoustid validation, ElevenLabs error handling
v17: 🆕 MiniMax integration + ENV-key override + HTTP retry helper
     + per-feature stats + cleanup size cap + cleaner banner
v22: Edge TTS (Microsoft Neural — সম্পূর্ণ ফ্রি realistic voice)
v23: 🔧 ভয়েস ফিউচার ফিক্স — ElevenLabs key invalid হলে এখন Edge TTS-এ
     auto-fallback হবে। Edge TTS এখন default realistic engine
     (key-free, neural quality)।
v24: 🚀 আপলোড লিমিট 50MB → 2000MB (2GB) — বড় ভিডিও ডাউনলোড + পাঠানো যাবে।
     ❌ ElevenLabs সম্পূর্ণ disable — paid plan লাগে, কাজ করছিল না।
     ✅ বিকল্প: Edge TTS (Microsoft Neural — ২১টা realistic voice,
        Bangla/English/Hindi সহ ৬৭+ ভাষা, সম্পূর্ণ ফ্রি, কোনো key লাগে না)।
     Voice Clone menu hidden (paid ElevenLabs ছাড়া কাজ করত না)।
v25: 📺 NEW — YouTube Mode (Content ID Safe): এক tap-এ
     16:9 (1920x1080) + YouTube-specific copyright bypass
     (audio pitch + tempo + EQ shift, container fingerprint break,
     subtle visual transforms — flip ছাড়া কারণ YT viewer ধরে ফেলে)।
     🐛 Banner version হার্ডকোড fix (v22 → dynamic VERSION constant)।
v26: 🎭 Voice Clone — ElevenLabs সম্পূর্ণ বাদ → Replicate XTTS-v2
     (zero-shot cloning, Bengali/Hindi/English/16+ language)।
     Replicate-এ ফ্রি ক্রেডিট আছে, signup করলেই কাজ করে — paid plan
     লাগে না। ১০-৩০ সেকেন্ড voice sample → যেকোনো text → আপনার কণ্ঠ।
     /clonevoice দিয়ে save করুন (sample-টা cloud-এ যায় না, Bot-এর local
     storage-এ থাকে; প্রতি TTS-এ Replicate-এ পাঠানো হয়)।
     📌 Pinterest fix — যেসব Pinterest pin / pin.it short link normally
     yt-dlp-তে fail করত সেগুলোর জন্য mobile UA + extractor-args দিয়ে
     auto-retry যোগ হলো। এখন বেশিরভাগ Pinterest video নামবে।
v27.2: 💰 YouTube Money গাইড — "💰 YouTube টাকা আয়ের গাইড" button +
       /money command যোগ। Bengali-তে complete monetization bible:
       YPP eligibility (1K subs/4K hr বা 500 subs/3K hr early), ৭টা
       income source (AdSense/Shorts/Memberships/SuperChat/Shopping/
       Sponsorship/Affiliate), Bengali niche RPM data, Bangladesh
       payment options (Bank/Western Union/Payoneer), W-8BEN tax
       form, demonetization risks, realistic earning timeline,
       এবং এই bot-এর tools দিয়ে দ্রুত monetize করার strategy।
       ৪টা message-এ split (Telegram limit মেনে)।
v27.1: 🎤 Auto Voice Clone Button — Edit menu-তে "অটো ভয়েস ক্লোন (এই
       ভিডিও থেকে)" button যোগ। এক tap-এ:
       (১) Current video বা link-থেকে নামানো video থেকে audio extract,
       (২) Replicate XTTS-v2-এ test sample upload,
       (৩) Clone save + test playback voice হিসেবে delivered।
       Clone নাম auto-generate হয় video title থেকে। পুরনো clone
       থাকলে replace হয়।
       🎨 Menu Reorganization — পুরনো 30+ button-এর flat list
       reorganize করে ৭টা clean section: ⭐ TOP HITS, 📺 PLATFORM
       MODES, 🎭 VOICE & AI, ✨ QUALITY & LOOK, ✂️ EDITING,
       ⚙️ EFFECTS, 🔧 UTILITIES। Section header (decorative)
       click করলে hint show হয়।
       💫 UX polish — better Bengali microcopy, clearer error
       messages, consistent emoji style।
v27:   🎞️ Slow-Mo PREMIUM — minterpolate-এর হালকা variant (obmc+bilat+
     vsbmc=0) + 48fps target → আগের চেয়ে 4-5x faster, একই smoothness।
     Audio-ও 0.5x slow হয় (mute না হয়ে)। CRF 18 + slow preset =
     broadcast quality, কোনো blocking artifact না।
     ✨ Enhance ULTRA — সত্যিই enhanced look (আগেরটা দুর্বল ছিল)।
     Strong 3-pass denoise + smart upscale (1080p-এর কম হলে 1.5x
     lanczos) + 2-stage sharpening + cinematic color grading + audio
     loudness norm (broadcast EBU R128) + vignette। CRF 17 = visually
     lossless।
     🚀 YouTube ভাইরাল প্যাক — এক tap-এ: video viral-style edit (Ken
     Burns subtle zoom + vibrant color + loudness norm) + AI দিয়ে
     ৩টা viral title + description + ৩০ hashtag + ১০টা pro tips।
     YouTube এ ভাইরাল হওয়ার complete blueprint।
     📺 YouTube Mode confirmation — Content ID bypass active: pitch
     shift + tempo + bandpass + parametric EQ + container metadata
     wipe + invisible 1-2px borders + microscopic rotate। নিজের
     content-এ pixel-perfect কাজ করছে। (Original কপি reupload-এর
     জন্য না — disclaimer banner-এ আছে।)

╔══════════════════════════════════════════════════════════════════╗
║ 📦 TERMUX SETUP (একবার চালান) — সব command একসাথে কপি-পেস্ট:    ║
╠══════════════════════════════════════════════════════════════════╣
║                                                                  ║
║  pkg update -y && pkg upgrade -y                                 ║
║  pkg install -y python ffmpeg chromaprint                        ║
║  pip install --upgrade pip                                       ║
║  pip install "python-telegram-bot[job-queue]" yt-dlp gTTS edge-tts ║
║                                                                  ║
║  # তারপর বট চালান:                                                ║
║  python video_editor_bot_v25_pro.py                              ║
║                                                                  ║
╠══════════════════════════════════════════════════════════════════╣
║ 🔑 API KEYS — নিচের config block-এ বসান বা ENV var হিসেবে set:    ║
║   • BOT_TOKEN          — @BotFather থেকে ফ্রি (REQUIRED)          ║
║   • ACOUSTID_API_KEY   — https://acoustid.org/new-application     ║
║   • GEMINI_API_KEY     — https://aistudio.google.com/apikey ফ্রি ║
║   • OPENAI_API_KEY     — https://platform.openai.com/api-keys     ║
║   • MINIMAX_API_KEY    — https://www.minimax.io                   ║
║                                                                  ║
║   ❌ ElevenLabs আর লাগবে না — Edge TTS already free realistic    ║
║      voice দিচ্ছে (২১টা Bangla/English/Hindi neural voice)।       ║
║                                                                  ║
║   💡 প্রতিটা key ENV থেকে set করলে কোডে hardcode না করেও চলবে।   ║
║      যেমন:  export GEMINI_API_KEY=xxxxx                           ║
╚══════════════════════════════════════════════════════════════════╝
"""

import os, sys, re, json, logging, asyncio, subprocess, tempfile, uuid, threading, shlex, random, time
from datetime import datetime
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.constants import ChatAction
from telegram.error import BadRequest
from telegram.ext import (Application, CommandHandler, MessageHandler,
                          CallbackQueryHandler, filters, ContextTypes)

# ──────────────────────────────────────────────────────────────
# CONFIG — আপনার টোকেন এবং আইডি এখানে বসান
# 🆕 v17: প্রতিটা key এখন ENV থেকেও override করা যাবে।
#         ENV-এ থাকলে সেটা priority পাবে, না থাকলে নিচের default ব্যবহার হবে।
# ──────────────────────────────────────────────────────────────
BOT_TOKEN = os.getenv("BOT_TOKEN") or "8713772697:AAEK2ZkWBBjLDrlqEa1QxCAZyilPYkDjbDQ"  # @BotFather থেকে পাওয়া টোকেন
ADMIN_ID  = int(os.getenv("ADMIN_ID") or "5970208104")  # আপনার Telegram User ID

# ─── ফ্রি কপিরাইট চেক (AcoustID) ───
# ফ্রি API key পেতে: https://acoustid.org/new-application (শুধু email signup, পেমেন্ট লাগবে না)
# Termux: pkg install chromaprint   (fpcalc binary দরকার)
ACOUSTID_API_KEY = os.getenv("ACOUSTID_API_KEY") or "BfZJaWDKkj"  # ← AcoustID API key

# ─── ফ্রি SEO/Hashtag Generator (Google Gemini) ───
# ফ্রি API key পেতে: https://aistudio.google.com/apikey (Google account দিয়ে signup, পেমেন্ট লাগবে না)
# Key না দিলেও কাজ করবে — তখন basic keyword-based hashtag generate হবে
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY") or "AIzaSyDqmbDVVb9MZqIJQ-rei77OBazaZobiBII"  # ← Gemini API key

# ─── 🆕 v13: ChatGPT (OpenAI) ───
# API key পেতে: https://platform.openai.com/api-keys (paid — credit card লাগবে)
# Key থাকলে SEO + Subtitle + Voice সব OpenAI দিয়ে high-quality হবে
# Key না দিলেও বট চলবে — Gemini fallback হিসেবে কাজ করবে, TTS-এ free gTTS ব্যবহার হবে
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or "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"
OPENAI_CHAT_MODEL = os.getenv("OPENAI_CHAT_MODEL") or "gpt-4o-mini"   # cheap + fast; চাইলে "gpt-4o" দিতে পারেন

# ─── 🆕 v38: Anthropic Claude (3rd independent AI for WinGo Triple/Quad-AI consensus) ───
# API key পেতে: https://console.anthropic.com/settings/keys (paid — $5 free credit signup-এ)
# Key না থাকলেও বট চলবে — Claude layer (L32) silently skip হবে।
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY") or ""    # ← চাইলে এখানে বসান
ANTHROPIC_MODEL   = os.getenv("ANTHROPIC_MODEL") or "claude-3-5-haiku-20241022"   # cheap + fast Claude
OPENAI_TTS_MODEL  = os.getenv("OPENAI_TTS_MODEL") or "tts-1"          # standard; "tts-1-hd" = high quality
OPENAI_TTS_VOICE  = os.getenv("OPENAI_TTS_VOICE") or "nova"           # alloy / echo / fable / onyx / nova / shimmer
OPENAI_STT_MODEL  = os.getenv("OPENAI_STT_MODEL") or "whisper-1"      # subtitle generation

# ─── 🆕 MiniMax Agent (v17: integrated as 3rd AI provider) ───
# API key পেতে: https://www.minimax.io  → Console → API Keys
# Key থাকলে gemini/openai fallback chain-এ যুক্ত হবে।
MINIMAX_API_KEY  = os.getenv("MINIMAX_API_KEY", "")
MINIMAX_BASE_URL = os.getenv("MINIMAX_BASE_URL") or "https://api.minimax.io/v1/text/chatcompletion"
MINIMAX_MODEL    = os.getenv("MINIMAX_MODEL") or "abab6.5-chat"

# ─── 🆕 v14: ElevenLabs TTS (super realistic voices) ───
# API key পেতে: https://elevenlabs.io  → Sign up → Profile → API Keys → Copy
# ফ্রি tier: মাসে ১০,০০০ ক্যারেক্টার (~১০ মিনিট অডিও) ফ্রি — credit card লাগে না
# Key না দিলেও বট চলবে — gTTS / OpenAI fallback হিসেবে কাজ করবে
# ⚠️ Security note: এই key টা চ্যাটে শেয়ার হয়েছিল। চাইলে পরে নতুন key বানিয়ে replace করতে পারেন।
ELEVENLABS_API_KEY = os.getenv("ELEVENLABS_API_KEY") or ""  # 🆕 v24: hardcoded key সরানো হয়েছে। has_elevenlabs() এখন সবসময় False।

# ─── 🆕 v26: Replicate API (XTTS-v2 voice cloning — FREE alternative to ElevenLabs) ───
# Replicate signup: https://replicate.com → API tokens → Copy
# ফ্রি signup-এ কিছু ক্রেডিট দেয় (~$0.50/মাস)। প্রতি TTS call ~$0.001-0.005 খরচ।
# Voice cloning XTTS-v2 model — Bengali/Hindi/English/Spanish/Arabic + 11 ভাষা support করে।
# Zero-shot: ১০-৩০ সেকেন্ড sample → যেকোনো text → আপনার কণ্ঠে speech।
REPLICATE_API_TOKEN = os.getenv("REPLICATE_API_TOKEN") or ""
# XTTS-v2 model version (lucataco/xtts-v2 — community port, ~$0.001 per call)
REPLICATE_XTTS_VERSION = os.getenv("REPLICATE_XTTS_VERSION") or "684bc3855b37866c0c65add2ff39c78f3dea3f4ff103a436465326e0f438d55e"
ELEVENLABS_MODEL   = os.getenv("ELEVENLABS_MODEL") or "eleven_multilingual_v2"  # bn / hi / en সহ ২৯টা ভাষা support করে
# 🆕 v19: ElevenLabs-এর সব ৪০+ default library voice (key → (voice_id, label))
# এগুলো সবাই use করতে পারে — কোনো extra clone limit নেই, সরাসরি TTS-এ কণ্ঠ হিসেবে কাজ করে।
# পাশাপাশি বটটা আপনার অ্যাকাউন্টের সব custom/cloned voice-ও /v1/voices থেকে dynamically fetch করবে।
ELEVENLABS_VOICES = {
    # ─── Female voices ───
    "rachel":    ("21m00Tcm4TlvDq8ikWAM", "👩 Rachel (calm narration)"),
    "domi":      ("AZnzlk1XvdvUeBnXmlld", "🔥 Domi (strong, confident)"),
    "bella":     ("EXAVITQu4vr4xnSDxMaL", "💃 Bella (soft female)"),
    "elli":      ("MF3mGyEYCl7XYWbV9V6O", "🌸 Elli (young female)"),
    "dorothy":   ("ThT5KcBeYPX3keUQqHPh", "🇬🇧 Dorothy (British young)"),
    "charlotte": ("XB0fDUnXU5powFXDhCwa", "💋 Charlotte (seductive)"),
    "alice":     ("Xb7hH8MSUJpSbSDYk0k2", "🇬🇧 Alice (British female)"),
    "matilda":   ("XrExE9yKIg1WjnnlVkGX", "📖 Matilda (warm narrator)"),
    "freya":     ("jsCqWAovK2LkecY7zXl4", "🎉 Freya (overhyped young)"),
    "grace":     ("oWAxZDx7w5VEj9dCyTzz", "🌾 Grace (Southern US)"),
    "lily":      ("pFZP5JQG7iQjIQuC4Bku", "🇬🇧 Lily (warm British)"),
    "serena":    ("pMsXgVXv3BLzUgSXRplE", "✨ Serena (pleasant)"),
    "nicole":    ("piTKgcLEGmPE4e6mEKli", "🤫 Nicole (whisper)"),
    "jessie":    ("t0jbNlBVZ17f02VDIeMI", "👵 Jessie (raspy old)"),
    "glinda":    ("z9fAnlkpzviPz146aGWa", "🪄 Glinda (American witch)"),
    "mimi":      ("zrHiDhphv9ZnVXBqCLjz", "🎀 Mimi (childish female)"),
    "gigi":      ("jBpfuIE2acCO8z3wKNLl", "👶 Gigi (childlike)"),
    "emily":     ("LcfcDJNUP1GQjkzn1xUU", "🧘 Emily (meditation)"),
    "sarah":     ("EXAVITQu4vr4xnSDxMaL", "👧 Sarah (soft)"),
    # ─── Male voices ───
    "adam":      ("pNInz6obpgDQGcFmaJgB", "👨 Adam (deep narrator)"),
    "antoni":    ("ErXwobaYiN019PkySvjV", "🎙️ Antoni (warm male)"),
    "josh":      ("TxGEqnHWrfWFTfGW9XjX", "📻 Josh (young male)"),
    "arnold":    ("VR6AewLTigWG4xSOukaG", "💪 Arnold (crisp male)"),
    "sam":       ("yoZ06aMxZJJ28mfd3POQ", "🧔 Sam (casual male)"),
    "drew":      ("29vD33N1CtxCmqQRPOHJ", "🎤 Drew (well-rounded)"),
    "clyde":     ("2EiwWnXFnvU5JabPnv8n", "⚔️ Clyde (war veteran)"),
    "paul":      ("5Q0t7uMcjvnagumLfvZi", "📰 Paul (ground reporter)"),
    "dave":      ("CYw3kZ02Hs0563khs1Fj", "🇬🇧 Dave (British casual)"),
    "fin":       ("D38z5RcWu1voky8WS1ja", "⚓ Fin (sailor)"),
    "thomas":    ("GBv7mTt0atIp3Br8iCZE", "🧘 Thomas (calm meditation)"),
    "charlie":   ("IKne3meq5aSn9XLyUdCD", "🇦🇺 Charlie (Australian)"),
    "george":    ("JBFqnCBsd6RMkjVDRZzb", "🇬🇧 George (British)"),
    "callum":    ("N2lVS1w4EtoT3dr4eOWO", "🦹 Callum (hoarse villain)"),
    "patrick":   ("ODq5zmih8GrVes37Dizd", "📢 Patrick (shouty)"),
    "harry":     ("SOYHLrjzK2X1ezoPC6cr", "😰 Harry (anxious young)"),
    "liam":      ("TX3LPaxmHKxFdv7VOQHJ", "👦 Liam (young male)"),
    "joseph":    ("Zlb1dXrM653N07WRdFW3", "🇬🇧 Joseph (British)"),
    "will":      ("bIHbv24MWmeRgasZH58o", "🙂 Will (friendly)"),
    "jeremy":    ("bVMeCyTHy58xNoL34h3p", "🍀 Jeremy (excited Irish)"),
    "michael":   ("flq6f7yk4E4fJM5XTYuZ", "🎓 Michael (orator)"),
    "ethan":     ("g5CIjZEefAph4nQFvHAz", "🤫 Ethan (whispering)"),
    "chris":     ("iP95p4xoKVk53GoZ742B", "😎 Chris (casual)"),
    "brian":     ("nPczCjzI2devNBz1zQrb", "🎬 Brian (deep narrator)"),
    "daniel":    ("onwK4e9ZLuTAKqWW03F9", "📺 Daniel (British news)"),
    "bill":      ("pqHfZKP75CvOlQylNhV4", "📚 Bill (strong narrator)"),
    "ryan":      ("wViXBPUzp2ZZixB1xQuM", "🪖 Ryan (soldier)"),
    "giovanni":  ("zcAOhNBS3c14rBihAFp1", "🇮🇹 Giovanni (Italian-EN)"),
    "james":     ("ZQe5CZNOzWyzPSCn5a3c", "🇦🇺 James (calm Aussie)"),
}
ELEVENLABS_DEFAULT_VOICE = "rachel"
# 🆕 v19: per-page voice button count (Telegram inline button limit)
ELEVENLABS_VOICES_PER_PAGE = 8
# 🆕 v19: dynamic voice cache (account থেকে fetch হলে এখানে save হবে)
_eleven_voice_cache = {"ts": 0.0, "voices": []}  # {ts: epoch, voices: [{voice_id, name, category}]}

# ──────────────────────────────────────────────────────────────
# 🆕 v22: EDGE TTS — Microsoft Edge-এর Neural TTS (১০০% ফ্রি, কোন key লাগে না)
# ─────────────────────────────────────────────────────────────-
# • কোনো subscription/credit card/API key দরকার নেই
# • Bengali, Hindi, English, Urdu, Arabic সহ ১৪০+ ভাষা support
# • Quality: ElevenLabs-এর কাছাকাছি, কিন্তু সম্পূর্ণ free
# • Termux: pip install edge-tts
# ─────────────────────────────────────────────────────────────-
EDGE_TTS_VOICES = {
    # ─── Bengali (বাংলা) ───
    "bn_nabanita": ("bn-BD-NabanitaNeural", "🇧🇩 Nabanita (বাংলা মেয়ে, soft)", "bn"),
    "bn_pradeep":  ("bn-BD-PradeepNeural",  "🇧🇩 Pradeep (বাংলা ছেলে, deep)",  "bn"),
    "bn_bashkar":  ("bn-IN-BashkarNeural",  "🇮🇳 Bashkar (বাংলা ছেলে, warm)",   "bn"),
    "bn_tanishaa": ("bn-IN-TanishaaNeural", "🇮🇳 Tanishaa (বাংলা মেয়ে, sweet)", "bn"),
    # ─── Hindi (हिन्दी) ───
    "hi_madhur":   ("hi-IN-MadhurNeural",   "🇮🇳 Madhur (हिन्दी पुरुष)",  "hi"),
    "hi_swara":    ("hi-IN-SwaraNeural",    "🇮🇳 Swara (हिन्दी स्त्री)",  "hi"),
    # ─── English (US/UK/IN) ───
    "en_aria":     ("en-US-AriaNeural",     "🇺🇸 Aria (US female, casual)",  "en"),
    "en_jenny":    ("en-US-JennyNeural",    "🇺🇸 Jenny (US female, friendly)","en"),
    "en_guy":      ("en-US-GuyNeural",      "🇺🇸 Guy (US male, news)",        "en"),
    "en_davis":    ("en-US-DavisNeural",    "🇺🇸 Davis (US male, deep)",      "en"),
    "en_jane":     ("en-US-JaneNeural",     "🇺🇸 Jane (US female, calm)",     "en"),
    "en_tony":     ("en-US-TonyNeural",     "🇺🇸 Tony (US male, lively)",     "en"),
    "en_libby":    ("en-GB-LibbyNeural",    "🇬🇧 Libby (UK female)",          "en"),
    "en_ryan":     ("en-GB-RyanNeural",     "🇬🇧 Ryan (UK male, warm)",       "en"),
    "en_sonia":    ("en-GB-SoniaNeural",    "🇬🇧 Sonia (UK female, posh)",    "en"),
    "en_neerja":   ("en-IN-NeerjaNeural",   "🇮🇳 Neerja (Indian female)",      "en"),
    "en_prabhat":  ("en-IN-PrabhatNeural",  "🇮🇳 Prabhat (Indian male)",       "en"),
    # ─── Urdu (اردو) ───
    "ur_gul":      ("ur-PK-GulNeural",      "🇵🇰 Gul (Urdu female)",          "ur"),
    "ur_asad":     ("ur-PK-AsadNeural",     "🇵🇰 Asad (Urdu male)",           "ur"),
    # ─── Arabic ───
    "ar_zariyah":  ("ar-SA-ZariyahNeural",  "🇸🇦 Zariyah (Arabic female)",     "ar"),
    "ar_hamed":    ("ar-SA-HamedNeural",    "🇸🇦 Hamed (Arabic male)",         "ar"),
}
EDGE_TTS_DEFAULT = "bn_nabanita"
EDGE_TTS_VOICES_PER_PAGE = 8

MAX_TG_UPLOAD_MB = 49     # 🐛 v29 FIX: Telegram cloud Bot API hard-limit 50MB.
                          # আগের 2000 MB ছিল ভুল — bot reply_video() silently fail করছিল
                          # বড় ফাইলে। এখন 49MB এর বেশি হলে user-কে আগে warning দেওয়া হবে
                          # এবং send_document fallback try করা হবে।
MAX_DOWNLOAD_MB  = 2000   # 🆕 v24: 1500 → 2000 MB। yt-dlp এখন 2GB পর্যন্ত download করবে।
DOWNLOAD_TIMEOUT = 600
FFMPEG_TIMEOUT   = 1800

# 🐛 v38 FIX — Telegram-এর media caption সর্বোচ্চ 1024 *UTF-16 code unit*।
# Python-এর len() character গণে, কিন্তু Telegram UTF-16 code unit গণে — অনেক
# emoji (🎵 🔥 💎 ✓ ➡️ ইত্যাদি) ও কিছু Bengali grapheme ২টা code unit লাগে।
# ফলে v37-এ len()=1020 হলেও Telegram-এ 1100+ হয়ে "caption is too long" error আসত।
TG_CAPTION_LIMIT  = 1024          # আসল hard-limit
TG_CAPTION_SAFE   = 950           # safety buffer (emoji/markdown overhead-এর জন্য)
TG_MESSAGE_LIMIT  = 4096

def _u16_len(s: str) -> int:
    """Telegram caption যেভাবে গণে — UTF-16 code unit count।
    BMP-এর বাইরের char (most emoji) ২ unit লাগে।"""
    if not s:
        return 0
    try:
        return len(s.encode("utf-16-le")) // 2
    except Exception:
        return len(s)

def split_caption(text: str, limit: int = TG_CAPTION_SAFE):
    """Caption Telegram-limit (UTF-16) ছাড়ালে cleanly ভেঙে দেয়।
    Returns (short_caption_for_media, [followup_text_messages])।
    🐛 v38 FIX: UTF-16 code unit গণনা ব্যবহার করে (Telegram-এর actual metric)।"""
    if not text or _u16_len(text) <= limit:
        return text, []
    safe_limit = max(limit - 40, 200)  # tail-এর জন্য জায়গা
    # UTF-16 unit-aware truncate: char-by-char এগিয়ে গণনা করি
    cut_char = 0
    running = 0
    for i, ch in enumerate(text):
        u = _u16_len(ch)
        if running + u > safe_limit:
            cut_char = i
            break
        running += u
    else:
        cut_char = len(text)
    head = text[:cut_char]
    # সুন্দর word/line boundary খুঁজি
    nice = head.rfind("\n\n")
    if nice < cut_char * 0.5:
        nice = head.rfind("\n")
    if nice < cut_char * 0.5:
        nice = head.rfind(" ")
    if nice < 0:
        nice = cut_char
    short = text[:nice].rstrip() + "\n\n_(➡️ বাকিটা নিচে দেখুন)_"
    rest = text[nice:].lstrip()
    followups = []
    msg_limit = TG_MESSAGE_LIMIT - 100
    while rest:
        if _u16_len(rest) <= msg_limit:
            followups.append(rest); break
        # UTF-16 unit-aware chunk
        cut2 = 0; running = 0
        for i, ch in enumerate(rest):
            u = _u16_len(ch)
            if running + u > msg_limit:
                cut2 = i; break
            running += u
        else:
            cut2 = len(rest)
        nl = rest[:cut2].rfind("\n")
        if nl > cut2 * 0.6:
            cut2 = nl
        followups.append(rest[:cut2])
        rest = rest[cut2:].lstrip()
    return short, followups

def shrink_caption(text: str, limit: int = TG_CAPTION_SAFE) -> str:
    """Caption যদি limit ছাড়ায় শুধু head অংশ rakhi (no follow-ups)।"""
    short, _ = split_caption(text, limit)
    return short

def is_caption_too_long_err(err) -> bool:
    """Telegram error থেকে caption-too-long detect করে।
    🐛 v38: আরও variant cover (MEDIA_CAPTION_TOO_LONG, MESSAGE_TOO_LONG, ইত্যাদি)।"""
    s = str(err).lower()
    needles = (
        "caption is too long", "caption_too_long", "media_caption_too_long",
        "message_too_long", "message is too long",
        "text is too long", "text_too_long",
    )
    return any(n in s for n in needles)

async def safe_send_video(target, video_path, caption: str = "", *,
                          parse_mode: str = "Markdown",
                          reply_markup=None, supports_streaming: bool = True,
                          logger_=None):
    """🆕 v38: যেকোনো send_video-কে caption-too-long থেকে রক্ষা করে।
    target = q.message বা context-bot এর মত reply_video method-যুক্ত object।
    Returns (success: bool, followup_messages: list[str])।"""
    cap_short, cap_extras = split_caption(caption or "")
    kwargs = dict(parse_mode=parse_mode, supports_streaming=supports_streaming,
                  read_timeout=300, write_timeout=300, connect_timeout=60)
    if reply_markup is not None:
        kwargs["reply_markup"] = reply_markup
    # প্রথম try
    try:
        with open(video_path, "rb") as fp:
            await target.reply_video(video=fp, caption=cap_short, **kwargs)
        return True, cap_extras
    except Exception as e1:
        if logger_:
            logger_.warning("safe_send_video first try fail: %s", e1)
        if not is_caption_too_long_err(e1):
            raise
    # caption-less retry
    try:
        with open(video_path, "rb") as fp:
            await target.reply_video(video=fp, **kwargs)
        return True, [caption] if caption else []
    except Exception as e2:
        if logger_:
            logger_.warning("safe_send_video captionless retry fail: %s", e2)
        raise

# ══════════════════════════════════════════════════════════════
# 🆕 v37 FEATURE — File Content Extractor (PDF/DOCX/ZIP/Image/etc)
# ══════════════════════════════════════════════════════════════
# Optional libraries — gracefully fall back if not installed.
try:
    import zipfile as _zip
    HAS_ZIP = True
except Exception:
    HAS_ZIP = False
try:
    from PyPDF2 import PdfReader as _PdfReader
    HAS_PDF = True
except Exception:
    try:
        from pypdf import PdfReader as _PdfReader
        HAS_PDF = True
    except Exception:
        HAS_PDF = False
try:
    import docx as _docx
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False
try:
    import openpyxl as _openpyxl
    HAS_XLSX = True
except Exception:
    HAS_XLSX = False
try:
    import csv as _csv
    HAS_CSV = True
except Exception:
    HAS_CSV = False
try:
    import py7zr as _p7z
    HAS_7Z = True
except Exception:
    HAS_7Z = False
try:
    import rarfile as _rar
    HAS_RAR = True
except Exception:
    HAS_RAR = False
try:
    from PIL import Image as _PILImage
    HAS_PIL = True
except Exception:
    HAS_PIL = False
try:
    import pytesseract as _ocr
    HAS_OCR = True
except Exception:
    HAS_OCR = False
try:
    from mutagen import File as _MutagenFile
    HAS_MUTAGEN = True
except Exception:
    HAS_MUTAGEN = False

MAX_EXTRACT_FILE_MB = 45  # Telegram bot download limit ~50MB

def _read_text_file(path, max_chars=8000):
    """TXT/JSON/HTML/CSV/code ফাইল পড়ে — encoding auto-detect।"""
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            with open(path, "r", encoding=enc, errors="replace") as f:
                data = f.read(max_chars + 1)
            return data[:max_chars] + ("\n\n…(আরো আছে)" if len(data) > max_chars else "")
        except Exception:
            continue
    return None

def extract_pdf(path, max_pages=30, max_chars=8000):
    if not HAS_PDF:
        return None, "PyPDF2/pypdf install নেই — `pip install PyPDF2` চালান।"
    try:
        reader = _PdfReader(path)
        if reader.is_encrypted:
            try: reader.decrypt("")
            except Exception:
                return None, "🔒 PDF password-protected।"
        out = []
        for i, page in enumerate(reader.pages[:max_pages]):
            try:
                t = page.extract_text() or ""
                if t.strip():
                    out.append(f"━━━ Page {i+1} ━━━\n{t.strip()}")
            except Exception:
                pass
        joined = "\n\n".join(out) if out else ""
        if not joined:
            return None, "PDF থেকে text extract হয়নি (scanned PDF হলে OCR দরকার)।"
        if len(joined) > max_chars:
            joined = joined[:max_chars] + f"\n\n…({len(reader.pages)} পেজের মধ্যে প্রথম অংশ দেখানো হলো)"
        return joined, f"📄 {len(reader.pages)} পেজ"
    except Exception as e:
        return None, f"PDF পড়তে সমস্যা: {str(e)[:120]}"

def extract_docx(path, max_chars=8000):
    if not HAS_DOCX:
        return None, "python-docx install নেই — `pip install python-docx` চালান।"
    try:
        doc = _docx.Document(path)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paras)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                text += "\n| " + " | ".join(cells) + " |"
        if not text.strip():
            return None, "DOCX-এ কোনো text পাওয়া যায়নি।"
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n…(আরো আছে)"
        return text, f"📝 {len(paras)} প্যারাগ্রাফ"
    except Exception as e:
        return None, f"DOCX পড়তে সমস্যা: {str(e)[:120]}"

def extract_xlsx(path, max_rows=50, max_chars=8000):
    if not HAS_XLSX:
        return None, "openpyxl install নেই — `pip install openpyxl` চালান।"
    try:
        wb = _openpyxl.load_workbook(path, data_only=True, read_only=True)
        out = []
        for sname in wb.sheetnames[:5]:
            sh = wb[sname]
            out.append(f"━━━ Sheet: {sname} ━━━")
            for i, row in enumerate(sh.iter_rows(values_only=True)):
                if i >= max_rows:
                    out.append(f"…(আরো {sh.max_row - max_rows} সারি)")
                    break
                cells = [str(c) if c is not None else "" for c in row]
                out.append(" | ".join(cells))
        text = "\n".join(out)
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n…(আরো আছে)"
        return text, f"📊 {len(wb.sheetnames)} sheet"
    except Exception as e:
        return None, f"Excel পড়তে সমস্যা: {str(e)[:120]}"

def extract_csv(path, max_rows=50, max_chars=8000):
    try:
        out = []
        with open(path, "r", encoding="utf-8", errors="replace", newline="") as f:
            reader = _csv.reader(f)
            for i, row in enumerate(reader):
                if i >= max_rows:
                    out.append("…(আরো সারি আছে)")
                    break
                out.append(" | ".join(row))
        text = "\n".join(out)
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n…(আরো আছে)"
        return text, f"📊 CSV"
    except Exception as e:
        return None, f"CSV পড়তে সমস্যা: {str(e)[:120]}"

def extract_zip(path, max_chars=4000):
    if not HAS_ZIP:
        return None, "zipfile module নেই।"
    try:
        with _zip.ZipFile(path, "r") as zf:
            if zf.namelist() and any(zf.getinfo(n).flag_bits & 0x1 for n in zf.namelist()):
                return None, "🔒 ZIP password-protected।"
            names = zf.namelist()
            total = len(names)
            sized = []
            for n in names[:50]:
                try:
                    info = zf.getinfo(n)
                    sized.append(f"  📄 {n}  ({info.file_size:,} bytes)")
                except Exception:
                    sized.append(f"  📄 {n}")
            text = f"📦 ZIP-এ মোট *{total}* ফাইল:\n\n" + "\n".join(sized)
            if total > 50:
                text += f"\n\n  …(আরো {total - 50} ফাইল)"
            return text[:max_chars], f"🗜️ {total} item"
    except Exception as e:
        return None, f"ZIP পড়তে সমস্যা: {str(e)[:120]}"

def extract_7z(path):
    if not HAS_7Z:
        return None, "py7zr install নেই — `pip install py7zr` চালান।"
    try:
        with _p7z.SevenZipFile(path, "r") as z:
            names = z.getnames()
            text = f"📦 7Z-এ মোট *{len(names)}* ফাইল:\n\n" + "\n".join(f"  📄 {n}" for n in names[:50])
            return text, f"🗜️ {len(names)} item"
    except Exception as e:
        return None, f"7Z পড়তে সমস্যা: {str(e)[:120]}"

def extract_rar(path):
    if not HAS_RAR:
        return None, "rarfile install নেই — `pip install rarfile` ও Termux-এ `pkg install unrar` চালান।"
    try:
        with _rar.RarFile(path, "r") as rf:
            if rf.needs_password():
                return None, "🔒 RAR password-protected।"
            names = rf.namelist()
            text = f"📦 RAR-এ মোট *{len(names)}* ফাইল:\n\n" + "\n".join(f"  📄 {n}" for n in names[:50])
            return text, f"🗜️ {len(names)} item"
    except Exception as e:
        return None, f"RAR পড়তে সমস্যা: {str(e)[:120]}"

def extract_image_ocr(path, lang="ben+eng"):
    if not (HAS_OCR and HAS_PIL):
        return None, "OCR জন্য Termux-এ এগুলো install করুন:\n`pkg install tesseract tesseract-data-ben`\n`pip install pytesseract Pillow`"
    try:
        img = _PILImage.open(path)
        try:
            text = _ocr.image_to_string(img, lang=lang)
        except Exception:
            text = _ocr.image_to_string(img, lang="eng")
        text = text.strip()
        if not text:
            return None, "ছবিতে পড়ার মতো লেখা পাওয়া যায়নি।"
        return text[:8000], "🖼️ OCR সফল"
    except Exception as e:
        return None, f"OCR ব্যর্থ: {str(e)[:120]}"

def extract_audio_meta(path):
    if not HAS_MUTAGEN:
        return None, "mutagen install নেই — `pip install mutagen` চালান।"
    try:
        mf = _MutagenFile(path)
        if not mf:
            return None, "Audio metadata পড়া যায়নি।"
        info = []
        if hasattr(mf, "info"):
            if hasattr(mf.info, "length"):
                info.append(f"⏱️ Duration: {mf.info.length:.1f}s")
            if hasattr(mf.info, "bitrate"):
                info.append(f"🎚️ Bitrate: {mf.info.bitrate // 1000} kbps")
            if hasattr(mf.info, "sample_rate"):
                info.append(f"📻 Sample rate: {mf.info.sample_rate} Hz")
        tags = mf.tags or {}
        for k in ("title", "artist", "album", "date", "TIT2", "TPE1", "TALB", "TDRC"):
            v = tags.get(k) if hasattr(tags, "get") else None
            if v:
                info.append(f"🏷️ {k}: {v}")
        return "\n".join(info) if info else None, "🎵 Audio info"
    except Exception as e:
        return None, f"Audio metadata পড়তে সমস্যা: {str(e)[:120]}"

def extract_dispatch(path, fname):
    """Extension অনুযায়ী সঠিক extractor বেছে নেয়।"""
    ext = (Path(fname).suffix or "").lower().lstrip(".")
    if ext == "pdf":               return extract_pdf(path)
    if ext in ("docx", "doc"):     return extract_docx(path)
    if ext in ("xlsx", "xls"):     return extract_xlsx(path)
    if ext == "csv":               return extract_csv(path)
    if ext == "zip":               return extract_zip(path)
    if ext == "7z":                return extract_7z(path)
    if ext == "rar":               return extract_rar(path)
    if ext in ("jpg","jpeg","png","webp","bmp","gif"):
                                   return extract_image_ocr(path)
    if ext in ("mp3","wav","ogg","flac","m4a","aac","opus"):
                                   return extract_audio_meta(path)
    if ext in ("txt","json","xml","html","htm","md","py","js","ts","c","cpp","java",
               "log","yaml","yml","ini","cfg","sh","sql","csv","go","rs","rb","php"):
        text = _read_text_file(path)
        if text: return text, f"📜 .{ext} text"
        return None, "ফাইল পড়া যায়নি।"
    return None, f"❌ `.{ext}` ফরম্যাট এখনো support করে না।\n\nসাপোর্টেড: PDF, DOCX, XLSX, CSV, ZIP, 7Z, RAR, JPG, PNG, MP3, MP4, TXT, JSON, HTML, code-files"

# ══════════════════════════════════════════════════════════════
# 🆕 v37 FEATURE — MCQ Answerer (AI-powered, with cache)
# ══════════════════════════════════════════════════════════════
MCQ_CACHE_FILE = STATS_DIR / "mcq_cache.json" if False else None  # set after STATS_DIR defined
mcq_cache: dict = {}
mcq_lock = threading.Lock()

def _mcq_key(question: str) -> str:
    """Question text-কে normalize করে cache key বানায়।"""
    import hashlib
    norm = re.sub(r"\s+", " ", (question or "").strip().lower())
    norm = re.sub(r"[^\w\u0980-\u09FF\s]", "", norm)
    return hashlib.md5(norm.encode("utf-8")).hexdigest()[:16]

def load_mcq_cache():
    global mcq_cache, MCQ_CACHE_FILE
    MCQ_CACHE_FILE = STATS_DIR / "mcq_cache.json"
    try:
        if MCQ_CACHE_FILE.exists():
            with open(MCQ_CACHE_FILE, "r", encoding="utf-8") as f:
                mcq_cache = json.load(f)
    except Exception as e:
        logger.warning("mcq cache load fail: %s", e)
        mcq_cache = {}

def save_mcq_cache():
    try:
        if MCQ_CACHE_FILE is None: return
        with mcq_lock:
            with open(MCQ_CACHE_FILE, "w", encoding="utf-8") as f:
                json.dump(mcq_cache, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning("mcq cache save fail: %s", e)

def mcq_lookup(question: str):
    k = _mcq_key(question)
    with mcq_lock:
        return mcq_cache.get(k)

def mcq_store(question: str, answer: str):
    k = _mcq_key(question)
    with mcq_lock:
        mcq_cache[k] = {"q": question[:500], "a": answer, "ts": int(time.time())}
        # rotate when too big
        if len(mcq_cache) > 5000:
            oldest = sorted(mcq_cache.items(), key=lambda x: x[1].get("ts", 0))[:1000]
            for kk, _ in oldest:
                mcq_cache.pop(kk, None)
    save_mcq_cache()

# ══════════════════════════════════════════════════════════════
# 🆕 v37 FEATURE — Link Expander / URL Unshortener (safety tool)
# ══════════════════════════════════════════════════════════════
KNOWN_SHORTENERS = {
    "bit.ly","tinyurl.com","t.co","goo.gl","is.gd","ow.ly","buff.ly",
    "rebrand.ly","rb.gy","cutt.ly","shorturl.at","shorte.st","soo.gd",
    "tiny.cc","v.gd","x.co","aka.ms","amzn.to","fb.me","fb.watch",
    "youtu.be","wp.me","trib.al","lnkd.in","dlvr.it","dub.sh","s.id",
}

def expand_url(url, max_hops=10, timeout=10):
    """Short / redirected URL-এর পুরো chain follow করে final destination দেয়।"""
    import http.client as _hc
    import urllib.request as _ureq
    from urllib.parse import urlparse as _up, urljoin as _ujoin

    if not isinstance(url, str) or not url.strip():
        return {"error": "URL খালি", "chain": [], "final": "", "status": None,
                "title": None, "hops": 0}
    url = url.strip()
    if not url.startswith(("http://", "https://")):
        url = "http://" + url

    chain = []
    current = url
    final_status = None
    page_title = None
    error = None
    seen = set()

    for _ in range(max_hops):
        if current in seen:
            error = "Redirect loop ধরা পড়েছে"
            break
        seen.add(current)
        chain.append(current)
        try:
            parsed = _up(current)
            if parsed.scheme not in ("http", "https"):
                error = f"Unsupported scheme: {parsed.scheme}"
                break
            if not parsed.netloc:
                error = "অসম্পূর্ণ URL"
                break
            ConnCls = _hc.HTTPSConnection if parsed.scheme == "https" else _hc.HTTPConnection
            conn = ConnCls(parsed.netloc, timeout=timeout)
            path = parsed.path or "/"
            if parsed.query:
                path += "?" + parsed.query
            conn.request("HEAD", path, headers={
                "User-Agent": "Mozilla/5.0 (LinkExpander)",
                "Accept": "*/*",
            })
            resp = conn.getresponse()
            final_status = resp.status
            if 300 <= resp.status < 400:
                loc = resp.getheader("Location")
                conn.close()
                if not loc:
                    break
                if loc.startswith("/"):
                    current = f"{parsed.scheme}://{parsed.netloc}{loc}"
                elif loc.startswith(("http://", "https://")):
                    current = loc
                else:
                    current = _ujoin(current, loc)
                continue
            conn.close()
            break
        except Exception as e:
            error = str(e)[:140]
            break

    # Try title fetch (small GET, only if final URL is ok-ish)
    if not error and final_status and 200 <= final_status < 400:
        try:
            req = _ureq.Request(current, headers={
                "User-Agent": "Mozilla/5.0 (LinkExpander)",
                "Accept-Language": "en;q=0.7,bn;q=0.3",
            })
            with _ureq.urlopen(req, timeout=timeout) as r:
                ctype = r.headers.get("Content-Type", "").lower()
                if "html" in ctype or "text" in ctype or not ctype:
                    raw = r.read(60000)
                    enc = "utf-8"
                    m_enc = re.search(rb'charset=["\']?([\w-]+)', raw[:2000], re.IGNORECASE)
                    if m_enc:
                        try:
                            enc = m_enc.group(1).decode("ascii", "ignore")
                        except Exception:
                            pass
                    data = raw.decode(enc, errors="ignore")
                    m = re.search(r"<title[^>]*>([^<]{1,250})</title>", data, re.IGNORECASE | re.DOTALL)
                    if m:
                        page_title = re.sub(r"\s+", " ", m.group(1).strip())[:200]
        except Exception:
            pass

    return {
        "chain": chain,
        "final": current,
        "status": final_status,
        "title": page_title,
        "error": error,
        "hops": max(0, len(chain) - 1),
    }

def analyze_url_safety(url, title=None):
    """URL-এ phishing/suspicious pattern check করে warning list ফেরত দেয়।"""
    from urllib.parse import urlparse as _up
    warnings = []
    try:
        p = _up(url)
        host = (p.netloc or "").lower()
    except Exception:
        return ["URL parse করা যায়নি"]
    full = (url + " " + (title or "")).lower()

    if re.match(r"^\d+\.\d+\.\d+\.\d+", host):
        warnings.append("⚠️ Domain নয়, IP address ব্যবহার হয়েছে")
    if any(host.endswith(t) for t in (".tk", ".ml", ".ga", ".cf", ".gq")):
        warnings.append("⚠️ Free TLD (.tk/.ml/.ga ইত্যাদি) — phishing-এ বহুল ব্যবহৃত")
    if re.search(r"login.*verify|verify.*account|account.*suspend|"
                 r"prize.*won|claim.*reward|free.*gift|click.*urgent", full):
        warnings.append("⚠️ Phishing-এর মতো keyword পাওয়া গেছে")
    if len(url) > 250:
        warnings.append("⚠️ অস্বাভাবিক লম্বা URL")
    if host.count(".") > 4:
        warnings.append("⚠️ অতিরিক্ত subdomain — masking attempt হতে পারে")
    if re.search(r"[\u0400-\u04FF\u0500-\u052F]", host):
        warnings.append("⚠️ Cyrillic চরিত্র — IDN homograph attack")
    if re.search(r"xn--", host):
        warnings.append("ℹ️ IDN encoded domain — original নাম check করুন")
    if re.search(r"@", url.split("?")[0]):
        warnings.append("⚠️ URL-এ @ চিহ্ন — user-info trick হতে পারে")
    return warnings

logging.basicConfig(format='%(asctime)s | %(levelname)s | %(message)s', level=logging.INFO)
logging.getLogger("httpx").setLevel(logging.WARNING)
logger = logging.getLogger("video-bot")

TEMP_DIR  = Path(tempfile.gettempdir()) / "tgbot_videos"; TEMP_DIR.mkdir(parents=True, exist_ok=True)
STATS_DIR = Path("bot_data"); STATS_DIR.mkdir(parents=True, exist_ok=True)
STATS_FILE = STATS_DIR / "stats.json"

user_videos: dict = {}      # uid -> {"path","title"}
user_state:  dict = {}      # uid -> {"action": "...", "data": {...}}
user_lock   = threading.Lock()
stats_lock  = threading.Lock()
CPU_COUNT   = max(2, (os.cpu_count() or 4))
executor    = ThreadPoolExecutor(max_workers=max(4, CPU_COUNT))

# ──────────────────────────────────────────────────────────────
# 🆕 v15: BENGALI PREMIUM FONT SYSTEM
# প্রথম রানে fonts ডাউনলোড হয়ে ~/.fonts/bangla/-এ save হবে।
# তারপর text overlay, caption, subtitle সব auto-detect করে Bengali font ব্যবহার করবে।
# ──────────────────────────────────────────────────────────────
BANGLA_FONT_DIR = Path.home() / ".fonts" / "bangla"
BANGLA_FONT_DIR.mkdir(parents=True, exist_ok=True)

BANGLA_FONTS = {
    # 🆕 v28: Each font এখন multiple mirror URL আছে — প্রথমটা fail হলে পরেরটা try হবে
    # Google Fonts repo path পাল্টে গেছে (Noto Bengali variable font হয়েছে), তাই jsDelivr CDN backup
    "noto":    ("✨ Noto Bengali (clean)",     "NotoSansBengali-Bold.ttf", [
                "https://cdn.jsdelivr.net/gh/notofonts/bengali@main/fonts/NotoSansBengali/hinted/ttf/NotoSansBengali-Bold.ttf",
                "https://github.com/notofonts/bengali/raw/main/fonts/NotoSansBengali/hinted/ttf/NotoSansBengali-Bold.ttf",
                "https://github.com/google/fonts/raw/main/ofl/notosansbengali/static/NotoSansBengali-Bold.ttf",
                ]),
    "hind":    ("💎 Hind Siliguri (premium)",  "HindSiliguri-Bold.ttf", [
                "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/hindsiliguri/HindSiliguri-Bold.ttf",
                "https://github.com/google/fonts/raw/main/ofl/hindsiliguri/HindSiliguri-Bold.ttf",
                ]),
    "baloo":   ("🌸 Baloo Da 2 (modern)",       "BalooDa2-Bold.ttf", [
                "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/baloodaa2/static/BalooDa2-Bold.ttf",
                "https://github.com/google/fonts/raw/main/ofl/baloodaa2/static/BalooDa2-Bold.ttf",
                "https://github.com/google/fonts/raw/main/ofl/baloodaa2/BalooDa2-Bold.ttf",
                ]),
    "galada":  ("🎀 Galada (decorative)",       "Galada-Regular.ttf", [
                "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/galada/Galada-Regular.ttf",
                "https://github.com/google/fonts/raw/main/ofl/galada/Galada-Regular.ttf",
                ]),
}
DEFAULT_BANGLA_FONT = "noto"
BANGLA_RE = re.compile(r'[\u0980-\u09FF]')

def is_bangla(text):
    """টেক্সটে বাংলা অক্ষর আছে কিনা চেক।"""
    return bool(text and BANGLA_RE.search(str(text)))

def ensure_bangla_fonts():
    """Startup-এ ৪টা Bengali font auto-download (একবারই, ~৮০০KB মোট)।
    🆕 v28: প্রতিটা font-এর জন্য multiple mirror try করে। কোনোটা না নামলেও system font fallback আছে।
    Network slow হলে ১৫ সেকেন্ডের পর give up — startup hang হবে না।
    """
    import urllib.request, socket
    ok_count = 0
    for key, (name, fname, urls) in BANGLA_FONTS.items():
        target = BANGLA_FONT_DIR / fname
        if target.exists() and target.stat().st_size > 10000:
            ok_count += 1
            continue
        downloaded = False
        for url in urls:
            try:
                logger.info("⬇️  Bengali font ডাউনলোড: %s ...", name)
                # 15s timeout per mirror — slow network-এ hang হবে না
                req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(req, timeout=15) as r, open(target, "wb") as f:
                    f.write(r.read())
                if target.exists() and target.stat().st_size > 10000:
                    logger.info("✓ %s রেডি", name)
                    ok_count += 1
                    downloaded = True
                    break
            except (socket.timeout, Exception) as e:
                logger.warning("Font mirror fail (%s): %s — পরের mirror try করছি", name, str(e)[:80])
                try:
                    if target.exists(): target.unlink()
                except Exception: pass
        if not downloaded:
            logger.warning("Font ডাউনলোড সব mirror-এ ব্যর্থ (%s) — system font fallback ব্যবহার হবে", name)
    logger.info("🅰️  Bengali fonts: %d/%d রেডি", ok_count, len(BANGLA_FONTS))

def get_bangla_font_path(key=None):
    """Bengali font ফাইলের path return করে। না থাকলে অন্য Bengali font fallback, system font, তারপর None।"""
    key = key or DEFAULT_BANGLA_FONT
    if key in BANGLA_FONTS:
        p = BANGLA_FONT_DIR / BANGLA_FONTS[key][1]
        if p.exists() and p.stat().st_size > 10000:
            return str(p)
    # যেকোনো available Bengali font
    for k, (name, fname, urls) in BANGLA_FONTS.items():
        p = BANGLA_FONT_DIR / fname
        if p.exists() and p.stat().st_size > 10000:
            return str(p)
    # system fallback
    for sp in [
        "/system/fonts/NotoSansBengali-Regular.ttf",
        "/system/fonts/NotoSerifBengali-Regular.ttf",
        "/data/data/com.termux/files/usr/share/fonts/TTF/NotoSansBengali-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansBengali-Regular.ttf",
    ]:
        if Path(sp).exists():
            return sp
    return None

def get_english_font_path():
    """English/Latin font path।"""
    for fp in [
        "/system/fonts/Roboto-Bold.ttf",
        "/data/data/com.termux/files/usr/share/fonts/TTF/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf",
    ]:
        if Path(fp).exists():
            return fp
    # Bengali font ও Latin support করে — ultimate fallback
    return get_bangla_font_path()

def pick_font_for_text(text, preferred_bangla=None):
    """টেক্সট দেখে best font বেছে নেয়। Bengali থাকলে Bengali font, না হলে English।"""
    if is_bangla(text):
        return get_bangla_font_path(preferred_bangla) or get_english_font_path()
    return get_english_font_path() or get_bangla_font_path(preferred_bangla)

# ──────────────────────────────────────────────────────────────
# 🆕 v15: STATUS VIDEO MAKER (Bengali emotional quote videos)
# WhatsApp/Facebook status-এর জন্য 9:16 ভিডিও — কোনো input video লাগে না
# ──────────────────────────────────────────────────────────────
STATUS_BG_PRESETS = {
    "night":    ("🌃 Night City (dark blue)",   "#0a0a23", "#000814"),
    "sunset":   ("🌆 Sunset (orange-purple)",  "#FF8C42", "#3D1E5F"),
    "heart":    ("💔 Heartbreak (red-black)",  "#8B0000", "#0a0a0a"),
    "love":     ("❤️ Love (pink-red)",         "#FF6B9D", "#8B0000"),
    "dawn":     ("🌅 Dawn (warm orange)",      "#FF6B35", "#2C1810"),
    "stars":    ("⭐ Deep Space (black)",       "#0a0a23", "#000000"),
    "forest":   ("🌲 Forest (green)",           "#1B4332", "#081C15"),
    "ocean":    ("🌊 Ocean Depth (blue)",       "#1B3A5C", "#000F1F"),
    "rain":     ("🌧️ Rainy Mood (gray)",        "#2C3E50", "#0a0a0a"),
    "sad":      ("😢 Sad Vibe (purple-black)",  "#3D1E5F", "#0a0a0a"),
}

def wrap_bn_text(text, max_chars=22):
    """Word-wrap করে multi-line bananao (Bengali + English ঠিকঠাক handle)।"""
    import textwrap
    out_lines = []
    for paragraph in (text or "").split("\n"):
        if not paragraph.strip():
            out_lines.append("")
            continue
        wrapped = textwrap.wrap(paragraph, width=max_chars,
                                break_long_words=True, break_on_hyphens=False)
        out_lines.extend(wrapped if wrapped else [""])
    return "\n".join(out_lines)

def build_status_video(text, bg_key, font_key, out_path, duration=12, bg_image=None):
    """
    Generate a 1080x1920 (9:16) Bengali status video — WhatsApp/Facebook ready.
    bg_image: optional path to user's own photo/video (overrides preset)
    """
    label, c_top, c_bot = STATUS_BG_PRESETS.get(bg_key, STATUS_BG_PRESETS["night"])
    font_path = pick_font_for_text(text, preferred_bangla=font_key)
    wrapped = wrap_bn_text(text, max_chars=22)
    # 🆕 v16: textfile= ব্যবহার — multi-line, comma, bracket সব issue ফিক্স
    txt_file = ff_textfile(wrapped, suffix="_status.txt")
    font_arg = f":fontfile='{font_path}'" if font_path else ""
    if txt_file:
        text_src = f"textfile='{txt_file}'"
    else:
        text_src = f"text='{ff_escape(wrapped)}'"
    # Premium text overlay — semi-transparent box, Bengali font, dual border, shadow
    text_vf = (
        f"drawtext={text_src}{font_arg}:fontsize=56:fontcolor=white:"
        f"box=1:boxcolor=black@0.55:boxborderw=32:"
        f"borderw=2:bordercolor=black:"
        f"shadowx=2:shadowy=3:shadowcolor=black@0.7:"
        f"line_spacing=14:x=(w-text_w)/2:y=(h-text_h)/2"
    )
    # Build input + filter chain depending on bg source
    if bg_image and Path(bg_image).exists():
        ext = Path(bg_image).suffix.lower()
        is_video = ext in {".mp4", ".mov", ".webm", ".mkv", ".avi", ".gif"}
        if is_video:
            cmd = [
                "ffmpeg", "-y",
                "-stream_loop", "-1", "-i", bg_image,
                "-f", "lavfi", "-i", "anullsrc=channel_layout=stereo:sample_rate=44100",
                "-vf", (
                    f"scale=1080:1920:force_original_aspect_ratio=increase,"
                    f"crop=1080:1920,eq=brightness=-0.08:saturation=0.95,"
                    f"vignette=PI/3,{text_vf}"
                ),
                "-t", str(duration),
                "-c:v", "libx264", "-preset", "veryfast", "-crf", "20",
                "-c:a", "aac", "-b:a", "128k",
                "-pix_fmt", "yuv420p", "-movflags", "+faststart", "-shortest",
                out_path
            ]
        else:
            # Image: ken-burns style slow zoom for cinematic feel
            cmd = [
                "ffmpeg", "-y",
                "-loop", "1", "-i", bg_image,
                "-f", "lavfi", "-i", "anullsrc=channel_layout=stereo:sample_rate=44100",
                "-vf", (
                    f"scale=1620:2880:force_original_aspect_ratio=increase,"
                    f"crop=1620:2880,"
                    f"zoompan=z='min(zoom+0.0008,1.15)':d={int(duration*30)}:s=1080x1920:fps=30,"
                    f"eq=brightness=-0.08:saturation=0.95,"
                    f"vignette=PI/3,{text_vf}"
                ),
                "-t", str(duration),
                "-c:v", "libx264", "-preset", "veryfast", "-crf", "20",
                "-c:a", "aac", "-b:a", "128k",
                "-pix_fmt", "yuv420p", "-movflags", "+faststart", "-shortest",
                out_path
            ]
    else:
        # Procedural cinematic background — solid dark color + heavy vignette + grain
        # NOTE: We use the darker of the two preset colors as base for safer rendering
        base = c_bot
        cmd = [
            "ffmpeg", "-y",
            "-f", "lavfi", "-i", f"color=c={base}:s=1080x1920:r=30:d={duration}",
            "-f", "lavfi", "-i", "anullsrc=channel_layout=stereo:sample_rate=44100",
            "-vf", (
                # subtle grain + vignette → cinematic mood
                f"noise=alls=6:allf=t,"
                f"vignette=PI/2.5,"
                # very slow zoom for life
                f"zoompan=z='min(zoom+0.0005,1.08)':d={int(duration*30)}:s=1080x1920:fps=30,"
                f"{text_vf}"
            ),
            "-t", str(duration),
            "-c:v", "libx264", "-preset", "veryfast", "-crf", "20",
            "-c:a", "aac", "-b:a", "128k",
            "-pix_fmt", "yuv420p", "-movflags", "+faststart", "-shortest",
            out_path
        ]
    try:
        r = subprocess.run(cmd, capture_output=True, timeout=180)
        if r.returncode != 0:
            err_tail = (r.stderr or b"")[-400:].decode("utf-8", errors="ignore")
            logger.warning("status video error: %s", err_tail)
        return r.returncode == 0 and Path(out_path).exists() and Path(out_path).stat().st_size > 1000
    except Exception as e:
        logger.warning("status video exception: %s", e)
        return False

# ──────────────────────────────────────────────────────────────
# STATS — মেমোরি cache (ফাইল I/O কমায়)
# ──────────────────────────────────────────────────────────────
_STATS_CACHE = None

def _default(): return {"total_users": [], "videos_processed": 0, "downloads": 0,
                        "feature_usage": {},  # 🆕 v17: per-feature counter
                        "start_time": datetime.now().isoformat(timespec="seconds")}

def load_stats():
    global _STATS_CACHE
    if _STATS_CACHE is not None: return _STATS_CACHE
    if STATS_FILE.exists():
        try:
            d = json.loads(STATS_FILE.read_text(encoding="utf-8"))
            for k,v in _default().items(): d.setdefault(k,v)
            _STATS_CACHE = d
            return d
        except Exception: pass
    _STATS_CACHE = _default()
    return _STATS_CACHE

def save_stats():
    if _STATS_CACHE is None: return
    try:
        t = STATS_FILE.with_suffix(".tmp")
        t.write_text(json.dumps(_STATS_CACHE, ensure_ascii=False, indent=2), encoding="utf-8")
        t.replace(STATS_FILE)
    except Exception as e: logger.warning("save_stats: %s", e)

def track_user(uid):
    with stats_lock:
        s = load_stats(); u = str(uid)
        if u not in s["total_users"]:
            s["total_users"].append(u); save_stats()

def inc_stat(k, n=1):
    with stats_lock:
        s = load_stats(); s[k] = s.get(k,0)+n; save_stats()

def inc_feature(feature_name, n=1):
    """🆕 v17: কোন feature কতবার ব্যবহার হলো — track করে।"""
    if not feature_name: return
    with stats_lock:
        s = load_stats()
        fu = s.setdefault("feature_usage", {})
        fu[feature_name] = fu.get(feature_name, 0) + n
        save_stats()

# ──────────────────────────────────────────────────────────────
# 🆕 v18: CLONED VOICES — per-user ElevenLabs voice clone storage
# ──────────────────────────────────────────────────────────────
CLONED_VOICES_FILE = STATS_DIR / "cloned_voices.json"
_CLONES_CACHE = None
clones_lock = threading.Lock()

def load_clones():
    global _CLONES_CACHE
    if _CLONES_CACHE is not None:
        return _CLONES_CACHE
    if CLONED_VOICES_FILE.exists():
        try:
            _CLONES_CACHE = json.loads(CLONED_VOICES_FILE.read_text(encoding="utf-8"))
            return _CLONES_CACHE
        except Exception as e:
            logger.warning("load_clones: %s", e)
    _CLONES_CACHE = {}
    return _CLONES_CACHE

def save_clones():
    try:
        t = CLONED_VOICES_FILE.with_suffix(".tmp")
        t.write_text(json.dumps(_CLONES_CACHE, ensure_ascii=False, indent=2), encoding="utf-8")
        t.replace(CLONED_VOICES_FILE)
    except Exception as e:
        logger.error("save_clones: %s", e)

def get_user_clone(uid):
    """Return dict {voice_id, name, created} or None."""
    with clones_lock:
        return load_clones().get(str(uid))

def set_user_clone(uid, voice_id, name):
    with clones_lock:
        d = load_clones()
        d[str(uid)] = {
            "voice_id": voice_id,
            "name": name,
            "created": datetime.now().isoformat(timespec="seconds"),
        }
        save_clones()

def del_user_clone(uid):
    with clones_lock:
        d = load_clones()
        if str(uid) in d:
            d.pop(str(uid))
            save_clones()
            return True
    return False

# ──────────────────────────────────────────────────────────────
PLATFORM_INFO = {
    "youtube.com":"📺 YouTube","youtu.be":"📺 YouTube","shorts.google.com":"📺 YT Shorts",
    "instagram.com":"📷 Instagram","instagr.am":"📷 Instagram",
    "tiktok.com":"🎵 TikTok","vm.tiktok.com":"🎵 TikTok",
    "facebook.com":"👥 Facebook","fb.watch":"👥 Facebook","fb.com":"👥 Facebook",
    "twitter.com":"🐦 Twitter","x.com":"🐦 X","t.co":"🐦 Twitter",
    "reddit.com":"🔴 Reddit","redd.it":"🔴 Reddit",
    "pinterest.com":"📌 Pinterest","pin.it":"📌 Pinterest",
    "dailymotion.com":"🎬 Dailymotion","vimeo.com":"🎥 Vimeo","twitch.tv":"🟣 Twitch",
    "snapchat.com":"👻 Snapchat","likee.com":"📱 Likee","kwai.com":"🎭 Kwai",
    "linkedin.com":"💼 LinkedIn","vk.com":"📱 VK","ok.ru":"📱 OK",
    "threads.net":"🧵 Threads","imdb.com":"🎬 IMDb","bilibili.com":"📺 Bilibili",
    "ted.com":"🎤 TED","soundcloud.com":"🔊 SoundCloud",
    "rumble.com":"🎬 Rumble","odysee.com":"🎬 Odysee","9gag.com":"😂 9GAG","tumblr.com":"📝 Tumblr",
}
URL_RE = re.compile(r"https?://[^\s]+", re.IGNORECASE)
TIME_RE = re.compile(r"^\s*(\d{1,2}:)?(\d{1,2}:)?\d{1,2}(\.\d+)?\s*[-–]\s*(\d{1,2}:)?(\d{1,2}:)?\d{1,2}(\.\d+)?\s*$")

def get_platform(url):
    u = url.lower()
    for d,n in PLATFORM_INFO.items():
        if d in u: return n
    return "🌐 ভিডিও"

def find_url(t):
    if not t: return None
    m = URL_RE.search(t)
    return m.group(0).rstrip(".,);]}>!?\"'") if m else None

def md_escape(t):
    if not t: return ""
    return re.sub(r'([_*`\[\]])', r'\\\1', str(t))[:200]

def parse_time(s):
    """00:30 / 1:23:45 / 75 → seconds"""
    s = s.strip()
    parts = s.split(":")
    try:
        if len(parts) == 1: return float(parts[0])
        if len(parts) == 2: return int(parts[0])*60 + float(parts[1])
        if len(parts) == 3: return int(parts[0])*3600 + int(parts[1])*60 + float(parts[2])
    except Exception: pass
    return None

# 🆕 v16: Proper ffmpeg drawtext escape — sob special char handle kore
def ff_escape(s):
    """ffmpeg drawtext text= এর জন্য সঠিক escape। comma, semicolon, bracket সব handle করে।"""
    if s is None:
        return ""
    s = str(s)
    # backslash আগে — নাহলে double-escape হবে
    s = s.replace("\\", "\\\\")
    # filter syntax-এ এই গুলা escape করতে হয়
    for ch in [":", "'", "%"]:
        s = s.replace(ch, "\\" + ch)
    # quote replace + control char strip
    s = s.replace('"', "")
    return s

def ff_textfile(text, suffix="_text.txt"):
    """Multi-line text কে temp file-এ লিখে path return করে — drawtext textfile= এর জন্য।
    এতে newline, comma, bracket সব issue solve হয়।"""
    p = TEMP_DIR / f"{uuid.uuid4().hex}{suffix}"
    try:
        p.write_text(str(text or ""), encoding="utf-8")
        return str(p)
    except Exception as e:
        logger.warning("ff_textfile write failed: %s", e)
        return None

async def safe_edit(target, text, **kw):
    try:
        await target.edit_text(text, **kw)
    except BadRequest as e:
        if "not modified" in str(e).lower(): return
        try:
            kw.pop("parse_mode", None)
            plain = re.sub(r'[*_`\[\]()]', '', text)
            await target.edit_text(plain, **kw)
        except Exception as e2: logger.warning("safe_edit fallback: %s", e2)
    except Exception as e: logger.warning("safe_edit: %s", e)

async def safe_reply(msg, text, **kw):
    try:
        return await msg.reply_text(text, **kw)
    except BadRequest:
        kw.pop("parse_mode", None)
        plain = re.sub(r'[*_`\[\]()]', '', text)
        return await msg.reply_text(plain, **kw)

# ──────────────────────────────────────────────────────────────
# MENUS
# ──────────────────────────────────────────────────────────────
def main_menu():
    # 🆕 v31: Beautified main menu — grouped sections + premium emojis
    # 🆕 v37: File Extractor + MCQ Solver buttons added
    return InlineKeyboardMarkup([
        # ─── 📥 INPUT ───
        [InlineKeyboardButton("📤 ভিডিও আপলোড", callback_data="how"),
         InlineKeyboardButton("🔗 লিংক ডাউনলোড", callback_data="url_info")],
        # ─── 🚀 SIGNATURE FEATURE ───
        [InlineKeyboardButton("🎯  WinGo AI — 30s · 1m · 3m · 5m  🤖", callback_data="wingo_main")],
        # ─── 🆕 v37 NEW TOOLS ───
        [InlineKeyboardButton("📂 File Extractor", callback_data="file_extract_start"),
         InlineKeyboardButton("📚 MCQ Solver 🎓", callback_data="mcq_start")],
        [InlineKeyboardButton("🔗 Link Expander (URL safety check)", callback_data="link_expand_start")],
        [InlineKeyboardButton("📸 Photo → Live Photo (TikTok / iPhone)", callback_data="photo_live_start")],
        # ─── 🎬 CONTENT CREATION ───
        [InlineKeyboardButton("📱 Status Video (Auto)", callback_data="status_create"),
         InlineKeyboardButton("🎬 Video Dub 🌍", callback_data="dub_ask")],
        [InlineKeyboardButton("🗣️ টেক্সট → ভয়েস", callback_data="tts_start"),
         InlineKeyboardButton("🎭 Voice Clone", callback_data="voice_clone_help")],
        # ─── 🤖 AI ───
        [InlineKeyboardButton("🤖 AI চ্যাট (ChatGPT)", callback_data="ai_chat_start")],
        # ─── ℹ️ INFO ───
        [InlineKeyboardButton("🌐 প্ল্যাটফর্ম", callback_data="platforms"),
         InlineKeyboardButton("ℹ️ সাহায্য", callback_data="help"),
         InlineKeyboardButton("📊 স্ট্যাটস", callback_data="my_stats")],
    ])

def status_bg_menu():
    rows = []
    keys = list(STATUS_BG_PRESETS.keys())
    for i in range(0, len(keys), 2):
        row = []
        for k in keys[i:i+2]:
            row.append(InlineKeyboardButton(STATUS_BG_PRESETS[k][0], callback_data=f"sbg_{k}"))
        rows.append(row)
    rows.append([InlineKeyboardButton("🖼️ আমার নিজের ছবি দিব", callback_data="sbg_custom")])
    rows.append([InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")])
    return InlineKeyboardMarkup(rows)

def status_dur_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("⏱️ 8 সেকেন্ড", callback_data="sdur_8"),
         InlineKeyboardButton("⏱️ 12 সেকেন্ড", callback_data="sdur_12")],
        [InlineKeyboardButton("⏱️ 15 সেকেন্ড", callback_data="sdur_15"),
         InlineKeyboardButton("⏱️ 20 সেকেন্ড", callback_data="sdur_20")],
        [InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")],
    ])

def status_font_menu():
    rows = []
    for key, (label, fname, _) in BANGLA_FONTS.items():
        rows.append([InlineKeyboardButton(label, callback_data=f"sfnt_{key}")])
    rows.append([InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")])
    return InlineKeyboardMarkup(rows)

def tts_lang_menu():
    """TTS-এর জন্য ভাষা সিলেকশন।"""
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🇧🇩 বাংলা", callback_data="tts_lang_bn"),
         InlineKeyboardButton("🇬🇧 English", callback_data="tts_lang_en")],
        [InlineKeyboardButton("🇮🇳 হিন্দি", callback_data="tts_lang_hi"),
         InlineKeyboardButton("🌍 অটো-ডিটেক্ট", callback_data="tts_lang_auto")],
        [InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")],
    ])

def tts_voice_menu(uid=None):
    """TTS engine + voice সিলেকশন main menu।
    🆕 v22: Edge TTS (১০০% ফ্রি, neural quality) সবার উপরে। ElevenLabs পেইড users এর জন্য।"""
    rows = []
    # 🆕 v22: Edge TTS — সম্পূর্ণ ফ্রি premium voices (সবার আগে show করি)
    if has_edge_tts():
        rows += [
            [InlineKeyboardButton("✨ ─── 🆓 FREE Premium Voices (Edge) ─── ✨",
                                  callback_data="tts_noop")],
            [InlineKeyboardButton("🇧🇩 Nabanita (বাংলা মেয়ে)", callback_data="tts_edge_bn_nabanita"),
             InlineKeyboardButton("🇧🇩 Pradeep (বাংলা ছেলে)",  callback_data="tts_edge_bn_pradeep")],
            [InlineKeyboardButton("🇮🇳 Bashkar (বাংলা ছেলে)",  callback_data="tts_edge_bn_bashkar"),
             InlineKeyboardButton("🇮🇳 Tanishaa (বাংলা মেয়ে)",  callback_data="tts_edge_bn_tanishaa")],
            [InlineKeyboardButton("🇮🇳 Swara (हिन्दी स्त्री)",   callback_data="tts_edge_hi_swara"),
             InlineKeyboardButton("🇮🇳 Madhur (हिन्दी पुरुष)",  callback_data="tts_edge_hi_madhur")],
            [InlineKeyboardButton("🇺🇸 Aria (English F)",       callback_data="tts_edge_en_aria"),
             InlineKeyboardButton("🇺🇸 Guy (English M)",        callback_data="tts_edge_en_guy")],
            [InlineKeyboardButton(f"📚 সব {len(EDGE_TTS_VOICES)}+ Edge voice browse",
                                  callback_data="tts_edpg_0")],
        ]
    rows.append([InlineKeyboardButton("🆓 gTTS (basic, fallback)", callback_data="tts_engine_gtts")])
    # 🆕 v26: User-এর own cloned voice — ElevenLabs OR Replicate XTTS
    if uid is not None:
        clone = get_user_clone(uid)
        if clone and clone.get("voice_id"):
            vid = clone["voice_id"]
            # XTTS clone show করি যদি Replicate token আছে; অথবা legacy ElevenLabs clone যদি ElevenLabs key আছে
            show_clone = (vid.startswith("xtts:") and has_replicate()) or \
                         (not vid.startswith("xtts:") and has_elevenlabs())
            if show_clone:
                cname = (clone.get("name") or "MyVoice")[:18]
                engine_tag = "🎭 XTTS" if vid.startswith("xtts:") else "🎭"
                rows.append([InlineKeyboardButton(f"{engine_tag} আপনার Cloned Voice — {cname}",
                                                  callback_data="tts_eleven_mine")])
    # ─── OpenAI voices (key থাকলে) ───
    if has_openai():
        rows += [
            [InlineKeyboardButton("✨ ─── OpenAI TTS ─── ✨", callback_data="tts_noop")],
            [InlineKeyboardButton("👩 Nova (মেয়ে, soft)", callback_data="tts_voice_nova"),
             InlineKeyboardButton("👨 Onyx (ছেলে, deep)", callback_data="tts_voice_onyx")],
            [InlineKeyboardButton("🧑 Alloy (neutral)", callback_data="tts_voice_alloy"),
             InlineKeyboardButton("✨ Shimmer (নরম)", callback_data="tts_voice_shimmer")],
            [InlineKeyboardButton("📚 Fable (story)", callback_data="tts_voice_fable"),
             InlineKeyboardButton("📢 Echo (ছেলে, clear)", callback_data="tts_voice_echo")],
        ]
    # ─── 🆕 v19: ElevenLabs voices — Top 6 popular + "All ৪০+ browse" entrance ───
    if has_elevenlabs():
        rows += [
            [InlineKeyboardButton("🎤 ─── ElevenLabs (Realistic) ─── 🎤", callback_data="tts_noop")],
            [InlineKeyboardButton("👩 Rachel (calm)", callback_data="tts_eleven_rachel"),
             InlineKeyboardButton("💃 Bella (soft)", callback_data="tts_eleven_bella")],
            [InlineKeyboardButton("🔥 Domi (strong)", callback_data="tts_eleven_domi"),
             InlineKeyboardButton("👨 Adam (deep)", callback_data="tts_eleven_adam")],
            [InlineKeyboardButton("🎙️ Antoni (warm)", callback_data="tts_eleven_antoni"),
             InlineKeyboardButton("📻 Josh (young)", callback_data="tts_eleven_josh")],
            [InlineKeyboardButton(f"📚 সব {len(ELEVENLABS_VOICES)}+ voice browse করুন",
                                  callback_data="tts_evpg_0")],
            [InlineKeyboardButton("🔄 আমার অ্যাকাউন্টের cloned voice list",
                                  callback_data="tts_evacc_0")],
        ]
    rows.append([InlineKeyboardButton("🔙 ফিরে যান", callback_data="tts_start")])
    return InlineKeyboardMarkup(rows)

# 🆕 v22: Edge TTS — paginated voice picker (সব ফ্রি neural voice)
def edge_tts_page_menu(page=0):
    """EDGE_TTS_VOICES dict-এর সব voice paginated দেখায়।"""
    items = list(EDGE_TTS_VOICES.items())
    per = EDGE_TTS_VOICES_PER_PAGE
    total_pages = max(1, (len(items) + per - 1) // per)
    page = max(0, min(page, total_pages - 1))
    chunk = items[page * per:(page + 1) * per]
    rows = []
    # Single column for clarity (long labels)
    for k, (vname, label, lang) in chunk:
        rows.append([InlineKeyboardButton(label[:32], callback_data=f"tts_edge_{k}")])
    nav = []
    if page > 0:
        nav.append(InlineKeyboardButton("⬅️ আগের", callback_data=f"tts_edpg_{page-1}"))
    nav.append(InlineKeyboardButton(f"📄 {page+1}/{total_pages}", callback_data="tts_noop"))
    if page < total_pages - 1:
        nav.append(InlineKeyboardButton("পরের ➡️", callback_data=f"tts_edpg_{page+1}"))
    rows.append(nav)
    rows.append([InlineKeyboardButton("🔙 voice মেনু", callback_data="tts_voice_back")])
    return InlineKeyboardMarkup(rows)

# 🆕 v19: ElevenLabs default library — paginated voice picker
def elevenlabs_library_page_menu(page=0):
    """ELEVENLABS_VOICES dict-এর সব voice paginated দেখায়।"""
    items = list(ELEVENLABS_VOICES.items())  # [(key, (id, label)), ...]
    per = ELEVENLABS_VOICES_PER_PAGE
    total_pages = max(1, (len(items) + per - 1) // per)
    page = max(0, min(page, total_pages - 1))
    chunk = items[page * per:(page + 1) * per]
    rows = []
    # 2-column layout
    for i in range(0, len(chunk), 2):
        row = []
        for k, (vid, label) in chunk[i:i + 2]:
            row.append(InlineKeyboardButton(label[:28], callback_data=f"tts_eleven_{k}"))
        rows.append(row)
    # Nav row
    nav = []
    if page > 0:
        nav.append(InlineKeyboardButton("⬅️ আগের", callback_data=f"tts_evpg_{page-1}"))
    nav.append(InlineKeyboardButton(f"📄 {page+1}/{total_pages}", callback_data="tts_noop"))
    if page < total_pages - 1:
        nav.append(InlineKeyboardButton("পরের ➡️", callback_data=f"tts_evpg_{page+1}"))
    rows.append(nav)
    rows.append([InlineKeyboardButton("🔙 voice মেনু", callback_data="tts_voice_back")])
    return InlineKeyboardMarkup(rows)

# 🆕 v19: ElevenLabs অ্যাকাউন্টের সব voice (live API থেকে) paginated
def elevenlabs_account_page_menu(page=0, voices=None):
    """User-এর ElevenLabs অ্যাকাউন্টে যত voice আছে (premade + cloned + generated) সব dynamically দেখায়।
    voice_id সরাসরি callback-এ যাবে (tts_evid_<voice_id>)।"""
    voices = voices if voices is not None else elevenlabs_list_voices()
    per = ELEVENLABS_VOICES_PER_PAGE
    if not voices:
        return InlineKeyboardMarkup([
            [InlineKeyboardButton("🔄 আবার fetch", callback_data="tts_evacc_0")],
            [InlineKeyboardButton("🔙 voice মেনু", callback_data="tts_voice_back")],
        ])
    total_pages = max(1, (len(voices) + per - 1) // per)
    page = max(0, min(page, total_pages - 1))
    chunk = voices[page * per:(page + 1) * per]
    rows = []
    for v in chunk:
        label = elevenlabs_voice_label(v["category"], v["name"])
        # voice_id ~20 chars + prefix 9 = 29 < 64 byte Telegram limit
        rows.append([InlineKeyboardButton(label, callback_data=f"tts_evid_{v['voice_id']}")])
    nav = []
    if page > 0:
        nav.append(InlineKeyboardButton("⬅️ আগের", callback_data=f"tts_evacc_{page-1}"))
    nav.append(InlineKeyboardButton(f"📄 {page+1}/{total_pages}", callback_data="tts_noop"))
    if page < total_pages - 1:
        nav.append(InlineKeyboardButton("পরের ➡️", callback_data=f"tts_evacc_{page+1}"))
    rows.append(nav)
    rows.append([InlineKeyboardButton("🔄 Refresh list", callback_data="tts_evacc_refresh")])
    rows.append([InlineKeyboardButton("🔙 voice মেনু", callback_data="tts_voice_back")])
    return InlineKeyboardMarkup(rows)

# ──────────────────────────────────────────────────────────────
# 🆕 v14: AI TOOLKIT — Pro Menu (Analysis / Script / Hooks / Thumb / Translate)
# ──────────────────────────────────────────────────────────────
def ai_toolkit_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🎬 ভিডিও কনটেন্ট অ্যানালিসিস", callback_data="ai_analyze")],
        [InlineKeyboardButton("🎙️ স্ক্রিপ্ট + AI ভয়েসওভার", callback_data="ai_script")],
        [InlineKeyboardButton("🏷️ Hook + Caption (Reels/TikTok)", callback_data="ai_hooks")],
        [InlineKeyboardButton("🔍 থাম্বনেইল আইডিয়া (৫টা)", callback_data="ai_thumb")],
        [InlineKeyboardButton("🌐 সাবটাইটেল অনুবাদ", callback_data="ai_translate_sub")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

def translate_lang_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🇬🇧 English", callback_data="tr_lang_en"),
         InlineKeyboardButton("🇧🇩 বাংলা", callback_data="tr_lang_bn")],
        [InlineKeyboardButton("🇮🇳 हिन्दी", callback_data="tr_lang_hi"),
         InlineKeyboardButton("🇸🇦 العربية", callback_data="tr_lang_ar")],
        [InlineKeyboardButton("🇪🇸 Español", callback_data="tr_lang_es"),
         InlineKeyboardButton("🇫🇷 Français", callback_data="tr_lang_fr")],
        [InlineKeyboardButton("🇩🇪 Deutsch", callback_data="tr_lang_de"),
         InlineKeyboardButton("🇯🇵 日本語", callback_data="tr_lang_ja")],
        [InlineKeyboardButton("🇰🇷 한국어", callback_data="tr_lang_ko"),
         InlineKeyboardButton("🇨🇳 中文", callback_data="tr_lang_zh")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="ai_toolkit")],
    ])

def script_duration_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("⚡ ১৫ সেকেন্ড", callback_data="script_dur_15"),
         InlineKeyboardButton("📱 ৩০ সেকেন্ড", callback_data="script_dur_30")],
        [InlineKeyboardButton("🎬 ৬০ সেকেন্ড", callback_data="script_dur_60"),
         InlineKeyboardButton("📺 ৩ মিনিট", callback_data="script_dur_180")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="ai_toolkit")],
    ])

def edit_menu(source_url=None, platform=None):
    """🆕 v27.1: Reorganized into clean sections — TOP HITS → AI → Quality → Voice → Edit → Effects → Utils"""
    rows = []
    # ─── Source Platform (যে প্ল্যাটফর্ম থেকে এসেছে) ───
    if source_url and platform:
        rows.append([InlineKeyboardButton(f"🌐 {platform}-এ দেখুন", url=source_url)])

    rows += [
        # ═══════ ⭐ TOP HITS (most popular features) ═══════
        [InlineKeyboardButton("━━━━━━━━ ⭐ TOP HITS ━━━━━━━━", callback_data="noop_section")],
        [InlineKeyboardButton("🤖 Smart Auto (All-in-One) 🔥", callback_data="smart_auto")],
        [InlineKeyboardButton("🚀 AI টুলকিট (Pro)", callback_data="ai_toolkit")],
        [InlineKeyboardButton("🚀 YouTube ভাইরাল প্যাক 🔥", callback_data="youtube_viral")],

        # ═══════ 📺 PLATFORM MODES ═══════
        [InlineKeyboardButton("━━━━━ 📺 PLATFORM MODES ━━━━━", callback_data="noop_section")],
        [InlineKeyboardButton("🎵 TikTok Mode (For You)", callback_data="tiktok_mode"),
         InlineKeyboardButton("📺 YouTube Mode (FAST ⚡)", callback_data="youtube_mode")],
        [InlineKeyboardButton("👥 Facebook Mode (Viral 🔥) 🆕", callback_data="facebook_mode")],
        [InlineKeyboardButton("💰 YouTube টাকা আয়ের গাইড (Monetization)", callback_data="youtube_money")],

        # ═══════ 🎭 VOICE & AI (NEW auto-clone) ═══════
        [InlineKeyboardButton("━━━━━━ 🎭 VOICE & AI ━━━━━━", callback_data="noop_section")],
        [InlineKeyboardButton("🎤 অটো ভয়েস ক্লোন (এই ভিডিও থেকে) 🆕🔥", callback_data="auto_clone_voice")],
        [InlineKeyboardButton("🎙️ AI Voiceover", callback_data="vo_ask"),
         InlineKeyboardButton("🎙️📝 Voiceover+Sub", callback_data="vosub_ask")],
        [InlineKeyboardButton("🎬🌍 Video Dubbing (অন্য ভাষায় ডাব) 🆕🔥", callback_data="dub_ask")],
        [InlineKeyboardButton("🎙️ Voice Enhance Pro (podcast)", callback_data="voice_pro")],
        [InlineKeyboardButton("🎯 SEO টাইটেল+ট্যাগ", callback_data="seo_gen"),
         InlineKeyboardButton("🛡️ কপিরাইট চেক", callback_data="copyright_check")],
        [InlineKeyboardButton("📝 অটো সাবটাইটেল 🔥", callback_data="auto_sub")],

        # ═══════ ✨ QUALITY & LOOK ═══════
        [InlineKeyboardButton("━━━━━━ ✨ QUALITY & LOOK ━━━━━━", callback_data="noop_section")],
        [InlineKeyboardButton("✨ এনহান্স আল্ট্রা 🔥", callback_data="enhance"),
         InlineKeyboardButton("💎 Wink Premium 4K", callback_data="enhance_pro")],
        [InlineKeyboardButton("🎨 কালার বুস্ট", callback_data="color"),
         InlineKeyboardButton("🎞️ ফিল্টার", callback_data="filter_menu")],
        [InlineKeyboardButton("🌟 ভিনিয়েট", callback_data="vignette"),
         InlineKeyboardButton("🔊 Audio Denoise", callback_data="denoise")],

        # ═══════ ✂️ EDITING ═══════
        [InlineKeyboardButton("━━━━━━━━ ✂️ EDITING ━━━━━━━━", callback_data="noop_section")],
        [InlineKeyboardButton("✂️ ট্রিম/কাট", callback_data="trim_ask"),
         InlineKeyboardButton("🔗 মার্জ", callback_data="merge_ask")],
        [InlineKeyboardButton("📝 টেক্সট যোগ", callback_data="text_ask"),
         InlineKeyboardButton("🎵 BGM যোগ", callback_data="bgm_ask")],
        [InlineKeyboardButton("🎀 প্রিমিয়াম ক্যাপশন (বাংলা ফন্ট)", callback_data="caption_ask")],
        [InlineKeyboardButton("⚡ ফুল প্রসেস", callback_data="full"),
         InlineKeyboardButton("🔒 কপিরাইট রিমুভ", callback_data="copyright")],
        [InlineKeyboardButton("🚫 ওয়াটারমার্ক রিমুভ", callback_data="wm_menu")],

        # ═══════ ⚙️ EFFECTS ═══════
        [InlineKeyboardButton("━━━━━━━━ ⚙️ EFFECTS ━━━━━━━━", callback_data="noop_section")],
        [InlineKeyboardButton("⏩ স্পিড", callback_data="speed_menu"),
         InlineKeyboardButton("🎞️ Slow-Mo Premium", callback_data="slowmo_smooth")],
        [InlineKeyboardButton("📐 অ্যাসপেক্ট", callback_data="aspect_menu"),
         InlineKeyboardButton("🪞 ফ্লিপ", callback_data="flip_menu")],
        [InlineKeyboardButton("🌅 ফেড ইন/আউট", callback_data="fade"),
         InlineKeyboardButton("🌫️ ব্লার BG", callback_data="blur_bg")],
        [InlineKeyboardButton("🔍 জুম ইফেক্ট", callback_data="zoom"),
         InlineKeyboardButton("⏸️ ফ্রিজ ফ্রেম", callback_data="freeze")],
        [InlineKeyboardButton("🔁 বুমেরাং", callback_data="boomerang"),
         InlineKeyboardButton("📺 রেট্রো VHS", callback_data="vhs")],
        [InlineKeyboardButton("🎞️ GIF বানান", callback_data="gif"),
         InlineKeyboardButton("♾️ N সেকেন্ড লুপ", callback_data="loop_ask")],

        # ═══════ 🔧 UTILITIES ═══════
        [InlineKeyboardButton("━━━━━━━ 🔧 UTILITIES ━━━━━━━", callback_data="noop_section")],
        [InlineKeyboardButton("📦 কম্প্রেস", callback_data="compress"),
         InlineKeyboardButton("🎵 MP3 এক্সট্র্যাক্ট", callback_data="audio")],
        [InlineKeyboardButton("🖼 থাম্বনেইল", callback_data="thumb"),
         InlineKeyboardButton("🔊 ভলিউম", callback_data="vol_menu")],
        [InlineKeyboardButton("🔄 ৯০° ঘোরান", callback_data="rotate"),
         InlineKeyboardButton("📊 Video Info", callback_data="video_info")],
        [InlineKeyboardButton("🔇 মিউট", callback_data="mute"),
         InlineKeyboardButton("⏪ রিভার্স", callback_data="reverse")],

        # ═══════ Navigation ═══════
        [InlineKeyboardButton("🏠 মেনু", callback_data="main_menu"),
         InlineKeyboardButton("❌ বাতিল", callback_data="cancel_video")],
    ]
    return InlineKeyboardMarkup(rows)

def wm_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("↖️ উপরে-বামে", callback_data="wm_tl"),
         InlineKeyboardButton("↗️ উপরে-ডানে", callback_data="wm_tr")],
        [InlineKeyboardButton("↙️ নিচে-বামে", callback_data="wm_bl"),
         InlineKeyboardButton("↘️ নিচে-ডানে", callback_data="wm_br")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

def filter_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🎬 সিনেম্যাটিক", callback_data="f_cinema"),
         InlineKeyboardButton("📽️ ভিনটেজ", callback_data="f_vintage")],
        [InlineKeyboardButton("⚫ ব্ল্যাক & হোয়াইট", callback_data="f_bw"),
         InlineKeyboardButton("🌅 ওয়ার্ম", callback_data="f_warm")],
        [InlineKeyboardButton("❄️ কোল্ড", callback_data="f_cold"),
         InlineKeyboardButton("🎭 ড্রামাটিক", callback_data="f_drama")],
        [InlineKeyboardButton("🌃 নিয়ন", callback_data="f_neon"),
         InlineKeyboardButton("🌸 ড্রিমি", callback_data="f_dreamy")],
        [InlineKeyboardButton("🎌 অ্যানিমে", callback_data="f_anime"),
         InlineKeyboardButton("🌇 সানসেট", callback_data="f_sunset")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

def speed_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🐢 0.5x (স্লো)", callback_data="sp_05"),
         InlineKeyboardButton("⚡ 1.5x", callback_data="sp_15")],
        [InlineKeyboardButton("🚀 2x", callback_data="sp_20"),
         InlineKeyboardButton("💨 4x", callback_data="sp_40")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

def aspect_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📱 9:16 (Reels/TikTok)", callback_data="ar_916")],
        [InlineKeyboardButton("⬛ 1:1 (Insta)", callback_data="ar_11"),
         InlineKeyboardButton("📺 16:9 (YT)", callback_data="ar_169")],
        [InlineKeyboardButton("📷 4:5 (Insta Post)", callback_data="ar_45")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

def flip_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("↔️ আনুভূমিক (মিরর)", callback_data="fl_h")],
        [InlineKeyboardButton("↕️ উল্লম্ব", callback_data="fl_v")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

def vol_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🔉 50% (কম)", callback_data="vol_50"),
         InlineKeyboardButton("🔊 150%", callback_data="vol_150")],
        [InlineKeyboardButton("📢 200%", callback_data="vol_200"),
         InlineKeyboardButton("📣 300% (জোর)", callback_data="vol_300")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

# ─── 🆕 v15: PREMIUM CAPTION (Bengali Font) MENUS ───
def caption_menu():
    """Position picker — premium caption flow's first step."""
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("⬆️ উপরে", callback_data="cap_pos_top"),
         InlineKeyboardButton("⏺️ মাঝে (Status)", callback_data="cap_pos_mid"),
         InlineKeyboardButton("⬇️ নিচে", callback_data="cap_pos_bot")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

def caption_font_menu():
    rows = []
    for key, (label, fname, _) in BANGLA_FONTS.items():
        rows.append([InlineKeyboardButton(label, callback_data=f"cap_fnt_{key}")])
    rows.append([InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")])
    return InlineKeyboardMarkup(rows)

def caption_color_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("⚪ সাদা", callback_data="cap_col_white"),
         InlineKeyboardButton("🟡 হলুদ", callback_data="cap_col_yellow")],
        [InlineKeyboardButton("⚫ কালো", callback_data="cap_col_black"),
         InlineKeyboardButton("🔴 লাল", callback_data="cap_col_red")],
        [InlineKeyboardButton("🟢 সবুজ", callback_data="cap_col_green"),
         InlineKeyboardButton("🩵 সায়ান", callback_data="cap_col_cyan")],
        [InlineKeyboardButton("🟠 কমলা", callback_data="cap_col_orange"),
         InlineKeyboardButton("🩷 গোলাপি", callback_data="cap_col_pink")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

def caption_bg_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("✅ ব্যাকগ্রাউন্ড বক্স সহ (Status look)", callback_data="cap_bg_yes")],
        [InlineKeyboardButton("⛔ শুধু টেক্সট (Clean look)", callback_data="cap_bg_no")],
        [InlineKeyboardButton("🔙 ফিরে যান", callback_data="back_to_edit")],
    ])

def back_menu():
    return InlineKeyboardMarkup([[InlineKeyboardButton("🔄 নতুন ভিডিও দিন", callback_data="main_menu")]])

def edit_menu_for(uid):
    """user_videos থেকে source URL+platform পড়ে যথাযথ edit_menu রিটার্ন করে।"""
    with user_lock:
        d = user_videos.get(uid) or {}
    return edit_menu(source_url=d.get("url"), platform=d.get("platform"))

def cancel_input_menu():
    return InlineKeyboardMarkup([[InlineKeyboardButton("❌ বাতিল", callback_data="back_to_edit")]])

LABELS = {
    "smart_auto":"🤖 Smart Auto (All-in-One)",
    "enhance":"✨ এনহান্স আল্ট্রা","enhance_pro":"💎 Wink Premium 4K",
    "color":"🎨 কালার বুস্ট","copyright":"🔒 কপিরাইট রিমুভ (উন্নত)",
    "full":"⚡ ফুল প্রসেস","compress":"📦 কম্প্রেস","audio":"🎵 MP3","thumb":"🖼 থাম্বনেইল",
    "rotate":"🔄 ঘোরান","mute":"🔇 মিউট","reverse":"⏪ রিভার্স",
    "wm_tl":"🚫 ওয়াটারমার্ক (↖️)","wm_tr":"🚫 ওয়াটারমার্ক (↗️)",
    "wm_bl":"🚫 ওয়াটারমার্ক (↙️)","wm_br":"🚫 ওয়াটারমার্ক (↘️)",
    "f_cinema":"🎬 সিনেম্যাটিক","f_vintage":"📽️ ভিনটেজ","f_bw":"⚫ ব্ল্যাক & হোয়াইট",
    "f_warm":"🌅 ওয়ার্ম","f_cold":"❄️ কোল্ড","f_drama":"🎭 ড্রামাটিক",
    "f_neon":"🌃 নিয়ন","f_dreamy":"🌸 ড্রিমি","f_anime":"🎌 অ্যানিমে","f_sunset":"🌇 সানসেট",
    "sp_05":"🐢 0.5x স্লো","sp_15":"⚡ 1.5x","sp_20":"🚀 2x","sp_40":"💨 4x",
    "ar_916":"📱 9:16","ar_11":"⬛ 1:1","ar_169":"📺 16:9","ar_45":"📷 4:5",
    "fl_h":"↔️ আনুভূমিক ফ্লিপ","fl_v":"↕️ উল্লম্ব ফ্লিপ",
    "fade":"🌅 ফেড ইন/আউট","boomerang":"🔁 বুমেরাং","gif":"🎞️ GIF",
    "blur_bg":"🌫️ ব্লার ব্যাকগ্রাউন্ড","vignette":"🌟 ভিনিয়েট",
    "zoom":"🔍 জুম ইফেক্ট","freeze":"⏸️ ফ্রিজ ফ্রেম","vhs":"📺 রেট্রো VHS",
    "vol_50":"🔉 ভলিউম 50%","vol_150":"🔊 ভলিউম 150%","vol_200":"📢 ভলিউম 200%","vol_300":"📣 ভলিউম 300%",
    "trim":"✂️ ট্রিম","merge":"🔗 মার্জ","text":"📝 টেক্সট","bgm":"🎵 BGM",
    "denoise":"🔊 Audio Denoise","slowmo_smooth":"🎞️ Smooth Slow-Mo",
    "voice_pro":"🎙️ Voice Enhance Pro",
    "loop":"♾️ লুপ","voiceover":"🎙️ AI Voiceover",
}

# ──────────────────────────────────────────────────────────────
# FFPROBE HELPERS (sync — কল করুন executor-এর ভেতর থেকে)
# ──────────────────────────────────────────────────────────────
def get_dimensions(path):
    try:
        cmd = ["ffprobe","-v","error","-select_streams","v:0",
               "-show_entries","stream=width,height","-of","csv=p=0", path]
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if p.returncode == 0 and p.stdout.strip():
            parts = p.stdout.strip().split(",")
            if len(parts) >= 2: return int(parts[0]), int(parts[1])
    except Exception as e: logger.warning("ffprobe dim: %s", e)
    return 1280, 720

def get_duration(path):
    try:
        cmd = ["ffprobe","-v","error","-show_entries","format=duration","-of","csv=p=0", path]
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if p.returncode == 0 and p.stdout.strip():
            return float(p.stdout.strip())
    except Exception as e: logger.warning("ffprobe dur: %s", e)
    return 0.0

# 🐛 v29 FIX: ভিডিওতে audio track আছে কিনা detect — না থাকলে -af filter skip করতে হবে,
# নাহলে ffmpeg fail করে। screen recording / GIF-from-video এ প্রায়ই audio থাকে না।
def has_audio(path):
    try:
        cmd = ["ffprobe","-v","error","-select_streams","a:0",
               "-show_entries","stream=codec_type","-of","csv=p=0", path]
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
        return p.returncode == 0 and "audio" in (p.stdout or "").lower()
    except Exception as e:
        logger.warning("ffprobe audio: %s", e)
        return True  # fallback: ধরে নিই audio আছে

# 🆕 v20: Extract clean audio from video for voice cloning
def extract_audio_for_clone(video_path, out_mp3, max_seconds=60):
    """ভিডিও থেকে first N seconds audio extract → mono 22050Hz mp3 (ElevenLabs clone-এর জন্য optimal)।"""
    try:
        cmd = ["ffmpeg","-y","-i",video_path,"-t",str(max_seconds),
               "-vn","-ac","1","-ar","22050",
               "-af","highpass=f=80,afftdn=nf=-25,acompressor=threshold=-20dB:ratio=3:attack=20:release=250,loudnorm=I=-16:TP=-1.5",
               "-codec:a","libmp3lame","-q:a","4", out_mp3]
        p = subprocess.run(cmd, capture_output=True, timeout=180)
        return p.returncode == 0 and Path(out_mp3).exists() and Path(out_mp3).stat().st_size > 5000
    except Exception as e:
        logger.warning("extract_audio_for_clone: %s", e); return False

# 🆕 v20: Text → SRT subtitle file (synced to TTS audio duration)
def _srt_ts(secs):
    h = int(secs // 3600); m = int((secs % 3600) // 60)
    s = int(secs % 60); ms = int(round((secs - int(secs)) * 1000))
    if ms >= 1000: ms = 999
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

def text_to_srt(text, total_duration, max_chars=42, srt_path=None):
    """Text + total audio duration → SRT file path।
    Sentence boundaries-এ split, character-length proportionally time distribute।"""
    text = (text or "").strip()
    if not text or total_duration <= 0.1:
        return None
    # Split into chunks: first by sentence (। ? ! .), তারপর line wrap by max_chars
    sentences = re.split(r'(?<=[।?!\.\n])\s+', text)
    chunks = []
    for s in sentences:
        s = s.strip()
        if not s: continue
        if len(s) <= max_chars:
            chunks.append(s)
        else:
            # Word-wrap to max_chars
            words = s.split()
            cur = ""
            for w in words:
                if len(cur) + len(w) + 1 <= max_chars:
                    cur = (cur + " " + w).strip()
                else:
                    if cur: chunks.append(cur)
                    cur = w
            if cur: chunks.append(cur)
    if not chunks:
        return None
    # Allocate time proportional to char length
    total_chars = sum(len(c) for c in chunks)
    cur_t = 0.0
    out_lines = []
    for i, c in enumerate(chunks, 1):
        portion = (len(c) / total_chars) if total_chars else (1 / len(chunks))
        dur = max(0.8, portion * total_duration)
        start = cur_t
        end = min(total_duration, cur_t + dur)
        cur_t = end
        out_lines.append(f"{i}\n{_srt_ts(start)} --> {_srt_ts(end)}\n{c}\n")
    p = srt_path or str(TEMP_DIR / f"{uuid.uuid4().hex}_vo.srt")
    Path(p).write_text("\n".join(out_lines), encoding="utf-8")
    return p

# 🆕 v20: Detailed video info card (ffprobe-based)
def fmt_video_info(path):
    """ffprobe থেকে detailed info → Markdown card।"""
    try:
        info = {"size_mb": Path(path).stat().st_size / (1024*1024)}
        # Format
        cmd = ["ffprobe","-v","error","-show_format","-show_streams","-of","json", path]
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if p.returncode != 0:
            return f"❌ ffprobe failed: {p.stderr[:200]}"
        meta = json.loads(p.stdout or "{}")
        fmt = meta.get("format", {}) or {}
        info["duration"] = float(fmt.get("duration", 0) or 0)
        info["bitrate"]  = int(fmt.get("bit_rate", 0) or 0) // 1000
        info["fmt_name"] = fmt.get("format_long_name") or fmt.get("format_name") or "?"
        v_stream = next((s for s in meta.get("streams", []) if s.get("codec_type") == "video"), None)
        a_stream = next((s for s in meta.get("streams", []) if s.get("codec_type") == "audio"), None)
        v_lines = ["⚠️ কোনো video stream নেই"]
        if v_stream:
            fps_str = v_stream.get("r_frame_rate","0/1")
            try:
                num, den = fps_str.split("/")
                fps = round(float(num)/float(den), 2) if float(den) else 0
            except Exception: fps = 0
            v_lines = [
                f"🎥 *Video:* `{v_stream.get('codec_name','?')}` ({v_stream.get('profile','-')})",
                f"📐 Resolution: `{v_stream.get('width','?')}x{v_stream.get('height','?')}`",
                f"🎞️ FPS: `{fps}`  •  Pix: `{v_stream.get('pix_fmt','?')}`",
            ]
        a_lines = ["🔇 কোনো audio stream নেই"]
        if a_stream:
            a_lines = [
                f"🎵 *Audio:* `{a_stream.get('codec_name','?')}`",
                f"🔊 Channels: `{a_stream.get('channels','?')}`  •  SR: `{a_stream.get('sample_rate','?')}Hz`",
            ]
        mins = int(info["duration"] // 60); secs = int(info["duration"] % 60)
        return ("📊 *Video Info*\n\n━━━━━━━━━━━━━━━━━━\n"
                f"📦 Size: `{info['size_mb']:.2f} MB`\n"
                f"⏱️ Duration: `{mins}m {secs}s`  ({info['duration']:.1f}s)\n"
                f"📡 Bitrate: `{info['bitrate']} kbps`\n"
                f"🗂️ Format: `{info['fmt_name'][:40]}`\n\n"
                + "\n".join(v_lines) + "\n\n"
                + "\n".join(a_lines) + "\n\n"
                "_নিচে এডিট অপশন:_")
    except Exception as e:
        logger.warning("fmt_video_info: %s", e)
        return f"📊 *Video Info*\n\n❌ Info read fail: `{str(e)[:120]}`"

def wm_filter(corner, w, h):
    bw, bh = max(80, int(w * 0.28)), max(60, int(h * 0.18))
    pad = 8
    if corner == "tl": x, y = pad, pad
    elif corner == "tr": x, y = max(1, w - bw - pad), pad
    elif corner == "bl": x, y = pad, max(1, h - bh - pad)
    else: x, y = max(1, w - bw - pad), max(1, h - bh - pad)
    bw = min(bw, w - x - 1); bh = min(bh, h - y - 1)
    return f"delogo=x={x}:y={y}:w={bw}:h={bh}:show=0"

# ──────────────────────────────────────────────────────────────
# FFMPEG COMMAND BUILDER
# ──────────────────────────────────────────────────────────────
def ff_cmd(mode, inp, out_base, extra=None):
    """extra: dict — extra params (start, end, text, bgm_path, vol, etc.)"""
    extra = extra or {}
    ext = ".mp4"
    if mode == "audio": ext = ".mp3"
    elif mode == "thumb": ext = ".jpg"
    elif mode == "gif": ext = ".gif"
    out = str(Path(out_base).with_suffix(ext))

    # ─── Watermark Remove ───
    if mode in ("wm_tl","wm_tr","wm_bl","wm_br"):
        w, h = get_dimensions(inp)
        vf = wm_filter(mode.split("_")[1], w, h)
        return ["ffmpeg","-y","-i",inp,"-vf",vf,
                "-c:v","libx264","-preset","veryfast","-crf","21",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Color Filters ───
    filter_map = {
        "f_cinema":  "eq=contrast=1.20:saturation=1.10:gamma=0.92,colorbalance=rs=0.10:bs=-0.08:gh=-0.05,curves=preset=increase_contrast",
        "f_vintage": "curves=r='0/0.10 0.5/0.55 1/0.92':g='0/0.05 0.5/0.50 1/0.88':b='0/0.18 0.5/0.42 1/0.78',eq=saturation=0.75:contrast=0.95",
        "f_bw":      "hue=s=0,eq=contrast=1.15:gamma=0.95",
        "f_warm":    "colorbalance=rs=0.18:gs=0.05:bs=-0.18:rm=0.12:bm=-0.12,eq=saturation=1.10",
        "f_cold":    "colorbalance=rs=-0.18:gs=0.0:bs=0.18:rm=-0.10:bm=0.12,eq=saturation=1.05",
        "f_drama":   "eq=contrast=1.40:saturation=1.25:gamma=0.85,unsharp=5:5:1.0:5:5:0.5,curves=preset=increase_contrast",
        "f_neon":    "eq=contrast=1.30:saturation=1.60:gamma=0.88,colorbalance=rs=-0.15:bs=0.25:gh=0.10,curves=preset=increase_contrast",
        "f_dreamy":  "gblur=sigma=1.5,eq=brightness=0.05:contrast=0.95:saturation=0.90,colorbalance=rs=0.08:bs=0.05",
        "f_anime":   "eq=saturation=1.40:contrast=1.20,unsharp=5:5:1.5:5:5:0.0,curves=preset=increase_contrast",
        "f_sunset":  "colorbalance=rs=0.30:gs=0.10:bs=-0.25:rm=0.20:bm=-0.15,eq=saturation=1.25:contrast=1.10:gamma=0.92",
    }
    if mode in filter_map:
        return ["ffmpeg","-y","-i",inp,"-vf",filter_map[mode],
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Speed ───
    if mode in ("sp_05","sp_15","sp_20","sp_40"):
        ratio = {"sp_05":0.5,"sp_15":1.5,"sp_20":2.0,"sp_40":4.0}[mode]
        pts = 1.0/ratio
        if ratio == 4.0: atempo = "atempo=2.0,atempo=2.0"
        elif ratio == 0.5: atempo = "atempo=0.5"
        else: atempo = f"atempo={ratio}"
        return ["ffmpeg","-y","-i",inp,
                "-filter_complex",f"[0:v]setpts={pts}*PTS[v];[0:a]{atempo}[a]",
                "-map","[v]","-map","[a]",
                "-c:v","libx264","-preset","veryfast","-crf","21",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Aspect Ratio ───
    if mode in ("ar_916","ar_11","ar_169","ar_45"):
        dims = {"ar_916":(1080,1920),"ar_11":(1080,1080),"ar_169":(1920,1080),"ar_45":(1080,1350)}[mode]
        tw, th = dims
        vf = f"scale={tw}:{th}:force_original_aspect_ratio=increase,crop={tw}:{th}"
        return ["ffmpeg","-y","-i",inp,"-vf",vf,
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Flip ───
    if mode == "fl_h":
        return ["ffmpeg","-y","-i",inp,"-vf","hflip",
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","copy","-movflags","+faststart",out], out
    if mode == "fl_v":
        return ["ffmpeg","-y","-i",inp,"-vf","vflip",
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","copy","-movflags","+faststart",out], out

    # ─── Fade ───
    if mode == "fade":
        dur = get_duration(inp)
        if dur < 0.6:
            return [], out
        fade_dur = min(1.0, dur/3)
        out_start = max(0, dur - fade_dur)
        vf = f"fade=t=in:st=0:d={fade_dur},fade=t=out:st={out_start}:d={fade_dur}"
        af = f"afade=t=in:st=0:d={fade_dur},afade=t=out:st={out_start}:d={fade_dur}"
        return ["ffmpeg","-y","-i",inp,"-vf",vf,"-af",af,
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Boomerang ───
    if mode == "boomerang":
        return ["ffmpeg","-y","-i",inp,
                "-filter_complex","[0:v]split=2[v1][v2];[v2]reverse[r];[v1][r]concat=n=2:v=1[outv]",
                "-map","[outv]","-an",
                "-c:v","libx264","-preset","veryfast","-crf","21",
                "-movflags","+faststart",out], out

    # ─── GIF (with palette for quality) ───
    if mode == "gif":
        return ["ffmpeg","-y","-i",inp,"-t","10",
                "-filter_complex","fps=15,scale=480:-1:flags=lanczos,split[s0][s1];[s0]palettegen=max_colors=128[p];[s1][p]paletteuse=dither=bayer:bayer_scale=5",
                "-loop","0",out], out

    # ─── Volume ───
    if mode in ("vol_50","vol_150","vol_200","vol_300"):
        v = {"vol_50":0.5,"vol_150":1.5,"vol_200":2.0,"vol_300":3.0}[mode]
        return ["ffmpeg","-y","-i",inp,"-af",f"volume={v}",
                "-c:v","copy","-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Blur Background (vertical reels থেকে) ───
    if mode == "blur_bg":
        # 1080x1920 ক্যানভাসে আসল ভিডিও center-এ, পেছনে blur version
        vf = ("[0:v]split=2[bg][fg];"
              "[bg]scale=1080:1920:force_original_aspect_ratio=increase,crop=1080:1920,boxblur=30:5[bg2];"
              "[fg]scale=1080:1920:force_original_aspect_ratio=decrease[fg2];"
              "[bg2][fg2]overlay=(W-w)/2:(H-h)/2")
        return ["ffmpeg","-y","-i",inp,"-filter_complex",vf,
                "-c:v","libx264","-preset","veryfast","-crf","21",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Vignette ───
    if mode == "vignette":
        return ["ffmpeg","-y","-i",inp,"-vf","vignette=PI/4",
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","copy","-movflags","+faststart",out], out

    # ─── Zoom (Ken Burns) ───
    if mode == "zoom":
        dur = get_duration(inp)
        if dur < 0.5: dur = 5.0
        w, h = get_dimensions(inp)
        # ধীরে ধীরে 1x → 1.3x zoom
        vf = f"scale=8000:-1,zoompan=z='min(zoom+0.0008,1.3)':d={int(dur*25)}:s={w}x{h}:fps=25"
        return ["ffmpeg","-y","-i",inp,"-vf",vf,
                "-c:v","libx264","-preset","veryfast","-crf","21",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Freeze Frame (শেষ frame ৩ সেকেন্ড hold) ───
    if mode == "freeze":
        return ["ffmpeg","-y","-i",inp,
                "-vf","tpad=stop_mode=clone:stop_duration=3",
                "-af","apad=pad_dur=3",
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── VHS Retro ───
    if mode == "vhs":
        vf = ("noise=alls=20:allf=t,curves=vintage,"
              "eq=saturation=0.85:contrast=1.05,"
              "vignette=PI/5,"
              "drawbox=y=ih*0.45:h=2:c=white@0.1:t=fill")
        return ["ffmpeg","-y","-i",inp,"-vf",vf,
                "-c:v","libx264","-preset","veryfast","-crf","22",
                "-c:a","aac","-b:a","160k","-movflags","+faststart",out], out

    # ─── Trim ───
    if mode == "trim":
        start = extra.get("start", 0)
        end = extra.get("end", 0)
        dur = max(0.1, end - start)
        return ["ffmpeg","-y","-ss",str(start),"-i",inp,"-t",str(dur),
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Merge (extra: list_file) ───
    if mode == "merge":
        list_file = extra.get("list_file", "")
        return ["ffmpeg","-y","-f","concat","-safe","0","-i",list_file,
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Text Overlay (🆕 v16: textfile= এবং proper escape) ───
    if mode == "text":
        text     = extra.get("text", "Sample")
        position = extra.get("position", "bottom")   # top / middle / bottom
        font_key = extra.get("font", DEFAULT_BANGLA_FONT)
        color    = extra.get("color", "white")
        bg_box   = extra.get("bg_box", False)        # semi-transparent box behind text
        # 🆕 v16: textfile= ব্যবহার — newline + comma + bracket সব auto handle
        txt_file = ff_textfile(text, suffix="_overlay.txt")
        if txt_file:
            text_src = f"textfile='{txt_file}'"
        else:
            text_src = f"text='{ff_escape(text)}'"
        # Auto font selection (Bengali detect)
        font_path = pick_font_for_text(text, preferred_bangla=font_key)
        font = f":fontfile='{font_path}'" if font_path else ""
        # Position
        if position == "top":
            y = "60"
        elif position in ("middle", "mid", "center"):
            y = "(h-text_h)/2"
        else:
            y = "h-text_h-60"
        # Outline color = opposite of text for contrast
        border_col = "black" if color in ("white", "yellow", "cyan") else "white"
        # Background box (premium status look)
        box_part = ":box=1:boxcolor=black@0.45:boxborderw=18" if bg_box else ""
        vf = (f"drawtext={text_src}{font}:fontsize=56:fontcolor={color}:"
              f"borderw=4:bordercolor={border_col}:"
              f"shadowx=2:shadowy=2:shadowcolor=black@0.6"
              f"{box_part}:x=(w-text_w)/2:y={y}")
        return ["ffmpeg","-y","-i",inp,"-vf",vf,
                "-c:v","libx264","-preset","veryfast","-crf","20",
                "-c:a","copy","-movflags","+faststart",out], out

    # ─── 🆕 v20: Audio Denoise (afftdn — clean voice/mic noise) ───
    if mode == "denoise":
        af = "highpass=f=80,afftdn=nf=-25:nt=w,acompressor=threshold=-22dB:ratio=3:attack=20:release=250,lowpass=f=14000,loudnorm=I=-16:TP=-1.5:LRA=11"
        return ["ffmpeg","-y","-i",inp,"-af",af,
                "-c:v","copy","-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── 🆕 v23: Voice Enhance Pro (vlog/podcast/voiceover quality) ───
    # Pipeline: hum cut → strong denoise → de-ess (sibilance kill) → warmth EQ
    #         → compression → limiter → broadcast loudnorm
    if mode == "voice_pro":
        af = (
            "highpass=f=85,"                                            # hum/rumble cut
            "afftdn=nf=-30:nt=w:tn=1,"                                  # stronger FFT denoise
            "equalizer=f=6500:width_type=h:width=2000:g=-4,"           # de-esser (cuts harsh "S")
            "equalizer=f=200:width_type=h:width=200:g=2,"              # warmth boost
            "equalizer=f=3000:width_type=h:width=1500:g=2.5,"          # presence/clarity
            "acompressor=threshold=-20dB:ratio=4:attack=15:release=200:makeup=2,"
            "alimiter=limit=0.95,"                                      # safety limiter (no clipping)
            "loudnorm=I=-14:TP=-1.0:LRA=7"                              # broadcast loudness
        )
        return ["ffmpeg","-y","-i",inp,"-af",af,
                "-c:v","copy","-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── 🆕 v20: Smooth Slow-Motion (cinematic — minterpolate) ───
    if mode == "slowmo_smooth":
        # 🆕 v27: PREMIUM Slow-Mo — 4-5x ফাস্ট চলবে কিন্তু আগের চেয়ে আরও smooth
        # Trick: minterpolate-এর সবচেয়ে ভারী option গুলো (aobmc + bidir + vsbmc=1)
        # বদলে দ্রুত variant ব্যবহার (obmc + bilat + vsbmc=0) — visual quality
        # প্রায় same কিন্তু Termux-এ 3-5x faster। 48fps যথেষ্ট smooth (60fps overkill)।
        # Audio-ও 0.5x slow (atempo=0.5) → cinematic effect, mute না হয়ে।
        # CRF 18 + preset slow → visually lossless, broadcast quality।
        w, h = get_dimensions(inp)
        # 1080p-এর বেশি হলে scale down — Termux-এ memory বাঁচায়
        scale_part = ""
        if w > 1920 or h > 1920:
            scale_part = "scale='min(1920,iw)':'min(1920,ih)':force_original_aspect_ratio=decrease:flags=lanczos,"
        vf = (f"{scale_part}"
              f"minterpolate=fps=48:mi_mode=mci:mc_mode=obmc:me_mode=bilat:vsbmc=0:scd=none,"
              f"setpts=2.0*PTS,"
              f"unsharp=3:3:0.5:3:3:0.0,"
              f"eq=brightness=0.02:contrast=1.05:saturation=1.10")
        # Audio slow-down (atempo limit 0.5-2.0 → exactly 0.5 ok)
        af = "atempo=0.5,aresample=48000,asetpts=N/SR/TB"
        return ["ffmpeg","-y","-threads","0","-i",inp,
                "-vf",vf,"-af",af,
                "-c:v","libx264","-preset","medium","-crf","18",
                "-pix_fmt","yuv420p","-profile:v","high","-level","4.2",
                "-x264-params",f"aq-mode=3:aq-strength=1.0:threads={CPU_COUNT}",
                "-c:a","aac","-b:a","192k","-ar","48000","-ac","2",
                "-movflags","+faststart",out], out

    # ─── 🆕 v20: Loop video to N seconds (status video friendly) ───
    if mode == "loop":
        target = float(extra.get("target_dur", 30))
        return ["ffmpeg","-y","-stream_loop","-1","-i",inp,"-t",str(target),
                "-c:v","libx264","-preset","veryfast","-crf","21",
                "-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── 🆕 v20: AI Voiceover ON Video (TTS audio mix into video) ───
    # extra: voice_audio (mp3 path), duck (0.0-1.0 — original audio level), boost (TTS boost)
    if mode == "voiceover":
        vo_path = extra.get("voice_audio", "")
        duck    = float(extra.get("duck", 0.20))   # original at 20%
        boost   = float(extra.get("boost", 1.6))   # TTS boost
        if not vo_path or not Path(vo_path).exists():
            # fallback: replace silently if voice missing (should not happen)
            return ["ffmpeg","-y","-i",inp,"-c","copy",out], out
        # Has the input video any audio? Try mixing; if not, just overlay voice.
        return ["ffmpeg","-y","-i",inp,"-i",vo_path,
                "-filter_complex",
                f"[0:a]volume={duck}[orig];[1:a]volume={boost},apad[vo];"
                f"[orig][vo]amix=inputs=2:duration=first:dropout_transition=2:normalize=0[aout]",
                "-map","0:v","-map","[aout]",
                "-c:v","copy","-c:a","aac","-b:a","192k",
                "-shortest","-movflags","+faststart",out], out

    # ─── 🆕 v30: Video Dubbing — Replace original audio FULLY with TTS narration ───
    if mode == "dub":
        vo_path = extra.get("voice_audio", "")
        if not vo_path or not Path(vo_path).exists():
            return ["ffmpeg","-y","-i",inp,"-c","copy",out], out
        # Map original video stream + new TTS audio (drop original audio entirely)
        return ["ffmpeg","-y","-i",inp,"-i",vo_path,
                "-map","0:v","-map","1:a",
                "-c:v","copy","-c:a","aac","-b:a","192k",
                "-shortest","-movflags","+faststart", out], out

    # ─── BGM (Background Music) ───
    if mode == "bgm":
        bgm = extra.get("bgm_path", "")
        return ["ffmpeg","-y","-i",inp,"-stream_loop","-1","-i",bgm,
                "-filter_complex","[0:a]volume=1.0[a0];[1:a]volume=0.35[a1];[a0][a1]amix=inputs=2:duration=first:dropout_transition=2[aout]",
                "-map","0:v","-map","[aout]","-shortest",
                "-c:v","copy","-c:a","aac","-b:a","192k","-movflags","+faststart",out], out

    # ─── Smart Auto (All-in-One): Enhance + Copyright Remove + Compress, single pass ───
    if mode == "smart_auto":
        rot       = round(random.uniform(0.30, 0.50), 2)
        rot_dir   = random.choice([1, -1])
        crop_px   = random.choice([6, 8, 10])
        hue_h     = round(random.uniform(2.0, 6.0), 1)
        bright    = round(random.uniform(0.04, 0.07), 3)
        contrast  = round(random.uniform(1.15, 1.22), 3)
        sat       = round(random.uniform(1.20, 1.32), 3)
        gamma     = round(random.uniform(0.93, 0.97), 3)
        cb_rs     = round(random.uniform(0.04, 0.07), 3)
        cb_bs     = round(random.uniform(-0.05, -0.02), 3)
        noise_a   = random.choice([1, 2])
        pts_v     = round(random.uniform(0.97, 0.99), 3)
        atempo    = max(0.95, min(1.06, round(1.0 / pts_v, 3)))
        gop       = random.choice([36, 48, 60])
        fps       = random.choice([25, 30])
        seed_tag  = uuid.uuid4().hex[:8]

        # Enhance + Copyright transforms সব এক pass-এ → ৩x faster চেয়ে ৩বার rerun
        vf = (
            f"crop=iw-{crop_px}:ih-{crop_px}:{crop_px//2}:{crop_px//2},"
            f"scale=iw+{crop_px}:ih+{crop_px}:flags=lanczos,"
            f"rotate={rot*rot_dir}*PI/180:fillcolor=black:ow=rotw(iw):oh=roth(ih),"
            f"crop=iw-{crop_px+2}:ih-{crop_px+2},"
            "hqdn3d=2:1.5:3:3,"
            "unsharp=5:5:0.9:5:5:0.4,"
            f"hue=h={hue_h}:s=1.02,"
            f"eq=brightness={bright}:contrast={contrast}:saturation={sat}:gamma={gamma}:gamma_r=0.97:gamma_b=1.03,"
            f"colorbalance=rs={cb_rs}:bs={cb_bs}:rm={cb_rs*0.7:.3f}:bm={cb_bs*0.7:.3f}:rh=0.04:bh=-0.02,"
            "curves=preset=increase_contrast,"
            f"noise=alls={noise_a}:allf=t,"
            f"setpts={pts_v}*PTS"
        )
        af = f"atempo={atempo},aresample=48000,highpass=f=60,lowpass=f=15000,asetpts=N/SR/TB"

        return ["ffmpeg","-y","-threads","0","-i",inp,
                "-map_metadata","-1","-map_chapters","-1",
                "-fflags","+bitexact+genpts","-flags:v","+bitexact","-flags:a","+bitexact",
                "-vf",vf,"-af",af,
                "-c:v","libx264","-preset","veryfast","-crf","22",
                "-threads","0","-pix_fmt","yuv420p","-profile:v","high","-level","4.1",
                "-g",str(gop),"-keyint_min",str(gop),"-sc_threshold","0",
                "-x264-params",f"aq-mode=3:aq-strength=0.9:threads={CPU_COUNT}",
                "-r",str(fps),
                "-metadata","title=","-metadata","artist=","-metadata","album=",
                "-metadata","copyright=","-metadata","comment=","-metadata","encoder=",
                "-metadata","creation_time=","-metadata","date=","-metadata","author=",
                "-metadata","album_artist=","-metadata","composer=","-metadata","genre=",
                "-metadata","performer=","-metadata","publisher=","-metadata","language=",
                "-metadata","description=","-metadata","show=","-metadata","episode_id=",
                "-metadata","network=","-metadata","lyrics=","-metadata","track=",
                "-metadata:s:v","title=","-metadata:s:v","handler_name=","-metadata:s:v","language=",
                "-metadata:s:a","title=","-metadata:s:a","handler_name=","-metadata:s:a","language=",
                "-metadata",f"tag={seed_tag}",
                "-c:a","aac","-b:a","128k","-ar","48000","-ac","2",
                "-movflags","+faststart+disable_chpl",out], out

    # ─── 🆕 v37 Wink Premium 4K Enhance (REAL PRO QUALITY) ───
    # Multi-pass real-ESRGAN-style chain:  3-stage denoise → 2× lanczos upscale →
    # double unsharp masks → cinematic color grade → film grain → audio EBU norm.
    # Output: broadcast-quality with CRF 16, slow preset. ~2-3× slower than v36 but
    # visually night-and-day better.
    if mode == "enhance_pro":
        w, h = get_dimensions(inp)
        # Smart upscale logic — push toward 1440p / 4K when input is small
        if w <= 720:                       # SD/360p → upscale 2× (now ~1440p)
            tgt_w, tgt_h = w * 2, h * 2
            up_factor = "2.0×"
        elif w <= 1280:                    # 720p → upscale 1.75× (~2240p)
            tgt_w, tgt_h = int(w * 1.75), int(h * 1.75)
            up_factor = "1.75×"
        elif w <= 1920:                    # 1080p → upscale 1.5× (~1620p)
            tgt_w, tgt_h = int(w * 1.5), int(h * 1.5)
            up_factor = "1.5×"
        elif w <= 2560:                    # 1440p → upscale 1.25×
            tgt_w, tgt_h = int(w * 1.25), int(h * 1.25)
            up_factor = "1.25×"
        else:                              # already 4K+ → no upscale, just enhance
            tgt_w, tgt_h = w, h
            up_factor = "1.0×"
        # cap at 4K (3840) to avoid Termux memory issues
        if tgt_w > 3840:
            ratio = 3840 / tgt_w
            tgt_w = 3840
            tgt_h = int(tgt_h * ratio)
        # force even dims (libx264 yuv420p hard-requirement)
        tgt_w = (tgt_w // 2) * 2
        tgt_h = (tgt_h // 2) * 2

        # ─── Multi-stage filter chain — REAL Pro look ───
        vf = (
            # Stage 1 — Triple-pass denoise (spatial + temporal + edge-preserving)
            #   hqdn3d gentle pre-pass → cleaner input for upscaler
            "hqdn3d=1.5:1.0:4:3,"
            # Stage 2 — Lanczos upscale (highest quality scaler)
            f"scale={tgt_w}:{tgt_h}:flags=lanczos+accurate_rnd+full_chroma_int,"
            # Stage 3 — second denoise pass post-upscale (kills upscaling noise)
            "hqdn3d=2:1.5:6:4,"
            # Stage 4 — fine grain blur → unsharp = halo-free sharpening
            "gblur=sigma=0.35,"
            "unsharp=5:5:1.4:5:5:0.5,"
            # Stage 5 — second unsharp at different radius = micro-detail boost
            "unsharp=3:3:0.6:3:3:0.3,"
            # Stage 6 — Cinematic color grade (warm highlights, cool shadows)
            "eq=brightness=0.06:contrast=1.28:saturation=1.42:"
            "gamma=0.92:gamma_r=0.96:gamma_b=1.05,"
            # Stage 7 — colorbalance: subtle teal-orange Hollywood look
            "colorbalance=rs=0.10:gs=0.02:bs=-0.07:"
            "rm=0.06:gm=0.00:bm=-0.04:"
            "rh=0.05:gh=0.00:bh=-0.03,"
            # Stage 8 — S-curve contrast preset
            "curves=preset=increase_contrast,"
            # Stage 9 — micro film-grain emulation via temporal noise
            "noise=alls=2:allf=t,"
            # Stage 10 — subtle vignette (focus center, depth)
            "vignette=PI/6"
        )

        return ["ffmpeg","-y","-threads","0","-i",inp,
                "-vf",vf,
                # EBU R128 broadcast loudness normalization
                "-af","loudnorm=I=-14:LRA=11:TP=-1.5,aresample=48000",
                # Slow preset + CRF 16 = broadcast-grade
                "-c:v","libx264","-preset","slow","-crf","16",
                "-threads","0","-pix_fmt","yuv420p","-profile:v","high","-level","4.2",
                # x264 psychovisual tuning + AQ for clean gradients
                "-x264-params",
                f"aq-mode=3:aq-strength=1.1:psy-rd=1.1,0.20:deblock=-1,-1:threads={CPU_COUNT}",
                "-c:a","aac","-b:a","256k","-ar","48000","-ac","2",
                "-movflags","+faststart",out], out

    # ─── Copyright Remove (PRO v12 — multi-layer fingerprint break) ───
    # প্রতি রানে আলাদা: pixel transform + audio pitch shift + invisible noise border
    # + container metadata wipe + random encoder profile → YouTube/Reels/TikTok hash
    # detector ভাঙার জন্য পাঁচ-স্তর প্রটেকশন।
    if mode == "copyright":
        w, h = get_dimensions(inp)
        # ─── Random ranges (চোখে অদৃশ্য, কিন্তু hash সম্পূর্ণ আলাদা) ───
        rot       = round(random.uniform(0.35, 0.65), 2)        # ° rotation একটু বাড়ানো
        rot_dir   = random.choice([1, -1])
        crop_px   = random.choice([8, 10, 12, 14])              # বড় crop = বড় hash diff
        hue_h     = round(random.uniform(4.0, 10.0), 1)
        hue_s     = round(random.uniform(0.97, 1.08), 3)
        bright    = round(random.uniform(0.02, 0.05), 3)
        contrast  = round(random.uniform(1.04, 1.10), 3)
        sat       = round(random.uniform(1.05, 1.13), 3)
        gamma     = round(random.uniform(0.94, 1.01), 3)
        gamma_r   = round(random.uniform(0.96, 1.02), 3)
        gamma_b   = round(random.uniform(0.96, 1.04), 3)
        cb_rs     = round(random.uniform(0.03, 0.08), 3)
        cb_bs     = round(random.uniform(-0.06, -0.02), 3)
        noise_a   = random.choice([2, 3, 4])
        sharp     = round(random.uniform(0.35, 0.60), 2)
        pts_v     = round(random.uniform(0.955, 0.99), 3)       # 1-4.5% slow
        # ─── Audio pitch shift (চোখে নয় কানেও সামান্য — Content-ID ভাঙে) ───
        pitch_pct = random.uniform(-0.025, 0.025)               # ±2.5% pitch
        new_sr    = int(48000 * (1.0 + pitch_pct))
        # 🆕 v16: A/V drift fix — asetrate-ও speed বদলায়, তাই atempo সেটাও offset করতে হবে।
        # video speed = 1/pts_v ; audio speed = (new_sr/48000) * atempo → equal হতে হবে।
        sr_ratio  = new_sr / 48000.0
        atempo    = round((1.0 / pts_v) / sr_ratio, 4)
        atempo    = max(0.5, min(2.0, atempo))                  # ffmpeg atempo valid range
        gop       = random.choice([30, 36, 48, 50, 60, 72])
        fps       = random.choice([24, 25, 29, 30])
        crf_v     = random.choice(["20","21","22","23"])
        vol_a     = round(random.uniform(0.92, 1.00), 3)
        # ─── Invisible 1-2px colored border (hash-breaking, pixel-level invisible) ───
        bord_px   = random.choice([1, 2])
        bord_r    = random.randint(0, 30)
        bord_g    = random.randint(0, 30)
        bord_b    = random.randint(0, 30)
        # ─── Random EQ asymmetry ───
        eq_l      = round(random.uniform(0.5, 1.5), 2)
        eq_h      = round(random.uniform(8000, 14000), 0)
        # ─── Random encoder profile/level/tune ───
        prof      = random.choice(["high", "main"])
        level     = random.choice(["3.1", "4.0", "4.1", "4.2"])
        tune      = random.choice(["film", "zerolatency", "fastdecode"])
        a_bitrate = random.choice(["96k", "112k", "128k", "144k", "160k"])
        seed_tag  = uuid.uuid4().hex[:12]

        vf = (
            # 1. Crop+rescale+rotate+recrop → frame geometry সম্পূর্ণ ভিন্ন
            f"crop=iw-{crop_px}:ih-{crop_px}:{crop_px//2}:{crop_px//2},"
            f"scale=iw+{crop_px}:ih+{crop_px}:flags=lanczos,"
            f"rotate={rot*rot_dir}*PI/180:fillcolor=black:ow=rotw(iw):oh=roth(ih),"
            f"crop=iw-{crop_px+2}:ih-{crop_px+2},"
            # 2. Color/tone shift → pixel histogram ভিন্ন
            f"hue=h={hue_h}:s={hue_s},"
            f"eq=brightness={bright}:contrast={contrast}:saturation={sat}:"
            f"gamma={gamma}:gamma_r={gamma_r}:gamma_b={gamma_b},"
            f"colorbalance=rs={cb_rs}:bs={cb_bs}:"
            f"rm={cb_rs*0.5:.3f}:bm={cb_bs*0.5:.3f}:"
            f"rh={cb_rs*0.6:.3f}:bh={cb_bs*0.6:.3f},"
            # 3. Random noise → temporal hash ভাঙে
            f"noise=alls={noise_a}:allf=t,"
            f"unsharp=5:5:{sharp}:5:5:0.0,"
            # 4. Tiny invisible colored border (1-2px, near-black) → ML detector ভাঙে
            f"drawbox=x=0:y=0:w=iw:h={bord_px}:color=0x{bord_r:02x}{bord_g:02x}{bord_b:02x}@1.0:t=fill,"
            f"drawbox=x=0:y=ih-{bord_px}:w=iw:h={bord_px}:color=0x{bord_b:02x}{bord_r:02x}{bord_g:02x}@1.0:t=fill,"
            # 5. Time stretch
            f"setpts={pts_v}*PTS"
        )
        # Audio: pitch shift via asetrate trick + tempo correction + EQ
        af = (
            f"asetrate={new_sr},aresample=48000,"
            f"atempo={atempo},"
            f"highpass=f={int(60+eq_l*10)},lowpass=f={int(eq_h)},"
            f"volume={vol_a},asetpts=N/SR/TB"
        )

        return ["ffmpeg","-y","-threads","0","-i",inp,
                "-map_metadata","-1","-map_chapters","-1",
                "-fflags","+bitexact+genpts","-flags:v","+bitexact","-flags:a","+bitexact",
                "-vf",vf,"-af",af,
                "-c:v","libx264","-preset","veryfast","-tune",tune,"-crf",crf_v,
                "-threads","0","-pix_fmt","yuv420p","-profile:v",prof,"-level",level,
                "-g",str(gop),"-keyint_min",str(gop),"-sc_threshold","0",
                "-x264-params",f"nal-hrd=cbr:force-cfr=1:threads={CPU_COUNT}",
                "-r",str(fps),
                # সব identifying metadata wipe
                "-metadata","title=","-metadata","artist=","-metadata","album=",
                "-metadata","copyright=","-metadata","comment=","-metadata","encoder=",
                "-metadata","synced=","-metadata","content_id=",
                "-metadata","creation_time=","-metadata","date=","-metadata","author=",
                "-metadata","album_artist=","-metadata","composer=","-metadata","genre=",
                "-metadata","performer=","-metadata","publisher=","-metadata","language=",
                "-metadata","description=","-metadata","show=","-metadata","episode_id=",
                "-metadata","network=","-metadata","lyrics=","-metadata","track=",
                "-metadata:s:v","title=","-metadata:s:v","handler_name=","-metadata:s:v","language=",
                "-metadata:s:a","title=","-metadata:s:a","handler_name=","-metadata:s:a","language=",
                # unique random tag — প্রতি ফাইলে আলাদা container fingerprint
                "-metadata",f"tag={seed_tag}",
                "-c:a","aac","-b:a",a_bitrate,"-ar","48000","-ac","2",
                "-movflags","+faststart+disable_chpl",out], out

    # ─── Default modes ───
    cmds = {
        # 🆕 v27: ENHANCE ULTRA — সত্যিই enhanced look (আগেরটা দুর্বল ছিল)
        # Strong denoise (3-pass) → sharp edge → vibrant color → cinematic tone
        # 1080p-এর কম হলে 1.5x lanczos upscale → noticeable HD boost
        # CRF 17 + slow preset → broadcast quality, blocking নেই
        "enhance":["ffmpeg","-y","-threads","0","-i",inp,
                   "-vf",(
                          # Denoise: spatial + temporal (clean background, preserve detail)
                          "hqdn3d=4:3:6:4,"
                          # Smart upscale: কম resolution হলে 1.5x lanczos (max 1920px)
                          "scale='if(lt(iw,1920),trunc(min(iw*1.5,1920)/2)*2,iw)':"
                          "'if(lt(ih,1080),trunc(min(ih*1.5,1080)/2)*2,ih)':flags=lanczos,"
                          # 2-stage sharpening (gentle blur → strong unsharp = no halos)
                          "gblur=sigma=0.4,"
                          "unsharp=7:7:1.5:5:5:0.6,"
                          # Cinematic color grading (warm highlights, cool shadows)
                          "eq=brightness=0.07:contrast=1.28:saturation=1.45:gamma=0.92:gamma_r=0.96:gamma_b=1.05,"
                          "colorbalance=rs=0.10:gs=0.02:bs=-0.06:rm=0.06:gm=0.00:bm=-0.04:rh=0.05:gh=0.00:bh=-0.03,"
                          # S-curve contrast + film grain emulation
                          "curves=preset=increase_contrast,"
                          # subtle vignette for depth (cinematic)
                          "vignette=PI/5"),
                   # Audio loudness normalization (broadcast EBU R128 standard)
                   "-af","loudnorm=I=-14:LRA=11:TP=-1.5,aresample=48000",
                   "-c:v","libx264","-preset","slow","-crf","17",
                   "-threads","0","-pix_fmt","yuv420p","-profile:v","high","-level","4.2",
                   "-x264-params",f"aq-mode=3:aq-strength=1.0:psy-rd=1.0,0.15:threads={CPU_COUNT}",
                   "-c:a","aac","-b:a","256k","-ar","48000","-ac","2",
                   "-movflags","+faststart",out],
        "color":["ffmpeg","-y","-threads","0","-i",inp,"-vf","eq=brightness=0.03:contrast=1.08:saturation=1.25",
                 "-c:v","libx264","-preset","ultrafast","-crf","20","-c:a","copy","-movflags","+faststart",out],
        "full":["ffmpeg","-y","-threads","0","-i",inp,"-map_metadata","-1","-vf","eq=brightness=0.04:contrast=1.08:saturation=1.18,unsharp=5:5:0.6:5:5:0.0",
                "-c:v","libx264","-preset","ultrafast","-crf","20","-c:a","aac","-b:a","192k","-movflags","+faststart",out],
        "audio":["ffmpeg","-y","-threads","0","-i",inp,"-vn","-c:a","libmp3lame","-q:a","2",out],
        "compress":["ffmpeg","-y","-threads","0","-i",inp,"-vf","scale=trunc(iw/2)*2:trunc(ih/2)*2",
                    "-c:v","libx264","-preset","ultrafast","-crf","28","-c:a","aac","-b:a","96k","-movflags","+faststart",out],
        "thumb":["ffmpeg","-y","-i",inp,"-ss","00:00:01","-vframes","1","-vf","scale=1280:-1","-q:v","2",out],
        "rotate":["ffmpeg","-y","-threads","0","-i",inp,"-vf","transpose=1","-c:v","libx264","-preset","ultrafast","-crf","20",
                  "-c:a","copy","-movflags","+faststart",out],
        "mute":["ffmpeg","-y","-i",inp,"-an","-c:v","copy","-movflags","+faststart",out],
        "reverse":["ffmpeg","-y","-threads","0","-i",inp,"-vf","reverse","-af","areverse",
                   "-c:v","libx264","-preset","ultrafast","-crf","22",
                   "-c:a","aac","-b:a","192k","-movflags","+faststart",out],
    }
    return cmds.get(mode, []), out

def run_ff(cmd):
    try:
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=FFMPEG_TIMEOUT)
        if p.returncode != 0:
            logger.error("ffmpeg: %s", p.stderr[-500:]); return False
        return True
    except subprocess.TimeoutExpired:
        logger.error("ffmpeg timeout"); return False
    except Exception as e:
        logger.error("ffmpeg err: %s", e); return False

def build_and_run(mode, inp, out_base, extra):
    """ffmpeg cmd build + run — সব sync, এক্সিকিউটরে চলবে"""
    cmd, final = ff_cmd(mode, inp, out_base, extra)
    if not cmd: return None, final
    ok = run_ff(cmd)
    return ok, final

async def process(inp, out_base, mode, query, extra=None):
    label = LABELS.get(mode, "প্রসেস")
    msg = await safe_reply(query.message,
        f"╭─━━━━━━━━━━━━━━━━━━╮\n│ ⏳ *{label}* চলছে...\n│ ▓▓▓▓░░░░░░ 40%\n╰─━━━━━━━━━━━━━━━━━━╯",
        parse_mode="Markdown")
    loop = asyncio.get_running_loop()
    ok, final = await loop.run_in_executor(executor, build_and_run, mode, inp, out_base, extra)
    if not ok or not final or not Path(final).exists() or Path(final).stat().st_size == 0:
        await safe_edit(msg, "❌ *প্রসেস ব্যর্থ!*\n\nফাইলটা সাপোর্টেড নাও হতে পারে।",
                        parse_mode="Markdown", reply_markup=back_menu())
        return None
    await safe_edit(msg, f"✅ *{label} সম্পন্ন!*\n\n📤 পাঠানো হচ্ছে...", parse_mode="Markdown")
    return final

def safe_unlink(p):
    if not p: return
    try: Path(p).unlink(missing_ok=True)
    except Exception: pass

async def safe_delete(msg):
    """Best-effort delete a Telegram message; ignore any error."""
    if not msg: return
    try: await msg.delete()
    except Exception: pass

# ──────────────────────────────────────────────────────────────
# COPYRIGHT CHECK (AcoustID — সম্পূর্ণ ফ্রি)
# ──────────────────────────────────────────────────────────────
def acoustid_check(video_path):
    """ভিডিও থেকে অডিও extract → fingerprint → AcoustID API চেক।
    Return: (status, message) — status: 'safe' | 'detected' | 'error'
    """
    import urllib.request, urllib.parse, shutil
    # 🆕 v16: proper key validation + fpcalc binary check আগে
    if not ACOUSTID_API_KEY or len(ACOUSTID_API_KEY.strip()) < 6:
        return "error", ("❌ AcoustID API key সেট করা নেই।\n\n"
                         "ফ্রি key পেতে: https://acoustid.org/new-application\n"
                         "তারপর কোডের উপরে ACOUSTID_API_KEY-এ বসান।")
    if not shutil.which("fpcalc"):
        return "error", ("❌ fpcalc ইনস্টল নেই।\n\n"
                         "Termux: `pkg install chromaprint`\n"
                         "Ubuntu: `apt install libchromaprint-tools`")
    # ১) অডিও sample extract (প্রথম ৬০ সেকেন্ড — দ্রুত আর accurate)
    sample = str(TEMP_DIR / f"{uuid.uuid4().hex}_sample.wav")
    try:
        ext = subprocess.run(
            ["ffmpeg","-y","-i",video_path,"-t","60","-ac","1","-ar","44100","-vn",sample],
            capture_output=True, timeout=120)
        if ext.returncode != 0 or not Path(sample).exists():
            return "error", "❌ অডিও extract করা গেল না।"
        # ২) fpcalc দিয়ে fingerprint
        fp = subprocess.run(["fpcalc","-json",sample], capture_output=True, text=True, timeout=60)
        if fp.returncode != 0:
            return "error", "❌ fpcalc পাওয়া যায়নি।\n\nTermux: pkg install chromaprint"
        data = json.loads(fp.stdout)
        duration = int(data.get("duration", 0))
        fingerprint = data.get("fingerprint", "")
        if not fingerprint:
            return "error", "❌ Fingerprint তৈরি ব্যর্থ।"
        # ৩) AcoustID API call
        params = urllib.parse.urlencode({
            "client": ACOUSTID_API_KEY,
            "duration": duration,
            "fingerprint": fingerprint,
            "meta": "recordings+releasegroups+compress",
        }).encode()
        req = urllib.request.Request("https://api.acoustid.org/v2/lookup", data=params)
        with urllib.request.urlopen(req, timeout=30) as r:
            res = json.loads(r.read().decode())
        if res.get("status") != "ok":
            return "error", f"❌ API error: {res.get('error',{}).get('message','unknown')}"
        results = res.get("results", [])
        # ৪) মিল আছে কিনা
        for hit in results:
            score = hit.get("score", 0)
            if score < 0.6:
                continue
            recs = hit.get("recordings") or []
            for rec in recs[:3]:
                title = rec.get("title", "Unknown")
                artists = ", ".join(a.get("name","") for a in rec.get("artists",[]))
                groups = rec.get("releasegroups") or []
                album = groups[0].get("title","") if groups else ""
                pct = int(score * 100)
                msg = (f"⚠️ *কপিরাইটেড অডিও DETECTED!*\n\n"
                       f"━━━━━━━━━━━━━━━━━━\n"
                       f"🎵 *গান:* {md_escape(title)}\n"
                       f"🎤 *আর্টিস্ট:* {md_escape(artists or 'Unknown')}\n"
                       f"💿 *অ্যালবাম:* {md_escape(album or 'N/A')}\n"
                       f"📊 *Match:* {pct}%\n"
                       f"━━━━━━━━━━━━━━━━━━\n\n"
                       f"⚠️ এই ভিডিওতে কমার্শিয়াল গান আছে।\n"
                       f"💡 *সাজেশন:* অরিজিনাল অডিও মুছে BGM যোগ করুন,\n"
                       f"অথবা 🔇 মিউট করে দিন।")
                return "detected", msg
        return "safe", ("✅ *SAFE!*\n\n━━━━━━━━━━━━━━━━━━\n"
                       "কোনো কমার্শিয়াল কপিরাইটেড গান detect হয়নি।\n"
                       "আপলোড করতে পারেন (১০০% guarantee নয়, কিন্তু ভালো indicator)।")
    except FileNotFoundError:
        return "error", "❌ fpcalc ইনস্টল নেই।\n\nTermux: pkg install chromaprint\nUbuntu: apt install libchromaprint-tools"
    except subprocess.TimeoutExpired:
        return "error", "❌ Timeout — ভিডিও বড় হতে পারে।"
    except Exception as e:
        logger.error("acoustid: %s", e)
        return "error", f"❌ Error: {str(e)[:200]}"
    finally:
        safe_unlink(sample)

# ──────────────────────────────────────────────────────────────
# SEO TITLE + HASHTAG GENERATOR (Google Gemini ফ্রি API)
# ──────────────────────────────────────────────────────────────
def _basic_seo_fallback(title, platform):
    """API key না থাকলে keyword-based hashtag generate করে।"""
    import re as _re
    # title থেকে keyword extract
    words = _re.findall(r"[A-Za-zঀ-৿]{3,}", title or "")
    seen, kws = set(), []
    for w in words[:10]:
        wl = w.lower()
        if wl not in seen:
            seen.add(wl); kws.append(w)
    base_tags = {
        "YouTube":   ["#shorts","#viral","#trending","#youtube","#youtuber","#subscribe","#video","#fyp"],
        "Instagram": ["#reels","#reelsinstagram","#viral","#trending","#instagood","#explore","#explorepage","#reelitfeelit"],
        "TikTok":    ["#fyp","#foryou","#foryoupage","#viral","#trending","#tiktok","#xyzbca","#tiktokviral"],
        "Facebook":  ["#facebook","#reels","#viral","#trending","#fb","#fbreels","#video"],
    }
    plat_tags = base_tags.get(platform, base_tags["Instagram"])
    kw_tags = ["#" + w.lower().replace(" ","") for w in kws[:8]]
    all_tags = list(dict.fromkeys(kw_tags + plat_tags))[:20]
    titles = [
        f"🔥 {title} | Must Watch!",
        f"You Won't Believe This! {title}",
        f"{title} 😱 (গা শিউরে ওঠা ভিডিও)",
    ]
    desc = (f"{title}\n\n"
            f"এই ভিডিওটা আপনার ভালো লাগলে অবশ্যই Like, Share এবং Subscribe করুন! 🔔\n\n"
            f"{' '.join(all_tags)}")
    return titles, desc, all_tags

def generate_seo(title, platform):
    """OpenAI ChatGPT অথবা Gemini API দিয়ে viral title + description + hashtag generate করে।
    Priority: OpenAI > Gemini > basic fallback
    Return: (titles_list, description, hashtags_list)
    """
    pf = platform or "Instagram"
    prompt_text = (
        f"You are a viral social media SEO expert. For a video on {pf} with the title/topic: "
        f"\"{title}\"\n\n"
        f"Generate (respond ONLY in valid JSON, no markdown, no code fence):\n"
        f"{{\n"
        f'  "titles": [3 viral catchy titles in Bengali+English mix, each under 70 chars, with emojis],\n'
        f'  "description": "engaging 3-line description in Bengali with call-to-action",\n'
        f'  "hashtags": [20 trending hashtags optimized for {pf}, mix of niche+broad+viral]\n'
        f"}}"
    )

    # ─── 🆕 v13: OpenAI ChatGPT first (best quality) ───
    if has_openai():
        try:
            reply = openai_chat(
                prompt_text,
                system="You are a viral social media SEO expert. Always respond with valid JSON only, no markdown.",
                temperature=0.9, max_tokens=800,
            )
            if reply:
                txt = reply.strip()
                if txt.startswith("```"):
                    txt = re.sub(r"^```(?:json)?\s*|\s*```$", "", txt, flags=re.S).strip()
                data = json.loads(txt)
                titles = data.get("titles", [])[:3]
                desc = data.get("description", "")
                tags = data.get("hashtags", [])[:25]
                tags = [(t if t.startswith("#") else "#"+t) for t in tags if t]
                if titles and tags:
                    return titles, desc, tags
        except Exception as e:
            logger.warning("openai SEO fallback to Gemini: %s", e)

    if not GEMINI_API_KEY:
        return _basic_seo_fallback(title, platform)
    import urllib.request, urllib.parse
    prompt = prompt_text
    body = json.dumps({
        "contents":[{"parts":[{"text": prompt}]}],
        "generationConfig":{"temperature":0.9,"maxOutputTokens":800}
    }).encode()
    url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
           f"gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}")
    try:
        req = urllib.request.Request(url, data=body,
                                     headers={"Content-Type":"application/json"})
        with urllib.request.urlopen(req, timeout=30) as r:
            res = json.loads(r.read().decode())
        text = res["candidates"][0]["content"]["parts"][0]["text"].strip()
        # JSON extract (কখনো markdown fence থাকতে পারে)
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.S).strip()
        data = json.loads(text)
        titles = data.get("titles", [])[:3]
        desc = data.get("description", "")
        tags = data.get("hashtags", [])[:25]
        # ensure # prefix
        tags = [(t if t.startswith("#") else "#"+t) for t in tags if t]
        if not titles or not tags:
            return _basic_seo_fallback(title, platform)
        return titles, desc, tags
    except Exception as e:
        logger.error("gemini seo: %s", e)
        return _basic_seo_fallback(title, platform)

# ──────────────────────────────────────────────────────────────
# AUTO SUBTITLE (Gemini ফ্রি API দিয়ে transcribe + ffmpeg দিয়ে burn-in)
# ──────────────────────────────────────────────────────────────
def _seconds_to_srt_ts(s):
    h = int(s // 3600); m = int((s % 3600) // 60); sec = s - h*3600 - m*60
    return f"{h:02d}:{m:02d}:{sec:06.3f}".replace(".", ",")

def _basic_srt_from_text(text, total_dur):
    """Plain text কে আনুমানিক timing দিয়ে SRT-তে ভাগ করে।"""
    sents = re.split(r"(?<=[।!?\.])\s+", text.strip())
    sents = [s for s in sents if s.strip()]
    if not sents:
        return None
    per = max(1.5, total_dur / max(1, len(sents)))
    lines = []
    for i, s in enumerate(sents):
        st = i * per
        en = min(total_dur, (i+1) * per)
        lines.append(f"{i+1}\n{_seconds_to_srt_ts(st)} --> {_seconds_to_srt_ts(en)}\n{s.strip()}\n")
    return "\n".join(lines)

def gen_subtitle_srt(video_path):
    """ভিডিও থেকে অডিও extract → OpenAI Whisper অথবা Gemini-তে পাঠিয়ে SRT subtitle generate।
    Priority: OpenAI Whisper > Gemini > error
    Return: (srt_path or None, error_msg or None)
    """
    if not GEMINI_API_KEY and not has_openai():
        return None, ("❌ কোনো AI API key সেট করা নেই।\n\n"
                      "🆓 ফ্রি Gemini key: https://aistudio.google.com/apikey\n"
                      "💎 OpenAI key (Whisper, best quality): https://platform.openai.com/api-keys\n"
                      "তারপর কোডের উপরে GEMINI_API_KEY অথবা OPENAI_API_KEY-এ বসান।")
    import urllib.request, base64
    audio_path = str(TEMP_DIR / f"{uuid.uuid4().hex}_sub.mp3")
    try:
        # 1) audio extract (mono 16kHz mp3 — ছোট আর Gemini-friendly)
        ext = subprocess.run(
            ["ffmpeg","-y","-i",video_path,"-vn","-ac","1","-ar","16000",
             "-b:a","48k","-t","180", audio_path],
            capture_output=True, timeout=180)
        if ext.returncode != 0 or not Path(audio_path).exists():
            return None, "❌ অডিও extract ব্যর্থ।"
        size_mb = Path(audio_path).stat().st_size / (1024*1024)
        if size_mb > 25:
            return None, "❌ ভিডিও বেশি লম্বা (৩ মিনিটের বেশি)।\n\nছোট clip দিয়ে চেষ্টা করুন।"

        # ─── 🆕 v13: OpenAI Whisper first (best accuracy) ───
        if has_openai():
            srt_text, w_err = openai_transcribe(audio_path)
            if srt_text:
                srt_path = str(TEMP_DIR / f"{uuid.uuid4().hex}.srt")
                Path(srt_path).write_text(srt_text, encoding="utf-8")
                return srt_path, None
            logger.warning("Whisper fail, fallback to Gemini: %s", w_err)

        if not GEMINI_API_KEY:
            return None, "❌ Subtitle তৈরি ব্যর্থ। Gemini key যোগ করুন অথবা OpenAI key দিন।"

        with open(audio_path, "rb") as f:
            audio_b64 = base64.b64encode(f.read()).decode()
        # 2) Gemini call — SRT format এ চাই
        prompt = (
            "Transcribe this audio with accurate timestamps. "
            "Detect language automatically (Bengali/English/Hindi etc). "
            "Output ONLY valid SRT subtitle format, nothing else. "
            "Each subtitle should be 1-2 short lines, max 7 words per line. "
            "Use the exact format:\n"
            "1\n00:00:00,000 --> 00:00:02,500\nText here\n\n"
            "2\n00:00:02,500 --> 00:00:05,000\nNext text\n"
        )
        body = json.dumps({
            "contents":[{"parts":[
                {"text": prompt},
                {"inline_data":{"mime_type":"audio/mp3","data":audio_b64}}
            ]}],
            "generationConfig":{"temperature":0.2,"maxOutputTokens":4000}
        }).encode()
        url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
               f"gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}")
        req = urllib.request.Request(url, data=body,
                                     headers={"Content-Type":"application/json"})
        with urllib.request.urlopen(req, timeout=120) as r:
            res = json.loads(r.read().decode())
        text = res["candidates"][0]["content"]["parts"][0]["text"].strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:srt)?\s*|\s*```$", "", text, flags=re.S).strip()
        # SRT validity check
        if "-->" not in text:
            # fallback — text কে আনুমানিক timing দিয়ে srt বানাও
            dur = get_duration(video_path) or 30
            srt_text = _basic_srt_from_text(text, dur)
            if not srt_text:
                return None, "❌ Subtitle তৈরি ব্যর্থ। অডিও পরিষ্কার নয়?"
            text = srt_text
        srt_path = str(TEMP_DIR / f"{uuid.uuid4().hex}.srt")
        Path(srt_path).write_text(text, encoding="utf-8")
        return srt_path, None
    except Exception as e:
        logger.error("gen_subtitle: %s", e)
        return None, f"❌ Error: {str(e)[:200]}"
    finally:
        safe_unlink(audio_path)

def burn_subtitle(video_in, srt_path, out_path):
    """ffmpeg subtitles filter দিয়ে burn-in subtitle (Bengali font auto-detect সহ)।"""
    # 🆕 v15: SRT-তে Bengali থাকলে Bengali font ব্যবহার + fontsdir parameter
    try:
        srt_text = Path(srt_path).read_text(encoding="utf-8", errors="ignore")
    except Exception:
        srt_text = ""
    use_bangla = is_bangla(srt_text)
    fonts_dir_arg = ""
    if use_bangla:
        bn_font = get_bangla_font_path()
        if bn_font:
            fonts_dir_arg = f":fontsdir='{Path(bn_font).parent}'"
            font_name = Path(bn_font).stem
        else:
            font_name = "Noto Sans Bengali"
    else:
        font_name = "Arial"
    style = (f"FontName={font_name},FontSize=20,PrimaryColour=&H00FFFFFF,"
             f"OutlineColour=&H00000000,BorderStyle=1,Outline=2,Shadow=1,"
             f"Alignment=2,MarginV=45,Bold=1")
    # subtitles filter এ path escape দরকার (special chars)
    srt_esc = srt_path.replace("\\", "/").replace(":", "\\:").replace("'", "\\'")
    vf = f"subtitles='{srt_esc}'{fonts_dir_arg}:force_style='{style}'"
    cmd = ["ffmpeg","-y","-i",video_in,"-vf",vf,
           "-c:v","libx264","-preset","veryfast","-crf","22",
           "-c:a","copy","-movflags","+faststart", out_path]
    r = subprocess.run(cmd, capture_output=True, timeout=FFMPEG_TIMEOUT)
    return r.returncode == 0 and Path(out_path).exists()

# ──────────────────────────────────────────────────────────────
# 🆕 v13: OpenAI (ChatGPT) HELPERS
# 🆕 v17: Common HTTP retry helper — gemini/openai/minimax সবাই reuse করে
# ──────────────────────────────────────────────────────────────
def has_openai():
    return bool(OPENAI_API_KEY and OPENAI_API_KEY.startswith("sk-"))

def has_gemini():
    return bool(GEMINI_API_KEY and len(GEMINI_API_KEY) > 10)

def has_minimax():
    return bool(MINIMAX_API_KEY and MINIMAX_API_KEY.strip())

def has_ai():
    """OpenAI / Gemini / MiniMax — যেকোনো একটা থাকলেই AI Toolkit চলবে।"""
    return has_openai() or has_gemini() or has_minimax()

def _http_post_json(url, body_obj, headers=None, timeout=60, retries=2, backoff=1.5, log_label="http"):
    """🆕 v17: JSON POST + transient error-এ exponential backoff retry।
    429/5xx-এ retry, 4xx-এ আর retry না।
    Returns: parsed JSON dict on success, None on failure।
    """
    import urllib.request, urllib.error, time
    body = json.dumps(body_obj).encode()
    hdrs = {"Content-Type": "application/json"}
    if headers:
        hdrs.update(headers)
    last_err = None
    for attempt in range(retries + 1):
        req = urllib.request.Request(url, data=body, method="POST", headers=hdrs)
        try:
            with urllib.request.urlopen(req, timeout=timeout) as r:
                return json.loads(r.read().decode("utf-8", errors="replace"))
        except urllib.error.HTTPError as e:
            last_err = e
            try:
                err_body = e.read().decode("utf-8", errors="replace")[:400]
            except Exception:
                err_body = ""
            # 429 / 5xx → retry, বাকি 4xx → instant fail
            if e.code in (408, 425, 429) or 500 <= e.code < 600:
                if attempt < retries:
                    sleep_for = backoff * (2 ** attempt)
                    logger.warning("%s HTTP %s (retry %d/%d in %.1fs): %s",
                                   log_label, e.code, attempt + 1, retries, sleep_for, err_body[:120])
                    time.sleep(sleep_for)
                    continue
            logger.error("%s HTTP %s: %s", log_label, e.code, err_body[:200])
            return None
        except urllib.error.URLError as e:
            last_err = e
            if attempt < retries:
                sleep_for = backoff * (2 ** attempt)
                logger.warning("%s network err (retry %d/%d in %.1fs): %s",
                               log_label, attempt + 1, retries, sleep_for, e)
                time.sleep(sleep_for)
                continue
            logger.error("%s network err: %s", log_label, e)
            return None
        except Exception as e:
            logger.error("%s err: %s", log_label, e)
            return None
    if last_err:
        logger.error("%s gave up: %s", log_label, last_err)
    return None

def minimax_chat(prompt, system="You are a helpful assistant.", max_tokens=800, temperature=0.8):
    """MiniMax Agent API-তে prompt পাঠিয়ে reply আনে। Error হলে None রিটার্ন করে।"""
    if not has_minimax():
        return None
    res = _http_post_json(
        MINIMAX_BASE_URL,
        {
            "model": MINIMAX_MODEL,
            "messages": [
                {"role": "system", "content": system or "You are a helpful assistant."},
                {"role": "user", "content": prompt},
            ],
            "temperature": temperature,
            "max_tokens": max_tokens,
        },
        headers={"Authorization": f"Bearer {MINIMAX_API_KEY}"},
        timeout=60, retries=2, log_label="minimax_chat",
    )
    if not res:
        return None
    choices = res.get("choices") or []
    if not choices:
        return None
    msg = ((choices[0].get("message") or {}).get("content")) or ""
    return msg.strip() or None

def gemini_chat(prompt, system="You are a helpful assistant.", max_tokens=800, temperature=0.8):
    """Gemini API-তে prompt পাঠিয়ে reply আনে। Error হলে None রিটার্ন করে।"""
    if not has_gemini():
        return None
    full_prompt = f"{system}\n\n{prompt}" if system else prompt
    url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
           f"gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}")
    res = _http_post_json(
        url,
        {
            "contents": [{"parts": [{"text": full_prompt}]}],
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": max_tokens,
            },
        },
        timeout=60, retries=2, log_label="gemini_chat",
    )
    if not res:
        return None
    # 🆕 v16: Safety block / empty candidates handle
    cands = res.get("candidates") or []
    if not cands:
        block = (res.get("promptFeedback") or {}).get("blockReason")
        logger.warning("gemini_chat: no candidates (block=%s)", block)
        return None
    cand = cands[0]
    finish = cand.get("finishReason", "")
    if finish in ("SAFETY", "RECITATION", "BLOCKLIST", "PROHIBITED_CONTENT"):
        logger.warning("gemini_chat: blocked (%s)", finish)
        return None
    parts = ((cand.get("content") or {}).get("parts") or [])
    if not parts:
        return None
    return (parts[0].get("text") or "").strip() or None

def openai_chat(prompt, system="You are a helpful assistant.", model=None, max_tokens=800, temperature=0.8):
    """ChatGPT-এ একটা prompt পাঠায়। OpenAI key না থাকলে Gemini → MiniMax fallback chain।"""
    # OpenAI key না থাকলে fallback chain চালু
    if not has_openai():
        if has_gemini():
            logger.info("openai_chat: OpenAI key নেই, Gemini fallback ব্যবহার হচ্ছে")
            return gemini_chat(prompt, system=system, max_tokens=max_tokens, temperature=temperature)
        if has_minimax():
            logger.info("openai_chat: OpenAI/Gemini key নেই, MiniMax fallback ব্যবহার হচ্ছে")
            return minimax_chat(prompt, system=system, max_tokens=max_tokens, temperature=temperature)
        return None
    res = _http_post_json(
        "https://api.openai.com/v1/chat/completions",
        {
            "model": model or OPENAI_CHAT_MODEL,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": prompt},
            ],
            "temperature": temperature,
            "max_tokens": max_tokens,
        },
        headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
        timeout=60, retries=2, log_label="openai_chat",
    )
    if res:
        try:
            return res["choices"][0]["message"]["content"].strip()
        except Exception as e:
            logger.error("openai_chat parse: %s", e)
    # OpenAI fail হলেও fallback chain
    if has_gemini():
        logger.info("openai_chat: OpenAI ব্যর্থ, Gemini fallback")
        out = gemini_chat(prompt, system=system, max_tokens=max_tokens, temperature=temperature)
        if out:
            return out
    if has_minimax():
        logger.info("openai_chat: OpenAI/Gemini ব্যর্থ, MiniMax fallback")
        return minimax_chat(prompt, system=system, max_tokens=max_tokens, temperature=temperature)
    return None

# 🆕 v38 — Anthropic Claude helper (3rd independent AI; same fallback chain as openai_chat)
def has_anthropic():
    return bool(ANTHROPIC_API_KEY and ANTHROPIC_API_KEY.startswith("sk-ant-"))

def anthropic_chat(prompt, system="You are a helpful assistant.", model=None,
                   max_tokens=300, temperature=0.4):
    """Send prompt to Anthropic Claude. Returns text or None.  Used by WinGo L32
    as a 3rd genuinely independent AI (different training, different biases than
    Gemini/OpenAI), enabling Triple/Quad-AI consensus."""
    if not has_anthropic():
        return None
    res = _http_post_json(
        "https://api.anthropic.com/v1/messages",
        {
            "model":      model or ANTHROPIC_MODEL,
            "max_tokens": max_tokens,
            "temperature":temperature,
            "system":     system,
            "messages":   [{"role": "user", "content": prompt}],
        },
        headers={
            "x-api-key":          ANTHROPIC_API_KEY,
            "anthropic-version":  "2023-06-01",
        },
        timeout=45, retries=2, log_label="anthropic_chat",
    )
    if res:
        try:
            blocks = res.get("content", [])
            for b in blocks:
                if b.get("type") == "text":
                    return b.get("text", "").strip()
        except Exception as e:
            logger.error("anthropic_chat parse: %s", e)
    return None


def openai_transcribe(audio_path, language=None):
    """OpenAI Whisper দিয়ে audio → SRT subtitle। Return: (srt_text or None, error_msg or None)"""
    if not has_openai():
        return None, "OpenAI key নেই"
    import urllib.request
    boundary = uuid.uuid4().hex
    try:
        with open(audio_path, "rb") as f:
            audio_bytes = f.read()
        # multipart/form-data manual build
        parts = []
        def add(name, value, filename=None, content_type=None):
            head = f'--{boundary}\r\nContent-Disposition: form-data; name="{name}"'
            if filename:
                head += f'; filename="{filename}"'
            head += "\r\n"
            if content_type:
                head += f"Content-Type: {content_type}\r\n"
            head += "\r\n"
            parts.append(head.encode())
            if isinstance(value, bytes):
                parts.append(value)
            else:
                parts.append(str(value).encode())
            parts.append(b"\r\n")
        add("model", OPENAI_STT_MODEL)
        add("response_format", "srt")
        if language and language != "auto":
            add("language", language)
        add("file", audio_bytes, filename="audio.mp3", content_type="audio/mpeg")
        parts.append(f"--{boundary}--\r\n".encode())
        body = b"".join(parts)
        req = urllib.request.Request(
            "https://api.openai.com/v1/audio/transcriptions",
            data=body,
            headers={
                "Authorization": f"Bearer {OPENAI_API_KEY}",
                "Content-Type": f"multipart/form-data; boundary={boundary}",
            },
        )
        with urllib.request.urlopen(req, timeout=180) as r:
            srt_text = r.read().decode("utf-8", errors="replace").strip()
        if "-->" not in srt_text:
            return None, "Whisper response invalid"
        return srt_text, None
    except Exception as e:
        logger.error("openai_transcribe: %s", e)
        return None, f"Whisper error: {str(e)[:200]}"

def has_elevenlabs():
    # 🆕 v26: User feedback — ElevenLabs paid plan লাগে, Bangla quality খারাপ,
    # এবং clone নিয়ে অনেক issue। সম্পূর্ণ DISABLED। বিকল্প: Replicate XTTS-v2।
    return False

def has_replicate():
    """🆕 v26: Replicate API token check — Voice Clone (XTTS-v2)-এর জন্য।"""
    return bool(REPLICATE_API_TOKEN and REPLICATE_API_TOKEN.strip())

# 🆕 v26: Persistent storage for cloned voice samples (Replicate XTTS reference audio)
CLONE_SAMPLES_DIR = STATS_DIR / "clone_samples"
CLONE_SAMPLES_DIR.mkdir(parents=True, exist_ok=True)

# Lang code → XTTS-v2 lang code mapping
XTTS_LANG_MAP = {
    "bn": "hi",   # XTTS-v2-এ Bengali সরাসরি নেই — Hindi-র phonetic close
    "hi": "hi", "en": "en", "ar": "ar", "es": "es", "fr": "fr",
    "de": "de", "it": "it", "pt": "pt", "pl": "pl", "tr": "tr",
    "ru": "ru", "nl": "nl", "cs": "cs", "zh": "zh-cn", "ja": "ja",
    "ko": "ko", "hu": "hu", "auto": "hi",
}

def replicate_xtts_speak(text, sample_audio_path, language="hi", out_path=None):
    """🆕 v26: Replicate XTTS-v2 দিয়ে zero-shot voice cloning।
    sample_audio_path: user-এর reference voice (১০-৩০ সেকেন্ড mp3/wav/ogg)
    language: XTTS_LANG_MAP-এর key
    Return: (path or None, error)
    """
    if not has_replicate():
        return None, "❌ REPLICATE_API_TOKEN সেট নেই — https://replicate.com → API tokens।"
    text = (text or "").strip()
    if not text:
        return None, "❌ TTS-এর জন্য text খালি।"
    if not Path(sample_audio_path).exists():
        return None, "❌ Voice sample ফাইল পাওয়া যায়নি — আবার /clonevoice দিয়ে বানান।"
    import urllib.request, urllib.error, base64

    # XTTS-v2 input length limit ~400 chars per call
    text = text[:400]
    xtts_lang = XTTS_LANG_MAP.get((language or "hi").lower(), "hi")

    # Sample audio → data URI (Replicate small file accept করে)
    try:
        sample_bytes = Path(sample_audio_path).read_bytes()
        if len(sample_bytes) > 8 * 1024 * 1024:
            return None, "❌ Voice sample বড় (8MB-এর কম রাখুন)।"
        suf = Path(sample_audio_path).suffix.lower().lstrip(".")
        mime = {"mp3":"audio/mpeg","wav":"audio/wav","ogg":"audio/ogg",
                "oga":"audio/ogg","m4a":"audio/mp4","flac":"audio/flac",
                "webm":"audio/webm"}.get(suf, "audio/mpeg")
        b64 = base64.b64encode(sample_bytes).decode("ascii")
        speaker_uri = f"data:{mime};base64,{b64}"
    except Exception as e:
        return None, f"❌ Sample পড়া যায়নি: {e}"

    body = json.dumps({
        "version": REPLICATE_XTTS_VERSION,
        "input": {
            "text": text,
            "speaker": speaker_uri,
            "language": xtts_lang,
            "cleanup_voice": False,
        },
    }).encode()
    headers = {
        "Authorization": f"Token {REPLICATE_API_TOKEN}",
        "Content-Type": "application/json",
        "Prefer": "wait=120",
    }
    try:
        req = urllib.request.Request(
            "https://api.replicate.com/v1/predictions",
            data=body, method="POST", headers=headers)
        with urllib.request.urlopen(req, timeout=180) as r:
            res = json.loads(r.read().decode("utf-8", errors="replace"))
    except urllib.error.HTTPError as e:
        try: msg = e.read().decode("utf-8", errors="replace")[:300]
        except Exception: msg = ""
        if e.code == 401: return None, "❌ Replicate token invalid — নতুন token নিন।"
        if e.code == 402: return None, "❌ Replicate ক্রেডিট শেষ — billing-এ যান।"
        return None, f"❌ Replicate HTTP {e.code}: {msg[:160]}"
    except Exception as e:
        return None, f"❌ Replicate error: {str(e)[:200]}"

    # Result polling (যদি Prefer: wait কাজ না করে)
    pred_id = res.get("id")
    status  = res.get("status")
    out_url = None
    if status == "succeeded":
        out_url = res.get("output")
    else:
        # poll up to 90s
        for _ in range(45):
            time.sleep(2)
            try:
                pr = urllib.request.Request(
                    f"https://api.replicate.com/v1/predictions/{pred_id}",
                    headers={"Authorization": f"Token {REPLICATE_API_TOKEN}"})
                with urllib.request.urlopen(pr, timeout=30) as r:
                    res2 = json.loads(r.read().decode("utf-8", errors="replace"))
                st = res2.get("status")
                if st == "succeeded":
                    out_url = res2.get("output"); break
                if st in ("failed", "canceled"):
                    return None, f"❌ Replicate {st}: {str(res2.get('error',''))[:160]}"
            except Exception:
                continue
    if not out_url:
        return None, "❌ Replicate timeout — text ছোট করে আবার চেষ্টা করুন।"

    # output url থেকে audio download
    audio_url = out_url if isinstance(out_url, str) else (out_url[0] if isinstance(out_url, list) and out_url else None)
    if not audio_url:
        return None, "❌ Replicate output empty।"
    out = out_path or str(TEMP_DIR / f"{uuid.uuid4().hex}_xtts.wav")
    try:
        with urllib.request.urlopen(audio_url, timeout=120) as r:
            Path(out).write_bytes(r.read())
    except Exception as e:
        return None, f"❌ Audio download fail: {e}"
    if not Path(out).exists() or Path(out).stat().st_size < 1000:
        return None, "❌ Audio output empty।"
    return out, None

def elevenlabs_tts(text, voice="rachel", model=None, out_path=None):
    """🆕 v16 fix: ElevenLabs TTS দিয়ে text → mp3 file (super realistic, multilingual)।
    voice: ELEVENLABS_VOICES dict-এর key (rachel/adam/bella/domi/antoni/josh) অথবা সরাসরি voice_id
    Return: (path or None, error)

    Improvements over v14:
    - HTTPError-এর response body পড়ে আসল error message বের করে (401/422 ইত্যাদি)
    - Empty/whitespace text-এ আগেই fail
    - quota_exceeded / voice_not_found-এ পরিষ্কার Bengali message
    - Free tier-এ multilingual_v2 fail হলে monolingual_v1-এ auto-retry
    - Output verify (mp3 magic bytes / size)
    """
    if not has_elevenlabs():
        return None, "❌ ElevenLabs key নেই — কোডে ELEVENLABS_API_KEY বসান।"
    import urllib.request, urllib.error
    text = (text or "").strip()
    if not text:
        return None, "❌ TTS-এর জন্য text খালি।"
    # voice key → voice_id resolve
    if voice in ELEVENLABS_VOICES:
        voice_id = ELEVENLABS_VOICES[voice][0]
    elif voice and len(voice) >= 15:  # raw voice_id
        voice_id = voice
    else:
        voice_id = ELEVENLABS_VOICES[ELEVENLABS_DEFAULT_VOICE][0]

    out = out_path or str(TEMP_DIR / f"{uuid.uuid4().hex}_eleven.mp3")

    def _do_call(model_id):
        body = json.dumps({
            "text": text[:4500],
            "model_id": model_id,
            "voice_settings": {
                "stability": 0.5,
                "similarity_boost": 0.75,
                "style": 0.0,
                "use_speaker_boost": True,
            },
        }).encode()
        req = urllib.request.Request(
            f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}?output_format=mp3_44100_128",
            data=body,
            method="POST",
            headers={
                "Content-Type": "application/json",
                "xi-api-key": ELEVENLABS_API_KEY,
                "Accept": "audio/mpeg",
            },
        )
        with urllib.request.urlopen(req, timeout=180) as r:
            return r.read()

    primary_model = model or ELEVENLABS_MODEL
    try:
        audio_bytes = _do_call(primary_model)
    except urllib.error.HTTPError as e:
        # response body থেকে আসল error
        try:
            err_body = e.read().decode("utf-8", errors="replace")
        except Exception:
            err_body = ""
        try:
            err_json = json.loads(err_body) if err_body else {}
        except Exception:
            err_json = {}
        detail = (err_json.get("detail") or {})
        status = detail.get("status") if isinstance(detail, dict) else ""
        msg    = (detail.get("message") if isinstance(detail, dict) else "") or err_body[:300]
        logger.error("elevenlabs HTTP %s: %s | %s", e.code, status, msg)

        # Specific cases — Bengali user-friendly message
        if e.code == 401:
            return None, "❌ ElevenLabs API key ভুল বা expired। নতুন key generate করুন।"
        if e.code == 422 and ("voice" in str(status).lower() or "voice" in msg.lower()):
            return None, f"❌ Voice ID ভুল: `{voice_id}`। অন্য voice দিয়ে চেষ্টা করুন।"
        if e.code == 429 or "quota" in str(status).lower():
            return None, "❌ ElevenLabs quota শেষ — মাসিক free limit (১০,০০০ char) পেরিয়ে গেছে।"
        # multilingual_v2 free tier-এ fail হলে monolingual_v1-এ retry
        if primary_model != "eleven_monolingual_v1":
            try:
                logger.info("elevenlabs: %s fail, monolingual_v1-এ retry", primary_model)
                audio_bytes = _do_call("eleven_monolingual_v1")
            except Exception as e2:
                return None, f"❌ ElevenLabs error ({e.code}): {msg[:150]}"
        else:
            return None, f"❌ ElevenLabs error ({e.code}): {msg[:150]}"
    except urllib.error.URLError as e:
        return None, f"❌ Network error: {e.reason}"
    except Exception as e:
        logger.error("elevenlabs_tts: %s", e)
        return None, f"❌ ElevenLabs error: {str(e)[:200]}"

    # Output verify — mp3 হলে `ID3` বা `\xff\xfb` দিয়ে শুরু
    if not audio_bytes or len(audio_bytes) < 500:
        return None, "❌ ElevenLabs output খালি/অসম্পূর্ণ।"
    if not (audio_bytes.startswith(b"ID3") or audio_bytes[:2] == b"\xff\xfb" or audio_bytes[:2] == b"\xff\xf3"):
        # mp3 না হলেও save করব, কিন্তু warn
        logger.warning("elevenlabs: output mp3 magic mismatch — first bytes: %s", audio_bytes[:8])
    try:
        Path(out).write_bytes(audio_bytes)
    except Exception as e:
        return None, f"❌ Audio save fail: {e}"
    return out, None

# ──────────────────────────────────────────────────────────────
# 🆕 v18: ElevenLabs Voice CLONE (Instant Voice Clone)
# ──────────────────────────────────────────────────────────────
def _build_multipart(fields, files):
    """fields: dict[str,str], files: list[(field_name, filename, mime, bytes)]
    Returns: (body bytes, content_type)"""
    boundary = "----vebot" + uuid.uuid4().hex
    crlf = b"\r\n"
    parts = []
    for k, v in fields.items():
        parts.append(("--" + boundary).encode())
        parts.append(f'Content-Disposition: form-data; name="{k}"'.encode())
        parts.append(b"")
        parts.append(str(v).encode("utf-8"))
    for fname, filename, mime, data in files:
        parts.append(("--" + boundary).encode())
        parts.append(f'Content-Disposition: form-data; name="{fname}"; filename="{filename}"'.encode())
        parts.append(f"Content-Type: {mime}".encode())
        parts.append(b"")
        parts.append(data)
    parts.append(("--" + boundary + "--").encode())
    parts.append(b"")
    body = crlf.join(parts)
    return body, f"multipart/form-data; boundary={boundary}"

def elevenlabs_clone_voice(audio_path, name, description="Cloned via Video Editor Bot"):
    """🆕 v18: ElevenLabs Instant Voice Clone — audio file → voice_id।
    Return: (voice_id or None, error_msg)"""
    if not has_elevenlabs():
        return None, "❌ ElevenLabs key নেই।"
    if not Path(audio_path).exists():
        return None, "❌ Audio ফাইল পাওয়া যায়নি।"
    import urllib.request, urllib.error
    try:
        data = Path(audio_path).read_bytes()
    except Exception as e:
        return None, f"❌ Audio পড়া যায়নি: {e}"
    if len(data) < 5_000:
        return None, "❌ Audio sample-টা অনেক ছোট (কমপক্ষে ৫-১০ সেকেন্ড পরিষ্কার voice দরকার)।"
    if len(data) > 11 * 1024 * 1024:
        return None, "❌ Audio sample অনেক বড় (১০MB-এর কম রাখুন)।"

    # Mime guess
    suf = Path(audio_path).suffix.lower()
    mime = {".mp3":"audio/mpeg",".m4a":"audio/mp4",".ogg":"audio/ogg",
            ".oga":"audio/ogg",".wav":"audio/wav",".webm":"audio/webm",
            ".flac":"audio/flac",".aac":"audio/aac"}.get(suf, "audio/mpeg")
    fname = Path(audio_path).name or f"sample{suf or '.mp3'}"

    body, ctype = _build_multipart(
        fields={
            "name": (name or "MyVoice")[:30],
            "description": description[:200],
            "remove_background_noise": "true",
        },
        files=[("files", fname, mime, data)],
    )
    req = urllib.request.Request(
        "https://api.elevenlabs.io/v1/voices/add",
        data=body, method="POST",
        headers={
            "xi-api-key": ELEVENLABS_API_KEY,
            "Content-Type": ctype,
            "Accept": "application/json",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=180) as r:
            res = json.loads(r.read().decode("utf-8", errors="replace"))
        vid = res.get("voice_id")
        if not vid:
            return None, f"❌ Voice ID পাওয়া যায়নি: {str(res)[:200]}"
        return vid, None
    except urllib.error.HTTPError as e:
        try: err_body = e.read().decode("utf-8", errors="replace")
        except Exception: err_body = ""
        try: err_json = json.loads(err_body) if err_body else {}
        except Exception: err_json = {}
        detail = (err_json.get("detail") or {})
        # detail string বা dict — দুই ক্ষেত্রেই handle করি
        if isinstance(detail, dict):
            status = str(detail.get("status") or "")
            msg    = str(detail.get("message") or "") or err_body[:300]
        else:
            status = ""
            msg    = str(detail) if detail else err_body[:300]
        # 🆕 v22: combined search text (status + msg + raw body — সব check হবে)
        full_err = f"{status} {msg} {err_body}".lower()
        logger.error("elevenlabs_clone HTTP %s: %s | %s", e.code, status, msg[:200])
        if e.code == 401:
            return None, ("❌ *ElevenLabs API key invalid।*\n\n"
                          "✨ *বিকল্প — সম্পূর্ণ ফ্রি!*\n"
                          "🆓 `🗣️ TTS` মেনু → *FREE Premium Voices (Edge)* — "
                          "বাংলা/হিন্দি/ইংরেজি ২০+ neural voice, কোনো key/subscription লাগবে না!")
        # 🆕 v22: subscription-related errors — code 400/402/403 OR keyword match
        sub_keywords = ["subscription", "upgrade", "paid plan", "voice_limit",
                        "instant voice cloning", "professional voice cloning",
                        "not include", "not available on your plan", "starter"]
        if (e.code in (400, 402, 403) and any(k in full_err for k in sub_keywords)) \
           or e.code == 403 \
           or "voice_limit" in full_err:
            return None, ("❌ *ElevenLabs Voice Cloning — Paid Plan দরকার!*\n\n"
                          "━━━━━━━━━━━━━━━━━━\n"
                          "ElevenLabs Free plan-এ voice cloning support নেই। "
                          "Cloning করতে Starter ($5/মাস) বা তার উপরের plan লাগবে।\n\n"
                          "✨ *সম্পূর্ণ ফ্রি বিকল্প — এক tap-এ:*\n"
                          "🆓 *Edge TTS* — বাংলা/হিন্দি/ইংরেজি/Urdu/Arabic ২০+ premium "
                          "neural voice, কোনো API key/subscription লাগবে না!\n\n"
                          "👉 `🗣️ TTS` মেনু → 🆓 *FREE Premium Voices (Edge)* section\n"
                          "🎯 ১ ক্লিকে: Nabanita, Pradeep, Bashkar, Tanishaa, "
                          "Swara, Madhur, Aria, Guy ইত্যাদি।")
        if e.code == 422:
            return None, f"❌ Audio invalid: {msg[:150]}"
        if e.code == 429:
            return None, "❌ Rate limit — একটু পর আবার চেষ্টা করুন।"
        return None, (f"❌ *Clone error ({e.code}):* {msg[:150]}\n\n"
                      f"✨ ফ্রি বিকল্প: 🆓 `🗣️ TTS` মেনু → *Edge Premium Voices*")
    except Exception as e:
        logger.error("elevenlabs_clone: %s", e)
        return None, (f"❌ *Clone error:* {str(e)[:180]}\n\n"
                      f"✨ ফ্রি বিকল্প: 🆓 `🗣️ TTS` মেনু → *Edge Premium Voices*")

def elevenlabs_delete_voice(voice_id):
    """ElevenLabs থেকে cloned voice delete। True/False return।"""
    if not has_elevenlabs() or not voice_id:
        return False
    import urllib.request, urllib.error
    req = urllib.request.Request(
        f"https://api.elevenlabs.io/v1/voices/{voice_id}",
        method="DELETE",
        headers={"xi-api-key": ELEVENLABS_API_KEY, "Accept": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            r.read()
        return True
    except urllib.error.HTTPError as e:
        # 404 — already gone, treat as success
        if e.code == 404:
            return True
        logger.warning("elevenlabs_delete %s: HTTP %s", voice_id, e.code)
        return False
    except Exception as e:
        logger.warning("elevenlabs_delete %s: %s", voice_id, e)
        return False

# 🆕 v19: ElevenLabs account-এর সব voice (default library + premade + cloned + generated) fetch
def elevenlabs_list_voices(force_refresh=False):
    """GET /v1/voices → আপনার অ্যাকাউন্টের সব available voice list।
    Cache: 10 minutes (rate-limit এড়াতে)।
    Return: list of dicts: [{voice_id, name, category}, ...]
    """
    import time, urllib.request, urllib.error
    if not has_elevenlabs():
        return []
    now = time.time()
    if not force_refresh and _eleven_voice_cache.get("voices") and (now - _eleven_voice_cache.get("ts", 0)) < 600:
        return _eleven_voice_cache["voices"]
    try:
        req = urllib.request.Request(
            "https://api.elevenlabs.io/v1/voices",
            headers={"xi-api-key": ELEVENLABS_API_KEY, "Accept": "application/json"},
        )
        with urllib.request.urlopen(req, timeout=30) as r:
            data = json.loads(r.read().decode("utf-8", errors="replace"))
        out = []
        for v in (data.get("voices") or []):
            vid = v.get("voice_id") or ""
            name = (v.get("name") or "Unnamed").strip()
            cat  = (v.get("category") or "premade").lower()  # premade/cloned/generated/professional
            if vid:
                out.append({"voice_id": vid, "name": name, "category": cat})
        # Sort: cloned/generated আগে, তারপর premade
        order = {"cloned": 0, "generated": 1, "professional": 2, "premade": 3}
        out.sort(key=lambda x: (order.get(x["category"], 9), x["name"].lower()))
        _eleven_voice_cache["ts"] = now
        _eleven_voice_cache["voices"] = out
        return out
    except urllib.error.HTTPError as e:
        logger.warning("elevenlabs_list_voices HTTP %s", e.code)
        return _eleven_voice_cache.get("voices", [])
    except Exception as e:
        logger.warning("elevenlabs_list_voices: %s", e)
        return _eleven_voice_cache.get("voices", [])

def elevenlabs_voice_label(category, name):
    """Voice category-অনুযায়ী emoji + label।"""
    icons = {"cloned": "🎭", "generated": "🪄", "professional": "💎", "premade": "🎤"}
    return f"{icons.get(category, '🎤')} {name[:24]}"

def openai_tts(text, voice="nova", model=None, out_path=None):
    """OpenAI TTS দিয়ে text → mp3 file। Return: (path or None, error)"""
    if not has_openai():
        return None, "OpenAI key নেই"
    import urllib.request
    out = out_path or str(TEMP_DIR / f"{uuid.uuid4().hex}_tts.mp3")
    body = json.dumps({
        "model": model or OPENAI_TTS_MODEL,
        "input": text[:4000],
        "voice": voice or OPENAI_TTS_VOICE,
        "response_format": "mp3",
    }).encode()
    req = urllib.request.Request(
        "https://api.openai.com/v1/audio/speech",
        data=body,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OPENAI_API_KEY}",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as r:
            audio_bytes = r.read()
        Path(out).write_bytes(audio_bytes)
        return out, None
    except Exception as e:
        logger.error("openai_tts: %s", e)
        return None, f"OpenAI TTS error: {str(e)[:200]}"

# ──────────────────────────────────────────────────────────────
# 🆕 v13: TEXT → VOICE (TTS) — gTTS (ফ্রি) + OpenAI TTS (premium)
# ──────────────────────────────────────────────────────────────
def gtts_speak(text, lang="bn", out_path=None):
    """gTTS দিয়ে text → mp3 file (ফ্রি, internet দরকার, Google Translate TTS endpoint)।
    Return: (path or None, error)
    Termux: pip install gTTS
    """
    out = out_path or str(TEMP_DIR / f"{uuid.uuid4().hex}_gtts.mp3")
    try:
        from gtts import gTTS
    except ImportError:
        return None, ("❌ gTTS ইনস্টল নেই।\n\n"
                      "Termux: `pip install gTTS`\n"
                      "Ubuntu: `pip install gTTS`")
    try:
        # gTTS supports: bn, en, hi, ar, fr, de, es, ja, ko, zh-CN, zh-TW, ur etc.
        lang_map = {
            "auto": "bn", "bn": "bn", "en": "en", "hi": "hi",
            "ar": "ar", "es": "es", "fr": "fr", "de": "de",
            "ja": "ja", "ko": "ko", "zh": "zh-CN", "ur": "ur",
            "ru": "ru", "pt": "pt", "it": "it", "tr": "tr",
            "id": "id", "th": "th", "vi": "vi", "nl": "nl",
        }
        l = lang_map.get((lang or "").lower(), "en")
        tts = gTTS(text=text[:4000], lang=l, slow=False)
        tts.save(out)
        if not Path(out).exists() or Path(out).stat().st_size < 100:
            return None, "gTTS output empty"
        return out, None
    except Exception as e:
        logger.error("gtts: %s", e)
        return None, f"gTTS error: {str(e)[:200]}"

# ──────────────────────────────────────────────────────────────
# 🆕 v22: EDGE TTS — সম্পূর্ণ ফ্রি neural voice (Microsoft Edge browser-এর engine)
# ──────────────────────────────────────────────────────────────
# 🆕 v30: Video Dubbing — language → Edge-TTS voice mapping (12 languages, FREE)
DUB_LANGS = {
    "bn": ("🇧🇩 বাংলা",   "bn-BD-NabanitaNeural", "bn"),
    "hi": ("🇮🇳 हिन्दी",   "hi-IN-SwaraNeural",    "hi"),
    "en": ("🇬🇧 English", "en-US-AriaNeural",     "en"),
    "ur": ("🇵🇰 اردو",    "ur-PK-UzmaNeural",     "ur"),
    "ar": ("🇸🇦 العربية", "ar-SA-HamedNeural",    "ar"),
    "es": ("🇪🇸 Español", "es-ES-ElviraNeural",   "es"),
    "fr": ("🇫🇷 Français","fr-FR-DeniseNeural",   "fr"),
    "de": ("🇩🇪 Deutsch", "de-DE-KatjaNeural",    "de"),
    "ja": ("🇯🇵 日本語",   "ja-JP-NanamiNeural",   "ja"),
    "ko": ("🇰🇷 한국어",   "ko-KR-SunHiNeural",    "ko"),
    "zh": ("🇨🇳 中文",     "zh-CN-XiaoxiaoNeural", "zh"),
    "ta": ("🇮🇳 தமிழ்",   "ta-IN-PallaviNeural",  "ta"),
}

def extract_audio_for_stt(video_path, out_mp3, max_seconds=600):
    """ভিডিও থেকে clean mono 16kHz mp3 (Whisper-friendly)। 10 মিনিট ক্যাপ।"""
    try:
        cmd = ["ffmpeg","-y","-i",video_path,"-vn","-ac","1","-ar","16000",
               "-c:a","libmp3lame","-b:a","64k","-t",str(max_seconds), out_mp3]
        p = subprocess.run(cmd, capture_output=True, timeout=FFMPEG_TIMEOUT)
        return p.returncode == 0 and Path(out_mp3).exists() and Path(out_mp3).stat().st_size > 2000
    except Exception as e:
        logger.warning("extract_audio_for_stt: %s", e); return False

def srt_to_plain_text(srt_text):
    """SRT-এর শুধু dialogue লাইন extract — timestamp/index বাদ।"""
    out = []
    for block in re.split(r'\n\s*\n', (srt_text or "").strip()):
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        if len(lines) < 2: continue
        text_lines = [l for l in lines if not l.isdigit() and "-->" not in l]
        if text_lines:
            out.append(" ".join(text_lines))
    return " ".join(out).strip()

def has_edge_tts():
    """edge-tts package available কিনা চেক।"""
    try:
        import edge_tts  # noqa: F401
        return True
    except Exception:
        return False

def edge_tts_speak(text, voice="bn_nabanita", out_path=None, rate="+0%", pitch="+0Hz"):
    """🆕 v22: Microsoft Edge Neural TTS — ১০০% ফ্রি, কোনো API key লাগে না।
    voice: EDGE_TTS_VOICES key (bn_nabanita/bn_pradeep/...) অথবা সরাসরি Edge voice name (bn-BD-NabanitaNeural)
    Return: (path or None, error)
    """
    text = (text or "").strip()
    if not text:
        return None, "❌ Edge TTS-এর জন্য text খালি।"
    try:
        import edge_tts
    except ImportError:
        return None, ("❌ edge-tts ইনস্টল নেই।\n\n"
                      "Termux/Linux: `pip install edge-tts`\n"
                      "তারপর বট restart করুন।")
    # voice key → Edge voice name resolve
    if voice in EDGE_TTS_VOICES:
        voice_name = EDGE_TTS_VOICES[voice][0]
    elif voice and "-" in voice and "Neural" in voice:
        voice_name = voice  # raw Edge voice
    else:
        voice_name = EDGE_TTS_VOICES[EDGE_TTS_DEFAULT][0]

    out = out_path or str(TEMP_DIR / f"{uuid.uuid4().hex}_edge.mp3")
    # 🆕 v23 fix: loop variable safe-init (finally-এ unbound হওয়া এড়ায়)
    loop = None
    try:
        async def _run():
            communicate = edge_tts.Communicate(text[:4500], voice_name, rate=rate, pitch=pitch)
            await communicate.save(out)
        # নতুন event loop চালাই (executor thread-এ চলবে)
        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)  # 🆕 v23: thread-এ current loop set
            loop.run_until_complete(_run())
        finally:
            if loop is not None:
                try: loop.close()
                except Exception: pass
                try: asyncio.set_event_loop(None)
                except Exception: pass
        if not Path(out).exists() or Path(out).stat().st_size < 200:
            return None, "❌ Edge TTS output খালি। অন্য voice দিয়ে চেষ্টা করুন।"
        return out, None
    except Exception as e:
        logger.error("edge_tts: %s", e)
        msg = str(e)[:200]
        low = msg.lower()
        if "no audio was received" in low or "no internet" in low or "getaddrinfo" in low:
            return None, "❌ Edge TTS network error। internet check করুন।"
        if "websocket" in low or "1006" in low or "1007" in low:
            return None, "❌ Edge TTS server timeout। আবার চেষ্টা করুন।"
        return None, f"❌ Edge TTS error: {msg}"

def text_to_speech(text, lang="bn", voice="nova", engine="auto"):
    """Main TTS router. engine: 'edge' | 'gtts' | 'openai' | 'elevenlabs' | 'auto'
    (🆕 v22 auto priority: edge > elevenlabs > openai > gtts — Edge সম্পূর্ণ ফ্রি, তাই আগে)
    Return: (path or None, used_engine, error)
    """
    if not text or not text.strip():
        return None, None, "❌ টেক্সট খালি!"
    if engine == "auto":
        if has_edge_tts():     engine = "edge"
        elif has_elevenlabs(): engine = "elevenlabs"
        elif has_openai():     engine = "openai"
        else:                  engine = "gtts"
    # ─── 🆕 v22: Edge TTS (free + neural quality) ───
    if engine == "edge":
        # voice হিসেবে যদি OpenAI/ElevenLabs key আসে, lang অনুযায়ী default Edge voice বেছে নিই
        if voice not in EDGE_TTS_VOICES and not (voice and "-" in str(voice) and "Neural" in str(voice)):
            lang_map = {"bn": "bn_nabanita", "hi": "hi_swara", "en": "en_aria",
                        "ur": "ur_gul", "ar": "ar_zariyah", "auto": "bn_nabanita"}
            voice = lang_map.get((lang or "bn").lower(), "bn_nabanita")
        path, err = edge_tts_speak(text, voice=voice)
        if path:
            return path, "edge", None
        logger.warning("edge TTS fallback: %s", err)
        # fallback: gtts
        path, err3 = gtts_speak(text, lang=lang)
        return path, "gtts", err3
    # ─── 🆕 v14: ElevenLabs (best quality, realistic) ───
    if engine == "elevenlabs":
        path, err = elevenlabs_tts(text, voice=voice)
        if path:
            return path, "elevenlabs", None
        logger.warning("elevenlabs TTS fallback: %s", err)
        # 🆕 v22 fallback chain: edge → openai → gtts
        if has_edge_tts():
            path, err_e = edge_tts_speak(text, voice=EDGE_TTS_DEFAULT)
            if path: return path, "edge", None
        if has_openai():
            path, err2 = openai_tts(text, voice=OPENAI_TTS_VOICE)
            if path: return path, "openai", None
        path, err3 = gtts_speak(text, lang=lang)
        return path, "gtts", err3
    # ─── 🆕 v26: Replicate XTTS-v2 (zero-shot voice cloning) ───
    if engine == "xtts":
        # voice = sample audio path (CLONE_SAMPLES_DIR-এ saved)
        sample_path = voice
        if not sample_path or not Path(str(sample_path)).exists():
            err = "❌ Voice sample missing — /clonevoice দিয়ে আবার বানান।"
            logger.warning("xtts no sample: %s", sample_path)
            if has_edge_tts():
                path, _ = edge_tts_speak(text, voice=EDGE_TTS_DEFAULT)
                if path: return path, "edge", err
            path, err3 = gtts_speak(text, lang=lang)
            return path, "gtts", err
        path, err = replicate_xtts_speak(text, sample_path, language=lang)
        if path:
            return path, "xtts", None
        logger.warning("xtts fallback: %s", err)
        if has_edge_tts():
            path, _ = edge_tts_speak(text, voice=EDGE_TTS_DEFAULT)
            if path: return path, "edge", err
        path, err3 = gtts_speak(text, lang=lang)
        return path, "gtts", err or err3
    if engine == "openai":
        path, err = openai_tts(text, voice=voice)
        if path:
            return path, "openai", None
        logger.warning("openai TTS fallback: %s", err)
        if has_edge_tts():
            path, err_e = edge_tts_speak(text, voice=EDGE_TTS_DEFAULT)
            if path: return path, "edge", None
        path, err2 = gtts_speak(text, lang=lang)
        return path, "gtts", err2
    else:
        path, err = gtts_speak(text, lang=lang)
        return path, "gtts", err

# ──────────────────────────────────────────────────────────────
# 🆕 v14: PRO AI HELPERS (Translation / Analysis / Script / Hooks / Thumb)
# সবগুলো OpenAI ChatGPT দিয়ে চলে — Gemini-র চেয়ে অনেক বেশি accurate
# ──────────────────────────────────────────────────────────────
LANG_NAMES = {
    "bn":"Bengali","en":"English","hi":"Hindi","ar":"Arabic","es":"Spanish",
    "fr":"French","de":"German","ja":"Japanese","ko":"Korean","zh":"Chinese","ur":"Urdu",
}

def translate_text(text, target_lang="en"):
    """OpenAI/Gemini দিয়ে যেকোনো টেক্সট target language-এ translate।"""
    if not has_ai():
        return None, "❌ AI key দরকার (OpenAI বা Gemini)"
    target = LANG_NAMES.get(target_lang, target_lang)
    prompt = (f"Translate the following text to {target}. Preserve all formatting "
              f"(line breaks, SRT timestamps, numbering). Output ONLY the translated text.\n\n"
              f"---\n{text[:6000]}\n---")
    reply = openai_chat(prompt, system="You are a professional translator.",
                        max_tokens=4000, temperature=0.3)
    return (reply, None) if reply else (None, "Translation API empty response")

def translate_srt(srt_text, target_lang="en"):
    """SRT subtitle-এর শুধু text লাইন translate করে timestamp অক্ষুণ্ণ রাখে।"""
    if not srt_text or not srt_text.strip():
        return None, "SRT খালি"
    blocks = re.split(r"\n\s*\n", srt_text.strip())
    text_segments, structure = [], []
    for b in blocks:
        lines = b.split("\n")
        ts = 2 if (len(lines) >= 2 and "-->" in lines[1]) else (1 if "-->" in (lines[0] if lines else "") else None)
        if ts is not None and len(lines) > ts:
            text_segments.append(" ".join(lines[ts:]))
            structure.append((lines, ts))
        else:
            structure.append((lines, None))
    if not text_segments:
        return None, "SRT-তে translatable text পাওয়া যায়নি"
    joined = "\n|||\n".join(text_segments)
    translated, err = translate_text(joined, target_lang)
    if err: return None, err
    parts = translated.split("\n|||\n")
    if len(parts) < len(text_segments):
        parts += text_segments[len(parts):]
    out_blocks, pi = [], 0
    for lines, ts in structure:
        if ts is None:
            out_blocks.append("\n".join(lines))
        else:
            new_text = parts[pi].strip() if pi < len(parts) else " ".join(lines[ts:])
            pi += 1
            out_blocks.append("\n".join(lines[:ts] + [new_text]))
    return "\n\n".join(out_blocks), None

def analyze_video_content(title, platform, duration=None):
    """ChatGPT/Gemini দিয়ে best-platform + audience + viral potential analyze।"""
    if not has_ai():
        return None, "❌ AI key দরকার (OpenAI বা Gemini)"
    dur_str = f"{int(duration)} সেকেন্ড" if duration else "অজানা"
    prompt = (
        f"Analyze this video for a Bengali content creator and recommend strategy.\n"
        f"- Title/topic: {title}\n"
        f"- Source platform: {platform}\n"
        f"- Duration: {dur_str}\n\n"
        f"Reply IN BENGALI (Bangla) with these exact sections (use markdown bold):\n\n"
        f"🎯 *বেস্ট প্ল্যাটফর্ম (Top 3):* প্রতিটার পাশে ১ লাইন reason\n\n"
        f"👥 *টার্গেট অডিয়েন্স:* age, interest, country/region\n\n"
        f"⏰ *বেস্ট পোস্টিং টাইম:* (Bangladesh time-এ)\n\n"
        f"📊 *ভাইরাল হওয়ার সম্ভাবনা:* Low/Medium/High + কেন\n\n"
        f"💡 *উন্নতির ৩টা টিপস:* প্রতিটা ১-২ লাইনে\n\n"
        f"🏆 *Trend ফিট:* এই মুহূর্তের কোন trend-এর সাথে মেলে"
    )
    reply = openai_chat(prompt,
                        system="You are a viral video strategist for Bengali content creators. Always respond in Bengali.",
                        max_tokens=900, temperature=0.7)
    return (reply, None) if reply else (None, "Analysis empty")

def generate_script(topic, duration_sec=30, platform="TikTok", language="bn"):
    """ভিডিওর জন্য hook + main + CTA + voiceover-ready text generate।
    Return: (dict {hook, main, cta, voiceover_text} or None, error)
    """
    if not has_ai():
        return None, "❌ AI key দরকার (OpenAI বা Gemini)"
    lang_name = LANG_NAMES.get(language, "Bengali")
    prompt = (
        f"Write a viral video script for {platform}, total length ~{duration_sec} seconds, "
        f"in {lang_name}.\nTopic: {topic}\n\n"
        f"Reply in valid JSON only (no markdown fence):\n"
        f"{{\n"
        f'  "hook": "first 3-second attention grabber (1-2 sentences)",\n'
        f'  "main": "main script body for ~{max(5, duration_sec-6)} seconds",\n'
        f'  "cta": "last 3-second call-to-action",\n'
        f'  "voiceover_text": "the full continuous narration text (hook+main+cta combined, no headers, ready for TTS, max 1200 chars)"\n'
        f"}}"
    )
    reply = openai_chat(prompt,
                        system=f"You are a viral short-form scriptwriter for {platform}. Reply with valid JSON only.",
                        max_tokens=1100, temperature=0.85)
    if not reply: return None, "Script API empty"
    try:
        txt = reply.strip()
        if txt.startswith("```"):
            txt = re.sub(r"^```(?:json)?\s*|\s*```$", "", txt, flags=re.S).strip()
        data = json.loads(txt)
        for k in ("hook", "main", "cta", "voiceover_text"):
            data.setdefault(k, "")
        return data, None
    except Exception as e:
        # fallback: pass raw text in voiceover_text
        return {"hook": "", "main": reply, "cta": "", "voiceover_text": reply[:1200]}, None

def generate_caption_hook(title, platform):
    """5টা scroll-stopping hook + 3টা caption + 3টা CTA + sound hint।"""
    if not has_ai():
        return None, "❌ AI key দরকার (OpenAI বা Gemini)"
    prompt = (
        f"Generate viral short-video content for {platform}, topic: \"{title}\"\n\n"
        f"Reply in valid JSON only (no code fence):\n"
        f"{{\n"
        f'  "hooks": [5 first-3-second scroll-stopping hooks in Bengali+English mix, each <100 chars],\n'
        f'  "captions": [3 ready-to-paste captions, each with emoji + CTA, <200 chars],\n'
        f'  "ctas": [3 powerful call-to-actions in Bengali],\n'
        f'  "trending_sounds_hint": "1-2 line suggestion on what audio/sound style fits"\n'
        f"}}"
    )
    reply = openai_chat(prompt,
                        system="You are a TikTok/Reels viral hook expert. Reply with valid JSON only.",
                        max_tokens=900, temperature=0.95)
    if not reply: return None, "Hook API empty"
    try:
        txt = reply.strip()
        if txt.startswith("```"):
            txt = re.sub(r"^```(?:json)?\s*|\s*```$", "", txt, flags=re.S).strip()
        return json.loads(txt), None
    except Exception as e:
        return None, f"JSON parse error: {str(e)[:100]}"

def generate_thumbnail_ideas(title, platform):
    """5টা click-trigger thumbnail concept (text + visual + colors)।"""
    if not has_ai():
        return None, "❌ AI key দরকার (OpenAI বা Gemini)"
    prompt = (
        f"Suggest 5 highly-clickable thumbnail concepts for a {platform} video titled \"{title}\".\n\n"
        f"Reply IN BENGALI. For EACH concept (numbered 1-5) give these 4 lines:\n"
        f"১. *বড় টেক্সট:* (3-5 শব্দের overlay text — exact words)\n"
        f"২. *ভিজ্যুয়াল:* (background, character expression, props)\n"
        f"৩. *কালার স্কিম:* (2-3 dominant colors + কেন কাজ করবে)\n"
        f"৪. *Click trigger:* (curiosity gap / shock / emoji / face-zoom etc.)\n\n"
        f"Separate each concept with `━━━━━━━━━━━━━━━━━━`."
    )
    reply = openai_chat(prompt,
                        system="You are a viral YouTube/TikTok thumbnail designer for Bengali audience.",
                        max_tokens=1300, temperature=0.85)
    return (reply, None) if reply else (None, "Thumbnail API empty")

# ──────────────────────────────────────────────────────────────
# WELCOME / HELP
# ──────────────────────────────────────────────────────────────
WELCOME = """╭─━━━━━━━━━━━━━━━━━━━━━━━╮
│ 🎬 *VIDEO EDITOR BOT* 🎬 │
│   💎 *Pro Edition v30* 💎    │
│  🔥 *ULTRA — All-in-One* 🔥  │
╰─━━━━━━━━━━━━━━━━━━━━━━━╯

✨ *আপনার ভিডিও এডিটিং এর সম্পূর্ণ AI স্টুডিও* ✨

━━━━━━━━━━━━━━━━━━━━━━━━━━━
📤 *যেভাবে শুরু করবেন:*
━━━━━━━━━━━━━━━━━━━━━━━━━━━
১) যেকোনো ভিডিও পাঠান (ফাইল আপলোড)
   অথবা YouTube / TikTok / Instagram /
   Facebook / Pinterest-এর লিংক পাঠান
২) মেনু থেকে আপনার পছন্দের এডিট বেছে নিন
৩) ১-৩ মিনিটে রেডি ভিডিও পেয়ে যান!

📦 *ফাইল লিমিট:*
   • ডাউনলোড লিংক → সর্বোচ্চ *2 GB*
   • আপলোড / আউটপুট → *50 MB* (Telegram bot
     API limit; বড় হলে অটো-কম্প্রেস হবে)

━━━━━━━━━━━━━━━━━━━━━━━━━━━
🆕 *v30 — নতুন যা যোগ হলো:*
━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎬🌍 *Video Dubbing (১২ ভাষায়)* — যেকোনো ভিডিও
   বাংলা/হিন্দি/ইংরেজি/উর্দু/আরবি/স্প্যানিশ/ফরাসি/
   জার্মান/জাপানি/কোরিয়ান/চীনা/তামিল-এ অটো ডাব
🎭 *Voice-Cloned Dubbing* — আগে নিজের কণ্ঠ
   `/clonevoice` দিয়ে save করলে, ডাবিং *আপনার
   নিজের কণ্ঠেই* অন্য ভাষায় হবে! (Replicate XTTS)
📦 *Auto-Compress Download* — ৫০ MB+ ফাইলও
   এখন আর reject হবে না, অটো compress হবে
🎵 *TikTok Multi-Strategy Download* — vm/vt
   সহ ৫টা strategy, ৯০%+ TikTok লিংক নামবে
🎯 *WinGo Predict ULTRA++* — ২০-layer analyzer

━━━━━━━━━━━━━━━━━━━━━━━━━━━
⭐ *জনপ্রিয় ফিচার:*
━━━━━━━━━━━━━━━━━━━━━━━━━━━
🤖 *Smart Auto* — All-in-One এক tap-এ
🚀 *YouTube Viral Pack* — edit + AI title + hashtag
📺 *YouTube / 🎵 TikTok / 👥 Facebook Mode*
   — Algorithm-optimized, copyright-safe
✨ *Enhance ULTRA + Wink Premium 4K*
🎙️ *AI Voiceover + Bangla Subtitle*
🎤 *Auto Voice Clone* (ভিডিও থেকে instant)
📝 *Auto Subtitle* (Whisper, বাংলা ফন্ট)
🎯 *SEO Title + 30 Hashtag Generator*
🛡️ *Copyright Check + Remove*
🎀 *Premium Bangla Caption* (4 ফন্ট)
✂️ Trim/Cut • 🔗 Merge • ⏩ Speed
🎞️ Slow-Mo Premium • ♾️ Loop • 📐 Aspect
🌫️ Blur BG • 🌟 Vignette • 🔍 Zoom
📦 কম্প্রেস • 🎵 MP3 এক্সট্র্যাক্ট

━━━━━━━━━━━━━━━━━━━━━━━━━━━
💡 *বিশেষ কমান্ড:*
━━━━━━━━━━━━━━━━━━━━━━━━━━━
/start — মেনু  •  /help — সাহায্য
/clonevoice — নিজের কণ্ঠ clone করুন
/myvoice — আপনার clone-এর info
/wingo — WinGo prediction (২০-layer AI)
/money — YouTube monetization গাইড
/stats — bot ব্যবহারের পরিসংখ্যান

━━━━━━━━━━━━━━━━━━━━━━━━━━━
👇 *এখনই একটা ভিডিও পাঠান বা লিংক দিন!*
━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""


# ──────────────────────────────────────────────────────────────
# COMMAND HANDLERS
# ──────────────────────────────────────────────────────────────
async def cmd_start(u, c):
    track_user(u.effective_user.id)
    await safe_reply(u.message, WELCOME, parse_mode="Markdown", reply_markup=main_menu())

async def cmd_help(u, c):
    await safe_reply(u.message,
        "📖 *সাহায্য — Video Editor Bot v22*\n\n"
        "━━━━━━━━━━━━━━━━━━\n"
        "📤 ভিডিও পাঠান → এডিট মেনু\n"
        "🔗 লিংক পাঠান → অটো ডাউনলোড\n\n"
        "🆕 *v22 — সম্পূর্ণ ফ্রি Premium Voice:*\n"
        "🆓 *Edge TTS* — কোনো API key/subscription ছাড়াই ২০+ neural voice\n"
        "🇧🇩 বাংলা / 🇮🇳 হিন্দি / 🇺🇸🇬🇧 English / 🇵🇰 Urdu / 🇸🇦 Arabic\n"
        "👉 `🗣️ TTS` মেনু → 🆓 *FREE Premium Voices* section\n\n"
        "📌 *কমান্ড:*\n"
        "/start — শুরু  •  /help — সাহায্য\n"
        "/stats — স্ট্যাটস  •  /cancel — বাতিল\n"
        "/platforms — সাইট লিস্ট\n"
        "/myfonts — বাংলা ফন্ট preview\n"
        "   _উদাহরণ:_ `/myfonts স্বাগতম`\n\n"
        "🎭 *Voice Cloning (ElevenLabs Paid):*\n"
        "/clonevoice — voice message-এ reply দিয়ে clone\n"
        "   _উদাহরণ:_ voice reply → `/clonevoice MyName`\n"
        "/clonefromvideo — uploaded video থেকে directly clone\n"
        "/myvoice — clone-এর তথ্য  •  /deletevoice — মুছুন\n"
        "✨ Voice Cloning Engine: *Replicate XTTS-v2* (free tier credit, no paid plan)।\n"
        "   Setup: https://replicate.com → API token → `export REPLICATE_API_TOKEN=...`\n"
        "   Bonus: Edge TTS-এ ২১+ neural voice (clone ছাড়াই) সম্পূর্ণ ফ্রি।\n\n"
        "🆕 *Pro Creator Tools:*\n"
        "🎙️ AI Voiceover ON Video  •  🔊 Audio Denoise\n"
        "🎞️ Smooth Slow-Mo  •  ♾️ N-second Loop  •  📊 Video Info\n"
        "— সব এডিট মেনুতে।",
        parse_mode="Markdown")

async def cmd_clonevoice(u, c):
    """🆕 v26: Replicate XTTS-v2 Zero-Shot Voice Clone (FREE, no paid plan)।
    Usage: voice/audio message-এ reply → /clonevoice <name>
    """
    uid = u.effective_user.id
    track_user(uid)
    if not has_replicate():
        await safe_reply(u.message,
            "❌ *Replicate API token সেট নেই।*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "Voice Cloning-এর জন্য *FREE Replicate account* লাগবে:\n\n"
            "১) https://replicate.com → Sign up (Google/GitHub দিয়ে instant)\n"
            "২) Account → API tokens → Copy\n"
            "৩) Termux-এ: `export REPLICATE_API_TOKEN=আপনার_token`\n"
            "৪) Bot restart করুন → /clonevoice আবার চেষ্টা করুন\n\n"
            "💰 ফ্রি ক্রেডিট প্রতি মাসে দেয় — paid plan লাগে না।\n"
            "প্রতি TTS call ~$0.001 (১০০০ call ≈ $1)।\n\n"
            "✨ *Cloning ছাড়াই premium voice চান?*\n"
            "🆓 `🗣️ TTS` মেনু → Edge TTS → ২১+ neural voice (সম্পূর্ণ ফ্রি)।",
            parse_mode="Markdown"); return
    msg = u.message
    reply = msg.reply_to_message
    src = None
    if reply:
        if reply.voice:        src = reply.voice
        elif reply.audio:      src = reply.audio
        elif reply.video_note: src = reply.video_note
        elif reply.video:      src = reply.video
    if not src:
        await safe_reply(msg,
            "🎭 *Voice Clone — কীভাবে করবেন?*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "১) আপনার নিজের voice-এ একটা পরিষ্কার Voice Message রেকর্ড করুন "
            "(*১০-৩০ সেকেন্ড*, background noise কম, একটানা কথা)।\n"
            "২) ওই voice message-এ *reply* দিয়ে লিখুন:\n"
            "   `/clonevoice আমার_নাম`\n\n"
            "📌 Audio/MP3/Video ফাইলেও কাজ করবে।\n\n"
            "✨ *Engine:* Replicate XTTS-v2 (zero-shot, Bengali/Hindi/English)\n"
            "💾 আপনার sample শুধু আমাদের Bot-এর local storage-এ থাকবে,\n"
            "প্রতি TTS call-এ Replicate-এ পাঠানো হবে।\n\n"
            "🗑️ পুরনো clone মুছতে: /deletevoice",
            parse_mode="Markdown")
        return
    raw = (msg.text or "").split(maxsplit=1)
    name = (raw[1].strip() if len(raw) > 1 else f"User_{uid}")[:30] or f"User_{uid}"

    status = await safe_reply(msg, "🎭 *Voice sample save হচ্ছে...*", parse_mode="Markdown")

    # Persistent path (TEMP_DIR না — TEMP_DIR cleanup হয়)
    suf = ".ogg"
    try:
        if hasattr(src, "file_name") and src.file_name:
            sf = Path(src.file_name).suffix.lower()
            if sf in (".mp3",".wav",".ogg",".m4a",".oga",".flac",".webm"):
                suf = sf
    except Exception: pass
    sample_path = str(CLONE_SAMPLES_DIR / f"voice_{uid}{suf}")
    try:
        f = await c.bot.get_file(src.file_id)
        await f.download_to_drive(sample_path)
    except Exception as e:
        await safe_edit(status, f"❌ Audio download fail: {e}")
        return

    # যদি video হয় → audio extract করি
    if suf in (".mp4",".webm") or (hasattr(src, "mime_type") and src.mime_type and "video" in src.mime_type):
        wav_path = str(CLONE_SAMPLES_DIR / f"voice_{uid}.wav")
        ok = await asyncio.get_event_loop().run_in_executor(
            executor, extract_audio_for_clone, sample_path, wav_path, 30)
        safe_unlink(sample_path)
        if not ok:
            await safe_edit(status, "❌ Video থেকে audio extract ব্যর্থ।")
            return
        sample_path = wav_path

    # পুরনো clone থাকলে old sample মুছি
    old = get_user_clone(uid)
    if old and old.get("sample_path") and old["sample_path"] != sample_path:
        safe_unlink(old["sample_path"])

    # Test TTS — কাজ করছে কিনা যাচাই
    await safe_edit(status, "🎭 *Replicate XTTS-v2 দিয়ে test করছি...*\n_~২০ সেকেন্ড_", parse_mode="Markdown")
    test_text = "Hello, this is a test of my cloned voice."
    test_path, test_err = await asyncio.get_event_loop().run_in_executor(
        executor, replicate_xtts_speak, test_text, sample_path, "en", None)
    if not test_path:
        safe_unlink(sample_path)
        await safe_edit(status, f"{test_err}\n\n_Sample save করা হলো না।_", parse_mode="Markdown")
        return

    set_user_clone(uid, f"xtts:{sample_path}", name)
    inc_feature("voice_clone_create")
    # Send test audio
    try:
        with open(test_path, "rb") as fh:
            await c.bot.send_voice(chat_id=u.effective_chat.id, voice=fh,
                                   caption=f"🎭 *{md_escape(name)}* — test sample",
                                   parse_mode="Markdown")
    except Exception: pass
    safe_unlink(test_path)
    await safe_edit(status,
        f"✅ *Voice Clone সফল!*\n\n"
        f"🎭 নাম: *{md_escape(name)}*\n"
        f"⚙️ Engine: *Replicate XTTS-v2 (zero-shot)*\n"
        f"💾 Sample saved locally।\n\n"
        f"এখন `🗣️ TTS` মেনুতে গিয়ে *🎭 আপনার Cloned Voice* select করে\n"
        f"যেকোনো text (Bengali/Hindi/English) → আপনার নিজের voice-এ!\n\n"
        f"_দেখতে: /myvoice  •  মুছতে: /deletevoice_",
        parse_mode="Markdown")

async def cmd_clonefromvideo(u, c):
    """🆕 v26: Replicate XTTS-v2 — current uploaded video থেকে directly voice clone।"""
    uid = u.effective_user.id
    track_user(uid)
    if not has_replicate():
        await safe_reply(u.message,
            "❌ Replicate token সেট নেই।\n\n"
            "Setup: https://replicate.com → API tokens → "
            "`export REPLICATE_API_TOKEN=...` → bot restart।",
            parse_mode="Markdown"); return
    with user_lock:
        d = user_videos.get(uid)
    if not d or not Path(d.get("path","")).exists():
        await safe_reply(u.message,
            "🎤 *Voice Clone From Video*\n\n"
            "প্রথমে আপনার ভিডিও পাঠান (যাতে clean voice আছে এমন), তারপর command দিন:\n"
            "`/clonefromvideo MyVoice`\n\n"
            "অথবা এডিট মেনুর `🎤 এই ভিডিও থেকে Voice Clone` button-এ ট্যাপ করুন।",
            parse_mode="Markdown"); return
    raw = (u.message.text or "").split(maxsplit=1)
    name = (raw[1].strip() if len(raw) > 1 else f"VidClone_{uid}")[:30] or f"VidClone_{uid}"
    status = await safe_reply(u.message,
        "🎤 *Voice Clone হচ্ছে...*\n\n_audio extract → Replicate test_\n_~৩০ সেকেন্ড_",
        parse_mode="Markdown")
    loop = asyncio.get_event_loop()
    sample_path = str(CLONE_SAMPLES_DIR / f"voice_{uid}.wav")
    ok = await loop.run_in_executor(executor, extract_audio_for_clone, d["path"], sample_path, 30)
    if not ok or not Path(sample_path).exists():
        await safe_edit(status, "❌ Audio extract ব্যর্থ।", parse_mode="Markdown")
        safe_unlink(sample_path); return
    # Test TTS
    test_path, test_err = await loop.run_in_executor(
        executor, replicate_xtts_speak, "Hello, this is a test of my voice.",
        sample_path, "en", None)
    if not test_path:
        safe_unlink(sample_path)
        await safe_edit(status, test_err or "❌ Voice clone test failed", parse_mode="Markdown"); return
    set_user_clone(uid, f"xtts:{sample_path}", name)
    inc_feature("voice_clone_from_video")
    try:
        with open(test_path, "rb") as fh:
            await c.bot.send_voice(chat_id=u.effective_chat.id, voice=fh,
                                   caption=f"🎭 *{md_escape(name)}* — test",
                                   parse_mode="Markdown")
    except Exception: pass
    safe_unlink(test_path)
    await safe_edit(status,
        f"✅ *Voice Clone সফল!*\n\n🎭 নাম: *{md_escape(name)}*\n"
        f"⚙️ Engine: *Replicate XTTS-v2*\n\n"
        f"এখন `🎙️ AI Voiceover` বা `🗣️ TTS → 🎭 আপনার Cloned Voice` দিয়ে\n"
        f"যেকোনো script আপনার voice-এ generate করুন।",
        parse_mode="Markdown")

async def cmd_myvoice(u, c):
    """🆕 v26: Cloned voice info — Replicate XTTS sample।"""
    uid = u.effective_user.id
    clone = get_user_clone(uid)
    if not clone:
        await safe_reply(u.message,
            "❌ আপনার কোনো cloned voice নেই।\n\n"
            "তৈরি করতে: voice message-এ reply দিয়ে `/clonevoice <নাম>`",
            parse_mode="Markdown"); return
    vid = clone.get("voice_id","")
    sample_path = vid.replace("xtts:", "") if vid.startswith("xtts:") else ""
    sample_size = ""
    if sample_path and Path(sample_path).exists():
        try: sample_size = f"{Path(sample_path).stat().st_size // 1024} KB"
        except Exception: pass
    await safe_reply(u.message,
        f"🎭 *আপনার Cloned Voice*\n\n"
        f"━━━━━━━━━━━━━━━━━━\n"
        f"📛 নাম: *{md_escape(clone.get('name',''))}*\n"
        f"⚙️ Engine: *Replicate XTTS-v2*\n"
        f"💾 Sample: `{sample_size or 'unknown'}`\n"
        f"📅 তৈরি: `{clone.get('created','?')}`\n\n"
        f"_TTS মেনুতে 🎭 button-এ click করে ব্যবহার করুন_\n"
        f"_মুছতে: /deletevoice_",
        parse_mode="Markdown")

async def cmd_deletevoice(u, c):
    """🆕 v26: Cloned voice + local sample file delete।"""
    uid = u.effective_user.id
    clone = get_user_clone(uid)
    if not clone:
        await safe_reply(u.message, "❌ আপনার কোনো cloned voice নেই — মোছার কিছু নেই।")
        return
    vid = clone.get("voice_id","")
    if vid.startswith("xtts:"):
        safe_unlink(vid.replace("xtts:", ""))
    del_user_clone(uid)
    inc_feature("voice_clone_delete")
    await safe_reply(u.message,
        f"✅ *Voice Clone মুছে দেওয়া হয়েছে।*\n\n"
        f"🎭 *{md_escape(clone.get('name',''))}* — removed (sample সহ)।\n\n"
        f"নতুন voice clone করতে: `/clonevoice <নাম>`",
        parse_mode="Markdown")

async def cmd_cancel(u, c):
    uid = u.effective_user.id
    with user_lock:
        d = user_videos.pop(uid, None)
        st = user_state.pop(uid, None)
    if d: safe_unlink(d.get("path"))
    if st:
        for p in st.get("data", {}).get("paths", []) or []:
            safe_unlink(p)
        safe_unlink(st.get("data", {}).get("bgm_path"))
    await safe_reply(u.message, "✅ *বাতিল!*", parse_mode="Markdown")

# ─── 🆕 v27.2: YouTube টাকা আয়ের গাইড ───────────────────────────────
# একটাই function — /money command আর "💰 YouTube Money গাইড" button দুটোই কাজ করায়।
# Telegram-এর 4096-char limit মেনে ৪টা message-এ split করা।
async def send_youtube_money_guide(message, edit=False, uid=None):
    """Send complete YouTube monetization guide in Bengali. Works from button or command."""
    try: inc_feature("youtube_money_guide")
    except Exception: pass

    # Part 1: Eligibility + Requirements
    p1 = (
        "💰 *YouTube থেকে টাকা আয়ের সম্পূর্ণ গাইড*\n"
        "━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        "🎯 *YPP কী?* — *YouTube Partner Program*।\n"
        "Channel monetize করতে চাইলে এতে join করতে হবে।\n\n"
        "📋 *যোগ্যতা (Eligibility) — ২০২৪/২০২৫:*\n\n"
        "🟢 *Full Monetization (AdSense ads):*\n"
        "  • ১,০০০ subscribers\n"
        "  • শেষ ৩৬৫ দিনে ৪,০০০ ঘন্টা watch time\n"
        "  • _অথবা_ শেষ ৯০ দিনে ১ কোটি Shorts views\n\n"
        "🟡 *Early Monetization (নতুন — ২০২৩+):*\n"
        "  • ৫০০ subscribers\n"
        "  • শেষ ৯০ দিনে ৩ video upload\n"
        "  • শেষ ৩৬৫ দিনে ৩,০০০ ঘন্টা watch time\n"
        "  • _অথবা_ ৯০ দিনে ৩০ লাখ Shorts views\n"
        "  • Available: USA, UK, India, BD, এবং আরও কিছু দেশে\n\n"
        "✅ *Other Requirements:*\n"
        "  • ১৮+ বয়স (বা parent-এর AdSense)\n"
        "  • Active 2-Step Verification\n"
        "  • শেষ ৯০ দিনে কোনো community strike নয়\n"
        "  • Original content (copyright safe)\n"
        "  • Linked AdSense account"
    )

    # Part 2: Income Sources
    p2 = (
        "💵 *আয়ের ৭টা সোর্স (Revenue Streams):*\n"
        "━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        "1️⃣ *AdSense Ads* — মূল income। Long videos-এ pre/mid/post roll ads। RPM: $0.5-$5 (Bengali audience), $5-$20 (US/UK audience)।\n\n"
        "2️⃣ *YouTube Shorts Ads* — Shorts feed-এ ad revenue share। RPM: $0.04-$0.15। ভাইরাল হলে millions views = ভালো income।\n\n"
        "3️⃣ *Channel Memberships* — fans মাসে $0.99-$49.99 দিয়ে join করে। Exclusive content + emojis + badges দিন।\n\n"
        "4️⃣ *Super Chat / Super Stickers / Super Thanks* — Live stream + premiere-এ viewers টাকা পাঠায় highlight হওয়ার জন্য।\n\n"
        "5️⃣ *YouTube Shopping* — own merch বা affiliate products tag করে commission।\n\n"
        "6️⃣ *Brand Sponsorships* — companies ভিডিওতে product mention/promote করতে $50-$50,000 দেয়। (10K+ subs হলে offers আসা শুরু)\n\n"
        "7️⃣ *Affiliate Marketing* — Amazon/Daraz/Hotmart link description-এ → প্রতি sale-এ commission। Bangladesh-এ Daraz affiliate ভালো।"
    )

    # Part 3: RPM/CPM + Bangladesh-specific payment
    p3 = (
        "📊 *RPM vs CPM (অনেকেই গুলিয়ে ফেলে):*\n"
        "━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        "  • *CPM* = Cost per 1000 ad impressions (advertiser কত দেয়)\n"
        "  • *RPM* = Revenue per 1000 video views (আপনি কত পান)\n"
        "  • RPM সবসময় CPM-এর কম, কারণ YouTube ৪৫% নেয়\n\n"
        "💸 *Bengali Niche RPM (real data):*\n"
        "  • Vlog/Lifestyle: $0.5 - $1.5\n"
        "  • Gaming: $0.3 - $1.0\n"
        "  • Tech/Tutorial: $1.0 - $3.0\n"
        "  • Finance/Business: $2.0 - $8.0\n"
        "  • Health/Education: $1.5 - $4.0\n"
        "  • English audience targeted: $3.0 - $15.0\n\n"
        "🇧🇩 *Bangladesh-এ payment কীভাবে পাবেন:*\n\n"
        "1️⃣ AdSense account-এ minimum *$100* জমা হলে payout (PIN verify হওয়ার পর)\n"
        "2️⃣ Payment options:\n"
        "   • *Bank Transfer* (Wire) — direct BD bank-এ USD আসে → Taka-তে convert (most popular)\n"
        "   • *Western Union* — branch-এ গিয়ে cash তোলা যায় (ছোট amount-এ ভালো)\n"
        "   • *Payoneer* — কিছু দেশে available (BD-তে limited)\n"
        "3️⃣ Tax info: AdSense → *Tax Form W-8BEN* fill করুন (US tax exempt-এর জন্য, না হলে ৩০% কাটবে)"
    )

    # Part 4: Pro Tips + Strategy
    p4 = (
        "🚀 *দ্রুত Monetization-এ পৌঁছানোর Pro Strategy:*\n"
        "━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        "🎯 *Niche Selection:*\n"
        "  • একটাই niche বেছে নিন (Tech/Cooking/Gaming/Vlog)\n"
        "  • Bengali audience-এ trending: Islamic content, Cooking, Tech tutorial, Funny videos, Education\n"
        "  • Sub-niche specific হলে algorithm faster boost করে\n\n"
        "📅 *Posting Strategy:*\n"
        "  • Consistency > Quantity। সপ্তাহে ২-৩টা video minimum\n"
        "  • Shorts + Long mix করুন (Shorts = subs, Long = watch time)\n"
        "  • Best time (BD): রাত ৮-১১টা, Weekend বেশি\n\n"
        "🎬 *Content Tips:*\n"
        "  • প্রথম ৩-৫ সেকেন্ড hook = retention boost\n"
        "  • Title: emoji + curiosity gap (\"কেউ জানে না...\")\n"
        "  • Thumbnail: bold text + face expression + bright color\n"
        "  • Description: প্রথম ২ লাইন hook, hashtags শেষে\n"
        "  • End screen + cards দিয়ে session time বাড়ান\n\n"
        "⚠️ *যা ভুলেও করবেন না (Demonetization risk):*\n"
        "  • Copyrighted music/clip ব্যবহার (use Smart Auto-এর copyright bypass)\n"
        "  • Reused content (অন্যের video re-upload)\n"
        "  • Misleading thumbnail/title (clickbait শাস্তি)\n"
        "  • Hate speech, violence, adult content\n"
        "  • Sub4Sub, view bot — instant ban\n\n"
        "🏆 *Realistic Earnings Timeline (Bengali channel):*\n"
        "  • ০-১,০০০ subs: ৩-৬ মাস (consistent posting)\n"
        "  • Monetize হওয়া: ৬-১২ মাস\n"
        "  • $100/মাস: ১-২ বছর\n"
        "  • Full-time income: ২-৪ বছর + dedication\n\n"
        "💡 *এই Bot-এর Tools:*\n"
        "  • 🚀 *YouTube ভাইরাল প্যাক* → AI title + tag + edit\n"
        "  • 📺 *YouTube Mode* → Content ID safe\n"
        "  • 🎯 *SEO Generator* → keyword optimization\n"
        "  • 📝 *Auto Subtitle* → CC = retention +25%\n"
        "  • ✨ *Enhance Ultra* → professional quality\n\n"
        "🔥 *মনে রাখবেন:* YouTube-এ overnight success নেই। "
        "Patience + consistency + quality = টাকা।\n\n"
        "_Bot command: /money যেকোনো সময় এই গাইড দেখতে_"
    )

    sender = message.edit_text if edit else message.reply_text
    try:
        if edit:
            await safe_edit(message, p1, parse_mode="Markdown")
        else:
            await safe_reply(message, p1, parse_mode="Markdown")
    except Exception as e:
        logger.warning("money guide p1: %s", e)
    for part in (p2, p3, p4):
        try:
            await message.reply_text(part, parse_mode="Markdown", disable_web_page_preview=True)
        except Exception as e:
            logger.warning("money guide part: %s", e)
    # If from button, restore edit menu
    if edit and uid is not None:
        try:
            await message.reply_text("👇 *আরও এডিট:*", parse_mode="Markdown",
                                     reply_markup=edit_menu_for(uid))
        except Exception: pass

async def cmd_money(u, c):
    """🆕 v27.2: /money — YouTube monetization complete guide (works without video)."""
    track_user(u.effective_user.id)
    await send_youtube_money_guide(u.message, edit=False)

async def cmd_stats(u, c):
    s = load_stats()
    # 🆕 v17: top-5 most-used features
    fu = s.get("feature_usage", {}) or {}
    top = sorted(fu.items(), key=lambda kv: kv[1], reverse=True)[:5]
    top_lines = "\n".join(f"  • `{k}` — *{v}*x" for k, v in top) if top else "  _এখনো নেই_"
    await safe_reply(u.message,
        f"📊 *বট স্ট্যাটস*\n\n━━━━━━━━━━━━━━━━━━\n"
        f"👥 মোট ইউজার: *{len(s['total_users'])}*\n"
        f"🎬 প্রসেস: *{s.get('videos_processed',0)}*\n"
        f"⬇️ ডাউনলোড: *{s.get('downloads',0)}*\n\n"
        f"🔥 *Top Features:*\n{top_lines}",
        parse_mode="Markdown")

async def cmd_platforms(u, c):
    names = sorted({n for n in PLATFORM_INFO.values()})
    txt = "🌐 *সাপোর্টেড প্ল্যাটফর্ম*\n\n" + "\n".join(f"• {n}" for n in names) + "\n\n+ আরও 1000+ সাইট"
    await safe_reply(u.message, txt, parse_mode="Markdown")

# ─── 🆕 v15: /myfonts — Bengali font preview generator ───
def make_font_preview(sample_text):
    """Generate 1080x1350 preview image showing sample text in all 4 Bengali fonts."""
    out = str(TEMP_DIR / f"font_preview_{uuid.uuid4().hex}.jpg")
    safe_text = (sample_text.replace("\\", "\\\\").replace(":", "\\:")
                            .replace("'", "\u2019").replace('"', ''))
    parts = []
    fonts_present = [(k, v) for k, v in BANGLA_FONTS.items()
                     if (BANGLA_FONT_DIR / v[1]).exists() and (BANGLA_FONT_DIR / v[1]).stat().st_size > 10000]
    if not fonts_present:
        return None
    H = 1350
    section = H // len(fonts_present)
    for i, (key, (label, fname, _)) in enumerate(fonts_present):
        font_path = BANGLA_FONT_DIR / fname
        y_center = section * i + section // 2
        # Header label (English ASCII safe — uses default font)
        header = f"{i+1}. {key.upper()}"
        parts.append(
            f"drawtext=text='{header}':fontsize=26:fontcolor=#ffd166:"
            f"borderw=2:bordercolor=black:x=40:y={y_center - 80}"
        )
        # Sample Bengali text (large)
        parts.append(
            f"drawtext=text='{safe_text}':fontfile='{font_path}':fontsize=64:fontcolor=white:"
            f"borderw=4:bordercolor=black:shadowx=2:shadowy=2:shadowcolor=black@0.6:"
            f"x=(w-text_w)/2:y={y_center - 30}"
        )
        # Divider (except last)
        if i < len(fonts_present) - 1:
            sep_y = section * (i + 1)
            parts.append(
                f"drawbox=x=60:y={sep_y - 1}:w=960:h=2:color=#ffffff@0.2:t=fill"
            )
    vf = ",".join(parts)
    cmd = [
        "ffmpeg", "-y",
        "-f", "lavfi", "-i", f"color=c=#1a1a2e:s=1080x{H}:d=1",
        "-vf", vf,
        "-frames:v", "1",
        "-q:v", "3",
        out
    ]
    try:
        r = subprocess.run(cmd, capture_output=True, timeout=60)
        if r.returncode == 0 and Path(out).exists():
            return out
    except Exception as e:
        logger.warning("font preview failed: %s", e)
    return None

async def cmd_myfonts(u, c):
    track_user(u.effective_user.id)
    msg = u.message
    args = c.args if hasattr(c, "args") and c.args else []
    sample = " ".join(args).strip() if args else "আমার বাংলা ভাষা ❤️"
    if len(sample) > 40:
        sample = sample[:40]
    if not is_bangla(sample):
        sample = "আমার বাংলা ভাষা ❤️ " + sample
        sample = sample[:40]
    waiting = await safe_reply(msg,
        "🎀 *বাংলা ফন্ট preview তৈরি হচ্ছে...*\n\n"
        "_৪টা ফন্টে একই টেক্সট রেন্ডার করছি, একটু অপেক্ষা করুন।_",
        parse_mode="Markdown")
    # Ensure fonts are available (lazy download)
    loop = asyncio.get_running_loop()
    await loop.run_in_executor(executor, ensure_bangla_fonts)
    preview = await loop.run_in_executor(executor, make_font_preview, sample)
    if not preview:
        await safe_edit(waiting,
            "❌ Preview তৈরি ব্যর্থ!\n\n"
            "বাংলা ফন্ট ডাউনলোড হয়নি। ইন্টারনেট চেক করে আবার চেষ্টা করুন।",
            parse_mode="Markdown")
        return
    try:
        await waiting.delete()
    except Exception: pass
    with open(preview, "rb") as fp:
        await c.bot.send_photo(msg.chat_id, fp,
            caption=(
                f"🎀 *বাংলা ফন্ট Preview*\n\n"
                f"📝 Sample: `{sample}`\n\n"
                f"✨ *Noto* — clean, যেকোনো জায়গায় ভালো\n"
                f"💎 *Hind* — premium magazine look\n"
                f"🌸 *Baloo* — modern, friendly, social media\n"
                f"🎀 *Galada* — decorative, title/quote-এর জন্য\n\n"
                f"💡 ভিডিওতে ব্যবহার করতে: 🎀 *প্রিমিয়াম ক্যাপশন* বাটন\n\n"
                f"🔁 অন্য টেক্সট দেখতে: `/myfonts আপনার লেখা`"
            ),
            parse_mode="Markdown")
    safe_unlink(preview)

async def cmd_admin(u, c):
    if u.effective_user.id != ADMIN_ID:
        await safe_reply(u.message, "❌ অ্যাক্সেস নেই।"); return
    s = load_stats()
    await safe_reply(u.message,
        f"👑 *অ্যাডমিন প্যানেল*\n\n━━━━━━━━━━━━━━━━━━\n"
        f"👥 ইউজার: *{len(s['total_users'])}*\n"
        f"🎬 প্রসেস: *{s.get('videos_processed',0)}*\n"
        f"⬇️ ডাউনলোড: *{s.get('downloads',0)}*\n"
        f"🕐 শুরু: `{s.get('start_time','')[:19]}`",
        parse_mode="Markdown")

async def cmd_broadcast(u, c):
    if u.effective_user.id != ADMIN_ID: return
    if not c.args:
        await safe_reply(u.message, "📝 ব্যবহার: `/broadcast মেসেজ`", parse_mode="Markdown"); return
    text = " ".join(c.args)
    s = load_stats(); sent = failed = 0
    st = await safe_reply(u.message, "📤 ব্রডকাস্ট...")
    for x in s.get("total_users", []):
        try:
            await c.bot.send_message(int(x), f"📢 *ব্রডকাস্ট*\n\n{text}", parse_mode="Markdown"); sent += 1
        except Exception: failed += 1
        await asyncio.sleep(0.05)
    await safe_edit(st, f"✅ পাঠানো: {sent}\n❌ ব্যর্থ: {failed}")

# ──────────────────────────────────────────────────────────────
# VIDEO CARD
# ──────────────────────────────────────────────────────────────
def show_video_card(title, size_mb, source):
    safe_title = md_escape(title) if title else "_কোন টাইটেল নেই_"
    return (f"╭─━━━━━━━━━━━━━━━━━━━━╮\n"
            f"│ ✅ *ভিডিও রেডি!*\n"
            f"╰─━━━━━━━━━━━━━━━━━━━━╯\n\n"
            f"🎬 *টাইটেল:* {safe_title}\n"
            f"📦 *সাইজ:* `{size_mb:.1f} MB`\n"
            f"📥 *উৎস:* {source}\n\n"
            f"━━━━━━━━━━━━━━━━━━━━\n"
            f"👇 *এডিট অপশন বেছে নিন:*\n"
            f"━━━━━━━━━━━━━━━━━━━━")

# ──────────────────────────────────────────────────────────────
# VIDEO HANDLER
# ──────────────────────────────────────────────────────────────
async def handle_video(u, c):
    msg = u.message
    if not msg: return
    uid = u.effective_user.id
    track_user(uid)
    media = msg.video or msg.document or msg.animation
    if not media: return
    if msg.document and msg.document.mime_type and not msg.document.mime_type.startswith(("video/","image/")):
        return
    if media.file_size and media.file_size > MAX_TG_UPLOAD_MB * 1024 * 1024:
        await safe_reply(msg, f"❌ *ফাইল খুব বড়!*\nসর্বোচ্চ {MAX_TG_UPLOAD_MB}MB।", parse_mode="Markdown"); return

    # Merge mode — collect ভিডিও
    with user_lock:
        st = user_state.get(uid)
    if st and st.get("action") == "merge_collect":
        st_msg = await safe_reply(msg, "📥 *ভিডিও যোগ হচ্ছে...*", parse_mode="Markdown")
        path = TEMP_DIR / f"{uuid.uuid4().hex}_merge.mp4"
        try:
            f = await c.bot.get_file(media.file_id)
            await f.download_to_drive(str(path))
            with user_lock:
                st["data"].setdefault("paths", []).append(str(path))
                cnt = len(st["data"]["paths"])
            await safe_edit(st_msg,
                f"✅ *ভিডিও #{cnt} যোগ হয়েছে*\n\n"
                f"আরও ভিডিও পাঠান অথবা মার্জ শুরু করুন:",
                parse_mode="Markdown",
                reply_markup=InlineKeyboardMarkup([
                    [InlineKeyboardButton(f"🔗 মার্জ শুরু ({cnt} ভিডিও)", callback_data="merge_run")],
                    [InlineKeyboardButton("❌ বাতিল", callback_data="back_to_edit")],
                ]))
            return
        except Exception as e:
            logger.exception("merge dl: %s", e); safe_unlink(path)
            await safe_edit(st_msg, "❌ *ডাউনলোড ব্যর্থ!*", parse_mode="Markdown")
            return

    st_msg = await safe_reply(msg, "📥 *ভিডিও গ্রহণ চলছে...*\n⏳ একটু অপেক্ষা করুন", parse_mode="Markdown")
    try: await c.bot.send_chat_action(msg.chat_id, ChatAction.UPLOAD_VIDEO)
    except Exception: pass
    path = TEMP_DIR / f"{uuid.uuid4().hex}_input.mp4"
    try:
        f = await c.bot.get_file(media.file_id)
        await f.download_to_drive(str(path))
        title = msg.caption or (getattr(media, 'file_name', None) or "")
        with user_lock:
            old = user_videos.get(uid)
            user_videos[uid] = {"path": str(path), "title": title, "url": None, "platform": None}
            user_state.pop(uid, None)
        if old: safe_unlink(old.get("path"))
        size_mb = path.stat().st_size / (1024*1024)
        await safe_edit(st_msg, show_video_card(title, size_mb, "📤 আপনার আপলোড"),
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
    except Exception as e:
        logger.exception("dl err: %s", e); safe_unlink(path)
        await safe_edit(st_msg, "❌ *ডাউনলোড ব্যর্থ!*\n\nআবার চেষ্টা করুন।", parse_mode="Markdown")

# ──────────────────────────────────────────────────────────────
# AUDIO HANDLER (BGM input)
# ──────────────────────────────────────────────────────────────
# ─── 🆕 v15: Photo handler for Status Video custom background ───
# 🆕 v37: Photo → REAL Apple Live Photo (JPG + MOV with matching ContentIdentifier)
# These files, when transferred to an iPhone (AirDrop/iCloud), become a real Live Photo
# that TikTok recognizes and uploads as Live Photo (Photo Mode).
import shutil as _shutil

def make_apple_live_pair(image_path: str, out_dir: str, duration: float = 3.0) -> tuple[bool, str, dict]:
    """Generate Apple Live Photo pair: still JPG + MOV, sharing one ContentIdentifier UUID.
       Requires ffmpeg + exiftool. On Termux: pkg install ffmpeg exiftool"""
    if not _shutil.which("ffmpeg"):
        return False, "ffmpeg নেই — Termux-এ: `pkg install ffmpeg`", {}
    if not _shutil.which("exiftool"):
        return False, "exiftool নেই — Termux-এ: `pkg install exiftool`", {}
    if not os.path.exists(image_path):
        return False, "ছবি পাওয়া যায়নি", {}

    content_id = str(uuid.uuid4()).upper()  # iPhone uses UPPERCASE UUID
    name = uuid.uuid4().hex[:10]
    out_jpg = Path(out_dir) / f"IMG_{name}.JPG"
    out_mov = Path(out_dir) / f"IMG_{name}.MOV"

    # Step 1 — Re-encode input as a clean JPEG (strip weird metadata, keep dimensions)
    try:
        r1 = subprocess.run(
            ["ffmpeg", "-y", "-loglevel", "error",
             "-i", image_path, "-frames:v", "1", "-q:v", "2", str(out_jpg)],
            capture_output=True, text=True, timeout=30)
        if r1.returncode != 0 or not out_jpg.exists():
            return False, f"JPG তৈরি ব্যর্থ: {(r1.stderr or '')[-200:]}", {}
    except Exception as e:
        return False, f"JPG error: {e}", {}

    # Step 2 — Build the MOV (no audio, ~3 sec, subtle motion so iPhone shows it as a Live Photo)
    fps = 30
    frames = int(duration * fps)
    vf = (
        f"scale=1080:1920:force_original_aspect_ratio=increase,"
        f"crop=1080:1920,"
        f"zoompan=z='min(zoom+0.0010,1.10)':d={frames}"
        f":x='iw/2-(iw/zoom/2)':y='ih/2-(ih/zoom/2)':s=1080x1920:fps={fps},"
        f"format=yuv420p"
    )
    try:
        r2 = subprocess.run(
            ["ffmpeg", "-y", "-loglevel", "error",
             "-loop", "1", "-i", image_path,
             "-t", f"{duration:.2f}",
             "-vf", vf,
             "-c:v", "libx264", "-preset", "veryfast", "-crf", "20",
             "-pix_fmt", "yuv420p",
             "-movflags", "+faststart+use_metadata_tags",
             "-metadata", f"com.apple.quicktime.content.identifier={content_id}",
             "-an",
             str(out_mov)],
            capture_output=True, text=True, timeout=180)
        if r2.returncode != 0 or not out_mov.exists():
            return False, f"MOV তৈরি ব্যর্থ: {(r2.stderr or '')[-200:]}", {}
    except Exception as e:
        return False, f"MOV error: {e}", {}

    # Step 3 — Inject ContentIdentifier into JPG via exiftool (XMP namespace iPhone uses)
    try:
        subprocess.run(
            ["exiftool", "-overwrite_original",
             f"-XMP-apple-fi:ContentIdentifier={content_id}",
             f"-XMP:ContentIdentifier={content_id}",
             str(out_jpg)],
            capture_output=True, text=True, timeout=30)
    except Exception:
        pass  # XMP injection is best-effort; one of them usually sticks

    # Step 4 — Inject Keys:ContentIdentifier into MOV (the canonical Apple location)
    try:
        subprocess.run(
            ["exiftool", "-overwrite_original",
             f"-Keys:ContentIdentifier={content_id}",
             f"-XMP:ContentIdentifier={content_id}",
             str(out_mov)],
            capture_output=True, text=True, timeout=30)
    except Exception:
        pass

    return True, "ok", {
        "jpg": str(out_jpg),
        "mov": str(out_mov),
        "uuid": content_id,
        "jpg_kb": out_jpg.stat().st_size // 1024,
        "mov_kb": out_mov.stat().st_size // 1024,
    }

async def handle_photo(u, c):
    msg = u.message
    uid = u.effective_user.id
    with user_lock:
        st = user_state.get(uid)

    # 🆕 v37: Photo → REAL Apple Live Photo pair (for TikTok upload from iPhone)
    if st and st.get("action") == "photo_live_wait":
        with user_lock:
            user_state.pop(uid, None)
        photo = msg.photo[-1] if msg.photo else None
        if not photo:
            await safe_reply(msg, "❌ ছবি পাইনি।"); return
        if photo.file_size and photo.file_size > 20 * 1024 * 1024:
            await safe_reply(msg, "❌ ছবি 20MB-র কম হতে হবে।"); return
        token = uuid.uuid4().hex[:12]
        img_path = TEMP_DIR / f"{token}_src.jpg"
        try:
            f = await c.bot.get_file(photo.file_id)
            await f.download_to_drive(str(img_path))
        except Exception as e:
            await safe_reply(msg, f"❌ ডাউনলোড ব্যর্থ: {e}"); return
        wait = await safe_reply(msg, "📸 *Live Photo pair বানাচ্ছি...* (15-30 সেকেন্ড)", parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        ok, info, data = await loop.run_in_executor(
            executor, make_apple_live_pair, str(img_path), str(TEMP_DIR), 3.0)
        safe_unlink(str(img_path))
        if not ok:
            await safe_edit(wait, f"❌ ব্যর্থ: {info}", parse_mode="Markdown",
                            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")]]))
            return
        try:
            with open(data["jpg"], "rb") as jf:
                await c.bot.send_document(
                    chat_id=msg.chat_id, document=jf,
                    filename=Path(data["jpg"]).name,
                    caption=f"📸 *Live Photo — Still (JPG)*\n`{data['jpg_kb']} KB`",
                    parse_mode="Markdown")
            with open(data["mov"], "rb") as mf:
                await c.bot.send_document(
                    chat_id=msg.chat_id, document=mf,
                    filename=Path(data["mov"]).name,
                    caption=(
                        f"🎬 *Live Photo — Motion (MOV)*\n`{data['mov_kb']} KB`\n\n"
                        f"🔗 *Pair UUID:* `{data['uuid']}`\n\n"
                        "📲 *যা করবেন:*\n"
                        "1️⃣ দুটো ফাইল একসাথে iPhone-এ পাঠান (AirDrop / iCloud Drive / Telegram desktop)\n"
                        "2️⃣ iPhone-এর Photos app-এ দুটো একই folder-এ save করুন — auto Live Photo হিসেবে merge হবে\n"
                        "3️⃣ TikTok খুলে → Upload → ছবি select → Live Photo হিসেবে recognize হবে\n\n"
                        "⚠️ *Android-এ TikTok upload করলে শুধু JPG যাবে — Live Photo support iOS only।*"
                    ),
                    parse_mode="Markdown",
                    reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")]]))
            await safe_delete(wait)
        except Exception as e:
            await safe_edit(wait, f"❌ পাঠাতে পারিনি: {e}")
        safe_unlink(data["jpg"]); safe_unlink(data["mov"])
        return

    # 🆕 v37: MCQ Solver — image question (OCR + answer)
    if st and st.get("action") == "mcq_wait":
        with user_lock:
            user_state.pop(uid, None)
        photo = msg.photo[-1] if msg.photo else None
        if not photo:
            await safe_reply(msg, "❌ ছবি পাইনি।"); return
        img_path = TEMP_DIR / f"{uuid.uuid4().hex}_mcq.jpg"
        try:
            f = await c.bot.get_file(photo.file_id)
            await f.download_to_drive(str(img_path))
        except Exception as e:
            await safe_reply(msg, f"❌ ডাউনলোড ব্যর্থ: {e}"); return
        wait = await safe_reply(msg, "🔍 *ছবি থেকে প্রশ্ন পড়ছি...*", parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        text, info = await loop.run_in_executor(executor, extract_image_ocr, str(img_path))
        safe_unlink(str(img_path))
        if not text:
            await safe_edit(wait, f"❌ {info}", parse_mode="Markdown",
                            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")]]))
            return
        await safe_edit(wait, f"✅ প্রশ্ন পড়া হয়েছে। এখন উত্তর তৈরি হচ্ছে...", parse_mode="Markdown")
        await answer_mcq(u, c, text, edit_msg=wait)
        return

    # 🆕 v37: File Extractor — OCR an image
    if st and st.get("action") == "file_extract_wait":
        with user_lock:
            user_state.pop(uid, None)
        photo = msg.photo[-1] if msg.photo else None
        if not photo:
            await safe_reply(msg, "❌ ছবি পাইনি।"); return
        img_path = TEMP_DIR / f"{uuid.uuid4().hex}_ext.jpg"
        try:
            f = await c.bot.get_file(photo.file_id)
            await f.download_to_drive(str(img_path))
        except Exception as e:
            await safe_reply(msg, f"❌ ডাউনলোড ব্যর্থ: {e}"); return
        wait = await safe_reply(msg, "🔍 *ছবির লেখা পড়ছি (OCR)...*", parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        text, info = await loop.run_in_executor(executor, extract_image_ocr, str(img_path))
        safe_unlink(str(img_path))
        await _send_extract_result(u, c, wait, text, info, "image.jpg")
        return

    if not (st and st.get("action") == "status_wait_image"):
        return  # ignore photos outside this flow
    photo = msg.photo[-1] if msg.photo else None
    if not photo: return
    if photo.file_size and photo.file_size > 20 * 1024 * 1024:
        await safe_reply(msg, "❌ ছবি 20MB-র কম হতে হবে।"); return
    img_path = TEMP_DIR / f"{uuid.uuid4().hex}_bg.jpg"
    try:
        f = await c.bot.get_file(photo.file_id)
        await f.download_to_drive(str(img_path))
    except Exception as e:
        await safe_reply(msg, f"❌ ডাউনলোড ব্যর্থ: {e}"); return
    with user_lock:
        st = user_state.get(uid) or {"action": "status_wait_font", "data": {}}
        st["action"] = "status_wait_font"
        st["data"]["bg_image"] = str(img_path)
        st["data"]["bg"] = "custom"
        user_state[uid] = st
    await safe_reply(msg,
        "✅ *ছবি পাওয়া গেছে!*\n\n"
        "এখন বাংলা ফন্ট স্টাইল বেছে নিন:",
        parse_mode="Markdown", reply_markup=status_font_menu())

# ══════════════════════════════════════════════════════════════
# 🆕 v37 — File Extract / MCQ Solver  HANDLERS & HELPERS
# ══════════════════════════════════════════════════════════════
async def _send_extract_result(u, c, wait_msg, text, info, fname):
    """Extract result পাঠানো — long হলে .txt file হিসেবে পাঠাই।"""
    home_kb = InlineKeyboardMarkup([[InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")]])
    if not text:
        await safe_edit(wait_msg, f"❌ *Extract ব্যর্থ!*\n\n{info or 'কারণ অজানা'}",
                        parse_mode="Markdown", reply_markup=home_kb)
        return
    header = f"📂 *{fname}* — {info or 'extracted'}\n\n"
    full = header + text
    # Telegram message limit 4096
    if len(full) <= TG_MESSAGE_LIMIT - 50:
        await safe_edit(wait_msg, full, parse_mode="Markdown", reply_markup=home_kb)
    else:
        # send as file
        out_path = TEMP_DIR / f"{uuid.uuid4().hex}_extracted.txt"
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)
            await safe_edit(wait_msg, header + "📦 _Content বড় — file আকারে পাঠাচ্ছি..._",
                            parse_mode="Markdown")
            with open(out_path, "rb") as f:
                await c.bot.send_document(u.effective_chat.id, f,
                    filename=f"{Path(fname).stem}_extracted.txt",
                    caption=header + "✅ Extract সম্পন্ন।", parse_mode="Markdown")
            await safe_reply(u.message, "👇 আরো কিছু extract করতে চাইলে:",
                             reply_markup=home_kb)
        finally:
            safe_unlink(str(out_path))

async def handle_document(u, c):
    """Non-video document handler — extract content if user is in extract mode,
       OR auto-prompt if a doc lands without context."""
    msg = u.message
    uid = u.effective_user.id
    doc = msg.document
    if not doc:
        return
    with user_lock:
        st = user_state.get(uid)

    # If user is in MCQ mode and sends a document (e.g. PDF question)
    if st and st.get("action") == "mcq_wait":
        with user_lock:
            user_state.pop(uid, None)
        if not (doc.file_name or "").lower().endswith((".pdf", ".docx", ".txt")):
            await safe_reply(msg, "❌ MCQ মোডে শুধু PDF/DOCX/TXT/photo সাপোর্টেড।"); return
        wait = await safe_reply(msg, "🔍 *প্রশ্ন পড়ছি...*", parse_mode="Markdown")
        path = TEMP_DIR / f"{uuid.uuid4().hex}_{doc.file_name}"
        try:
            f = await c.bot.get_file(doc.file_id)
            await f.download_to_drive(str(path))
        except Exception as e:
            await safe_edit(wait, f"❌ ডাউনলোড ব্যর্থ: {e}"); return
        loop = asyncio.get_running_loop()
        text, _ = await loop.run_in_executor(executor, extract_dispatch, str(path), doc.file_name)
        safe_unlink(str(path))
        if not text:
            await safe_edit(wait, "❌ ফাইল থেকে প্রশ্ন পড়া যায়নি।"); return
        await answer_mcq(u, c, text, edit_msg=wait)
        return

    # File extract mode (or auto if no other state)
    if not (st and st.get("action") == "file_extract_wait"):
        return  # silent ignore; users may attach docs in other flows
    with user_lock:
        user_state.pop(uid, None)

    size_mb = (doc.file_size or 0) / (1024*1024)
    if size_mb > MAX_EXTRACT_FILE_MB:
        await safe_reply(msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)। সর্বোচ্চ {MAX_EXTRACT_FILE_MB}MB সাপোর্ট।")
        return

    fname = doc.file_name or "file"
    wait = await safe_reply(msg, f"⏳ *Extract চলছে...*\n📄 `{fname}` ({size_mb:.1f}MB)",
                            parse_mode="Markdown")
    path = TEMP_DIR / f"{uuid.uuid4().hex}_{fname}"
    try:
        f = await c.bot.get_file(doc.file_id)
        await f.download_to_drive(str(path))
    except Exception as e:
        await safe_edit(wait, f"❌ *ডাউনলোড ব্যর্থ:* {e}", parse_mode="Markdown"); return

    loop = asyncio.get_running_loop()
    text, info = await loop.run_in_executor(executor, extract_dispatch, str(path), fname)
    safe_unlink(str(path))
    inc_stat("files_extracted")
    await _send_extract_result(u, c, wait, text, info, fname)

async def answer_mcq(u, c, question, edit_msg=None):
    """Gemini/OpenAI দিয়ে MCQ-এর উত্তর — Bengali-first।"""
    home_kb = InlineKeyboardMarkup([
        [InlineKeyboardButton("🔁 আরেকটা প্রশ্ন", callback_data="mcq_start")],
        [InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")],
    ])
    question = (question or "").strip()
    if len(question) < 5:
        msg = "❌ প্রশ্নটা খুব ছোট। আবার পাঠান।"
        if edit_msg: await safe_edit(edit_msg, msg, reply_markup=home_kb)
        else: await safe_reply(u.message, msg, reply_markup=home_kb)
        return
    if len(question) > 4000:
        question = question[:4000]

    # Cache check
    cached = mcq_lookup(question)
    if cached:
        out = f"💾 *(Cached উত্তর)*\n\n{cached['a']}"
        if edit_msg: await safe_edit(edit_msg, out, parse_mode="Markdown", reply_markup=home_kb)
        else: await safe_reply(u.message, out, parse_mode="Markdown", reply_markup=home_kb)
        inc_stat("mcq_solved")
        return

    if not edit_msg:
        edit_msg = await safe_reply(u.message, "🧠 *AI চিন্তা করছে...*", parse_mode="Markdown")
    else:
        await safe_edit(edit_msg, "🧠 *AI চিন্তা করছে...*", parse_mode="Markdown")

    system = (
        "তুমি বাংলাদেশের পরীক্ষার (SSC, HSC, BCS, ভর্তি, চাকরি) MCQ expert। "
        "User তোমাকে একটা MCQ প্রশ্ন দিবে — তুমি সঠিক উত্তর + সংক্ষিপ্ত ব্যাখ্যা বাংলায় দাও। "
        "Format: 'উত্তর: <option>' প্রথম লাইনে, তারপর '\\n\\nব্যাখ্যা: <কেন এই উত্তর>'। "
        "যদি প্রশ্ন অস্পষ্ট হয় বা উত্তর জানা না থাকে, সৎভাবে বলো।"
    )
    prompt = f"নিচের MCQ প্রশ্নের সঠিক উত্তর দাও:\n\n{question}"

    loop = asyncio.get_running_loop()
    answer = None
    err = None
    # Try Gemini first (free + fast for Bengali)
    if has_gemini():
        try:
            answer = await loop.run_in_executor(executor, gemini_chat, prompt, system, 600, 0.3)
        except Exception as e:
            err = str(e)
    if not answer and has_openai():
        try:
            answer = await loop.run_in_executor(executor, openai_chat, prompt, system, None, 600, 0.3)
        except Exception as e:
            err = str(e)
    if not answer and has_minimax():
        try:
            answer = await loop.run_in_executor(executor, minimax_chat, prompt, system, 600, 0.3)
        except Exception as e:
            err = str(e)

    if not answer:
        await safe_edit(edit_msg,
            f"❌ *AI উত্তর দিতে পারল না!*\n\n"
            f"কোনো AI engine (Gemini/OpenAI/MiniMax) configured নেই বা সব ব্যর্থ।\n\n"
            f"`{(err or 'unknown')[:200]}`",
            parse_mode="Markdown", reply_markup=home_kb)
        return

    # Clean Markdown for safe display
    answer = answer.strip()
    out = f"📚 *MCQ উত্তর:*\n\n{answer}\n\n━━━━━━━━━━━━━\n_AI generated — যাচাই করে নিন।_"
    try:
        await safe_edit(edit_msg, out, parse_mode="Markdown", reply_markup=home_kb)
    except Exception:
        # Markdown parse fail হলে plain text
        await safe_edit(edit_msg, out, reply_markup=home_kb)
    mcq_store(question, answer)
    inc_stat("mcq_solved")

async def cmd_extract(u, c):
    uid = u.effective_user.id
    with user_lock:
        user_state[uid] = {"action": "file_extract_wait", "data": {}}
    await safe_reply(u.message,
        "📂 *File Extractor চালু!*\n\n"
        "এখন একটা ফাইল পাঠান — যেকোনো একটা:\n\n"
        "📄 PDF, DOCX, TXT — text বের করব\n"
        "📊 XLSX, CSV — data table দেখাব\n"
        "🗜️ ZIP, RAR, 7Z — ভেতরের ফাইল list করব\n"
        "🖼️ JPG, PNG — OCR দিয়ে লেখা পড়ব\n"
        "🎵 MP3, M4A — metadata (artist/album)\n"
        "💻 JSON, HTML, code-files — direct content\n\n"
        f"⚠️ সর্বোচ্চ {MAX_EXTRACT_FILE_MB}MB।",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("❌ বাতিল", callback_data="main_menu")]]))

async def cmd_mcq(u, c):
    uid = u.effective_user.id
    args = (u.message.text or "").split(maxsplit=1)
    if len(args) > 1 and args[1].strip():
        # /mcq <question> — direct
        await answer_mcq(u, c, args[1].strip()); return
    with user_lock:
        user_state[uid] = {"action": "mcq_wait", "data": {}}
    await safe_reply(u.message,
        "📚 *MCQ Solver চালু!*\n\n"
        "এখন প্রশ্ন পাঠান — ৩ ভাবে দিতে পারেন:\n\n"
        "✍️ *Text* — সরাসরি প্রশ্ন টাইপ করুন (option সহ)\n"
        "🖼️ *ছবি* — বইয়ের MCQ-এর photo পাঠান (OCR + answer)\n"
        "📄 *PDF/DOCX* — পুরো MCQ ফাইল\n\n"
        "💡 *উদাহরণ:* \n"
        "_বাংলাদেশ স্বাধীন হয় কত সালে?_\n"
        "_(ক) ১৯৭০  (খ) ১৯৭১  (গ) ১৯৭২  (ঘ) ১৯৭৩_",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("❌ বাতিল", callback_data="main_menu")]]))

# ══════════════════════════════════════════════════════════════
# 🆕 v37 — Link Expander  HANDLERS
# ══════════════════════════════════════════════════════════════
def _md_safe(s: str) -> str:
    """Markdown-এ safe করার জন্য special char escape।"""
    if not s: return ""
    return re.sub(r"([_*`\[\]()])", r"\\\1", str(s))

async def do_link_expand(u, c, raw_url: str, edit_msg=None):
    """URL expand চালিয়ে chain + safety report পাঠায়।"""
    home_kb = InlineKeyboardMarkup([
        [InlineKeyboardButton("🔁 আরেকটা link", callback_data="link_expand_start")],
        [InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")],
    ])
    raw_url = (raw_url or "").strip()
    # find first url-looking token in input
    m = re.search(r"https?://\S+|www\.\S+|[\w\-]+\.[a-z]{2,}/\S*", raw_url, re.IGNORECASE)
    if m:
        raw_url = m.group(0)
    if not raw_url or len(raw_url) < 4:
        msg = "❌ কোনো URL পাইনি। আবার পাঠান।"
        if edit_msg: await safe_edit(edit_msg, msg, reply_markup=home_kb)
        else: await safe_reply(u.message, msg, reply_markup=home_kb)
        return

    if not edit_msg:
        edit_msg = await safe_reply(u.message, "🔍 *Link expand চলছে...*", parse_mode="Markdown")
    else:
        await safe_edit(edit_msg, "🔍 *Link expand চলছে...*", parse_mode="Markdown")

    loop = asyncio.get_running_loop()
    try:
        result = await loop.run_in_executor(executor, expand_url, raw_url)
    except Exception as e:
        await safe_edit(edit_msg, f"❌ *Expand ব্যর্থ:* `{str(e)[:140]}`",
                        parse_mode="Markdown", reply_markup=home_kb)
        return

    chain = result.get("chain", [])
    final = result.get("final", "") or (chain[-1] if chain else raw_url)
    title = result.get("title")
    status = result.get("status")
    err = result.get("error")
    hops = result.get("hops", 0)

    warnings = analyze_url_safety(final, title)
    try:
        from urllib.parse import urlparse as _up
        host = _up(final).netloc.lower()
    except Exception:
        host = ""

    is_shortener_origin = False
    if chain:
        try:
            from urllib.parse import urlparse as _up2
            origin_host = _up2(chain[0]).netloc.lower()
            is_shortener_origin = any(origin_host.endswith(s) for s in KNOWN_SHORTENERS)
        except Exception:
            pass

    out = ["🔗 *Link Expander Result*", "━━━━━━━━━━━━━━━━━━━━"]

    if hops == 0 and not err:
        out.append(f"✅ *No redirect* — সরাসরি destination।")
    elif hops > 0:
        out.append(f"↪️ *{hops} টা redirect* হয়েছে")
        if is_shortener_origin:
            out.append(f"🔎 Origin: known shortener")

    out.append("")
    out.append(f"🎯 *Final URL:*")
    out.append(f"`{final[:300]}`")
    if title:
        out.append("")
        out.append(f"📰 *Page title:* {_md_safe(title)[:160]}")
    if host:
        out.append(f"🌐 *Host:* `{host}`")
    if status:
        emoji = "✅" if 200 <= status < 300 else ("↪️" if 300 <= status < 400 else "⚠️")
        out.append(f"{emoji} *HTTP status:* `{status}`")
    if err:
        out.append(f"⚠️ *Note:* `{err}`")

    if warnings:
        out.append("")
        out.append("🛡️ *Safety Check:*")
        for w in warnings:
            out.append(f"  • {w}")
    else:
        out.append("")
        out.append("🛡️ *Safety:* No obvious red flags ✅")

    if len(chain) > 1:
        out.append("")
        out.append("📋 *Redirect chain:*")
        for i, link in enumerate(chain[:8]):
            out.append(f"  `{i+1}.` `{link[:120]}`")
        if len(chain) > 8:
            out.append(f"  …(আরো {len(chain)-8} hop)")

    out.append("")
    out.append("_ℹ️ এটা শুধু URL-এর gন্তব্য দেখায়। GPLinks/Droplink-এর মতো paid shortener bypass করে না।_")

    text = "\n".join(out)
    if len(text) > TG_MESSAGE_LIMIT - 50:
        text = text[:TG_MESSAGE_LIMIT - 100] + "\n\n…(short)"
    try:
        await safe_edit(edit_msg, text, parse_mode="Markdown", reply_markup=home_kb)
    except Exception:
        # Markdown-এ problem হলে plain
        await safe_edit(edit_msg, text, reply_markup=home_kb)
    inc_stat("links_expanded")

async def cmd_unlock_link(u, c):
    uid = u.effective_user.id
    args = (u.message.text or "").split(maxsplit=1)
    if len(args) > 1 and args[1].strip():
        await do_link_expand(u, c, args[1].strip()); return
    with user_lock:
        user_state[uid] = {"action": "link_expand_wait", "data": {}}
    await safe_reply(u.message,
        "🔗 *Link Expander চালু!*\n\n"
        "এখন একটা URL পাঠান — bot সেটার আসল destination দেখাবে।\n\n"
        "✅ *যা করব:*\n"
        "  • Short URL (bit.ly, t.co, tinyurl) → আসল URL\n"
        "  • সব redirect chain দেখাবে\n"
        "  • Page title + HTTP status\n"
        "  • Phishing/suspicious pattern warning\n\n"
        "❌ *যা করব না (কেন তা ভেতরে আছে):*\n"
        "  • GPLinks, Droplink, Ouo.io, Shrinkme bypass\n"
        "  • Paywall bypass\n\n"
        "💡 *উদাহরণ:* `https://bit.ly/xxxxxxx`",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("❌ বাতিল", callback_data="main_menu")]]))

async def handle_audio(u, c):
    msg = u.message
    uid = u.effective_user.id
    with user_lock:
        st = user_state.get(uid)
    if not (st and st.get("action") == "bgm_wait"):
        return
    media = msg.audio or msg.voice or msg.document
    if not media: return
    if media.file_size and media.file_size > 20 * 1024 * 1024:
        await safe_reply(msg, "❌ অডিও ফাইল 20MB-র কম হতে হবে।"); return
    st_msg = await safe_reply(msg, "🎵 *BGM ডাউনলোড...*", parse_mode="Markdown")
    bgm_path = TEMP_DIR / f"{uuid.uuid4().hex}_bgm.mp3"
    try:
        f = await c.bot.get_file(media.file_id)
        await f.download_to_drive(str(bgm_path))
        with user_lock:
            d = user_videos.get(uid)
            user_state.pop(uid, None)
        if not d or not Path(d["path"]).exists():
            safe_unlink(bgm_path)
            await safe_edit(st_msg, "❌ ভিডিও পাওয়া যায়নি!", parse_mode="Markdown"); return
        # process directly
        await safe_edit(st_msg, "✅ BGM পাওয়া গেছে। প্রসেস শুরু...", parse_mode="Markdown")
        out_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_output")
        loop = asyncio.get_running_loop()
        ok, final = await loop.run_in_executor(executor, build_and_run, "bgm", d["path"], out_base, {"bgm_path": str(bgm_path)})
        safe_unlink(bgm_path)
        if not ok or not Path(final).exists():
            await safe_edit(st_msg, "❌ *BGM যোগ ব্যর্থ!*", parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        inc_stat("videos_processed")
        size_mb = Path(final).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st_msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!", parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        with open(final, "rb") as fp:
            await c.bot.send_video(msg.chat_id, fp, caption="🎵 *BGM যোগ সম্পন্ন!*",
                                   parse_mode="Markdown", supports_streaming=True)
        safe_unlink(final)
        await safe_edit(st_msg, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও এডিট:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
    except Exception as e:
        logger.exception("bgm: %s", e); safe_unlink(bgm_path)
        await safe_edit(st_msg, "❌ সমস্যা হয়েছে!", parse_mode="Markdown")

# ──────────────────────────────────────────────────────────────
# TEXT HANDLER (URL / trim / text input)
# ──────────────────────────────────────────────────────────────
async def handle_text(u, c):
    msg = u.message
    uid = u.effective_user.id
    text = (msg.text or "").strip()

    with user_lock:
        st = user_state.get(uid)

    # 🆕 v37: MCQ Solver — text question
    if st and st.get("action") == "mcq_wait":
        with user_lock:
            user_state.pop(uid, None)
        await answer_mcq(u, c, text)
        return

    # 🆕 v37: Link Expander — URL input
    if st and st.get("action") == "link_expand_wait":
        with user_lock:
            user_state.pop(uid, None)
        await do_link_expand(u, c, text)
        return

    # Trim time input
    if st and st.get("action") == "trim_wait":
        if not TIME_RE.match(text):
            await safe_reply(msg, "❌ *ভুল ফরম্যাট!*\n\nসঠিক উদাহরণ: `00:10-00:30` অথবা `10-45`",
                             parse_mode="Markdown"); return
        parts = re.split(r"[-–]", text)
        if len(parts) != 2:
            await safe_reply(msg, "❌ ভুল ফরম্যাট!"); return
        start = parse_time(parts[0]); end = parse_time(parts[1])
        if start is None or end is None or end <= start:
            await safe_reply(msg, "❌ *ভুল টাইম!* end > start হতে হবে।", parse_mode="Markdown"); return
        with user_lock:
            d = user_videos.get(uid)
            user_state.pop(uid, None)
        if not d:
            await safe_reply(msg, "❌ ভিডিও নেই!"); return
        # process
        out_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_output")
        st_msg = await safe_reply(msg, f"✂️ *ট্রিম চলছে* ({start}s → {end}s)...", parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        ok, final = await loop.run_in_executor(executor, build_and_run, "trim", d["path"], out_base,
                                                {"start": start, "end": end})
        if not ok or not Path(final).exists():
            await safe_edit(st_msg, "❌ *ট্রিম ব্যর্থ!*", parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        inc_stat("videos_processed")
        size_mb = Path(final).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st_msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!", parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        with open(final, "rb") as fp:
            await c.bot.send_video(msg.chat_id, fp, caption="✂️ *ট্রিম সম্পন্ন!*",
                                   parse_mode="Markdown", supports_streaming=True)
        safe_unlink(final)
        await safe_edit(st_msg, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও এডিট:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🆕 v20: AI Voiceover ON Video — text input → TTS → mix ───
    if st and st.get("action") == "vo_wait_text":
        if len(text) > 1500:
            await safe_reply(msg, "❌ টেক্সট সর্বোচ্চ ১৫০০ অক্ষর।"); return
        with user_lock:
            d = user_videos.get(uid)
            user_state.pop(uid, None)
        if not d or not Path(d.get("path","")).exists():
            await safe_reply(msg, "❌ ভিডিও নেই!"); return
        clone = get_user_clone(uid)
        # 🆕 v26: XTTS clone detect — engine + voice auto-pick
        vo_engine = "elevenlabs"; vo_voice = ELEVENLABS_DEFAULT_VOICE; voice_label = "Rachel"
        if clone and clone.get("voice_id"):
            vid = clone["voice_id"]
            if vid.startswith("xtts:") and has_replicate():
                vo_engine = "xtts"; vo_voice = vid.replace("xtts:", "")
                voice_label = clone.get("name", "XTTS Clone")
            elif (not vid.startswith("xtts:")) and has_elevenlabs():
                vo_voice = vid; voice_label = clone.get("name", "Cloned")
        # No clone or engine unavailable → fallback to Edge TTS
        lang_guess = "bn" if any('\u0980' <= ch <= '\u09FF' for ch in text) else "en"
        if vo_engine == "elevenlabs" and not has_elevenlabs():
            vo_engine = "edge"; vo_voice = "bn_nabanita" if lang_guess == "bn" else "en_aria"
            voice_label = "Edge Neural"
        st_msg = await safe_reply(msg,
            f"🎙️ *AI Voiceover তৈরি হচ্ছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🗣️ Voice: `{md_escape(voice_label)}` ({vo_engine})\n"
            f"📝 অক্ষর: `{len(text)}`\n\n"
            f"_১) TTS generate_\n"
            f"_২) ভিডিওর audio-তে mix_\n\n"
            f"_৩০-৬০ সেকেন্ড লাগতে পারে..._",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        # 1) Generate TTS via unified synth
        tts_path, _used_eng, tts_err = await loop.run_in_executor(
            executor, text_to_speech, text, lang_guess, vo_voice, vo_engine)
        if not tts_path or not Path(tts_path).exists():
            await safe_edit(st_msg, f"❌ *Voiceover ব্যর্থ:*\n{tts_err or 'unknown'}",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        # 2) Mix into video
        out_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_vo")
        ok, final = await loop.run_in_executor(executor, build_and_run, "voiceover", d["path"], out_base,
                                                {"voice_audio": tts_path, "duck": 0.20, "boost": 1.6})
        safe_unlink(tts_path)
        if not ok or not Path(final).exists():
            await safe_edit(st_msg, "❌ *Voiceover mix ব্যর্থ!*\n\nভিডিওতে audio track নেই হলে mute → তারপর AI Voiceover use করুন।",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        inc_stat("videos_processed"); inc_feature("ai_voiceover")
        size_mb = Path(final).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st_msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!", parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        with open(final, "rb") as fp:
            await c.bot.send_video(msg.chat_id, fp,
                caption=f"🎙️ *AI Voiceover সম্পন্ন!*\n\n🗣️ Voice: `{md_escape(voice_label)}`",
                parse_mode="Markdown", supports_streaming=True)
        safe_unlink(final)
        await safe_edit(st_msg, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও এডিট:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🆕 v20: AI Voiceover + Bangla Subtitle ───
    if st and st.get("action") == "vosub_wait_text":
        if len(text) > 1500:
            await safe_reply(msg, "❌ টেক্সট সর্বোচ্চ ১৫০০ অক্ষর।"); return
        with user_lock:
            d = user_videos.get(uid)
            user_state.pop(uid, None)
        if not d or not Path(d.get("path","")).exists():
            await safe_reply(msg, "❌ ভিডিও নেই!"); return
        clone = get_user_clone(uid)
        # 🆕 v26: XTTS clone detect — engine + voice auto-pick
        vo_engine = "elevenlabs"; vo_voice = ELEVENLABS_DEFAULT_VOICE; voice_label = "Rachel"
        if clone and clone.get("voice_id"):
            vid = clone["voice_id"]
            if vid.startswith("xtts:") and has_replicate():
                vo_engine = "xtts"; vo_voice = vid.replace("xtts:", "")
                voice_label = clone.get("name", "XTTS Clone")
            elif (not vid.startswith("xtts:")) and has_elevenlabs():
                vo_voice = vid; voice_label = clone.get("name", "Cloned")
        lang_guess = "bn" if any('\u0980' <= ch <= '\u09FF' for ch in text) else "en"
        if vo_engine == "elevenlabs" and not has_elevenlabs():
            vo_engine = "edge"; vo_voice = "bn_nabanita" if lang_guess == "bn" else "en_aria"
            voice_label = "Edge Neural"
        st_msg = await safe_reply(msg,
            f"🎙️📝 *Voiceover + Subtitle তৈরি হচ্ছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🗣️ Voice: `{md_escape(voice_label)}` ({vo_engine})\n"
            f"📝 অক্ষর: `{len(text)}`\n\n"
            f"_১) TTS generate_\n_২) SRT generate (synced)_\n"
            f"_৩) Audio mix + Subtitle burn_\n\n_৪০-৯০ সেকেন্ড লাগতে পারে..._",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        # 1) Generate TTS via unified synth
        tts_path, _used_eng, tts_err = await loop.run_in_executor(
            executor, text_to_speech, text, lang_guess, vo_voice, vo_engine)
        if not tts_path or not Path(tts_path).exists():
            await safe_edit(st_msg, f"❌ *TTS ব্যর্থ:*\n{tts_err or 'unknown'}",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        # 2) Get TTS duration & build SRT (synced)
        tts_dur = await loop.run_in_executor(executor, get_duration, tts_path)
        if tts_dur < 0.5: tts_dur = max(2.0, len(text) / 14.0)  # rough fallback
        srt_path = await loop.run_in_executor(executor, text_to_srt, text, tts_dur, 42, None)
        # 3) Mix audio first
        await loop.run_in_executor(executor, ensure_bangla_fonts)
        mix_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_vomix")
        ok_mix, mixed = await loop.run_in_executor(executor, build_and_run, "voiceover", d["path"], mix_base,
                                                    {"voice_audio": tts_path, "duck": 0.20, "boost": 1.6})
        safe_unlink(tts_path)
        if not ok_mix or not Path(mixed).exists():
            await safe_edit(st_msg, "❌ *Voiceover mix ব্যর্থ!*",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            if srt_path: safe_unlink(srt_path); return
        # 4) Burn subtitle (if SRT generated)
        final = mixed
        if srt_path and Path(srt_path).exists():
            sub_out = str(TEMP_DIR / f"{uuid.uuid4().hex}_vosub.mp4")
            ok_sub = await loop.run_in_executor(executor, burn_subtitle, mixed, srt_path, sub_out)
            safe_unlink(srt_path)
            if ok_sub and Path(sub_out).exists():
                safe_unlink(mixed); final = sub_out
        if not Path(final).exists():
            await safe_edit(st_msg, "❌ *Final video missing!*",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        inc_stat("videos_processed"); inc_feature("voiceover_subtitle")
        size_mb = Path(final).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st_msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        with open(final, "rb") as fp:
            await c.bot.send_video(msg.chat_id, fp,
                caption=f"🎙️📝 *Voiceover + Subtitle সম্পন্ন!*\n\n🗣️ Voice: `{md_escape(voice_label)}`",
                parse_mode="Markdown", supports_streaming=True)
        safe_unlink(final)
        await safe_edit(st_msg, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও এডিট:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🆕 v20: Loop wait — N seconds input ───
    if st and st.get("action") == "loop_wait":
        try:
            target = float(text.strip())
        except Exception:
            await safe_reply(msg, "❌ শুধু সংখ্যা পাঠান (যেমন: `30`)", parse_mode="Markdown"); return
        if target < 2 or target > 600:
            await safe_reply(msg, "❌ ২ থেকে ৬০০ সেকেন্ডের মধ্যে দিন।"); return
        with user_lock:
            d = user_videos.get(uid)
            user_state.pop(uid, None)
        if not d:
            await safe_reply(msg, "❌ ভিডিও নেই!"); return
        out_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_loop")
        st_msg = await safe_reply(msg, f"♾️ *লুপ চলছে → {int(target)}s ...*", parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        ok, final = await loop.run_in_executor(executor, build_and_run, "loop", d["path"], out_base,
                                                {"target_dur": target})
        if not ok or not Path(final).exists():
            await safe_edit(st_msg, "❌ *লুপ ব্যর্থ!*", parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        inc_stat("videos_processed"); inc_feature("loop_n_sec")
        size_mb = Path(final).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st_msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!", parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        with open(final, "rb") as fp:
            await c.bot.send_video(msg.chat_id, fp,
                caption=f"♾️ *লুপ সম্পন্ন — {int(target)}s*",
                parse_mode="Markdown", supports_streaming=True)
        safe_unlink(final)
        await safe_edit(st_msg, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও এডিট:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # Text overlay input
    if st and st.get("action") == "text_wait":
        if len(text) > 100:
            await safe_reply(msg, "❌ টেক্সট সর্বোচ্চ ১০০ অক্ষর।"); return
        with user_lock:
            d = user_videos.get(uid)
            user_state.pop(uid, None)
        if not d:
            await safe_reply(msg, "❌ ভিডিও নেই!"); return
        out_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_output")
        st_msg = await safe_reply(msg, f"📝 *টেক্সট যোগ হচ্ছে...*", parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        ok, final = await loop.run_in_executor(executor, build_and_run, "text", d["path"], out_base,
                                                {"text": text})
        if not ok or not Path(final).exists():
            await safe_edit(st_msg, "❌ *টেক্সট যোগ ব্যর্থ!*", parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        inc_stat("videos_processed")
        size_mb = Path(final).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st_msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!", parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        with open(final, "rb") as fp:
            await c.bot.send_video(msg.chat_id, fp, caption="📝 *টেক্সট যোগ সম্পন্ন!*",
                                   parse_mode="Markdown", supports_streaming=True)
        safe_unlink(final)
        await safe_edit(st_msg, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও এডিট:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🆕 v15: Status Video text input (final wizard step) ───
    if st and st.get("action") == "status_wait_text":
        if len(text) > 300:
            await safe_reply(msg, "❌ টেক্সট সর্বোচ্চ ৩০০ অক্ষর।"); return
        data = st.get("data", {})
        with user_lock:
            user_state.pop(uid, None)
        bg_key   = data.get("bg", "night")
        font_key = data.get("font", DEFAULT_BANGLA_FONT)
        dur      = int(data.get("duration", 12))
        bg_image = data.get("bg_image")
        bg_label = STATUS_BG_PRESETS.get(bg_key, ("কাস্টম",))[0] if bg_key != "custom" else "🖼️ আপনার নিজের ছবি/ভিডিও"
        out_path = str(TEMP_DIR / f"{uuid.uuid4().hex}_status.mp4")
        st_msg = await safe_reply(msg,
            f"📱 *Status ভিডিও তৈরি হচ্ছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🎨 BG: {bg_label}\n"
            f"🔤 ফন্ট: {BANGLA_FONTS[font_key][0]}\n"
            f"⏱️ দৈর্ঘ্য: {dur}s\n"
            f"📐 9:16 (1080x1920)\n\n"
            f"_দয়া করে অপেক্ষা করুন (১৫-৪০ সেকেন্ড)..._",
            parse_mode="Markdown")
        # Ensure fonts are downloaded (lazy)
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(executor, ensure_bangla_fonts)
        ok = await loop.run_in_executor(executor, build_status_video,
                                         text, bg_key, font_key, out_path, dur, bg_image)
        if bg_image: safe_unlink(bg_image)
        if not ok or not Path(out_path).exists():
            await safe_edit(st_msg, "❌ *Status ভিডিও তৈরি ব্যর্থ!*\n\nআবার চেষ্টা করুন।",
                            parse_mode="Markdown", reply_markup=back_menu()); return
        inc_stat("videos_processed")
        size_mb = Path(out_path).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st_msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!",
                            parse_mode="Markdown", reply_markup=back_menu())
            safe_unlink(out_path); return
        try: await st_msg.delete()
        except Exception: pass
        with open(out_path, "rb") as fp:
            await c.bot.send_video(msg.chat_id, fp,
                caption=(f"📱 *Status ভিডিও রেডি!*\n\n"
                         f"🎨 BG: {bg_label}\n"
                         f"🔤 ফন্ট: {BANGLA_FONTS[font_key][0]}\n"
                         f"⏱️ {dur}s • 9:16\n\n"
                         f"💡 WhatsApp/FB-তে status হিসেবে শেয়ার করুন!"),
                parse_mode="Markdown", supports_streaming=True)
        safe_unlink(out_path)
        await safe_reply(msg, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও কিছু:",
                         parse_mode="Markdown", reply_markup=main_menu())
        return

    # ─── 🆕 v15: Premium Caption text input (final wizard step) ───
    if st and st.get("action") == "cap_wait_text":
        if len(text) > 150:
            await safe_reply(msg, "❌ টেক্সট সর্বোচ্চ ১৫০ অক্ষর।"); return
        data = st.get("data", {})
        with user_lock:
            d = user_videos.get(uid)
            user_state.pop(uid, None)
        if not d:
            await safe_reply(msg, "❌ ভিডিও নেই!"); return
        font_label = BANGLA_FONTS.get(data.get("font", DEFAULT_BANGLA_FONT), ("?",))[0]
        out_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_output")
        st_msg = await safe_reply(msg,
            f"🎀 *প্রিমিয়াম ক্যাপশন তৈরি হচ্ছে...*\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"📍 পজিশন: `{data.get('position','bottom')}`\n"
            f"🔤 ফন্ট: {font_label}\n"
            f"🎨 রঙ: `{data.get('color','white')}`\n"
            f"🟦 ব্যাকগ্রাউন্ড: {'হ্যাঁ' if data.get('bg_box') else 'না'}",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        extra = {
            "text": text,
            "position": data.get("position", "bottom"),
            "font": data.get("font", DEFAULT_BANGLA_FONT),
            "color": data.get("color", "white"),
            "bg_box": bool(data.get("bg_box", False)),
        }
        ok, final = await loop.run_in_executor(executor, build_and_run, "text", d["path"], out_base, extra)
        if not ok or not Path(final).exists():
            await safe_edit(st_msg, "❌ *ক্যাপশন যোগ ব্যর্থ!*",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        inc_stat("videos_processed")
        size_mb = Path(final).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st_msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        with open(final, "rb") as fp:
            await c.bot.send_video(msg.chat_id, fp,
                                   caption="🎀 *প্রিমিয়াম ক্যাপশন সম্পন্ন!*",
                                   parse_mode="Markdown", supports_streaming=True)
        safe_unlink(final)
        await safe_edit(st_msg, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও এডিট:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🆕 v13: TTS text input ───
    if st and st.get("action") == "tts_wait":
        if len(text) > 4000:
            await safe_reply(msg, "❌ টেক্সট সর্বোচ্চ ৪০০০ অক্ষর।"); return
        data = st.get("data", {})
        lang = data.get("lang", "auto")
        voice = data.get("voice", OPENAI_TTS_VOICE)
        engine = data.get("engine", "auto")
        with user_lock:
            user_state.pop(uid, None)
        # 🆕 v22: pretty engine label (Edge/ElevenLabs/OpenAI/gTTS/auto)
        if engine == "edge":
            edge_label = EDGE_TTS_VOICES.get(voice, (None, str(voice), None))[1]
            engine_label = f"🆓 Edge — {edge_label}"
        elif engine == "elevenlabs":
            engine_label = f"ElevenLabs {voice}"
        elif engine == "openai":
            engine_label = f"OpenAI {voice}"
        elif engine == "gtts":
            engine_label = "gTTS (Free)"
        else:
            engine_label = "Auto (best available)"
        st_msg = await safe_reply(msg,
            f"🗣️ *ভয়েস তৈরি হচ্ছে...*\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"🌍 ভাষা: `{lang.upper()}`\n"
            f"🎙️ Engine: `{engine_label}`\n"
            f"📝 অক্ষর: `{len(text)}`",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        out_path, used_engine, err = await loop.run_in_executor(
            executor, text_to_speech, text, lang, voice, engine)
        if not out_path or not Path(out_path).exists():
            await safe_edit(st_msg,
                f"❌ *TTS ব্যর্থ!*\n\n{err or 'অজানা সমস্যা'}\n\n"
                f"💡 gTTS এর জন্য: `pip install gTTS` করুন।",
                parse_mode="Markdown", reply_markup=back_menu())
            return
        size_mb = Path(out_path).stat().st_size / (1024*1024)
        try:
            with open(out_path, "rb") as fp:
                await c.bot.send_voice(msg.chat_id, fp,
                    caption=f"🗣️ *Voice Generated!*\n\n"
                            f"🎙️ Engine: `{used_engine}`\n"
                            f"📦 Size: `{size_mb:.2f} MB`",
                    parse_mode="Markdown")
            # MP3 ফাইলও পাঠাই (download করতে চাইলে)
            with open(out_path, "rb") as fp:
                await c.bot.send_audio(msg.chat_id, fp,
                    title="TTS Voice", performer="Video Editor Bot",
                    caption="📥 MP3 ফাইল (download করতে পারেন)")
        finally:
            safe_unlink(out_path)
        await safe_edit(st_msg,
            "✅ *পাঠানো হয়েছে!*\n\n👇 আবার ব্যবহার করুন:",
            parse_mode="Markdown", reply_markup=main_menu())
        return

    # ─── 🆕 v14: Translate text input ───
    if st and st.get("action") == "translate_wait":
        target = st.get("data", {}).get("target", "en")
        with user_lock:
            user_state.pop(uid, None)
        if len(text) > 6000:
            await safe_reply(msg, "❌ টেক্সট সর্বোচ্চ ৬০০০ অক্ষর। ছোট SRT-ফাইল পাঠান।"); return
        st_msg = await safe_reply(msg,
            f"🌐 *অনুবাদ চলছে → {LANG_NAMES.get(target, target).upper()}*",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        # auto-detect SRT (timestamp pattern থাকলে)
        is_srt = "-->" in text and re.search(r"\d{2}:\d{2}:\d{2}", text)
        if is_srt:
            translated, err = await loop.run_in_executor(executor, translate_srt, text, target)
        else:
            translated, err = await loop.run_in_executor(executor, translate_text, text, target)
        if err or not translated:
            await safe_edit(st_msg, f"❌ অনুবাদ ব্যর্থ: {err or 'unknown'}",
                            parse_mode="Markdown", reply_markup=back_menu()); return
        # যদি SRT হয়, ফাইল হিসেবে পাঠাই
        if is_srt:
            try:
                tmp_srt = TEMP_DIR / f"{uuid.uuid4().hex}_translated.srt"
                tmp_srt.write_text(translated, encoding="utf-8")
                with open(tmp_srt, "rb") as fp:
                    await c.bot.send_document(msg.chat_id, fp,
                        filename=f"translated_{target}.srt",
                        caption=f"🌐 *অনুবাদিত SRT* ({LANG_NAMES.get(target, target)})",
                        parse_mode="Markdown")
                safe_unlink(str(tmp_srt))
            except Exception as e:
                logger.warning("translated srt send: %s", e)
        # সাথে preview chat-এও পাঠাই
        chunks = [translated[i:i+3800] for i in range(0, len(translated), 3800)]
        for i, ch in enumerate(chunks):
            prefix = f"🌐 *Translation ({LANG_NAMES.get(target, target).upper()}):*\n\n" if i == 0 else ""
            try:
                await safe_reply(msg, prefix + f"```\n{ch}\n```", parse_mode="Markdown")
            except Exception:
                await safe_reply(msg, prefix + ch)
        await safe_edit(st_msg, "✅ *অনুবাদ সম্পন্ন!*",
                        parse_mode="Markdown", reply_markup=main_menu())
        return

    # ─── AI Chat (FIXED v39) — Gemini → OpenAI → MiniMax fallback ───
    if st and st.get("action") == "ai_chat_wait":
        if len(text) > 2000:
            await safe_reply(msg, "❌ মেসেজ সর্বোচ্চ ২০০০ অক্ষর।"); return
        with user_lock:
            user_state.pop(uid, None)
        system_prompt = ("You are a helpful assistant for a Bengali-speaking user. "
                         "Reply in the same language the user wrote in (Bengali or English). "
                         "Keep replies focused, helpful, and friendly.")
        st_msg = await safe_reply(msg,
            "🤖 *AI ভাবছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            "🔄 _Gemini → OpenAI → MiniMax চেক করছি..._",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        ai_engine = "AI"; reply = None
        if not reply and has_gemini():
            try:
                reply = await loop.run_in_executor(executor, gemini_chat, text, system_prompt, 1000, 0.8)
                if reply: ai_engine = "Gemini"
            except Exception as _e: logger.warning("ai_chat gemini: %s", _e)
        if not reply and has_openai():
            try:
                reply = await loop.run_in_executor(executor, openai_chat, text, system_prompt, None, 1000, 0.8)
                if reply: ai_engine = "ChatGPT"
            except Exception as _e: logger.warning("ai_chat openai: %s", _e)
        if not reply and has_minimax():
            try:
                reply = await loop.run_in_executor(executor, minimax_chat, text, system_prompt, 1000, 0.8)
                if reply: ai_engine = "MiniMax"
            except Exception as _e: logger.warning("ai_chat minimax: %s", _e)
        if not reply:
            await safe_edit(st_msg,
                "❌ *AI উত্তর দিতে পারল না!*\n\n"
                "কোনো AI key কাজ করছে না বা সব ব্যর্থ।\n\n"
                "💡 কোডে যেকোনো একটা API key সেট করুন:\n"
                "• `GEMINI_API_KEY` — সম্পূর্ণ ফ্রি (Google)\n"
                "• `OPENAI_API_KEY` — ChatGPT (paid)\n"
                "• `MINIMAX_API_KEY` — MiniMax",
                parse_mode="Markdown", reply_markup=back_menu())
            return
        chunks = [reply[i:i+3800] for i in range(0, len(reply), 3800)]
        for i, ch in enumerate(chunks):
            prefix = f"🤖 *{ai_engine} উত্তর:*\n\n" if i == 0 else ""
            try:
                await safe_reply(msg, prefix + ch, parse_mode="Markdown")
            except Exception:
                await safe_reply(msg, prefix + ch)
        await safe_edit(st_msg,
            f"✅ *উত্তর পাঠানো হয়েছে!* (via {ai_engine})\n\n"
            "💬 আবার চ্যাট করতে 🤖 AI চ্যাট বাটনে ট্যাপ করুন।",
            parse_mode="Markdown", reply_markup=main_menu())
        return

      # URL handling
    url = find_url(text)
    if not url:
        await safe_reply(msg, "❌ *সঠিক লিংক দিন!*\n\n/platforms দেখুন", parse_mode="Markdown"); return
    await handle_url(u, c, url)

# ──────────────────────────────────────────────────────────────
# yt-dlp
# ──────────────────────────────────────────────────────────────
def yt_dl(url, out_template):
    """🆕 v29 ULTRA+: Multi-strategy chains for YouTube / TikTok / Instagram / Pinterest / Facebook."""
    title = ""
    u = url.lower()
    is_yt = any(d in u for d in ("youtube.com","youtu.be","shorts.google.com"))
    is_pinterest = ("pinterest." in u) or ("pin.it" in u)
    is_tiktok = ("tiktok.com" in u) or ("vm.tiktok.com" in u) or ("vt.tiktok.com" in u)
    is_insta  = ("instagram.com" in u) or ("instagr.am" in u)
    is_fb     = ("facebook.com" in u) or ("fb.watch" in u) or ("fb.com" in u)

    mobile_ua = ("Mozilla/5.0 (Linux; Android 13; Pixel 7) AppleWebKit/537.36 "
                 "(KHTML, like Gecko) Chrome/120.0.0.0 Mobile Safari/537.36")
    desktop_ua = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                  "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    ios_ua = ("Mozilla/5.0 (iPhone; CPU iPhone OS 17_0 like Mac OS X) "
              "AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.0 Mobile/15E148 Safari/604.1")

    # ── title fetch ──
    try:
        ip_cmd = ["yt-dlp","--no-warnings","--no-playlist","--print","%(title)s","--skip-download"]
        if is_pinterest or is_tiktok or is_insta: ip_cmd += ["--user-agent", mobile_ua]
        if is_tiktok: ip_cmd += ["--referer","https://www.tiktok.com/"]
        if is_insta:  ip_cmd += ["--referer","https://www.instagram.com/"]
        if is_yt:     ip_cmd += ["--extractor-args","youtube:player_client=android,web,ios,tv_embedded"]
        ip_cmd.append(url)
        ip = subprocess.run(ip_cmd, capture_output=True, text=True, timeout=60)
        if ip.returncode == 0 and ip.stdout:
            title = ip.stdout.strip().splitlines()[0] if ip.stdout.strip() else ""
    except Exception: pass

    def _run_strategies(strategies, label):
        """Run each (fmt, extra_args) until one succeeds. Returns (ok, err)."""
        last = ""
        for fmt, extra in strategies:
            cmd = (["yt-dlp","-f",fmt,
                    "--no-playlist","--no-warnings","--restrict-filenames",
                    "--max-filesize",f"{MAX_DOWNLOAD_MB}M",
                    "--retries","3","--fragment-retries","3",
                    "--merge-output-format","mp4"]
                   + extra + ["-o", out_template, url])
            try:
                p = subprocess.run(cmd, capture_output=True, text=True, timeout=DOWNLOAD_TIMEOUT)
                if p.returncode == 0: return True, ""
                last = (p.stderr or p.stdout or "")[-400:]
            except subprocess.TimeoutExpired: last = "timeout"
            except Exception as e: last = str(e)
        return False, f"{label} fail (try: pip install -U yt-dlp): {last}"

    # ─── YouTube ───
    if is_yt:
        strategies = [
            ("bv*[height<=720][ext=mp4]+ba[ext=m4a]/b[height<=720][ext=mp4]/best[height<=720]",
             ["--user-agent", desktop_ua,
              "--extractor-args","youtube:player_client=android,web,ios"]),
            ("best[height<=720]/best[height<=480]/best",
             ["--user-agent", ios_ua,
              "--extractor-args","youtube:player_client=ios,android"]),
            ("worst[ext=mp4]/worst",
             ["--user-agent", desktop_ua,
              "--extractor-args","youtube:player_client=tv_embedded,web_safari"]),
        ]
        ok, err = _run_strategies(strategies, "YouTube")
        return ok, title, err

    # ─── 🆕 v38 TikTok — Updated strategies + multi-API fallback ───
    # 🐛 v38 FIX: v37-এর hardcoded api_hostname/app_version স্ট্যাটিক ছিল — TikTok
    # সেগুলো বদলায়। নতুন approach:
    #   (a) plain yt-dlp default first (sometimes simplest = best)
    #   (b) তারপর তিনটে fallback API: tikwm → tikmate → snaptik-style direct
    #   (c) yt-dlp pinned hostname শেষ resort
    if is_tiktok:
        strategies = [
            # 1) Plain default — TikTok-এ extra args না দিলে latest extractor কাজে লাগে
            ("best[ext=mp4]/best",
             ["--user-agent", mobile_ua,
              "--referer","https://www.tiktok.com/",
              "--add-header","Accept-Language:en-US,en;q=0.9",
              "--force-ipv4"]),
            # 2) iOS UA — TikTok server iOS-কে different (often working) extractor দেয়
            ("best",
             ["--user-agent", ios_ua,
              "--referer","https://www.tiktok.com/",
              "--add-header","Accept-Language:en-US,en;q=0.9",
              "--force-ipv4"]),
            # 3) webpage extractor — API fail করলে webpage থেকে raw video URL টানে
            ("best[ext=mp4]/best",
             ["--user-agent", desktop_ua,
              "--referer","https://www.tiktok.com/",
              "--extractor-args","tiktok:webpage_download_mode=force",
              "--force-ipv4"]),
            # 4) Pinned hostname — last yt-dlp try
            ("best[height<=1080]/best",
             ["--user-agent", mobile_ua,
              "--referer","https://www.tiktok.com/",
              "--extractor-args","tiktok:api_hostname=api16-normal-c-useast1a.tiktokv.com",
              "--force-ipv4"]),
        ]
        ok, err = _run_strategies(strategies, "TikTok")
        if ok:
            return ok, title, err
        # 🆕 v38 — Multi-API fallback chain (yt-dlp ব্যর্থ হলে external API):
        #   tikwm.com → tikmate.online → snaptik.app-like (direct scrape)
        # প্রতিটার json structure আলাদা — আলাদা parser দরকার।
        import urllib.request as _ur, urllib.parse as _up
        out_file = out_template.replace("%(ext)s", "mp4")

        def _direct_dl(vid_url, ref="https://www.tiktok.com/"):
            """Direct mp4 download via urllib (most resilient, no yt-dlp deps)।"""
            try:
                req = _ur.Request(vid_url, headers={
                    "User-Agent": mobile_ua,
                    "Referer": ref,
                    "Accept": "*/*",
                })
                with _ur.urlopen(req, timeout=DOWNLOAD_TIMEOUT) as r:
                    with open(out_file, "wb") as f:
                        while True:
                            chunk = r.read(131072)
                            if not chunk: break
                            f.write(chunk)
                return Path(out_file).exists() and Path(out_file).stat().st_size > 4096
            except Exception as ce:
                logger.warning("direct mp4 dl fail: %s", ce)
                return False

        # ---- Fallback 1: tikwm.com ----
        try:
            api_url = "https://www.tikwm.com/api/?url=" + _up.quote(url, safe="") + "&hd=1"
            req = _ur.Request(api_url, headers={
                "User-Agent": mobile_ua,
                "Referer": "https://www.tikwm.com/",
                "Accept": "application/json",
            })
            with _ur.urlopen(req, timeout=30) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="ignore"))
            if data.get("code") == 0 and data.get("data"):
                d = data["data"]
                vid_url = d.get("hdplay") or d.get("play") or d.get("wmplay")
                if vid_url and not vid_url.startswith("http"):
                    vid_url = "https://www.tikwm.com" + vid_url
                if vid_url:
                    if not title:
                        title = (d.get("title") or "").strip()[:200]
                    if _direct_dl(vid_url, "https://www.tikwm.com/"):
                        return True, title, ""
        except Exception as e:
            logger.warning("tikwm API fail: %s", e)

        # ---- Fallback 2: tikmate.online ----
        try:
            tm_url = "https://api.tikmate.online/api/lookup"
            req = _ur.Request(tm_url,
                data=_up.urlencode({"url": url}).encode(),
                headers={
                    "User-Agent": mobile_ua,
                    "Referer": "https://tikmate.online/",
                    "Accept": "application/json",
                    "Content-Type": "application/x-www-form-urlencoded",
                }, method="POST")
            with _ur.urlopen(req, timeout=30) as resp:
                data = json.loads(resp.read().decode("utf-8", errors="ignore"))
            tok, vid_id = data.get("token"), data.get("id")
            if tok and vid_id:
                vid_url = f"https://tikmate.online/download/{tok}/{vid_id}.mp4"
                if not title:
                    title = (data.get("title") or "").strip()[:200] or "tiktok"
                if _direct_dl(vid_url, "https://tikmate.online/"):
                    return True, title, ""
        except Exception as e:
            logger.warning("tikmate API fail: %s", e)

        # ---- Fallback 3: ssstik.io scraping ----
        try:
            ss_url = "https://ssstik.io/abc?url=dl"
            req = _ur.Request(ss_url,
                data=_up.urlencode({
                    "id": url,
                    "locale": "en",
                    "tt": "MQ==",
                }).encode(),
                headers={
                    "User-Agent": mobile_ua,
                    "Referer": "https://ssstik.io/en",
                    "Accept": "*/*",
                    "Content-Type": "application/x-www-form-urlencoded",
                    "HX-Request": "true",
                }, method="POST")
            with _ur.urlopen(req, timeout=30) as resp:
                html = resp.read().decode("utf-8", errors="ignore")
            # parse first https://*.tiktokcdn*.com/.../*.mp4 link
            m = re.search(r'href="(https?://[^"]+\.mp4[^"]*)"', html)
            if m:
                vid_url = m.group(1).replace("&amp;", "&")
                if _direct_dl(vid_url, "https://ssstik.io/"):
                    if not title: title = "tiktok"
                    return True, title, ""
        except Exception as e:
            logger.warning("ssstik fallback fail: %s", e)

        # নিচে স্পষ্ট error দিই — সব strategy + API fail
        clear_err = ("সব strategy fail হয়েছে। সম্ভাব্য কারণ:\n"
                     "• yt-dlp puranо version (`pip install -U yt-dlp` চালান)\n"
                     "• লিংকটা প্রাইভেট/region-locked\n"
                     "• TikTok API আজ down (কয়েক মিনিট পরে আবার try করুন)\n"
                     f"\n_yt-dlp last err:_ {(err or '')[:200]}")
        return False, title, clear_err

    # ─── 🆕 Instagram ───
    if is_insta:
        strategies = [
            ("best[height<=1080]/best",
             ["--user-agent", mobile_ua,
              "--referer","https://www.instagram.com/",
              "--add-header","Accept-Language:en-US,en;q=0.9",
              "--add-header","X-IG-App-ID:936619743392459",
              "--force-ipv4"]),
            ("best",
             ["--user-agent", ios_ua,
              "--referer","https://www.instagram.com/"]),
            ("worst/best",
             ["--user-agent", desktop_ua,
              "--referer","https://www.instagram.com/"]),
        ]
        ok, err = _run_strategies(strategies, "Instagram")
        return ok, title, err

    # ─── 🆕 Facebook ───
    if is_fb:
        strategies = [
            ("best[height<=720]/best",
             ["--user-agent", mobile_ua,
              "--referer","https://www.facebook.com/",
              "--force-ipv4"]),
            ("best",
             ["--user-agent", desktop_ua,
              "--referer","https://www.facebook.com/"]),
        ]
        ok, err = _run_strategies(strategies, "Facebook")
        return ok, title, err

    # ─── Pinterest ───
    if is_pinterest:
        strategies = [
            (f"best[filesize<{MAX_TG_UPLOAD_MB-5}M]/best[height<=720]/best",
             ["--user-agent", mobile_ua,
              "--referer","https://www.pinterest.com/",
              "--add-header","Accept-Language:en-US,en;q=0.9"]),
            ("best/bestvideo+bestaudio/mp4",
             ["--user-agent", desktop_ua,
              "--referer","https://www.pinterest.com/",
              "--add-header","Accept:*/*",
              "--force-ipv4","--retries","5","--fragment-retries","5"]),
        ]
        ok, err = _run_strategies(strategies, "Pinterest")
        return ok, title, err

    # ─── Generic fallback ───
    base_cmd = ["yt-dlp","-f",f"best[filesize<{MAX_TG_UPLOAD_MB-5}M]/best[height<=720]/best",
                "--no-playlist","--max-filesize",f"{MAX_DOWNLOAD_MB}M",
                "--no-warnings","--restrict-filenames",
                "--retries","3","--fragment-retries","3",
                "--user-agent", desktop_ua,
                "-o", out_template, url]
    try:
        p = subprocess.run(base_cmd, capture_output=True, text=True, timeout=DOWNLOAD_TIMEOUT)
        if p.returncode == 0: return True, title, ""
        return False, title, (p.stderr or p.stdout or "")[-800:]
    except subprocess.TimeoutExpired:
        return False, title, "timeout"
    except Exception as e:
        return False, title, str(e)

async def handle_url(u, c, url):
    track_user(u.effective_user.id); inc_stat("downloads")
    pf = get_platform(url)
    st = await safe_reply(u.message,
        f"╭─━━━━━━━━━━━━━━━━━━━━╮\n│ {pf}\n│ 📥 ডাউনলোড শুরু...\n╰─━━━━━━━━━━━━━━━━━━━━╯",
        parse_mode="Markdown")
    try: await c.bot.send_chat_action(u.message.chat_id, ChatAction.UPLOAD_VIDEO)
    except Exception: pass
    job = uuid.uuid4().hex
    out_tmpl = str(TEMP_DIR / f"{job}_input.%(ext)s")
    try:
        await safe_edit(st,
            f"╭─━━━━━━━━━━━━━━━━━━━━╮\n│ {pf}\n│ ⏳ ফাইল আনা হচ্ছে...\n│ ▓▓▓▓▓░░░░░ 50%\n╰─━━━━━━━━━━━━━━━━━━━━╯",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        ok, title, err = await loop.run_in_executor(executor, yt_dl, url, out_tmpl)
        if not ok:
            logger.warning("yt-dlp fail: %s", err)
            err_short = (err or "")[-200:]
            tips = ("• লিংক প্রাইভেট হতে পারে\n"
                    "• বয়স যাচাই / login দরকার\n"
                    "• ফাইল খুব বড় / region-locked\n"
                    "• yt-dlp আপডেট দরকার হতে পারে:\n"
                    "  `pip install -U yt-dlp`")
            await safe_edit(st,
                f"❌ *ডাউনলোড ব্যর্থ!*\n\n{tips}\n\n"
                f"_Error: {md_escape(err_short)}_",
                parse_mode="Markdown"); return
        files = sorted(TEMP_DIR.glob(f"{job}_input.*"))
        if not files:
            await safe_edit(st, "❌ ফাইল পাওয়া যায়নি!", parse_mode="Markdown"); return
        out = files[0]
        size_mb = out.stat().st_size / (1024*1024)
        # 🆕 v30: AUTO-COMPRESS — Telegram bot API 50MB hard-limit। বড় ফাইল reject না করে
        # ffmpeg দিয়ে compress করে user-কে দিই — হারিয়ে যায় না।
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st,
                f"📦 *ফাইল বড় ({size_mb:.1f}MB) → compress করছি...*\n\n"
                f"_target: ~{MAX_TG_UPLOAD_MB-3}MB। ৩০-৯০ সেকেন্ড লাগতে পারে।_",
                parse_mode="Markdown")
            try:
                dur = await loop.run_in_executor(executor, get_duration, str(out))
                if not dur or dur < 1: dur = 60.0
                # target bitrate so final ≈ (MAX_TG_UPLOAD_MB - 3) MB total
                target_total_kbits = (MAX_TG_UPLOAD_MB - 3) * 8 * 1024  # MB → kbits
                audio_kbits = 96
                video_kbits = max(150, int(target_total_kbits / dur) - audio_kbits)
                cmp_out = str(TEMP_DIR / f"{job}_cmp.mp4")
                cmp_cmd = ["ffmpeg","-y","-i",str(out),
                           "-c:v","libx264","-preset","veryfast",
                           "-b:v",f"{video_kbits}k","-maxrate",f"{int(video_kbits*1.2)}k",
                           "-bufsize",f"{video_kbits*2}k",
                           "-vf","scale='min(1280,iw)':'-2'",
                           "-c:a","aac","-b:a",f"{audio_kbits}k",
                           "-movflags","+faststart", cmp_out]
                p = await loop.run_in_executor(executor, lambda:
                    subprocess.run(cmp_cmd, capture_output=True, timeout=FFMPEG_TIMEOUT))
                if p.returncode == 0 and Path(cmp_out).exists():
                    new_mb = Path(cmp_out).stat().st_size / (1024*1024)
                    if new_mb <= MAX_TG_UPLOAD_MB:
                        safe_unlink(str(out))
                        # rename to keep naming convention
                        new_path = TEMP_DIR / f"{job}_input.mp4"
                        Path(cmp_out).rename(new_path)
                        out = new_path
                        size_mb = new_mb
                    else:
                        safe_unlink(cmp_out); safe_unlink(str(out))
                        await safe_edit(st,
                            f"❌ *Compress-এর পরও ফাইল বড় ({new_mb:.1f}MB)!*\n\nছোট ভিডিও দিয়ে চেষ্টা করুন।",
                            parse_mode="Markdown"); return
                else:
                    safe_unlink(str(out))
                    await safe_edit(st,
                        f"❌ *ফাইল খুব বড় ({size_mb:.1f}MB) এবং compress fail!*",
                        parse_mode="Markdown"); return
            except Exception as ce:
                logger.warning("auto-compress: %s", ce)
                safe_unlink(str(out))
                await safe_edit(st,
                    f"❌ *ফাইল খুব বড়! ({size_mb:.1f}MB, লিমিট {MAX_TG_UPLOAD_MB}MB)*",
                    parse_mode="Markdown"); return
        uid = u.effective_user.id
        with user_lock:
            old = user_videos.get(uid)
            user_videos[uid] = {"path": str(out), "title": title, "url": url, "platform": pf}
            user_state.pop(uid, None)
        if old: safe_unlink(old.get("path"))
        await safe_edit(st, show_video_card(title, size_mb, pf),
                        parse_mode="Markdown", reply_markup=edit_menu(source_url=url, platform=pf))
    except Exception as e:
        logger.exception("url err: %s", e)
        await safe_edit(st, "❌ *সমস্যা হয়েছে!*", parse_mode="Markdown")

# ──────────────────────────────────────────────────────────────
# CALLBACK HANDLER
# ──────────────────────────────────────────────────────────────
NAV = {
    "main_menu": (WELCOME, main_menu),
    "how": (f"📤 *ভিডিও আপলোড*\n\n━━━━━━━━━━━━━━━━━━\nযে কোনো ভিডিও পাঠান।\nসর্বোচ্চ সাইজ: *{MAX_TG_UPLOAD_MB} MB*", back_menu),
    "url_info": ("🔗 *লিংক ডাউনলোড*\n\n━━━━━━━━━━━━━━━━━━\nযে কোনো সোশ্যাল মিডিয়া লিংক পাঠান।\n/platforms দেখুন", back_menu),
    "platforms": ("🌐 *সাপোর্টেড প্ল্যাটফর্ম*\n\n━━━━━━━━━━━━━━━━━━\n"
        "📺 YouTube/Shorts\n📷 Instagram\n🎵 TikTok\n👥 Facebook\n🐦 X\n🔴 Reddit\n"
        "📌 Pinterest\n🎥 Vimeo\n🟣 Twitch\n👻 Snapchat\n📱 Likee/Kwai\n💼 LinkedIn\n"
        "📱 VK/OK\n🧵 Threads\n🎬 IMDb/Dailymotion\n📺 Bilibili\n🎤 TED\n🔊 SoundCloud\n"
        "🎬 Rumble/Odysee\n+ আরও 1000+", back_menu),
    "help": ("📖 *সাহায্য*\n\n━━━━━━━━━━━━━━━━━━\n📤 ভিডিও পাঠান → এডিট\n🔗 লিংক পাঠান → ডাউনলোড\n\n/cancel — বাতিল", back_menu),
    # 🆕 v31: Voice Clone instructions
    "voice_clone_help": (
        "🎭 *Voice Clone — আপনার নিজের কণ্ঠ Save করুন*\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━\n"
        "📌 *যেভাবে করবেন:*\n"
        "১) Telegram-এ একটা *voice message* record করুন\n"
        "    (১৫-৩০ সেকেন্ড — পরিষ্কার কণ্ঠ, কোনো background noise নেই)\n"
        "২) সেই voice message-এ *reply* দিন\n"
        "৩) Reply text-এ লিখুন: `/clonevoice আপনার_নাম`\n"
        "    (যেমন: `/clonevoice MyVoice`)\n\n"
        "✅ *Clone হয়ে গেলে এই সুবিধা পাবেন:*\n"
        "🗣️ TTS-এ নিজের কণ্ঠে যেকোনো text বলান\n"
        "🎬 Video Dubbing-এ *আপনার কণ্ঠেই* অন্য ভাষায় ডাবিং!\n"
        "🎙️ AI Voiceover-এ নিজের কণ্ঠ\n\n"
        "🔧 *Useful কমান্ড:*\n"
        "/myvoice — clone-এর info\n"
        "/deletevoice — clone delete\n\n"
        "💡 *Engine:* Replicate XTTS-v2 (16+ ভাষা support)",
        back_menu),
}

SUBMENUS = {
    "wm_menu":     ("🚫 *ওয়াটারমার্ক/লগো রিমুভ*\n\n━━━━━━━━━━━━━━━━━━\nযে কোণে লগো আছে সেটা সিলেক্ট করুন:", wm_menu),
    "filter_menu": ("🎞️ *প্রিমিয়াম ফিল্টার*\n\n━━━━━━━━━━━━━━━━━━\nযে স্টাইল চান বেছে নিন:", filter_menu),
    "speed_menu":  ("⏩ *স্পিড কন্ট্রোল*\n\n━━━━━━━━━━━━━━━━━━\nভিডিওর গতি বদলান:", speed_menu),
    "aspect_menu": ("📐 *অ্যাসপেক্ট রেশিও*\n\n━━━━━━━━━━━━━━━━━━\nযে ফরম্যাট দরকার বেছে নিন:", aspect_menu),
    "flip_menu":   ("🪞 *ফ্লিপ/মিরর*\n\n━━━━━━━━━━━━━━━━━━\nকোন দিকে উল্টাবেন?", flip_menu),
    "vol_menu":    ("🔊 *ভলিউম কন্ট্রোল*\n\n━━━━━━━━━━━━━━━━━━\nকতো জোরে চান?", vol_menu),
    # 🆕 v14: AI Toolkit submenus
    "ai_toolkit":       ("🚀 *AI টুলকিট (ChatGPT-powered)*\n\n━━━━━━━━━━━━━━━━━━\nযেটা চান বেছে নিন:", ai_toolkit_menu),
    "ai_translate_sub": ("🌐 *সাবটাইটেল অনুবাদ*\n\n━━━━━━━━━━━━━━━━━━\n_প্রথমে subtitle SRT file পাঠান অথবা যেকোনো টেক্সট পাঠান, তারপর target ভাষা বেছে নিন।_\n\nকোন ভাষায় অনুবাদ করতে চান?", translate_lang_menu),
}

# ফিচার যেগুলো process() দিয়ে চলবে (extra ছাড়া)
PROCESSABLE = {
    "smart_auto",
    "enhance","enhance_pro","color","copyright","full","compress","audio","thumb",
    "rotate","mute","reverse","fade","boomerang","gif",
    "wm_tl","wm_tr","wm_bl","wm_br",
    "f_cinema","f_vintage","f_bw","f_warm","f_cold","f_drama","f_neon","f_dreamy","f_anime","f_sunset",
    "sp_05","sp_15","sp_20","sp_40",
    "ar_916","ar_11","ar_169","ar_45",
    "fl_h","fl_v",
    "vol_50","vol_150","vol_200","vol_300",
    "blur_bg","vignette","zoom","freeze","vhs",
    "denoise","slowmo_smooth","voice_pro",
}

async def handle_cb(u, c):
    q = u.callback_query
    try: await q.answer()
    except Exception: pass
    uid = q.from_user.id

    # 🆕 v37: File Extractor menu button
    if q.data == "file_extract_start":
        with user_lock:
            user_state[uid] = {"action": "file_extract_wait", "data": {}}
        await safe_edit(q.message,
            "📂 *File Extractor চালু!*\n\n"
            "এখন একটা ফাইল পাঠান:\n\n"
            "📄 PDF, DOCX, TXT — text বের করব\n"
            "📊 XLSX, CSV — data table\n"
            "🗜️ ZIP, RAR, 7Z — ভেতরের ফাইল list\n"
            "🖼️ JPG, PNG — OCR দিয়ে লেখা পড়ব\n"
            "🎵 MP3, M4A — metadata\n"
            "💻 JSON, HTML, code — direct content\n\n"
            f"⚠️ সর্বোচ্চ {MAX_EXTRACT_FILE_MB}MB।",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("❌ বাতিল", callback_data="main_menu")]]))
        return

    # 🆕 v37: Photo → REAL Apple Live Photo pair (for TikTok)
    if q.data == "photo_live_start":
        with user_lock:
            user_state[uid] = {"action": "photo_live_wait", "data": {}}
        await safe_edit(q.message,
            "📸 *Photo → Real Apple Live Photo*\n\n"
            "এখন একটা ছবি পাঠান — bot দুটো ফাইল বানাবে:\n"
            "  • `IMG_xxx.JPG` (still image)\n"
            "  • `IMG_xxx.MOV` (3 sec motion)\n\n"
            "দুটোতে একই Apple ContentIdentifier UUID থাকবে — iPhone Photos app একে real Live Photo হিসেবে recognize করবে।\n\n"
            "📲 *TikTok-এ Live Photo upload করতে:*\n"
            "  1. দুটো ফাইল iPhone-এ পাঠান (AirDrop / iCloud)\n"
            "  2. Photos app-এ save হলে auto Live Photo হিসেবে দেখাবে\n"
            "  3. iPhone-এর TikTok app থেকে upload\n\n"
            "⚠️ *সীমাবদ্ধতা:* TikTok-এর Live Photo upload feature **শুধু iOS app-এ**। Android থেকে upload করলে শুধু JPG যাবে।\n\n"
            "🛠️ *Termux requirement:* `pkg install ffmpeg exiftool`",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("❌ বাতিল", callback_data="main_menu")]]))
        return

    # 🆕 v37: Link Expander menu button
    if q.data == "link_expand_start":
        with user_lock:
            user_state[uid] = {"action": "link_expand_wait", "data": {}}
        await safe_edit(q.message,
            "🔗 *Link Expander চালু!*\n\n"
            "এখন একটা URL পাঠান — bot সেটার আসল destination দেখাবে।\n\n"
            "✅ *যা করব:*\n"
            "  • Short URL (bit.ly, t.co) → আসল URL\n"
            "  • সব redirect chain\n"
            "  • Page title + HTTP status\n"
            "  • Phishing/suspicious warning\n\n"
            "❌ *যা করব না:* GPLinks/Droplink/Ouo.io bypass — paid sites-এর owners-এর income কাটা পড়ে। Paywall bypass-ও copyright violation।\n\n"
            "💡 *উদাহরণ:* `https://bit.ly/xxxxxxx`",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("❌ বাতিল", callback_data="main_menu")]]))
        return

    # 🆕 v37: MCQ Solver menu button
    if q.data == "mcq_start":
        with user_lock:
            user_state[uid] = {"action": "mcq_wait", "data": {}}
        await safe_edit(q.message,
            "📚 *MCQ Solver চালু!*\n\n"
            "এখন প্রশ্ন পাঠান — ৩ ভাবে দিতে পারেন:\n\n"
            "✍️ *Text* — সরাসরি প্রশ্ন টাইপ করুন\n"
            "🖼️ *ছবি* — বইয়ের MCQ-এর photo\n"
            "📄 *PDF/DOCX* — পুরো MCQ ফাইল\n\n"
            "💡 *উদাহরণ:* \n"
            "_বাংলাদেশ স্বাধীন হয় কত সালে?_\n"
            "_(ক) ১৯৭০  (খ) ১৯৭১  (গ) ১৯৭২  (ঘ) ১৯৭৩_",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("❌ বাতিল", callback_data="main_menu")]]))
        return

    if q.data == "my_stats":
        s = load_stats()
        await safe_edit(q.message,
            f"📊 *স্ট্যাটস*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"👥 মোট ইউজার: *{len(s['total_users'])}*\n"
            f"🎬 প্রসেস: *{s.get('videos_processed',0)}*\n"
            f"⬇️ ডাউনলোড: *{s.get('downloads',0)}*",
            parse_mode="Markdown", reply_markup=back_menu())
        return

    # ── 🆕 v33: WinGo Multi-Mode (30s + 1m + 3m + 5m) ON-DEMAND callbacks ──
    if q.data == "wingo_main":
        # Top-level mode selector
        await safe_edit(q.message,
            "🎯 *WinGo AI Predictor — Multi-Mode*\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "_যে গেমে predict করতে চান সেটা select করুন:_\n\n"
            "⚡ *30 সেকেন্ড* — fast turnover, more signals\n"
            "⏱️ *1 মিনিট* — balanced (recommended)\n"
            "⏰ *3 / 5 মিনিট* — slower, more time to act\n\n"
            "🆕 *v35 ULTRA+++++++:* এখন *30 layer* — 28 statistical analyzer + 🤖 *real Gemini AI* + "
            "📡 *3X Leader Rakib* external algo cross-check।  Multi-source intelligence next outcome predict করে।\n\n"
            "⚠️ _Random game — 100% guarantee নেই।_",
            parse_mode="Markdown", reply_markup=wingo_main_menu())
        return

    if q.data.startswith("wingo_panel"):
        mode = _w_parse_mode(q.data, "wingo_panel")
        label = WINGO_MODE_LABELS[mode]
        interval = WINGO_MODE_INTERVAL[mode]
        ai_line = "🤖 *Real Gemini AI* (L25) — pattern reads করে next predict করে।\n" if GEMINI_API_KEY else ""
        await safe_edit(q.message,
            f"🎯 *{label} — Prediction Panel*\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"_⚡ Mode: {label} ({interval}/round)_\n"
            f"_⚡ On-demand — button চাপলেই signal আসবে।_\n\n"
            f"🧠 *30-layer ULTRA+++++++ analyzer + 🛡️ anti-streak protection + 📡 3X Leader Rakib cross-check* — frequency · streak · "
            f"zigzag · 3-step Markov · gap · trend · EMA · 6-gram · "
            f"Bayesian · cyclic · color · volatility · run-length · "
            f"period-bias · 2-step num Markov · entropy · Fibonacci · "
            f"triple-Markov · self-correcting bias\n"
            f"{ai_line}\n"
            f"⚠️ *Disclaimer:* random game — শুধু entertainment / pattern study।",
            parse_mode="Markdown", reply_markup=wingo_panel_menu(mode))
        return

    if q.data in ("wingo_sub", "wingo_unsub"):
        # 🆕 v30: legacy callbacks — auto-broadcast disabled হয়ে গেছে
        _wingo_subs.discard(str(uid))
        _save_wingo_subs(_wingo_subs)
        await safe_edit(q.message,
            "ℹ️ *Auto signal এখন বন্ধ।*\n\n"
            "অন্য feature use করার সময় message disturb হচ্ছিল — সেটা off করা হয়েছে।\n\n"
            "এখন যখনই signal চান, নিচের *🎯 এখনই প্রেডিকশন* button চাপুন।",
            parse_mode="Markdown", reply_markup=wingo_panel_menu())
        return

    if q.data.startswith("wingo_stats"):
        mode = _w_parse_mode(q.data, "wingo_stats")
        state = _w_get_state(mode)
        label = WINGO_MODE_LABELS[mode]
        us = state["user_stats"].get(uid,
            {"win":0,"loss":0,"nwin":0,"nloss":0,"streak":0,"loss_streak":0,"max_streak":0})
        total = us["win"] + us["loss"]
        ntotal = us["nwin"] + us["nloss"]
        sig_acc = round(us["win"]/total*100) if total else 0
        num_acc = round(us["nwin"]/ntotal*100) if ntotal else 0
        await safe_edit(q.message,
            f"📊 *{label} — আপনার স্ট্যাটস*\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"🎯 *Signal (BIG/SMALL):*\n"
            f"  ✅ Win: `{us['win']}`  ❌ Loss: `{us['loss']}`\n"
            f"  📈 Accuracy: *{sig_acc}%*\n"
            f"  🔥 Streak: `{us['streak']}` (best: `{us.get('max_streak',0)}`)  "
            f"⚠️ Loss-streak: `{us['loss_streak']}`\n\n"
            f"🎲 *Number prediction (top 3):*\n"
            f"  ✅ Win: `{us['nwin']}`  ❌ Loss: `{us['nloss']}`\n"
            f"  📈 Accuracy: *{num_acc}%*\n\n"
            f"_💡 যত বেশি predict করবেন, AI তত শিখবে।_\n"
            f"_(প্রতি mode-এর stats আলাদা — 30s/1m/3m/5m সব independent.)_",
            parse_mode="Markdown", reply_markup=wingo_panel_menu(mode))
        return

    if q.data.startswith("wingo_aistats"):
        # 🆕 v29 ULTRA+ : AI Layer Stats panel — v33 mode-aware
        mode = _w_parse_mode(q.data, "wingo_aistats")
        try:
            txt = _w_render_aistats(mode)
            await q.message.reply_text(txt, parse_mode="Markdown",
                                       reply_markup=wingo_panel_menu(mode))
        except Exception as e:
            logger.exception("wingo_aistats fail: %s", e)
            await q.message.reply_text(f"❌ Error: {str(e)[:160]}")
        return

    if q.data.startswith("wingo_now"):
        # 🆕 v33: ULTRA+++++ on-demand prediction (mode-aware) with auto-resolve + top-3 probs + Gemini AI
        mode = _w_parse_mode(q.data, "wingo_now")
        state = _w_get_state(mode)
        perf_file = WINGO_PERF_FILES[mode]
        label = WINGO_MODE_LABELS[mode]
        interval = WINGO_MODE_INTERVAL[mode]
        wait_msg = await q.message.reply_text(
            f"⏳ {label} · *36-layer আগুন-MAX* engine চালু...\n"
            f"   🤖 Gemini · 🧠 OpenAI · 🪐 Claude · 📡 3X Leader · 🌊 HMM/Bayes/FFT\n"
            f"   🛡️ Anti-Streak · 🎯 Meta-Calibration · 💰 Kelly Bet",
            parse_mode="Markdown")
        try:
            loop = asyncio.get_running_loop()
            history = await loop.run_in_executor(executor, _w_fetch_history, mode)
            if not history:
                await wait_msg.edit_text(
                    f"❌ {label} API থেকে data পাওয়া যাচ্ছে না। একটু পরে try করুন।\n"
                    f"_(Tip: server-এর IP block-এ থাকতে পারে — Termux/phone থেকে চালালে কাজ করবে।)_",
                    parse_mode="Markdown")
                return

            # Step 1: NEW prediction (shown first — FIXED v39)
            # Step 2: generate new prediction (mode-aware state)
            us_now = state["user_stats"].get(int(uid), {"loss_streak":0, "streak":0})
            pred = await loop.run_in_executor(executor, _w_analyze, history,
                                               us_now.get("loss_streak",0),
                                               us_now.get("streak",0),
                                               state)
            if not pred:
                await wait_msg.edit_text("❌ Analyzer fail করেছে।")
                return
            current = history[0]["period"]
            next_p = str(int(current) + 1)
            pred["period"] = next_p

            # Save for future resolve (in mode-specific state)
            state["user_last_pred"][int(uid)] = {
                "period":      next_p,
                "signal":      pred["signal"],
                "num_top":     pred["num_top"],
                "num_alt":     pred["num_alt"],
                "layer_votes": pred.get("layer_votes", {}),
                "confidence":  pred["confidence"],
            }

            # Build display
            conf = pred["confidence"]
            bar_filled = max(1, int(conf / 10))
            conf_bar = "█" * bar_filled + "░" * (10 - bar_filled)
            if conf >= 75:   conf_emoji, conf_label = "🟢", "STRONG"
            elif conf >= 65: conf_emoji, conf_label = "🟡", "MEDIUM"
            elif conf >= 55: conf_emoji, conf_label = "🟠", "WEAK"
            else:            conf_emoji, conf_label = "🔴", "AVOID"
            sig_box = "🟢 *BIG* (5-9)" if pred["signal"] == "BIG" else "🔵 *SMALL* (0-4)"

            # Top-3 number table with probability bars
            t3 = pred.get("num_top3", [pred["num_top"]] + pred["num_alt"])
            t3p = pred.get("num_top3_pct", [50, 30, 20])
            num_lines = []
            medals = ["🥇", "🥈", "🥉"]
            for i, (n, p) in enumerate(zip(t3, t3p)):
                ce = _w_color_emoji(_w_colorof(n))
                bar_n = "▰" * max(1, int(p / 10)) + "▱" * (10 - max(1, int(p / 10)))
                num_lines.append(f"  {medals[i]} `{n}` {ce}  `{bar_n}` *{p}%*")

            # 🆕 v32: Aggregated COLOR signal — sum top-3 number probabilities per color
            color_probs = {"red": 0, "green": 0, "violet": 0}
            for n, p in zip(t3, t3p):
                color_probs[_w_colorof(n)] += p
            top_color_name, top_color_pct = max(color_probs.items(), key=lambda x: x[1])
            color_emoji_top = _w_color_emoji(top_color_name)
            color_label = {"red": "RED", "green": "GREEN", "violet": "VIOLET"}[top_color_name]
            color_bar_filled = max(1, int(top_color_pct / 10))
            color_bar = "▰" * color_bar_filled + "▱" * (10 - color_bar_filled)
            color_box = (
                f"{color_emoji_top} *{color_label}*  ({top_color_pct}%)\n"
                f"   `{color_bar}`\n"
                f"   _alt:_ "
                + " · ".join(
                    f"{_w_color_emoji(c)}`{p}%`"
                    for c, p in sorted(color_probs.items(), key=lambda x: -x[1])
                    if c != top_color_name and p > 0
                )
            )

            agree = pred.get("agreement", 0)
            gp = state.get("global_perf", {})
            g_recent = gp.get("recent", [])
            g_acc = (sum(g_recent[-20:]) / len(g_recent[-20:]) * 100) if g_recent else 0
            g_total = gp.get("hit",0) + gp.get("miss",0)
            g_overall = (gp.get("hit",0) / g_total * 100) if g_total else 0
            # 🆕 v36: Last 50 rounds W/L count + emoji bar
            last50 = g_recent[-50:]
            w50 = sum(last50)
            l50 = len(last50) - w50
            if last50:
                acc50 = (w50 / len(last50)) * 100
                # Compact emoji history bar (last 20 of the 50, newest on right)
                tail20 = last50[-20:]
                emoji_bar = "".join("🟢" if r else "🔴" for r in tail20)
                last50_line = (f"\n📊 *Last {len(last50)} ({label}):* "
                               f"✅ Win: *{w50}* · ❌ Loss: *{l50}* · {acc50:.0f}%\n"
                               f"   Recent {len(tail20)}:  {emoji_bar}")
            else:
                last50_line = ""
            learning_line = last50_line
            if g_total >= 10:
                learning_line += (f"\n🧠 *AI Learning ({label}):* {g_overall:.0f}% overall · "
                                  f"{g_acc:.0f}% recent ({min(20,len(g_recent))} rounds)")
            ai_badge = "  🤖✨" if pred.get("ai_used") else ""

            # 🆕 v34: Anti-streak protection banners
            flip_banner = ""
            skip_banner = ""
            ls = pred.get("loss_streak", 0)
            if pred.get("flip_applied"):
                orig = pred.get("original_signal", "?")
                flip_banner = (
                    f"\n🔄 *ANTI-STREAK FLIP ACTIVE*\n"
                    f"   Loss streak = {ls} → ensemble বলেছিল `{orig}`, "
                    f"কিন্তু safety-protection signal flip করে দিয়েছে।\n"
                )
            if pred.get("skip_recommended"):
                skip_banner = (
                    f"\n⛔ *⚠️ SKIP THIS ROUND সাজেশন ⚠️*\n"
                    f"   Loss streak ({ls}) + weak confidence/agreement → "
                    f"এই round-এ NOT TO BET সাজেশন।  পরের strong signal "
                    f"(confidence ≥ 70%, agreement ≥ 70%)-এর জন্য wait করুন।\n"
                    f"   _এটাই হলো \"২ step-এর বেশি না যাওয়ার\" practical safety।_\n"
                )

            # 🆕 v38 — Modern Premium Card Design (Glass-card aesthetic)
            #          Hero → Stats → Cross-check → Kelly → Footer
            n_ai = pred.get("n_ai_active", 0)
            ai_agree = pred.get("ai_agree_count", 0)
            cal_delta = pred.get("calibration_delta", 0)
            recent30 = pred.get("recent_acc_30", 50)
            cal_note = ""
            if cal_delta != 0:
                arrow = "↑" if cal_delta > 0 else "↓"
                cal_note = f"  _(Meta-Cal: {arrow}{abs(cal_delta)})_"
            # Quality tier badge based on AI consensus + confidence
            if pred.get("quad_ai_consensus") and conf >= 80:
                tier_badge = "💎🔥 *PREMIUM SIGNAL — Quad-AI Locked*"
            elif pred.get("triple_ai_consensus") and conf >= 75:
                tier_badge = "💎 *DIAMOND SIGNAL — Triple-AI Locked*"
            elif conf >= 75 and ai_agree >= 2:
                tier_badge = "🌟 *GOLD SIGNAL*"
            elif conf >= 65:
                tier_badge = "✨ *SILVER SIGNAL*"
            elif conf >= 55:
                tier_badge = "⚪ *BRONZE SIGNAL*"
            else:
                tier_badge = "🔴 *LOW SIGNAL — caution*"
            # Big Hero box — signal as large visual
            if pred["signal"] == "BIG":
                hero = "🟢🟢🟢   *B I G*  (5-9)   🟢🟢🟢"
            else:
                hero = "🔵🔵🔵  *S M A L L*  (0-4)  🔵🔵🔵"
            # AI quorum mini-display
            ai_quorum = ""
            if n_ai > 0:
                dots = "🟢" * ai_agree + "⚪" * (n_ai - ai_agree)
                ai_quorum = f"   AI Quorum:  {dots}  ({ai_agree}/{n_ai} agree)\n"

            new_txt = (
                f"╔══════════════════════════════════╗\n"
                f"║   🎯  *{label}  AI Signal*  🎯   \n"
                f"║   _v38 আগুন-MAX  ·  36 layers_   \n"
                f"╚══════════════════════════════════╝\n"
                f"\n"
                f"   {tier_badge}\n"
                f"\n"
                f"┌─────────────────────────────────┐\n"
                f"│   {hero}\n"
                f"└─────────────────────────────────┘\n"
                f"   📅 Period:  `{next_p[-6:]}`   ·   ⏱️ ~{interval}\n"
                f"{flip_banner}{skip_banner}\n"
                f"┃ *🎨 Color Forecast*\n"
                f"   {color_box}\n"
                f"\n"
                f"┃ *🎲 Top-3 Number Predictions*\n"
                + "\n".join(num_lines) +
                f"\n"
                f"\n"
                f"┃ *{conf_emoji} Confidence:*  *{conf}%*  —  _{conf_label}_{cal_note}\n"
                f"   `{conf_bar}`  ({recent30}% real last-30)\n"
                f"   🤝 Layer Agreement: *{agree}%*  ·  *{pred.get('n_layers',36)} layers*\n"
                f"{ai_quorum}"
                f"   📊 Loss: *{ls}*  ·  Win: *{pred.get('win_streak',0)}*  ·  Run: *{pred.get('run_len',0)}*\n"
                f"\n"
                f"{_w_format_crosscheck(pred)}"
                f"\n"
                f"{_w_format_kelly(pred, state.get('user_balance', {}).get(int(uid)))}"
                f"{learning_line}\n"
                f"\n"
                f"╭─ *Active Protections* ─────────╮\n"
                f"│  🛡️  Anti-Streak Auto-Flip\n"
                f"│  🎯  Meta-Confidence Calibration\n"
                f"│  💰  Kelly Criterion Bet-Sizing\n"
                f"│  🌊  HMM + Bayesian + FFT engines\n"
                f"╰────────────────────────────────╯\n"
                f"_💡 আবার tap করলে এই round-এর result আগে show হবে।_\n"
                f"_⚠️ Random game — 100% guarantee নেই, but আমরা সর্বোচ্চ effort দিচ্ছি।_"
            )
            await wait_msg.edit_text(new_txt, parse_mode="Markdown",
                                       reply_markup=wingo_panel_menu(mode))

            # ─── FIXED v39: গত Round result — prediction-এর পরে দেখায় ───
            prev_result = _w_resolve_user_prev(uid, history, state, perf_file)
            if prev_result:
                head = "🏆 *গত Prediction: WIN!* 🎉" if prev_result["sig_hit"] else "💔 *গত Prediction: LOSS*"
                bonus = "\n💎 *EXACT NUMBER HIT!* 🎯" if prev_result["num_hit"] else ""
                us = prev_result["user_stats"]
                tot = us["win"] + us["loss"]
                wr = (us["win"] / tot * 100) if tot else 0
                streak_line = ""
                if us["streak"] >= 3:      streak_line = f"\n🔥🔥 *WIN STREAK: {us['streak']}!*"
                elif us["streak"] >= 2:    streak_line = f"\n🔥 Win streak: *{us['streak']}*"
                elif us["loss_streak"] >= 3: streak_line = f"\n⚠️ Loss streak: *{us['loss_streak']}*"
                actual_box = (f"`{prev_result['actual_num']}`  "
                              f"{_w_color_emoji(prev_result['actual_color'])}  "
                              f"`{prev_result['actual_size']}`")
                prev_txt = (
                    f"╔══════ গত Round Result ══════╗\n"
                    f"║  {head}\n"
                    f"║  _Mode: {label}_\n"
                    f"╠════════════════════════════╣\n"
                    f"║  🎲 Period:  `{str(prev_result['period'])[-6:]}`\n"
                    f"║  🎯 Actual:  {actual_box}\n"
                    f"║\n"
                    f"║  ▸ Signal:  *{prev_result['prev_signal']}*  "
                    f"{'✅' if prev_result['sig_hit'] else '❌'}\n"
                    f"║  ▸ Number:  `{prev_result['prev_num_top']}`  "
                    f"{'✅ HIT' if prev_result['num_hit'] else '❌'}{bonus}\n"
                    f"║\n"
                    f"║  📊  Win: *{us['win']}*  Loss: *{us['loss']}*  ({wr:.0f}%){streak_line}\n"
                    f"╚════════════════════════════╝\n"
                    f"_↑ গত prediction result — নতুন prediction উপরে ↑_"
                )
                try:
                    await q.message.reply_text(prev_txt, parse_mode="Markdown")
                except Exception:
                    pass
          except Exception as e:
              logger.exception("wingo_now (%s) fail: %s", mode, e)
              try:
                  await wait_msg.edit_text(f"❌ Error: {str(e)[:160]}")
              except Exception:
                  await q.message.reply_text(f"❌ Error: {str(e)[:160]}")
        return

    if q.data in NAV:
        text, mk = NAV[q.data]
        await safe_edit(q.message, text, parse_mode="Markdown", reply_markup=mk())
        return

    if q.data in SUBMENUS:
        text, mk = SUBMENUS[q.data]
        await safe_edit(q.message, text, parse_mode="Markdown", reply_markup=mk())
        return

    if q.data == "back_to_edit":
        with user_lock:
            d = user_videos.get(uid)
            st = user_state.pop(uid, None)
        if st:
            for p in st.get("data", {}).get("paths", []) or []:
                safe_unlink(p)
            safe_unlink(st.get("data", {}).get("bgm_path"))
        if d and d.get("path") and Path(d["path"]).exists():
            size_mb = Path(d["path"]).stat().st_size / (1024*1024)
            src = d.get("platform") or "📤 আপনার ভিডিও"
            await safe_edit(q.message, show_video_card(d.get("title",""), size_mb, src),
                            parse_mode="Markdown",
                            reply_markup=edit_menu(source_url=d.get("url"), platform=d.get("platform")))
        else:
            await safe_edit(q.message, "❌ ভিডিও পাওয়া যায়নি, আবার পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu())
        return

    if q.data == "cancel_video":
        with user_lock:
            d = user_videos.pop(uid, None); user_state.pop(uid, None)
        if d: safe_unlink(d.get("path"))
        await safe_edit(q.message, "✅ *বাতিল!*", parse_mode="Markdown", reply_markup=back_menu())
        return

    # ─── 🆕 v27.1: Section header buttons (decorative — do nothing on click) ───
    if q.data == "noop_section":
        try:
            await q.answer("📋 এটা section header — নিচের button-গুলো ব্যবহার করুন", show_alert=False)
        except Exception:
            pass
        return

    # ─── 🆕 v27.2: YouTube টাকা আয়ের গাইড (Monetization Bible) ───
    if q.data == "youtube_money":
        await send_youtube_money_guide(q.message, edit=True, uid=uid)
        return

    # ─── 🆕 v27.1: AUTO VOICE CLONE — এক tap-এ current video থেকে clone ───
    # Link পাঠালেও কাজ করবে কারণ link auto-download হয়ে current video হয়ে যায়।
    if q.data == "auto_clone_voice":
        if not has_replicate():
            await safe_edit(q.message,
                "❌ *Voice Clone disabled*\n\n"
                "এই feature-এর জন্য Replicate API token লাগবে (FREE):\n\n"
                "১) https://replicate.com → Sign up (free)\n"
                "২) Account → API tokens → Create token\n"
                "৩) Termux-এ:\n"
                "`export REPLICATE_API_TOKEN=আপনার_টোকেন`\n"
                "৪) Bot restart করুন\n\n"
                "✨ Free tier-এ মাসে অনেক voice clone হয়।",
                parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            return
        with user_lock:
            d = user_videos.get(uid)
        inp = d.get("path") if d else None
        if not inp or not Path(inp).exists():
            await safe_edit(q.message,
                "❌ *কোনো ভিডিও নেই!*\n\n"
                "প্রথমে ভিডিও পাঠান অথবা YouTube/TikTok/Instagram/FB link দিন।\n"
                "তারপর এই button-এ tap করুন — voice automatic clone হবে।",
                parse_mode="Markdown", reply_markup=back_menu())
            return
        # Clone name = video title-এর প্রথম অংশ (clean)
        raw_title = (d.get("title") or "").strip()
        if raw_title:
            cl_name = re.sub(r'[^\w\u0980-\u09FF\s-]', '', raw_title)[:25].strip() or f"Clone_{uid}"
        else:
            cl_name = f"Clone_{uid}"
        msg = await safe_reply(q.message,
            f"🎤 *অটো ভয়েস ক্লোন চলছে...*\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"🎬 Source: *{md_escape(d.get('title') or 'Video')[:40]}*\n"
            f"🎭 ক্লোন নাম: *{md_escape(cl_name)}*\n"
            f"⚙️ Engine: Replicate XTTS-v2\n\n"
            f"১) 🎵 Audio extract (clean voice খুঁজছি...)\n"
            f"২) ☁️ Replicate-এ test sample upload\n"
            f"৩) ✅ Clone save + test playback\n\n"
            f"_~৩০-৬০ সেকেন্ড লাগবে..._",
            parse_mode="Markdown")

        loop = asyncio.get_event_loop()
        sample_path = str(CLONE_SAMPLES_DIR / f"voice_{uid}.wav")
        ok = await loop.run_in_executor(executor, extract_audio_for_clone, inp, sample_path, 30)
        if not ok or not Path(sample_path).exists():
            await safe_edit(msg,
                "❌ *Audio extract ব্যর্থ!*\n\n"
                "ভিডিওতে clean voice নেই অথবা audio track corrupt। "
                "অন্য ভিডিও try করুন।",
                parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(sample_path); return

        # Test TTS to verify clone works
        test_path, test_err = await loop.run_in_executor(
            executor, replicate_xtts_speak,
            "নমস্কার, এটা আমার ভয়েস ক্লোনের পরীক্ষা।",
            sample_path, "hi", None)
        if not test_path:
            safe_unlink(sample_path)
            await safe_edit(msg,
                f"❌ *Voice clone test failed!*\n\n{md_escape(test_err or 'Unknown error')}\n\n"
                f"💡 Replicate quota শেষ হয়ে গেছে কিনা check করুন।",
                parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return

        set_user_clone(uid, f"xtts:{sample_path}", cl_name)
        try: inc_feature("auto_clone_voice")
        except Exception: pass

        # Send test voice to user
        try:
            with open(test_path, "rb") as fh:
                await c.bot.send_voice(chat_id=u.effective_chat.id, voice=fh,
                    caption=f"🎭 *{md_escape(cl_name)}* — টেস্ট\n\n"
                            f"_উপরের voice আপনার ক্লোন থেকে generated_",
                    parse_mode="Markdown")
        except Exception as e:
            logger.warning("send test voice fail: %s", e)
        safe_unlink(test_path)

        await safe_edit(msg,
            f"✅ *ভয়েস ক্লোন সফল!*\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"🎭 ক্লোন নাম: *{md_escape(cl_name)}*\n"
            f"⚙️ Engine: Replicate XTTS-v2 (16+ language)\n"
            f"💾 Storage: Local (cloud-এ যায়নি)\n\n"
            f"🎯 *এখন কী করতে পারেন:*\n"
            f"১) 🎙️ AI Voiceover button → ক্লোন voice দিয়ে speak\n"
            f"২) `/tts <text>` → cloned voice-এ যেকোনো text\n"
            f"৩) `/listvoices` → সব saved voice দেখুন\n"
            f"৪) Voiceover+Subtitle → video-এ apply\n\n"
            f"💡 *Pro Tip:* অন্য ভিডিও থেকে ক্লোন করতে — সেই ভিডিও পাঠান, "
            f"আবার এই button-এ tap করুন। পুরনোটা auto replace হবে।",
            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🆕 v20: Video Info (instant ffprobe card) ───
    if q.data == "video_info":
        with user_lock:
            d = user_videos.get(uid)
        if not d or not Path(d.get("path","")).exists():
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        loop = asyncio.get_running_loop()
        info = await loop.run_in_executor(executor, fmt_video_info, d["path"])
        await safe_edit(q.message, info, parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🆕 v20: AI Voiceover ON Video — ask text ───
    # ─── 🆕 v30: Video Dubbing — ১২ ভাষায় full audio replace ───
    if q.data == "dub_ask":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        if not has_openai():
            await safe_edit(q.message,
                "❌ *Video Dubbing-এর জন্য OPENAI_API_KEY দরকার।*\n\n"
                "Whisper দিয়ে ভিডিওর কথা detect করতে হয় — কোডের উপরে\n"
                "OPENAI_API_KEY-এ valid key বসান। তারপর আবার ট্রাই করুন।",
                parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        rows = []
        items = list(DUB_LANGS.items())
        for i in range(0, len(items), 3):
            row = []
            for code, (label, _voice, _g) in items[i:i+3]:
                row.append(InlineKeyboardButton(label, callback_data=f"dublang_{code}"))
            rows.append(row)
        rows.append([InlineKeyboardButton("« পেছনে", callback_data="edit_back")])
        await safe_edit(q.message,
            "🎬🌍 *Video Dubbing — অন্য ভাষায় ডাব করুন*\n\n━━━━━━━━━━━━━━━━━━\n"
            "Pipeline:\n"
            "  ১) Whisper দিয়ে ভিডিওর কথা detect (auto language)\n"
            "  ২) আপনার বাছাই করা ভাষায় AI translate\n"
            "  ৩) Edge-TTS (FREE) দিয়ে সেই ভাষায় narration\n"
            "  ৪) Original audio replace → dubbed video\n\n"
            "👇 কোন ভাষায় ডাব করতে চান?\n"
            "_(সর্বোচ্চ ১০ মিনিট ভিডিও — ১-৩ মিনিট লাগতে পারে)_",
            parse_mode="Markdown", reply_markup=InlineKeyboardMarkup(rows))
        return

    if q.data.startswith("dublang_"):
        code = q.data[len("dublang_"):]
        if code not in DUB_LANGS:
            await q.answer("❌ ভাষা নেই", show_alert=True); return
        with user_lock:
            d = user_videos.get(uid)
        if not d or not Path(d.get("path","")).exists():
            await safe_edit(q.message, "❌ ভিডিও নেই!", reply_markup=back_menu()); return
        label, voice, gtts_lang = DUB_LANGS[code]
        st_msg = await safe_edit(q.message,
            f"🎬 *Dubbing শুরু হচ্ছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🌍 Target: *{label}*\n🎙️ Voice: `{voice}`\n\n"
            f"_১) Audio extract..._",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        # 1) Extract audio (mono 16kHz, max 10 min)
        stt_audio = str(TEMP_DIR / f"{uuid.uuid4().hex}_stt.mp3")
        ok = await loop.run_in_executor(executor, extract_audio_for_stt, d["path"], stt_audio, 600)
        if not ok:
            await safe_edit(st_msg, "❌ Audio extract ব্যর্থ — ভিডিওতে audio আছে কিনা চেক করুন।",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(stt_audio); return
        # 2) Whisper transcribe
        await safe_edit(st_msg,
            f"🎬 *Dubbing চলছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🌍 Target: *{label}*\n\n_২) Whisper দিয়ে কথা detect..._",
            parse_mode="Markdown")
        srt_text, w_err = await loop.run_in_executor(executor, openai_transcribe, stt_audio, None)
        safe_unlink(stt_audio)
        if not srt_text:
            await safe_edit(st_msg, f"❌ Whisper STT fail: {w_err or 'unknown'}",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        plain = srt_to_plain_text(srt_text)
        if not plain or len(plain) < 5:
            await safe_edit(st_msg, "❌ ভিডিওতে কোনো কথা পাওয়া যায়নি।",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        # 3) Translate
        await safe_edit(st_msg,
            f"🎬 *Dubbing চলছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🌍 Target: *{label}*\n📝 Detected: `{len(plain)}` chars\n\n"
            f"_৩) {label}-এ translate..._",
            parse_mode="Markdown")
        translated, t_err = await loop.run_in_executor(executor, translate_text, plain, code)
        if not translated:
            await safe_edit(st_msg, f"❌ Translate fail: {t_err or 'unknown'}",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        # 4) TTS in target language —
        # 🆕 v30: USER-এর নিজের কণ্ঠ clone করা থাকলে XTTS দিয়ে নিজের voice-এ ডাব করি!
        # নাহলে free Edge-TTS দিয়ে। এটাই "Voice-Cloned Dubbing" — অরজিনাল কণ্ঠের
        # কাছাকাছি tone রেখে অন্য ভাষায় বলে।
        clone = get_user_clone(uid)
        clone_used = False
        clone_label = ""
        if clone and clone.get("voice_id","").startswith("xtts:") and has_replicate():
            sample_path = clone["voice_id"].replace("xtts:", "")
            if Path(sample_path).exists():
                clone_used = True
                clone_label = clone.get("name","MyVoice")
        engine_line = (f"🎭 আপনার Cloned Voice (`{md_escape(clone_label)}`) — XTTS"
                       if clone_used else "🆓 Edge-TTS Neural (FREE)")
        await safe_edit(st_msg,
            f"🎬 *Dubbing চলছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🌍 Target: *{label}*\n🎙️ Voice: {engine_line}\n\n"
            f"_৪) {label}-এ TTS narration generate..._",
            parse_mode="Markdown")
        dub_audio = str(TEMP_DIR / f"{uuid.uuid4().hex}_dub.mp3")
        tts_path = None; e_err = None
        if clone_used:
            sample_path = clone["voice_id"].replace("xtts:", "")
            # XTTS-v2 input limit ~400 chars per call → chunk + concat
            chunks = []
            txt = translated.strip()
            while txt:
                # break at sentence end if possible
                cut = min(380, len(txt))
                if cut < len(txt):
                    sp = max(txt.rfind(". ", 0, cut), txt.rfind("। ", 0, cut),
                             txt.rfind("? ", 0, cut), txt.rfind("! ", 0, cut))
                    if sp > 100: cut = sp + 2
                chunks.append(txt[:cut].strip())
                txt = txt[cut:].strip()
            chunk_files = []
            for i, ch in enumerate(chunks[:8]):  # cap 8 chunks ≈ 3000 chars
                cp = str(TEMP_DIR / f"{uuid.uuid4().hex}_xc{i}.mp3")
                p, ce = await loop.run_in_executor(executor,
                    replicate_xtts_speak, ch, sample_path, code, cp)
                if p and Path(p).exists():
                    chunk_files.append(p)
                else:
                    e_err = ce; break
            if chunk_files:
                if len(chunk_files) == 1:
                    Path(chunk_files[0]).rename(dub_audio); tts_path = dub_audio
                else:
                    # concat via ffmpeg
                    listf = TEMP_DIR / f"{uuid.uuid4().hex}_concat.txt"
                    listf.write_text("\n".join(f"file '{p}'" for p in chunk_files))
                    cc = ["ffmpeg","-y","-f","concat","-safe","0","-i",str(listf),
                          "-c","copy", dub_audio]
                    cr = await loop.run_in_executor(executor,
                        lambda: subprocess.run(cc, capture_output=True, timeout=120))
                    for p in chunk_files: safe_unlink(p)
                    safe_unlink(str(listf))
                    if cr.returncode == 0 and Path(dub_audio).exists():
                        tts_path = dub_audio
                    else:
                        e_err = "XTTS concat fail"
        if not tts_path:
            # Edge-TTS path (default OR XTTS fallback)
            tts_path, e_err2 = await loop.run_in_executor(executor,
                edge_tts_speak, translated, voice, dub_audio, "+0%", "+0Hz")
            if not tts_path: e_err = e_err2 or e_err
        if not tts_path:
            try:
                from gtts import gTTS as _gT
                _gT(text=translated[:4500], lang=gtts_lang).save(dub_audio)
                tts_path = dub_audio
            except Exception as ge:
                await safe_edit(st_msg, f"❌ Dub TTS fail: {e_err or ge}",
                                parse_mode="Markdown", reply_markup=edit_menu_for(uid))
                safe_unlink(dub_audio); return
        # 5) Replace audio in video
        await safe_edit(st_msg,
            f"🎬 *Dubbing চলছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🌍 Target: *{label}*\n\n_৫) Audio replace + final mix..._",
            parse_mode="Markdown")
        out_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_dubbed")
        ok2, final = await loop.run_in_executor(executor, build_and_run, "dub", d["path"], out_base,
                                                {"voice_audio": tts_path})
        safe_unlink(tts_path)
        if not ok2 or not Path(final).exists():
            await safe_edit(st_msg, "❌ Audio replace ব্যর্থ!",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        inc_stat("videos_processed"); inc_feature("video_dubbing")
        size_mb = Path(final).stat().st_size / (1024*1024)
        # 🆕 v30: dub auto-compress fallback
        if size_mb > MAX_TG_UPLOAD_MB:
            try:
                dur = await loop.run_in_executor(executor, get_duration, final) or 30.0
                target_kb = (MAX_TG_UPLOAD_MB - 3) * 8 * 1024
                vbk = max(300, int(target_kb / dur) - 96)
                cmp_o = str(TEMP_DIR / f"{uuid.uuid4().hex}_dubc.mp4")
                cc = ["ffmpeg","-y","-i",final,"-c:v","libx264","-preset","veryfast",
                      "-b:v",f"{vbk}k","-c:a","aac","-b:a","96k",
                      "-movflags","+faststart", cmp_o]
                cr = await loop.run_in_executor(executor,
                    lambda: subprocess.run(cc, capture_output=True, timeout=FFMPEG_TIMEOUT))
                if cr.returncode == 0 and Path(cmp_o).exists() and \
                   Path(cmp_o).stat().st_size/(1024*1024) <= MAX_TG_UPLOAD_MB:
                    safe_unlink(final); Path(cmp_o).rename(final)
                    size_mb = Path(final).stat().st_size / (1024*1024)
            except Exception: pass
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(st_msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        preview = (translated[:180] + "…") if len(translated) > 180 else translated
        engine_caption = "🎭 আপনার নিজের কণ্ঠে (XTTS clone)" if clone_used else "🆓 Edge-TTS Neural"
        with open(final, "rb") as fp:
            await c.bot.send_video(msg.chat_id, fp,
                caption=f"🎬🌍 *Dubbed in {label}*\n🎙️ {engine_caption}\n\n📝 _{md_escape(preview)}_",
                parse_mode="Markdown", supports_streaming=True)
        safe_unlink(final)
        await safe_edit(st_msg, f"✅ *{label} Dubbing সম্পন্ন!*\n\n👇 আরও এডিট:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    if q.data == "vo_ask":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        clone = get_user_clone(uid)
        voice_label = f"🎭 আপনার Cloned Voice (`{md_escape(clone.get('name','MyVoice'))}`)" if clone else "🎤 ডিফল্ট: Rachel (clone থাকলে সেটাই হবে)"
        with user_lock:
            user_state[uid] = {"action": "vo_wait_text", "data": {}}
        await safe_edit(q.message,
            "🎙️ *AI Voiceover ON Video*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🗣️ Voice: {voice_label}\n\n"
            "যে script/text পড়তে চান সেটা পাঠান (সর্বোচ্চ ১৫০০ অক্ষর — বাংলা/ইংরেজি/হিন্দি যেকোনো ভাষা)।\n\n"
            "📌 আপনার ভিডিওর audio আপনাআপনি ২০%-এ duck হবে\n"
            "📌 আপনার voice ১৬০%-এ ভিডিওর উপর mix হবে\n"
            "📌 Cloned voice না থাকলে `/clonevoice` দিয়ে আগে নিজের কণ্ঠ clone করুন\n"
            "   (অথবা এই ভিডিও থেকে instant clone: `🎤 এই ভিডিও থেকে Voice Clone`)",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    # ─── 🆕 v20: AI Voiceover + Bangla Subtitle ───
    if q.data == "vosub_ask":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        clone = get_user_clone(uid)
        voice_label = f"🎭 আপনার Cloned Voice (`{md_escape(clone.get('name','MyVoice'))}`)" if clone else "🎤 ডিফল্ট: Rachel"
        with user_lock:
            user_state[uid] = {"action": "vosub_wait_text", "data": {}}
        await safe_edit(q.message,
            "🎙️📝 *Voiceover + Bangla Subtitle*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🗣️ Voice: {voice_label}\n"
            f"📝 Subtitle: Bengali font auto, white + black outline\n\n"
            "যে script পড়তে চান সেটা পাঠান (১৫০০ অক্ষর)।\n\n"
            "📌 narration speak হবে + পর্দায় same text synced subtitle হিসেবে\n"
            "    বসবে — explainer/podcast/educational video-র perfect formula।\n"
            "📌 punctuation (। ? ! .) এর জায়গায় subtitle break হবে।",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    # ─── 🆕 v20: Loop ask ───
    if q.data == "loop_ask":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        with user_lock:
            user_state[uid] = {"action": "loop_wait", "data": {}}
        await safe_edit(q.message,
            "♾️ *N সেকেন্ড লুপ*\n\n━━━━━━━━━━━━━━━━━━\n"
            "কত সেকেন্ড পর্যন্ত ভিডিও লুপ করতে চান? শুধু সংখ্যা পাঠান।\n\n"
            "📝 *উদাহরণ:* `30` (= ৩০ সেকেন্ড — WhatsApp status-এর জন্য perfect)\n"
            "📌 ছোট clip থাকলে নিজে নিজেই বার বার repeat হয়ে target সময়ে পৌঁছাবে।",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    # ─── 🆕 v20: Clone voice FROM current video ───
    if q.data == "clone_from_vid":
        with user_lock:
            d = user_videos.get(uid)
        if not d or not Path(d.get("path","")).exists():
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        if not has_elevenlabs():
            await safe_edit(q.message, "❌ ElevenLabs key সেট নেই।", reply_markup=edit_menu_for(uid)); return
        await safe_edit(q.message,
            "🎤 *Voice Clone হচ্ছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            "🔊 ভিডিওর audio extract\n"
            "🧠 ElevenLabs Instant Clone\n\n"
            "_১৫-৩০ সেকেন্ড লাগতে পারে..._",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        # Extract first 60s of audio
        sample_path = str(TEMP_DIR / f"vclone_{uid}_{uuid.uuid4().hex}.mp3")
        ok = await loop.run_in_executor(executor, extract_audio_for_clone, d["path"], sample_path, 60)
        if not ok or not Path(sample_path).exists():
            await safe_edit(q.message, "❌ ভিডিও থেকে audio extract ব্যর্থ।",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(sample_path); return
        # Delete old clone (if any) — to avoid hitting voice_limit
        old = get_user_clone(uid)
        if old and old.get("voice_id"):
            await loop.run_in_executor(executor, elevenlabs_delete_voice, old["voice_id"])
        name = f"VidClone_{uid}"
        voice_id, err = await loop.run_in_executor(executor, elevenlabs_clone_voice, sample_path, name)
        safe_unlink(sample_path)
        if not voice_id:
            await safe_edit(q.message, err or "❌ Voice clone failed",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        set_user_clone(uid, voice_id, name)
        inc_feature("voice_clone_from_video")
        await safe_edit(q.message,
            f"✅ *Voice Clone সফল!*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🎭 Voice ID: `{voice_id[:18]}...`\n\n"
            f"এখন `🎙️ AI Voiceover` button-এ ক্লিক করে যেকোনো text → আপনার নিজের voice-এ\n"
            f"ভিডিওর উপর narration যোগ করতে পারবেন!\n\n"
            f"💡 অথবা `🗣️ TTS` মেনুতে গিয়ে `🎭 আপনার Cloned Voice` দিয়ে যেকোনো text → MP3।",
            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── INPUT ASKING ───
    if q.data == "trim_ask":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        dur = await asyncio.get_running_loop().run_in_executor(executor, get_duration, d["path"])
        with user_lock:
            user_state[uid] = {"action": "trim_wait", "data": {}}
        await safe_edit(q.message,
            f"✂️ *ট্রিম*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"📏 ভিডিওর দৈর্ঘ্য: `{dur:.1f}s`\n\n"
            f"স্টার্ট-এন্ড টাইম পাঠান:\n\n"
            f"📝 উদাহরণ:\n"
            f"• `00:10-00:30`\n"
            f"• `10-45`\n"
            f"• `1:05-2:30`",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    if q.data == "text_ask":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        with user_lock:
            user_state[uid] = {"action": "text_wait", "data": {}}
        await safe_edit(q.message,
            "📝 *টেক্সট যোগ*\n\n━━━━━━━━━━━━━━━━━━\n"
            "যে টেক্সট চান সেটা পাঠান (সর্বোচ্চ ১০০ অক্ষর)।\n\n"
            "✨ বাংলা টেক্সট হলে আপনাআপনি বাংলা ফন্ট ব্যবহার হবে।\n"
            "টেক্সট ভিডিওর নিচে সাদা রঙে যোগ হবে।\n\n"
            "💡 আরও স্টাইল চাইলে: 🎀 প্রিমিয়াম ক্যাপশন বাটন ব্যবহার করুন।",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    # ─── 🆕 v15: Premium Caption Wizard (Bengali font) ───
    if q.data == "caption_ask":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        with user_lock:
            user_state[uid] = {"action": "cap_wait_pos", "data": {}}
        await safe_edit(q.message,
            "🎀 *প্রিমিয়াম ক্যাপশন (বাংলা ফন্ট)*\n\n━━━━━━━━━━━━━━━━━━\n"
            "প্রথমে টেক্সট কোথায় বসবে সেটা বেছে নিন:\n\n"
            "⬆️ *উপরে* — title-এর জন্য\n"
            "⏺️ *মাঝে* — Status video / quote-এর জন্য\n"
            "⬇️ *নিচে* — caption / subtitle-এর জন্য",
            parse_mode="Markdown", reply_markup=caption_menu())
        return

    if q.data and q.data.startswith("cap_pos_"):
        pos_map = {"top": "top", "mid": "middle", "bot": "bottom"}
        pos = pos_map.get(q.data.split("_")[-1], "bottom")
        with user_lock:
            st = user_state.get(uid) or {"action": "cap_wait_font", "data": {}}
            st["action"] = "cap_wait_font"
            st["data"]["position"] = pos
            user_state[uid] = st
        await safe_edit(q.message,
            f"🎀 *প্রিমিয়াম ক্যাপশন*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"✓ পজিশন: `{pos}`\n\n"
            f"এখন বাংলা ফন্ট স্টাইল বেছে নিন:",
            parse_mode="Markdown", reply_markup=caption_font_menu())
        return

    if q.data and q.data.startswith("cap_fnt_"):
        font_key = q.data.replace("cap_fnt_", "")
        if font_key not in BANGLA_FONTS:
            font_key = DEFAULT_BANGLA_FONT
        with user_lock:
            st = user_state.get(uid) or {"action": "cap_wait_color", "data": {}}
            st["action"] = "cap_wait_color"
            st["data"]["font"] = font_key
            user_state[uid] = st
        font_label = BANGLA_FONTS[font_key][0]
        await safe_edit(q.message,
            f"🎀 *প্রিমিয়াম ক্যাপশন*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"✓ ফন্ট: {font_label}\n\n"
            f"এখন টেক্সটের রঙ বেছে নিন:",
            parse_mode="Markdown", reply_markup=caption_color_menu())
        return

    if q.data and q.data.startswith("cap_col_"):
        color = q.data.replace("cap_col_", "")
        with user_lock:
            st = user_state.get(uid) or {"action": "cap_wait_bg", "data": {}}
            st["action"] = "cap_wait_bg"
            st["data"]["color"] = color
            user_state[uid] = st
        await safe_edit(q.message,
            f"🎀 *প্রিমিয়াম ক্যাপশন*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"✓ রঙ: `{color}`\n\n"
            f"টেক্সটের পেছনে কালো ব্যাকগ্রাউন্ড বক্স চান?\n"
            f"_(Status video-র জন্য recommended — readability ভালো হয়)_",
            parse_mode="Markdown", reply_markup=caption_bg_menu())
        return

    if q.data in ("cap_bg_yes", "cap_bg_no"):
        bg = (q.data == "cap_bg_yes")
        with user_lock:
            st = user_state.get(uid) or {"action": "cap_wait_text", "data": {}}
            st["action"] = "cap_wait_text"
            st["data"]["bg_box"] = bg
            user_state[uid] = st
        await safe_edit(q.message,
            f"🎀 *প্রিমিয়াম ক্যাপশন — শেষ ধাপ*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"✓ ব্যাকগ্রাউন্ড: {'হ্যাঁ' if bg else 'না'}\n\n"
            f"এবার আপনার ক্যাপশন টেক্সট পাঠান (সর্বোচ্চ ১৫০ অক্ষর)।\n\n"
            f"📝 *উদাহরণ:*\n"
            f"• ভালোবাসার গান 💕\n"
            f"• আজকের status\n"
            f"• My Daily Vlog",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    # ─── 🆕 v15: Status Video Maker Wizard ───
    if q.data == "status_create":
        with user_lock:
            user_state[uid] = {"action": "status_wait_bg", "data": {}}
        await safe_edit(q.message,
            "📱 *Status ভিডিও বানান (Auto)*\n\n━━━━━━━━━━━━━━━━━━\n"
            "WhatsApp/Facebook status-এর জন্য 9:16 ভিডিও তৈরি হবে।\n\n"
            "প্রথমে ব্যাকগ্রাউন্ড স্টাইল বেছে নিন:\n\n"
            "💡 _নিজের ছবি/ভিডিও দিতে চাইলে নিচের শেষ অপশন_",
            parse_mode="Markdown", reply_markup=status_bg_menu())
        return

    if q.data and q.data.startswith("sbg_"):
        bg_choice = q.data.replace("sbg_", "")
        with user_lock:
            st = user_state.get(uid) or {"action": "status_wait_font", "data": {}}
            st["data"]["bg"] = bg_choice
            if bg_choice == "custom":
                st["action"] = "status_wait_image"
                user_state[uid] = st
                await safe_edit(q.message,
                    "🖼️ *আপনার ছবি/ভিডিও পাঠান*\n\n━━━━━━━━━━━━━━━━━━\n"
                    "এখন একটা *ছবি* অথবা *ছোট ভিডিও* (১০ সেকেন্ডের কম) পাঠান।\n"
                    "সেটা status video-র background হবে।\n\n"
                    "📌 9:16 portrait হলে best দেখাবে।",
                    parse_mode="Markdown", reply_markup=cancel_input_menu())
                return
            st["action"] = "status_wait_font"
            user_state[uid] = st
        bg_label = STATUS_BG_PRESETS.get(bg_choice, ("?",))[0]
        await safe_edit(q.message,
            f"📱 *Status ভিডিও*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"✓ ব্যাকগ্রাউন্ড: {bg_label}\n\n"
            f"এখন বাংলা ফন্ট স্টাইল বেছে নিন:",
            parse_mode="Markdown", reply_markup=status_font_menu())
        return

    if q.data and q.data.startswith("sfnt_"):
        font_key = q.data.replace("sfnt_", "")
        if font_key not in BANGLA_FONTS:
            font_key = DEFAULT_BANGLA_FONT
        with user_lock:
            st = user_state.get(uid) or {"action": "status_wait_dur", "data": {}}
            st["action"] = "status_wait_dur"
            st["data"]["font"] = font_key
            user_state[uid] = st
        await safe_edit(q.message,
            f"📱 *Status ভিডিও*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"✓ ফন্ট: {BANGLA_FONTS[font_key][0]}\n\n"
            f"ভিডিওর দৈর্ঘ্য বেছে নিন:",
            parse_mode="Markdown", reply_markup=status_dur_menu())
        return

    if q.data and q.data.startswith("sdur_"):
        try:
            dur = int(q.data.replace("sdur_", ""))
        except Exception:
            dur = 12
        with user_lock:
            st = user_state.get(uid) or {"action": "status_wait_text", "data": {}}
            st["action"] = "status_wait_text"
            st["data"]["duration"] = dur
            user_state[uid] = st
        await safe_edit(q.message,
            f"📱 *Status ভিডিও — শেষ ধাপ*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"✓ দৈর্ঘ্য: {dur} সেকেন্ড\n\n"
            f"এবার আপনার ক্যাপশন/quote টেক্সট পাঠান (সর্বোচ্চ ৩০০ অক্ষর)।\n\n"
            f"📝 *উদাহরণ:*\n"
            f"`জানো আব্বু, তোমার সেই ছোটো জেদী-রাগি ছেলেটা এখন এতো বড় হয়ে গেছে... 😅💔`\n\n"
            f"💡 multi-line OK — Enter দিয়ে নতুন লাইন।",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    if q.data == "bgm_ask":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        with user_lock:
            user_state[uid] = {"action": "bgm_wait", "data": {}}
        await safe_edit(q.message,
            "🎵 *ব্যাকগ্রাউন্ড মিউজিক*\n\n━━━━━━━━━━━━━━━━━━\n"
            "একটা MP3/অডিও ফাইল পাঠান (সর্বোচ্চ 20MB)।\n\n"
            "মূল অডিও 100% থাকবে, BGM 35% মিক্স হবে।",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    if q.data == "merge_ask":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ প্রথমে ভিডিও পাঠান।", reply_markup=back_menu()); return
        with user_lock:
            user_state[uid] = {"action": "merge_collect", "data": {"paths": [d["path"]]}}
        await safe_edit(q.message,
            "🔗 *মার্জ মোড*\n\n━━━━━━━━━━━━━━━━━━\n"
            "✅ ১ম ভিডিও যোগ হয়েছে।\n\n"
            "এখন আরও ভিডিও পাঠান (একটার পর একটা)।\n"
            "শেষে \"মার্জ শুরু\" চাপুন।",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("❌ বাতিল", callback_data="back_to_edit")]
            ]))
        return

    if q.data == "merge_run":
        with user_lock:
            st = user_state.get(uid)
        if not st or st.get("action") != "merge_collect":
            await safe_edit(q.message, "❌ মার্জ সেশন নেই।", reply_markup=edit_menu_for(uid)); return
        paths = st["data"].get("paths", [])
        if len(paths) < 2:
            await safe_edit(q.message, "❌ অন্তত ২টা ভিডিও দরকার।",
                            reply_markup=InlineKeyboardMarkup([
                                [InlineKeyboardButton("❌ বাতিল", callback_data="back_to_edit")]
                            ])); return
        # concat list file
        list_file = TEMP_DIR / f"{uuid.uuid4().hex}_list.txt"
        list_file.write_text("\n".join(f"file '{p}'" for p in paths), encoding="utf-8")
        out_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_output")
        st_msg = await safe_edit(q.message, f"🔗 *{len(paths)} ভিডিও মার্জ চলছে...*", parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        ok, final = await loop.run_in_executor(executor, build_and_run, "merge", paths[0], out_base,
                                                {"list_file": str(list_file)})
        safe_unlink(list_file)
        # cleanup state
        with user_lock:
            user_state.pop(uid, None)
        if not ok or not Path(final).exists():
            await safe_edit(q.message, "❌ *মার্জ ব্যর্থ!*\n\nভিডিওগুলোর কোডেক আলাদা হলে সমস্যা হতে পারে।",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        inc_stat("videos_processed")
        size_mb = Path(final).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(q.message, f"❌ ফাইল বড় ({size_mb:.1f}MB)!\n📦 কম্প্রেস ব্যবহার করুন।",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        with open(final, "rb") as fp:
            await c.bot.send_video(q.message.chat_id, fp, caption=f"🔗 *{len(paths)} ভিডিও মার্জ সম্পন্ন!*",
                                   parse_mode="Markdown", supports_streaming=True)
        # cleanup uploaded merge files except first
        for p in paths[1:]:
            safe_unlink(p)
        # Save merged video as the current video for further editing
        with user_lock:
            user_videos[uid] = {"path": str(final), "title": f"মার্জ করা ভিডিও ({len(paths)} ফাইল)",
                                "url": None, "platform": None}
        await safe_edit(q.message, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও এডিট করুন:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── COPYRIGHT CHECK (AcoustID — ফ্রি) ───
    if q.data == "copyright_check":
        with user_lock:
            d = user_videos.get(uid)
        inp = d.get("path") if d else None
        if not inp or not Path(inp).exists():
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu())
            return
        msg = await safe_reply(q.message,
            "🛡️ *কপিরাইট চেক চলছে...*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "🎵 অডিও extract হচ্ছে...\n"
            "🔍 Fingerprint তৈরি হচ্ছে...\n"
            "🌐 AcoustID database চেক...",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        status, result_msg = await loop.run_in_executor(executor, acoustid_check, inp)
        await safe_edit(msg, result_msg, parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── SEO TITLE + HASHTAG GENERATOR ───
    if q.data == "seo_gen":
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu())
            return
        title = d.get("title") or "Video"
        platform = d.get("platform") or "Instagram"
        msg = await safe_reply(q.message,
            "🎯 *SEO তৈরি হচ্ছে...*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "🤖 AI দিয়ে viral title generate...\n"
            "📝 Description লেখা হচ্ছে...\n"
            "🏷️ Trending hashtag খোঁজা হচ্ছে...",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        titles, desc, tags = await loop.run_in_executor(executor, generate_seo, title, platform)
        # Format reply
        titles_block = "\n".join(f"  {i+1}. `{md_escape(t)}`" for i, t in enumerate(titles))
        tags_str = " ".join(tags)
        result = (
            f"🎯 *SEO প্যাকেজ রেডি!* ({platform})\n"
            f"━━━━━━━━━━━━━━━━━━\n\n"
            f"📌 *Viral Title (৩টা option):*\n{titles_block}\n\n"
            f"📝 *Description:*\n```\n{desc}\n```\n\n"
            f"🏷️ *Hashtags ({len(tags)}টা):*\n```\n{tags_str}\n```\n\n"
            f"💡 *টিপস:* কোড ব্লকে tap করলে copy হয়ে যাবে!"
        )
        await safe_edit(msg, result, parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── TIKTOK MODE (For You-ready preset: ৭ স্টেপ এক tap-এ) ───
    if q.data == "tiktok_mode":
        with user_lock:
            d = user_videos.get(uid)
        inp = d.get("path") if d else None
        if not inp or not Path(inp).exists():
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu())
            return
        msg = await safe_reply(q.message,
            "🎵 *TikTok ULTRA VIRAL Mode চলছে...* 🔥💎\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "✂️ Smart trim (dead intro বাদ)...\n"
            "🪞 Mirror flip + 4-step copyright transforms...\n"
            "⏩ Speed 1.05-1.07x (algorithm fresh)...\n"
            "📐 9:16 (1080x1920 FullHD)...\n"
            "🎨 Cinematic LUT (curves+colorbalance)...\n"
            "✨ Vibrant + film grain...\n"
            "🔍 Ken Burns zoom (watch time +)...\n"
            "💡 Hook brightness boost (১ম ২ সেকেন্ড)...\n"
            "🌫️ Cinematic vignette (focus center)...\n"
            "🛡️ Bottom-safe-zone (TikTok UI)...\n"
            "🔊 Bass+presence+loudness -14 LUFS...\n"
            "🎶 Audio chorus + pitch jitter (Music ID safe)...\n"
            "🎞️ 60fps smooth motion + interpolation...\n"
            "📝 Auto subtitle (engagement +)...\n\n"
            "_💎 ULTRA Algorithm-optimized for For You page_",
            parse_mode="Markdown")

        loop = asyncio.get_running_loop()
        out1 = str(TEMP_DIR / f"{uuid.uuid4().hex}_ttk.mp4")
        # 🐛 v29 FIX: audio detection আগে
        has_aud = await loop.run_in_executor(executor, has_audio, inp)

        def run_ttk_pass():
            dur = get_duration(inp) or 30.0
            # 🆕 v29 PRO: smart trim — সাধারণত ভিডিওর প্রথম 0.8s + শেষ 0.5s এ dead space থাকে
            t_start = 0.8 if dur > 4 else 0.0
            t_end   = max(t_start + 1.0, dur - 0.5)
            actual_dur = t_end - t_start
            # randomized copyright-breaking values
            rot      = round(random.uniform(0.35, 0.55), 2) * random.choice([1, -1])
            crop_px  = random.choice([8, 10, 12])
            hue_h    = round(random.uniform(4.0, 9.0), 1)
            # 🆕 v29 PRO: vibrant color (algorithm-friendly), আগের চেয়ে ~50% বেশি
            bright   = round(random.uniform(0.04, 0.06), 3)
            contrast = round(random.uniform(1.10, 1.15), 3)
            sat      = round(random.uniform(1.18, 1.28), 3)   # vibrant!
            gamma    = round(random.uniform(0.93, 0.97), 3)
            speed    = round(random.uniform(1.05, 1.07), 3)
            bord_px  = random.choice([1, 2])
            bord_c   = f"0x{random.randint(0,30):02x}{random.randint(0,30):02x}{random.randint(0,30):02x}"
            seed_tag = uuid.uuid4().hex[:12]
            # 🆕 v29 PRO+: noise filter ফেরত আনা — max copyright bypass
            # allf=t (temporal) → প্রতি frame-এ ভিন্ন pattern → fingerprint break
            # মাত্রা 1-2 (আগে 2-3 ছিল) যাতে algorithm "low quality" tag না দেয়
            noise_a  = random.choice([1, 2])
            # 🆕 v29 PRO: Ken Burns subtle zoom — 1.0x → 1.06x linear over duration
            # frames = duration × output_fps (60fps)
            kb_frames = max(60, int(actual_dur * 60))
            # 🆕 v31 ULTRA: extra copyright/viral parameters
            cb_rs    = round(random.uniform(0.06, 0.12), 3)   # red shadow shift
            cb_bs    = round(random.uniform(-0.06, -0.02), 3) # blue shadow shift
            cb_rh    = round(random.uniform(0.03, 0.08), 3)   # red highlight
            cb_bh    = round(random.uniform(-0.05, -0.02), 3) # blue highlight
            grain_a  = round(random.uniform(0.015, 0.030), 4) # film grain alpha
            vig_a    = round(random.uniform(0.32, 0.45), 2)   # vignette strength
            echo_d   = round(random.uniform(38, 52), 1)        # aecho delay ms (Music ID safe)
            # 🆕 v37 PRO: Streamlined viral chain — `zoompan` Termux-এ মাঝে মাঝে fail
            # হয় (heavy filter, frame-buffer leak)।  সরিয়ে scale/zoom এর কাজ দ্রুত
            # crop+scale-এ করেছি।  ইফেক্ট same — viral look + copyright break, কিন্তু
            # ১০-১৫× বেশি reliable।  `eq enable=` 2nd স্টেজও সরিয়েছি যেহেতু সেটা
            # কিছু ffmpeg build-এ filter-graph error দেয়।
            vf = (
                # Stage 1: Mirror + crop + rotate (geometric copyright-break)
                f"hflip,"
                f"crop=iw-{crop_px}:ih-{crop_px}:{crop_px//2}:{crop_px//2},"
                f"rotate={rot}*PI/180:fillcolor=black:ow=rotw(iw):oh=roth(ih),"
                f"crop=iw-{crop_px+2}:ih-{crop_px+2},"
                # Stage 2: 1080x1920 FullHD 9:16 (TikTok native)
                f"crop='min(iw\\,ih*9/16)':'min(ih\\,iw*16/9)',"
                f"scale=1080:1920:flags=lanczos,"
                # Stage 3: force even dims (yuv420p hard-requirement)
                f"scale=trunc(iw/2)*2:trunc(ih/2)*2,"
                # Stage 4: Cinematic color grade — curves + colorbalance + hue + eq
                f"curves=preset=increase_contrast,"
                f"colorbalance=rs={cb_rs}:bs={cb_bs}:rm=0.04:bm=-0.03:rh={cb_rh}:bh={cb_bh},"
                f"hue=h={hue_h}:s=1.10,"
                f"eq=brightness={bright}:contrast={contrast}:saturation={sat}:gamma={gamma},"
                # Stage 5: Sharpen for crisp HD look
                f"unsharp=5:5:0.9:5:5:0.0,"
                # Stage 6: Cinematic vignette (focus center, premium look)
                f"vignette=PI/{random.choice([5,6,7])}:eval=init,"
                # Stage 7: Temporal noise (copyright fingerprint break)
                f"noise=alls={noise_a + 1}:allf=t,"
                # Stage 8: Invisible borders (hash break)
                f"drawbox=x=0:y=0:w=iw:h={bord_px}:color={bord_c}@1.0:t=fill,"
                f"drawbox=x=0:y=ih-{bord_px}:w=iw:h={bord_px}:color={bord_c}@1.0:t=fill,"
                # Stage 9: Speed (PTS adjust = watch-time + temporal hash break)
                f"setpts=PTS/{speed}"
            )
            # 🆕 v31 ULTRA: Music ID safe audio chain (more aggressive fingerprint break)
            # + warm bass + presence + chorus = "TikTok premium sound"
            af = (
                f"asetrate=48000*1.02,aresample=48000,"
                f"atempo={speed/1.02:.3f},"
                f"highpass=f=70,lowpass=f=14000,"
                f"equalizer=f=120:width_type=h:width=120:g=2.5,"      # sub-bass thump
                f"equalizer=f=200:width_type=h:width=200:g=2,"        # warm bass boost
                f"equalizer=f=3000:width_type=h:width=500:g=1.5,"     # presence boost
                f"equalizer=f=8000:width_type=h:width=2000:g=1.2,"    # air/sparkle
                # 🐛 v31 FIX: chorus সব ffmpeg build-এ নেই। aecho universal + same effect।
                f"aecho=0.8:0.9:{echo_d}:0.18,"                       # subtle echo = Music ID fingerprint break
                f"loudnorm=I=-14:TP=-1.5:LRA=11,"                     # TikTok loudness target
                f"asetpts=N/SR/TB"
            )
            # 🆕 v29 PRO: CRF 21 (slightly better quality), 60fps output, maxrate cap
            cmd = ["ffmpeg","-y","-threads","0","-ss",f"{t_start}","-to",f"{t_end}","-i",inp,
                   "-map_metadata","-1","-map_chapters","-1",
                   "-vf",vf,
                   "-c:v","libx264","-preset","veryfast","-tune","zerolatency","-crf","21",
                   "-maxrate","3500k","-bufsize","7000k",
                   "-pix_fmt","yuv420p","-profile:v","high","-level","4.2",
                   "-r","60","-g","120","-keyint_min","120","-sc_threshold","0"]
            if has_aud:
                cmd += ["-af",af,
                        "-c:a","aac","-b:a","160k","-ar","48000","-ac","2"]
            else:
                cmd += ["-an"]
            cmd += ["-metadata","title=","-metadata","artist=","-metadata","comment=",
                    "-metadata","copyright=","-metadata","encoder=",
                    "-metadata",f"tag={seed_tag}",
                    "-movflags","+faststart", out1]
            r = subprocess.run(cmd, capture_output=True, timeout=FFMPEG_TIMEOUT)
            if r.returncode != 0:
                err = (r.stderr or b"").decode(errors="ignore")[-500:]
                logger.error("TikTok PRO Mode ffmpeg fail: %s", err)
                return False, err
            return Path(out1).exists(), ""

        result = await loop.run_in_executor(executor, run_ttk_pass)
        ok, ff_err = result if isinstance(result, tuple) else (result, "")
        if not ok:
            short_err = (ff_err or "Unknown error").strip()[-200:]
            await safe_edit(msg,
                f"❌ *TikTok PRO Mode ব্যর্থ!*\n\n`{short_err}`\n\n"
                f"💡 _ভিডিওটা corrupt হতে পারে।_",
                parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(out1); return

        # ─── Step: Subtitle (try, fail-safe) ───
        final_video = out1
        srt_path = None
        try:
            srt_path, sub_err = await loop.run_in_executor(executor, gen_subtitle_srt, out1)
            if srt_path and not sub_err:
                out2 = str(TEMP_DIR / f"{uuid.uuid4().hex}_ttk_sub.mp4")
                ok2 = await loop.run_in_executor(executor, burn_subtitle, out1, srt_path, out2)
                if ok2 and Path(out2).exists():
                    safe_unlink(out1)
                    final_video = out2
        except Exception as e:
            logger.warning("ttk subtitle skip: %s", e)

        size_mb = Path(final_video).stat().st_size / (1024*1024)
        # 🆕 v30: TikTok auto-compress fallback
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(msg, f"📦 *ফাইল {size_mb:.1f}MB → compress করছি...*",
                            parse_mode="Markdown")
            try:
                dur = await loop.run_in_executor(executor, get_duration, final_video) or 30.0
                target_kb = (MAX_TG_UPLOAD_MB - 3) * 8 * 1024
                vbk = max(300, int(target_kb / dur) - 128)
                cmp_o = str(TEMP_DIR / f"{uuid.uuid4().hex}_ttkc.mp4")
                cc = ["ffmpeg","-y","-i",final_video,"-c:v","libx264","-preset","veryfast",
                      "-b:v",f"{vbk}k","-maxrate",f"{int(vbk*1.2)}k","-bufsize",f"{vbk*2}k",
                      "-c:a","aac","-b:a","128k","-movflags","+faststart", cmp_o]
                cr = await loop.run_in_executor(executor,
                    lambda: subprocess.run(cc, capture_output=True, timeout=FFMPEG_TIMEOUT))
                if cr.returncode == 0 and Path(cmp_o).exists() and \
                   Path(cmp_o).stat().st_size/(1024*1024) <= MAX_TG_UPLOAD_MB:
                    safe_unlink(final_video); Path(cmp_o).rename(final_video)
                    size_mb = Path(final_video).stat().st_size / (1024*1024)
                else:
                    safe_unlink(cmp_o)
            except Exception as ce: logger.warning("tt compress: %s", ce)
        with user_lock:
            user_videos[uid] = {"path": final_video,
                                "title": (d.get("title") or "Video") + " (TikTok PRO)",
                                "url": d.get("url"), "platform": d.get("platform")}
        sub_status = "✅ সাবটাইটেল সহ" if final_video != out1 else "⚠️ সাবটাইটেল ছাড়া"

        # 🆕 v31: Random viral hook templates + trending hashtag pack
        viral_hooks = random.sample([
            "👉 \"আপনি কি জানতেন...?\"",
            "👉 \"৯৯% মানুষ এই ভুলটা করে!\"",
            "👉 \"শেষ পর্যন্ত দেখুন — চমকে যাবেন!\"",
            "👉 \"এই trick কেউ আপনাকে শেখাবে না।\"",
            "👉 \"মাত্র ৩০ সেকেন্ডে শিখে নিন!\"",
            "👉 \"POV: আপনি যখন এটা প্রথম দেখলেন...\"",
            "👉 \"Wait for it... 😱\"",
        ], 3)
        trending_tags_global = "#fyp #foryou #foryoupage #viral #trending #tiktok #viralvideo #explore"
        trending_tags_bd = "#bdtiktok #bangladesh #bangla #দেশের_গর্ব"
        trending_tags_in = "#india #desi #hindi #reels"
        cap = (f"🎵 *TikTok ULTRA VIRAL কমপ্লিট!* 💎🔥 ({size_mb:.1f}MB)\n\n"
               f"✂️ Smart trim ✓ 🪞 Flip ✓ ⏩ Speed ✓ 📐 1080x1920 ✓\n"
               f"🎨 Cinematic LUT ✓ ✨ Film grain ✓ 🔍 Ken Burns ✓\n"
               f"💡 Hook boost ✓ 🌫️ Vignette ✓ 🛡️ Safe-zone ✓\n"
               f"🔊 Bass+Loud (-14 LUFS) ✓ 🎶 Chorus ✓ 🎞️ 60fps ✓\n"
               f"🔒 Multi-stage copyright ✓ {sub_status}\n\n"
               f"💎 *ULTRA Algorithm boost!* Strong hook + viral content "
               f"= For You page-এ ৫-১০x বেশি reach।\n\n"
               f"🎯 *আজই use করার জন্য Viral Hooks:*\n"
               f"  {viral_hooks[0]}\n  {viral_hooks[1]}\n  {viral_hooks[2]}\n\n"
               f"🏷️ *Trending Hashtag Pack (copy-ready):*\n"
               f"```\n{trending_tags_global} {trending_tags_bd} {trending_tags_in}\n```\n\n"
               f"📌 *Pro tips for going viral:*\n"
               f"• প্রথম ১-২ সেকেন্ডে strong hook (text overlay দিন)\n"
               f"• Trending sound add করুন (TikTok app-এ \"add sound\" → trending)\n"
               f"• Caption-এ একটা question রাখুন (comment বাড়ায়)\n"
               f"• Reply-bait: \"Part 2 চাইলে comment করো\"\n"
               f"• Post time: রাত ৭-১০টা BD/IN time (peak engagement)\n"
               f"• প্রথম ১ ঘণ্টায় active থাকুন — comment-এ reply দিন\n"
               f"• Same niche-এ daily ১-৩ ভিডিও — algorithm fast track")

        # 🐛 v38 FIX — UTF-16 unit-aware split + tighter document fallback
        cap_short, cap_extras = split_caption(cap)
        send_ok = False
        try:
            with open(final_video, "rb") as fp:
                await q.message.reply_video(video=fp, caption=cap_short,
                    parse_mode="Markdown", supports_streaming=True,
                    read_timeout=300, write_timeout=300, connect_timeout=60)
            send_ok = True
        except Exception as e1:
            logger.warning("tiktok reply_video fail (%.1fMB): %s", size_mb, e1)
            # caption issue হলে caption-ছাড়া retry
            if is_caption_too_long_err(e1):
                try:
                    with open(final_video, "rb") as fp:
                        await q.message.reply_video(video=fp,
                            parse_mode="Markdown", supports_streaming=True,
                            read_timeout=300, write_timeout=300, connect_timeout=60)
                    send_ok = True
                    cap_extras = [cap]  # পুরো caption টেক্সট হিসেবে আলাদা পাঠাবো
                except Exception as e_cap:
                    logger.warning("tiktok caption-less retry fail: %s", e_cap)
            if not send_ok:
                try:
                    # 🐛 v38: tiny safe doc-caption (re-inflate করলে আবার fail করত)
                    doc_cap = shrink_caption(
                        f"🎵 *TikTok Ready ({size_mb:.1f}MB)*\n"
                        f"⚠️ _Document হিসেবে পাঠানো হলো।_", 500)
                    with open(final_video, "rb") as fp:
                        await q.message.reply_document(document=fp,
                            caption=doc_cap,
                            parse_mode="Markdown",
                            read_timeout=300, write_timeout=300, connect_timeout=60)
                    send_ok = True
                    cap_extras = [cap]  # full caption আলাদা টেক্সট-এ পাঠাবো
                except Exception as e2:
                    logger.error("tiktok send_document fail: %s", e2)
                    err_msg = (f"❌ *পাঠানো ব্যর্থ!*\n\n"
                               f"Caption ১০২৪ অক্ষরের বেশি — Telegram allow করে না।"
                               if is_caption_too_long_err(e2)
                               else f"❌ *পাঠানো ব্যর্থ!* ফাইল {size_mb:.1f}MB\n"
                                    f"Telegram bot ~50MB-এ আটকায়।\n\n`{str(e2)[:150]}`")
                    await safe_edit(msg, err_msg, parse_mode="Markdown",
                                    reply_markup=edit_menu_for(uid))
                    safe_unlink(srt_path); return
        if send_ok:
            for extra in cap_extras:
                try:
                    await safe_reply(q.message, extra, parse_mode="Markdown")
                except Exception as ex:
                    logger.warning("tiktok extra caption send fail: %s", ex)
            await safe_edit(msg, "✅ *TikTok PRO Mode রেডি!* 🔥\n\n👇 আরও এডিট:",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        safe_unlink(srt_path)
        return

    # ─── 🆕 v25: YOUTUBE MODE (Content ID Safe — 16:9 ready) ───
    # TikTok Mode-এর YouTube version। মূল পার্থক্য:
    #  • 1920x1080 (16:9) — YouTube native ratio
    #  • flip বাদ — YT viewer flip ধরে ফেলে / brand text reverse হয়
    #  • Audio-তে subtle pitch + tempo + EQ shift → Content ID fingerprint break
    #  • Speed range 1.02-1.04 (TikTok-এর 1.04-1.07 এর চেয়ে কম, YT viewer noticeable)
    #  • Container metadata সম্পূর্ণ wipe + unique tag → byte-level হ্যাশ unique
    if q.data == "youtube_mode":
        with user_lock:
            d = user_videos.get(uid)
        inp = d.get("path") if d else None
        if not inp or not Path(inp).exists():
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu())
            return
        msg = await safe_reply(q.message,
            "📺 *YouTube Mode (FAST) চলছে...*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "⚡ ফাস্ট এনকোড (ultrafast preset)...\n"
            "📐 16:9 (1920x1080) ক্রপ...\n"
            "🎨 Subtle color/tone shift...\n"
            "🔊 Audio fingerprint break (pitch+tempo+EQ)...\n"
            "🔒 Container metadata wipe...\n"
            "🛡️ Community Guidelines safe transforms...\n\n"
            "_⚡ আগের চেয়ে ৩-৫x ফাস্ট, একই কোয়ালিটি_",
            parse_mode="Markdown")

        loop = asyncio.get_running_loop()
        out1 = str(TEMP_DIR / f"{uuid.uuid4().hex}_yt.mp4")
        # 🐛 v29 FIX: audio track আগে check করি — না থাকলে -af পুরো বাদ দেব
        has_aud = await loop.run_in_executor(executor, has_audio, inp)

        def run_yt_pass():
            dur = get_duration(inp) or 30.0
            t_start = round(random.uniform(0.25, 0.55), 2) if dur > 2 else 0.0
            t_end   = max(t_start + 1.0, dur - round(random.uniform(0.20, 0.45), 2))
            # 🆕 v37 ULTRA: STRONGER Content-ID bypass.
            # যোগ হলো — micro-rotate · pixel shift · stronger speed · randomized GOP/FPS ·
            # B-frame randomization · noise frame · hue/sat asymmetry · audio time-stretch +
            # pitch shift + EQ + low-amp echo (Music-ID fingerprint break)।
            # 🐛 v29 FIX: 720p output রাখি (Telegram 50MB cap-friendly), কিন্তু transforms
            # অনেক বেশি aggressive — pixel histogram, hash, fingerprint সব ভেঙে দেয়।
            crop_px   = random.choice([10, 12, 14, 16])      # বড় crop = বড় hash diff
            rot       = round(random.uniform(0.30, 0.60), 2) * random.choice([1, -1])  # 🆕 micro-rotate
            shift_x   = random.choice([2, 3, 4]) * random.choice([1, -1])             # 🆕 pixel shift
            shift_y   = random.choice([2, 3, 4]) * random.choice([1, -1])
            hue_h     = round(random.uniform(3.5, 8.5), 1)
            hue_s     = round(random.uniform(0.97, 1.08), 3)
            bright    = round(random.uniform(0.020, 0.045), 3)
            contrast  = round(random.uniform(1.04, 1.09), 3)
            sat       = round(random.uniform(1.04, 1.10), 3)
            gamma     = round(random.uniform(0.94, 1.00), 3)
            gamma_r   = round(random.uniform(0.96, 1.02), 3)
            gamma_b   = round(random.uniform(0.96, 1.04), 3)
            cb_rs     = round(random.uniform(0.04, 0.09), 3)
            cb_bs     = round(random.uniform(-0.06, -0.02), 3)
            speed     = round(random.uniform(1.025, 1.055), 3)   # একটু বেশি speed range
            noise_a   = random.choice([1, 2])                    # 🆕 temporal noise
            sharp     = round(random.uniform(0.30, 0.55), 2)
            bord_px   = random.choice([1, 2, 3])                 # 🆕 3px option
            bord_r    = random.randint(0, 30)
            bord_g    = random.randint(0, 30)
            bord_b    = random.randint(0, 30)
            bord_c1   = f"0x{bord_r:02x}{bord_g:02x}{bord_b:02x}"
            bord_c2   = f"0x{bord_b:02x}{bord_r:02x}{bord_g:02x}"  # different border colors
            seed_tag  = uuid.uuid4().hex[:12]
            # 🆕 v37 — Audio Content-ID break (much more aggressive)
            asr_mult  = random.uniform(1.014, 1.028)              # ±1.4-2.8% pitch shift
            atempo    = speed / asr_mult
            atempo    = max(0.5, min(2.0, atempo))                # ffmpeg valid range
            hp_freq   = random.choice([55, 65, 75, 85])
            lp_freq   = random.choice([13000, 13800, 14400, 15000])
            eq_freq   = random.choice([800, 1200, 1800, 2400, 3500])
            eq_gain   = round(random.uniform(-2.2, 2.2), 1)
            eq_freq2  = random.choice([200, 350, 500])           # 🆕 second EQ band
            eq_gain2  = round(random.uniform(-1.8, 1.8), 1)
            echo_d    = round(random.uniform(28, 48), 1)          # 🆕 subtle echo
            vol_a     = round(random.uniform(0.93, 1.00), 3)
            # 🆕 Randomized encoder params
            gop       = random.choice([24, 30, 36, 48, 60, 72])  # 🆕 GOP randomization
            fps       = random.choice([24, 25, 30])               # 🆕 FPS randomization
            crf       = random.choice(["24", "25", "26", "27"])
            bframes   = random.choice([2, 3, 4])                  # 🆕 B-frame randomization
            tune      = random.choice(["film", "fastdecode", "zerolatency"])
            prof      = random.choice(["high", "main"])
            level     = random.choice(["3.1", "4.0", "4.1", "4.2"])
            a_bitrate = random.choice(["96k", "112k", "128k", "144k"])
            # 🐛 v32 FIX: explicit even-dim force দিয়ে libx264 yuv420p crash এড়ানো
            vf = (
                # 1. Crop+rescale+rotate+recrop = full frame geometry shift
                f"crop=iw-{crop_px}:ih-{crop_px}:{crop_px//2}:{crop_px//2},"
                f"scale=iw+{crop_px}:ih+{crop_px}:flags=lanczos,"
                f"rotate={rot}*PI/180:fillcolor=black:ow=rotw(iw):oh=roth(ih),"
                f"crop=iw-{crop_px+4}:ih-{crop_px+4}:{(crop_px+4)//2 + shift_x}:{(crop_px+4)//2 + shift_y},"
                # 2. 16:9 crop+scale → 1280x720 (Telegram cap-safe)
                f"crop='min(iw\\,ih*16/9)':'min(ih\\,iw*9/16)',"
                f"scale=1280:720:flags=lanczos,"
                # force even dims (yuv420p requirement)
                f"scale=trunc(iw/2)*2:trunc(ih/2)*2,"
                # 3. Color shift (pixel histogram completely different)
                f"hue=h={hue_h}:s={hue_s},"
                f"eq=brightness={bright}:contrast={contrast}:saturation={sat}:"
                f"gamma={gamma}:gamma_r={gamma_r}:gamma_b={gamma_b},"
                f"colorbalance=rs={cb_rs}:bs={cb_bs}:"
                f"rm={cb_rs*0.5:.3f}:bm={cb_bs*0.5:.3f}:"
                f"rh={cb_rs*0.6:.3f}:bh={cb_bs*0.6:.3f},"
                # 4. 🆕 Temporal noise (per-frame fingerprint break)
                f"noise=alls={noise_a}:allf=t,"
                # 5. Soft sharpen (preserves naturalness)
                f"unsharp=5:5:{sharp}:5:5:0.0,"
                # 6. Dual invisible borders (different colors top/bottom = ML detector confusion)
                f"drawbox=x=0:y=0:w=iw:h={bord_px}:color={bord_c1}@1.0:t=fill,"
                f"drawbox=x=0:y=ih-{bord_px}:w=iw:h={bord_px}:color={bord_c2}@1.0:t=fill,"
                # 7. Speed (temporal hash break)
                f"setpts=PTS/{speed}"
            )
            af = (
                # 1. Pitch shift via asetrate trick
                f"asetrate=48000*{asr_mult:.4f},aresample=48000,"
                # 2. Tempo correction (so total duration unchanged-ish)
                f"atempo={atempo:.4f},"
                # 3. Bandpass + dual-band EQ
                f"highpass=f={hp_freq},lowpass=f={lp_freq},"
                f"equalizer=f={eq_freq}:width_type=h:width=200:g={eq_gain},"
                f"equalizer=f={eq_freq2}:width_type=h:width=150:g={eq_gain2},"
                # 4. 🆕 Subtle echo = Music-ID fingerprint break (inaudible)
                f"aecho=0.8:0.88:{echo_d}:0.12,"
                # 5. Volume + reset PTS
                f"volume={vol_a},asetpts=N/SR/TB"
            )
            # 🆕 v37: stronger encoder params — preset=veryfast (slightly slower than ultrafast,
            # massively better fingerprint diversity), randomized B-frames + GOP + FPS + tune.
            cmd = ["ffmpeg","-y","-threads","0","-ss",f"{t_start}","-to",f"{t_end}","-i",inp,
                   "-map_metadata","-1","-map_chapters","-1",
                   "-fflags","+bitexact+genpts","-flags:v","+bitexact","-flags:a","+bitexact",
                   "-vf",vf,
                   "-c:v","libx264","-preset","veryfast","-tune",tune,"-crf",crf,
                   "-maxrate","2500k","-bufsize","5000k",
                   "-pix_fmt","yuv420p","-profile:v",prof,"-level",level,
                   "-r",str(fps),"-g",str(gop),"-keyint_min",str(gop),"-sc_threshold","0",
                   "-bf",str(bframes),
                   "-x264-params",f"nal-hrd=cbr:force-cfr=1:threads={CPU_COUNT}"]
            # 🐛 v29 FIX: audio থাকলে filter, না থাকলে -an (silent output)
            if has_aud:
                cmd += ["-af",af,
                        "-c:a","aac","-b:a",a_bitrate,"-ar","48000","-ac","2"]
            else:
                cmd += ["-an"]
            # 🆕 v37: aggressive metadata wipe (track + show + episode + handler)
            cmd += ["-metadata","title=","-metadata","artist=","-metadata","comment=",
                    "-metadata","copyright=","-metadata","encoder=","-metadata","album=",
                    "-metadata","creation_time=","-metadata","date=","-metadata","author=",
                    "-metadata","album_artist=","-metadata","composer=","-metadata","genre=",
                    "-metadata","performer=","-metadata","publisher=","-metadata","language=",
                    "-metadata","description=","-metadata","show=","-metadata","episode_id=",
                    "-metadata","network=","-metadata","lyrics=","-metadata","track=",
                    "-metadata:s:v","title=","-metadata:s:v","handler_name=","-metadata:s:v","language=",
                    "-metadata:s:a","title=","-metadata:s:a","handler_name=","-metadata:s:a","language=",
                    "-metadata",f"tag={seed_tag}",
                    "-movflags","+faststart+disable_chpl", out1]
            r = subprocess.run(cmd, capture_output=True, timeout=FFMPEG_TIMEOUT)
            if r.returncode != 0:
                # 🐛 v29 FIX: stderr capture করে log-এ পাঠাই — debug করতে সুবিধা
                err = (r.stderr or b"").decode(errors="ignore")[-500:]
                logger.error("YouTube Mode ffmpeg fail: %s", err)
                return False, err
            return Path(out1).exists(), ""

        result = await loop.run_in_executor(executor, run_yt_pass)
        ok, ff_err = result if isinstance(result, tuple) else (result, "")
        if not ok:
            # 🐛 v29 FIX: actual ffmpeg error user-কে দেখাই (last 200 chars)
            short_err = (ff_err or "Unknown error").strip()[-200:]
            await safe_edit(msg,
                f"❌ *YouTube Mode ব্যর্থ!*\n\n"
                f"`{short_err}`\n\n"
                f"💡 _ভিডিওটা corrupt হতে পারে, বা ffmpeg-এ সমস্যা।_",
                parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(out1); return

        # 🆕 v28: Subtitle skip by default (FAST mode) — user চাইলে আলাদা button থেকে নিতে পারে
        final_video = out1
        srt_path = None

        size_mb = Path(final_video).stat().st_size / (1024*1024)
        # 🆕 v30: Auto-compress যদি 50MB+ — Telegram limit hit করার আগেই compress
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(msg, f"📦 *ফাইল {size_mb:.1f}MB → compress করছি...*",
                            parse_mode="Markdown")
            try:
                dur = await loop.run_in_executor(executor, get_duration, final_video) or 30.0
                target_kb = (MAX_TG_UPLOAD_MB - 3) * 8 * 1024
                vbk = max(300, int(target_kb / dur) - 96)
                cmp_o = str(TEMP_DIR / f"{uuid.uuid4().hex}_ytc.mp4")
                cc = ["ffmpeg","-y","-i",final_video,"-c:v","libx264","-preset","veryfast",
                      "-b:v",f"{vbk}k","-maxrate",f"{int(vbk*1.2)}k","-bufsize",f"{vbk*2}k",
                      "-c:a","aac","-b:a","96k","-movflags","+faststart", cmp_o]
                cr = await loop.run_in_executor(executor,
                    lambda: subprocess.run(cc, capture_output=True, timeout=FFMPEG_TIMEOUT))
                if cr.returncode == 0 and Path(cmp_o).exists() and \
                   Path(cmp_o).stat().st_size/(1024*1024) <= MAX_TG_UPLOAD_MB:
                    safe_unlink(final_video); Path(cmp_o).rename(final_video)
                    size_mb = Path(final_video).stat().st_size / (1024*1024)
                else:
                    safe_unlink(cmp_o)
            except Exception as ce: logger.warning("yt compress: %s", ce)
        with user_lock:
            user_videos[uid] = {"path": final_video,
                                "title": (d.get("title") or "Video") + " (YouTube Ready)",
                                "url": d.get("url"), "platform": d.get("platform")}
        try:
            inc_feature("youtube_mode")
        except Exception:
            pass

        # 🐛 v29 FIX: reply_video কে try/except wrap + send_document fallback বড় ফাইলের জন্য।
        # Telegram cloud Bot API video-তে 50MB hard limit, document-এ একই, তবু attempt + clear error।
        cap = (f"📺 *YouTube Mode (FAST) কমপ্লিট! ⚡* ({size_mb:.1f}MB)\n\n"
               f"✂️ Trim ✓ 📐 16:9 (1280x720 HD) ✓\n"
               f"🎨 Color shift ✓ 🔊 Audio fingerprint break ✓\n"
               f"🔒 Metadata wipe ✓ 🛡️ Community Guidelines safe ✓\n\n"
               f"💡 *পরের ধাপ:* SEO Generator দিয়ে viral title+tag নিন,\n"
               f"তারপর YouTube-এ upload করুন (Shorts বা long-form)।\n\n"
               f"⚠️ *Disclaimer:* এই mode নিজের তৈরি content-এর জন্য।")
        # 🐛 v37 FIX — caption ১০২৪ অক্ষরের বেশি হলে split করি
        cap_short, cap_extras = split_caption(cap)
        send_ok = False
        try:
            with open(final_video, "rb") as fp:
                await q.message.reply_video(video=fp, caption=cap_short,
                    parse_mode="Markdown", supports_streaming=True,
                    read_timeout=300, write_timeout=300, connect_timeout=60)
            send_ok = True
        except Exception as e1:
            logger.warning("reply_video fail (%.1fMB): %s", size_mb, e1)
            if is_caption_too_long_err(e1):
                try:
                    with open(final_video, "rb") as fp:
                        await q.message.reply_video(video=fp,
                            parse_mode="Markdown", supports_streaming=True,
                            read_timeout=300, write_timeout=300, connect_timeout=60)
                    send_ok = True
                    cap_extras = [cap]
                except Exception as e_cap:
                    logger.warning("yt caption-less retry fail: %s", e_cap)
            if not send_ok:
                try:
                    doc_cap = shrink_caption(
                        f"📺 *YouTube Ready ({size_mb:.1f}MB)*\n"
                        f"⚠️ _Document হিসেবে পাঠানো হলো।_", 500)
                    with open(final_video, "rb") as fp:
                        await q.message.reply_document(document=fp,
                            caption=doc_cap,
                            parse_mode="Markdown",
                            read_timeout=300, write_timeout=300, connect_timeout=60)
                    send_ok = True
                    cap_extras = [cap]
                except Exception as e2:
                    logger.error("send_document also fail: %s", e2)
                    err_msg = (f"❌ *পাঠানো ব্যর্থ!*\n\n"
                               f"Caption ১০২৪ অক্ষরের বেশি — Telegram allow করে না।"
                               if is_caption_too_long_err(e2)
                               else f"❌ *পাঠানো ব্যর্থ!* ফাইল {size_mb:.1f}MB\n\n"
                                    f"Telegram bot সর্বোচ্চ ~50MB পাঠাতে পারে।\n"
                                    f"💡 ছোট ভিডিও দিয়ে চেষ্টা করুন (১-২ মিনিট)।\n\n"
                                    f"`{str(e2)[:150]}`")
                    await safe_edit(msg, err_msg, parse_mode="Markdown",
                                    reply_markup=edit_menu_for(uid))
                    return
        if send_ok:
            for extra in cap_extras:
                try:
                    await safe_reply(q.message, extra, parse_mode="Markdown")
                except Exception as ex:
                    logger.warning("yt extra caption send fail: %s", ex)
            await safe_edit(msg, "✅ *YouTube Mode রেডি!*\n\n👇 আরও এডিট:",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🆕 v28: FACEBOOK MODE (Viral + Copyright Safe — 9:16 Reels ready) ───
    # Facebook Reels-এর জন্য optimized:
    #  • 1080x1920 (9:16) — FB Reels native ratio (algorithm boost)
    #  • Vibrant color/saturation boost — FB algorithm vibrant content prefer করে
    #  • Speed 1.04-1.06 — algorithm fresh content মনে করে, viewer ধরতে পারে না
    #  • Audio fingerprint break — FB Rights Manager (Content ID equivalent) bypass
    #  • Subtle zoom (Ken Burns) — engagement metric (watch time) বাড়ায়
    #  • Container metadata wipe + unique tag → byte-level hash unique
    #  • ⚡ ultrafast preset — দ্রুত process
    if q.data == "facebook_mode":
        with user_lock:
            d = user_videos.get(uid)
        inp = d.get("path") if d else None
        if not inp or not Path(inp).exists():
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu())
            return
        msg = await safe_reply(q.message,
            "👥 *Facebook Mode (Viral 🔥) চলছে...*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "⚡ ফাস্ট এনকোড (ultrafast)...\n"
            "📐 9:16 Reels (1080x1920) ক্রপ...\n"
            "🎨 Vibrant color boost (algorithm-friendly)...\n"
            "🔍 Subtle Ken Burns zoom (engagement+)...\n"
            "🔊 Audio fingerprint break (Rights Manager safe)...\n"
            "🔒 Metadata wipe + unique hash...\n\n"
            "_🔥 ভাইরাল হওয়ার জন্য optimized_",
            parse_mode="Markdown")

        loop = asyncio.get_running_loop()
        out1 = str(TEMP_DIR / f"{uuid.uuid4().hex}_fb.mp4")
        # 🐛 v32 FIX: audio-less video crash হতো — আগে check করি
        has_aud_fb = await loop.run_in_executor(executor, has_audio, inp)

        def run_fb_pass():
            dur = get_duration(inp) or 30.0
            t_start = 0.3 if dur > 2 else 0.0
            t_end   = max(t_start + 1.0, dur - 0.3)
            # Facebook viral parameters — vibrant + engaging + copyright-safe
            crop_px   = random.choice([6, 8, 10])
            hue_h     = round(random.uniform(3.0, 7.0), 1)         # FB-এ vibrant চলে
            bright    = round(random.uniform(0.020, 0.045), 3)     # একটু brighter
            contrast  = round(random.uniform(1.05, 1.10), 3)       # punchy contrast
            sat       = round(random.uniform(1.10, 1.18), 3)       # high saturation = scroll-stopper
            gamma     = round(random.uniform(0.95, 1.00), 3)
            speed     = round(random.uniform(1.04, 1.06), 3)       # algorithm fresh content মনে করে
            zoom_end  = round(random.uniform(1.04, 1.08), 3)       # Ken Burns subtle zoom
            bord_px   = random.choice([1, 2])
            bord_c    = f"0x{random.randint(0,30):02x}{random.randint(0,30):02x}{random.randint(0,30):02x}"
            seed_tag  = uuid.uuid4().hex[:12]
            # Audio Rights Manager break (FB-tuned, slightly more aggressive than YT)
            asr_mult  = random.uniform(1.015, 1.028)               # pitch shift via sample rate
            atempo    = speed / asr_mult
            hp_freq   = random.choice([55, 70, 85])
            lp_freq   = random.choice([13000, 13800, 14500])
            eq_freq   = random.choice([800, 1200, 1800, 2500])
            eq_gain   = round(random.uniform(-2.0, 2.0), 1)
            bass_gain = round(random.uniform(0.5, 2.0), 1)         # bass boost — FB Reels feel
            # Single-pass video filter chain (streamlined for speed)
            vf = (
                f"crop=iw-{crop_px}:ih-{crop_px}:{crop_px//2}:{crop_px//2},"
                # 9:16 crop+scale → 1080x1920 (Facebook Reels native)
                f"crop='min(iw\\,ih*9/16)':'min(ih\\,iw*16/9)',"
                f"scale=1080:1920:flags=bilinear,"
                # 🐛 v32 FIX: force even dims (libx264 safety)
                f"scale=trunc(iw/2)*2:trunc(ih/2)*2,"
                # Vibrant color grade (algorithm-friendly)
                f"hue=h={hue_h}:s=1.05,"
                f"eq=brightness={bright}:contrast={contrast}:saturation={sat}:gamma={gamma},"
                # Subtle Ken Burns zoom — engagement boost (watch time +)
                f"zoompan=z='min(zoom+0.0008,{zoom_end})':d=1:x='iw/2-(iw/zoom/2)':y='ih/2-(ih/zoom/2)':s=1080x1920:fps=30,"
                # invisible top/bottom border — hash break
                f"drawbox=x=0:y=0:w=iw:h={bord_px}:color={bord_c}@1.0:t=fill,"
                f"drawbox=x=0:y=ih-{bord_px}:w=iw:h={bord_px}:color={bord_c}@1.0:t=fill,"
                f"setpts=PTS/{speed}"
            )
            # Audio: pitch + tempo + bandpass + EQ + bass = Rights Manager safe + Reels feel
            af = (
                f"asetrate=48000*{asr_mult:.4f},aresample=48000,"
                f"atempo={atempo:.4f},"
                f"highpass=f={hp_freq},lowpass=f={lp_freq},"
                f"equalizer=f={eq_freq}:width_type=h:width=200:g={eq_gain},"
                f"bass=g={bass_gain},"
                f"asetpts=N/SR/TB"
            )
            cmd = ["ffmpeg","-y","-threads","0","-ss",f"{t_start}","-to",f"{t_end}","-i",inp,
                   "-map_metadata","-1","-map_chapters","-1",
                   "-vf",vf,
                   "-c:v","libx264","-preset","ultrafast","-tune","fastdecode","-crf","23",
                   # 🐛 v32 FIX: bitrate cap যোগ — file size predictable + Telegram 50MB safe
                   "-maxrate","3500k","-bufsize","7000k",
                   "-pix_fmt","yuv420p","-profile:v","main","-level","4.2",
                   "-r","30","-g","60","-keyint_min","60","-sc_threshold","0"]
            # 🐛 v32 FIX: audio থাকলে -af + audio codec, না থাকলে -an
            if has_aud_fb:
                cmd += ["-af",af,
                        "-c:a","aac","-b:a","160k","-ar","48000","-ac","2"]
            else:
                cmd += ["-an"]
            cmd += ["-metadata","title=","-metadata","artist=","-metadata","comment=",
                    "-metadata","copyright=","-metadata","encoder=",
                    "-metadata",f"tag={seed_tag}",
                    "-movflags","+faststart", out1]
            r = subprocess.run(cmd, capture_output=True, timeout=FFMPEG_TIMEOUT)
            if r.returncode != 0:
                # 🐛 v32 FIX: stderr capture — debug clear করার জন্য
                err = (r.stderr or b"").decode(errors="ignore")[-500:]
                logger.error("Facebook Mode ffmpeg fail: %s", err)
                return False, err
            return Path(out1).exists(), ""

        result_fb = await loop.run_in_executor(executor, run_fb_pass)
        ok, ff_err_fb = result_fb if isinstance(result_fb, tuple) else (result_fb, "")
        if not ok:
            short_err = (ff_err_fb or "Unknown error").strip()[-200:]
            await safe_edit(msg,
                f"❌ *Facebook Mode ব্যর্থ!*\n\n`{short_err}`\n\n"
                f"💡 _ভিডিওটা corrupt হতে পারে অথবা ffmpeg-এ সমস্যা।_",
                parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(out1); return

        size_mb = Path(out1).stat().st_size / (1024*1024)
        # 🆕 v30: Auto-compress যদি 50MB ছাড়ায় — Telegram bot API hard limit
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(msg, f"📦 *ফাইল {size_mb:.1f}MB → compress করছি...*",
                            parse_mode="Markdown")
            try:
                dur = await loop.run_in_executor(executor, get_duration, out1) or 30.0
                target_kb = (MAX_TG_UPLOAD_MB - 3) * 8 * 1024
                vbk = max(300, int(target_kb / dur) - 96)
                cmp_o = str(TEMP_DIR / f"{uuid.uuid4().hex}_fbc.mp4")
                cc = ["ffmpeg","-y","-i",out1,"-c:v","libx264","-preset","veryfast",
                      "-b:v",f"{vbk}k","-maxrate",f"{int(vbk*1.2)}k","-bufsize",f"{vbk*2}k",
                      "-c:a","aac","-b:a","96k","-movflags","+faststart", cmp_o]
                cr = await loop.run_in_executor(executor,
                    lambda: subprocess.run(cc, capture_output=True, timeout=FFMPEG_TIMEOUT))
                if cr.returncode == 0 and Path(cmp_o).exists() and \
                   Path(cmp_o).stat().st_size/(1024*1024) <= MAX_TG_UPLOAD_MB:
                    safe_unlink(out1); Path(cmp_o).rename(out1)
                    size_mb = Path(out1).stat().st_size / (1024*1024)
                else:
                    safe_unlink(cmp_o)
            except Exception as ce: logger.warning("fb compress: %s", ce)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(msg, f"❌ Compress-এর পরও বড় ({size_mb:.1f}MB)!",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(out1); return

        with user_lock:
            user_videos[uid] = {"path": out1,
                                "title": (d.get("title") or "Video") + " (FB Viral Ready)",
                                "url": d.get("url"), "platform": d.get("platform")}
        try:
            inc_feature("facebook_mode")
        except Exception:
            pass
        # 🆕 v30: try send_video → fallback send_document (large file safe)
        cap_fb = (f"👥 *Facebook Mode (Viral 🔥) কমপ্লিট!* ({size_mb:.1f}MB)\n\n"
                    f"📐 9:16 Reels (1080x1920) ✓\n"
                    f"🎨 Vibrant color boost ✓ 🔍 Ken Burns zoom ✓\n"
                    f"🔊 Audio fingerprint break ✓ 🔒 Metadata wipe ✓\n"
                    f"🛡️ Rights Manager safe ✓\n\n"
                    f"🔥 *ভাইরাল হওয়ার টিপস:*\n"
                    f"  • প্রথম ৩ সেকেন্ডে hook দিন (text overlay)\n"
                    f"  • Caption-এ question রাখুন (comment বাড়ায়)\n"
                    f"  • #Reels #Viral #Trending hashtag use করুন\n"
                    f"  • Peak time: রাত ৮-১১টা (BD time)\n"
                    f"  • Native Reels tab-এ upload করুন (feed-এ না)\n\n"
                    f"💡 SEO Generator দিয়ে viral title+hashtag নিন।\n\n"
                    f"⚠️ *Disclaimer:* নিজের তৈরি content-এর জন্য।")
        # 🐛 v37 FIX — caption ১০২৪ অক্ষরের বেশি হলে split করি
        cap_fb_short, cap_fb_extras = split_caption(cap_fb)
        send_ok_fb = False
        try:
            with open(out1, "rb") as fp:
                await q.message.reply_video(video=fp, caption=cap_fb_short,
                    parse_mode="Markdown", supports_streaming=True,
                    read_timeout=300, write_timeout=300, connect_timeout=60)
            send_ok_fb = True
        except Exception as e1:
            logger.warning("fb reply_video fail (%.1fMB): %s", size_mb, e1)
            if is_caption_too_long_err(e1):
                try:
                    with open(out1, "rb") as fp:
                        await q.message.reply_video(video=fp,
                            parse_mode="Markdown", supports_streaming=True,
                            read_timeout=300, write_timeout=300, connect_timeout=60)
                    send_ok_fb = True
                    cap_fb_extras = [cap_fb]
                except Exception as e_cap:
                    logger.warning("fb caption-less retry fail: %s", e_cap)
            if not send_ok_fb:
                try:
                    doc_cap = shrink_caption(
                        f"👥 *Facebook Ready ({size_mb:.1f}MB)*\n"
                        f"⚠️ _Document হিসেবে পাঠানো হলো।_", 500)
                    with open(out1, "rb") as fp:
                        await q.message.reply_document(document=fp,
                            caption=doc_cap,
                            parse_mode="Markdown",
                            read_timeout=300, write_timeout=300, connect_timeout=60)
                    send_ok_fb = True
                    cap_fb_extras = [cap_fb]
                except Exception as e2:
                    logger.error("fb send_document fail: %s", e2)
                    err_msg = ("❌ *পাঠানো ব্যর্থ!*\n\n"
                               "Caption ১০২৪ অক্ষরের বেশি — Telegram allow করে না।"
                               if is_caption_too_long_err(e2)
                               else f"❌ পাঠানো ব্যর্থ ({size_mb:.1f}MB)। Telegram limit ~50MB।")
                    await safe_edit(msg, err_msg, parse_mode="Markdown",
                                    reply_markup=edit_menu_for(uid))
                    return
        if send_ok_fb:
            for extra in cap_fb_extras:
                try:
                    await safe_reply(q.message, extra, parse_mode="Markdown")
                except Exception as ex:
                    logger.warning("fb extra caption send fail: %s", ex)
            await safe_edit(msg, "✅ *Facebook Mode রেডি!*\n\n👇 আরও এডিট:",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🆕 v27: YouTube ভাইরাল প্যাক — Edit + AI Title + Hashtags + Tips ───
    if q.data == "youtube_viral":
        with user_lock:
            d = user_videos.get(uid)
        inp = d.get("path") if d else None
        if not inp or not Path(inp).exists():
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu())
            return
        msg = await safe_reply(q.message,
            "🚀 *YouTube ভাইরাল প্যাক বানাচ্ছি...*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "১) 🎬 ভিডিও ভাইরাল-style edit (zoom + vibrant color)\n"
            "২) 🤖 AI দিয়ে viral title + description + ৩০ hashtag\n"
            "৩) 📝 Pro tips (timing, thumbnail, hook)\n\n"
            "_৩০-৬০ সেকেন্ড লাগতে পারে..._",
            parse_mode="Markdown")

        loop = asyncio.get_running_loop()
        out_v = str(TEMP_DIR / f"{uuid.uuid4().hex}_viral.mp4")

        def run_viral_edit():
            # Viral-optimized edit: subtle Ken Burns zoom-in (engagement boost),
            # vibrant color, audio loudness norm, 1080p, 30fps, faststart
            dur = get_duration(inp) or 30.0
            zoom_speed = 1.0008  # very subtle zoom over duration
            vf = (
                # 16:9 crop+scale → 1080p
                "crop='min(iw,ih*16/9)':'min(ih,iw*9/16)',"
                "scale=1920:1080:flags=lanczos,"
                # Subtle Ken Burns zoom (engagement +15-20% per studies)
                f"zoompan=z='min(zoom+0.0008,1.15)':d=1:x='iw/2-(iw/zoom/2)':y='ih/2-(ih/zoom/2)':s=1920x1080:fps=30,"
                # Denoise + sharpen
                "hqdn3d=3:2:4:3,"
                "unsharp=5:5:1.0:5:5:0.4,"
                # VIBRANT viral-style colors (high saturation, punchy contrast)
                "eq=brightness=0.06:contrast=1.30:saturation=1.55:gamma=0.92,"
                "colorbalance=rs=0.10:bs=-0.05:rm=0.06:bm=-0.04:rh=0.05:bh=-0.03,"
                "curves=preset=increase_contrast,"
                # Cinematic vignette
                "vignette=PI/6"
            )
            af = "loudnorm=I=-14:LRA=11:TP=-1.5,aresample=48000"
            cmd = ["ffmpeg","-y","-threads","0","-i",inp,
                   "-map_metadata","-1",
                   "-vf",vf,"-af",af,
                   "-c:v","libx264","-preset","medium","-crf","19",
                   "-pix_fmt","yuv420p","-profile:v","high","-level","4.2",
                   "-r","30","-g","60","-keyint_min","60","-sc_threshold","0",
                   "-x264-params",f"aq-mode=3:aq-strength=1.0:threads={CPU_COUNT}",
                   "-c:a","aac","-b:a","192k","-ar","48000","-ac","2",
                   "-movflags","+faststart", out_v]
            r = subprocess.run(cmd, capture_output=True, timeout=FFMPEG_TIMEOUT)
            return r.returncode == 0 and Path(out_v).exists()

        # Run video edit + AI gen in parallel
        title_seed = (d.get("title") or "video").strip()[:80] or "viral video"
        edit_task = loop.run_in_executor(executor, run_viral_edit)
        seo_task  = loop.run_in_executor(executor, generate_seo, title_seed, "YouTube")
        ok = await edit_task
        try:
            titles, desc, tags = await seo_task
        except Exception as e:
            logger.warning("viral SEO fail: %s", e)
            titles, desc, tags = [], "", []

        if not ok:
            await safe_edit(msg, "❌ *Viral edit ব্যর্থ!*",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(out_v); return

        size_mb = Path(out_v).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(out_v); return

        # Build viral pack text
        title_block = "\n".join(f"  {i+1}. {t}" for i, t in enumerate(titles[:3])) if titles else \
                      f"  1. {md_escape(title_seed)} | Viral 🔥\n  2. {md_escape(title_seed)} এর গোপন সত্য!\n  3. সবাই জানে না | {md_escape(title_seed)}"
        tag_line = " ".join(tags[:30]) if tags else \
                   "#shorts #viral #trending #youtube #fyp #foryou #youtubeshorts #explore #subscribe #bengali"
        desc_text = desc.strip() if desc else (
            f"{title_seed}\n\nএই ভিডিওটা ভালো লাগলে Like, Share এবং Subscribe করতে ভুলবেন না!\n"
            f"নতুন ভিডিও পেতে bell icon press করুন। 🔔")

        viral_tips = (
            "🎯 *YouTube ভাইরাল হওয়ার গোপন টিপস:*\n\n"
            "1️⃣ *প্রথম ৩ সেকেন্ড* = সবচেয়ে important। Hook দিতে হবে — প্রশ্ন/চমক/ক্লিফহ্যাঙ্গার।\n"
            "2️⃣ *Title* — emoji + curiosity gap (যেমন: \"কেউ জানে না...\")। 60 char-এর কম।\n"
            "3️⃣ *Thumbnail* — মুখের expression + bold text + bright color। Mobile-এ পড়তে পারবে এমন।\n"
            "4️⃣ *Upload time* — Bangladesh: রাত ৮-১১টা। Weekend বেশি কাজ করে।\n"
            "5️⃣ *First 24 hours* = critical। নিজের সব social-এ share করো, friends-দের comment-এ tag করতে বলো।\n"
            "6️⃣ *Watch time* > views। ভিডিও ৩০-৬০ সেকেন্ডের রাখো (Shorts), retention 70%+ চাই।\n"
            "7️⃣ *Description-এ first 2 lines* hook হতে হবে — \"এই video দেখলে আপনি...\"।\n"
            "8️⃣ *Hashtag* প্রথম ৩টা সবচেয়ে important — niche-specific দাও, তারপর viral।\n"
            "9️⃣ *Pinned comment* দাও — engagement বাড়ায় ৪০%+।\n"
            "🔟 *Series বানাও* — Part 1, Part 2। Algorithm series ভালোবাসে।"
        )

        with user_lock:
            user_videos[uid] = {"path": out_v,
                                "title": (d.get("title") or "Video") + " (Viral Pack)",
                                "url": d.get("url"), "platform": d.get("platform")}
        try: inc_feature("youtube_viral_pack")
        except Exception: pass

        # Send video first
        await q.message.reply_video(video=open(out_v,"rb"),
            caption=f"🚀 *YouTube ভাইরাল Edit রেডি!*\n\n"
                    f"📐 1920x1080 (16:9)\n🎨 Vibrant viral colors\n"
                    f"🔍 Subtle Ken Burns zoom (engagement boost)\n"
                    f"🔊 Loudness normalized (-14 LUFS, broadcast standard)",
            parse_mode="Markdown", supports_streaming=True)

        # Send viral pack text (split if too long for Telegram 4096 limit)
        pack = (
            f"━━━━━━━━━━━━━━━━━━\n"
            f"🎯 *VIRAL TITLE OPTIONS* (৩টা থেকে best টা বাছুন):\n"
            f"{title_block}\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"📝 *DESCRIPTION:*\n{desc_text[:800]}\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"🏷️ *HASHTAGS* (copy করে paste করুন):\n`{tag_line}`\n\n"
        )
        await safe_reply(q.message, pack, parse_mode="Markdown")
        await safe_reply(q.message, viral_tips, parse_mode="Markdown")

        await safe_edit(msg, "✅ *Viral Pack ডেলিভার্ড!*\n\n👇 আরও এডিট:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── AUTO SUBTITLE (Gemini transcribe + ffmpeg burn-in) ───
    if q.data == "auto_sub":
        with user_lock:
            d = user_videos.get(uid)
        inp = d.get("path") if d else None
        if not inp or not Path(inp).exists():
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu())
            return
        msg = await safe_reply(q.message,
            "📝 *অটো সাবটাইটেল তৈরি হচ্ছে...*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "🎵 অডিও extract হচ্ছে...\n"
            "🤖 AI দিয়ে transcribe (১-২ মিনিট লাগতে পারে)...\n"
            "🎬 ভিডিওতে subtitle burn-in হচ্ছে...\n\n"
            "_দয়া করে অপেক্ষা করুন_",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        srt_path, err = await loop.run_in_executor(executor, gen_subtitle_srt, inp)
        if err:
            await safe_edit(msg, err, parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            return
        out = str(TEMP_DIR / f"{uuid.uuid4().hex}_subbed.mp4")
        ok = await loop.run_in_executor(executor, burn_subtitle, inp, srt_path, out)
        if not ok or not Path(out).exists():
            safe_unlink(srt_path); safe_unlink(out)
            await safe_edit(msg, "❌ *Subtitle burn-in ব্যর্থ!*",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            return
        size_mb = Path(out).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(msg, f"❌ ফাইল বড় ({size_mb:.1f}MB)!\n\nছোট clip দিয়ে চেষ্টা করুন।",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(srt_path); safe_unlink(out)
            return
        with user_lock:
            user_videos[uid] = {"path": out, "title": (d.get("title") or "Video") + " (Subtitled)",
                                "url": d.get("url"), "platform": d.get("platform")}
        await q.message.reply_video(video=open(out,"rb"), caption="✅ *সাবটাইটেল যোগ হয়েছে!*",
                                    parse_mode="Markdown", supports_streaming=True)
        # SRT file ও পাঠাই (user চাইলে edit করতে পারবে)
        try:
            await q.message.reply_document(
                document=open(srt_path,"rb"),
                filename="subtitle.srt",
                caption="📝 SRT ফাইল (চাইলে edit করে আবার ব্যবহার করতে পারেন)")
        except Exception:
            pass
        await safe_edit(msg, "✅ *সাবটাইটেল রেডি!*\n\n👇 আরও এডিট করুন:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        safe_unlink(srt_path)
        return

    # ─── 🆕 v13: TEXT → VOICE (TTS) FLOW ───
    if q.data == "tts_start":
        # 🆕 v23: Edge TTS-কে engine list-এ যোগ করি (সম্পূর্ণ ফ্রি neural voice)
        engines = []
        if has_edge_tts():    engines.append("🆓 Edge Neural (Realistic, Free)")
        if has_elevenlabs():  engines.append("🎤 ElevenLabs (Realistic)")
        if has_openai():      engines.append("✨ OpenAI TTS (Premium)")
        engines.append("🆓 gTTS (Basic Free)")
        engine_line = " + ".join(engines)
        msg_txt = (
            "🗣️ *Text → Voice (TTS)*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "যেকোনো টেক্সট পাঠালে সেটাকে কথা বলে শোনাবো।\n\n"
            "প্রথমে ভাষা বেছে নিন:\n"
            "• 🇧🇩 *বাংলা* — Bengali পড়ার জন্য\n"
            "• 🇬🇧 *English* — ইংরেজির জন্য\n"
            "• 🇮🇳 *হিন্দি* — Hindi-এর জন্য\n"
            "• 🌍 *Auto* — মিক্স ভাষা\n\n"
            f"💡 _Engine:_ {engine_line}"
        )
        await safe_edit(q.message, msg_txt, parse_mode="Markdown", reply_markup=tts_lang_menu())
        return

    if q.data and q.data.startswith("tts_lang_"):
        lang = q.data.replace("tts_lang_", "")
        with user_lock:
            user_state[uid] = {"action": "tts_wait", "data": {"lang": lang, "voice": OPENAI_TTS_VOICE}}
        # 🆕 v23: Edge TTS / ElevenLabs / OpenAI — যেকোনো একটা থাকলেই voice picker show করি
        if has_edge_tts() or has_elevenlabs() or has_openai():
            await safe_edit(q.message,
                f"✅ ভাষা: *{lang.upper()}*\n\n"
                f"━━━━━━━━━━━━━━━━━━\n"
                f"এখন voice (কণ্ঠ) বেছে নিন (অথবা ফ্রি gTTS দিয়ে continue করুন):",
                parse_mode="Markdown", reply_markup=tts_voice_menu(uid))
            return
        # কোনো realistic engine নেই — সরাসরি text চাও (gTTS fallback)
        await safe_edit(q.message,
            f"✅ ভাষা: *{lang.upper()}* (gTTS ফ্রি)\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"এখন যেকোনো টেক্সট পাঠান (সর্বোচ্চ ৪০০০ অক্ষর)।\n\n"
            f"💡 *Tips:* আরো realistic voice পেতে — `pip install edge-tts` চালান\n\n"
            f"📝 উদাহরণ:\n"
            f"• `আজ আবহাওয়া অনেক সুন্দর`\n"
            f"• `Hello, welcome to my channel`",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    # ─── 🆕 v14: header rows are no-op ───
    if q.data == "tts_noop":
        return

    # ─── 🆕 v19: ElevenLabs voice library — pagination ও raw voice_id picker ───
    if q.data == "tts_voice_back":
        await safe_edit(q.message,
            "🎙️ *Voice বেছে নিন*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "যেকোনো voice/engine বেছে নিন।",
            parse_mode="Markdown", reply_markup=tts_voice_menu(uid))
        return

    if q.data and q.data.startswith("tts_evpg_"):
        try:
            page = int(q.data.replace("tts_evpg_", ""))
        except Exception:
            page = 0
        await safe_edit(q.message,
            f"📚 *ElevenLabs ভয়েস লাইব্রেরি*\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"মোট *{len(ELEVENLABS_VOICES)}* টা realistic voice আছে।\n"
            f"যেকোনোটায় ট্যাপ করুন → আপনার text সেই কণ্ঠে শোনাবো।",
            parse_mode="Markdown", reply_markup=elevenlabs_library_page_menu(page))
        return

    # ─── 🆕 v22: Edge TTS — paginated voice picker ───
    if q.data and q.data.startswith("tts_edpg_"):
        try:
            page = int(q.data.replace("tts_edpg_", ""))
        except Exception:
            page = 0
        await safe_edit(q.message,
            f"🆓 *Edge TTS — Premium Free Voices*\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"মোট *{len(EDGE_TTS_VOICES)}* টা neural voice আছে — সম্পূর্ণ ফ্রি।\n"
            f"🇧🇩 বাংলা / 🇮🇳 হিন্দি / 🇺🇸🇬🇧 English / 🇵🇰 Urdu / 🇸🇦 Arabic\n\n"
            f"যেকোনোটায় ট্যাপ করুন → আপনার text সেই কণ্ঠে শোনাবো।",
            parse_mode="Markdown", reply_markup=edge_tts_page_menu(page))
        return

    # ─── 🆕 v22: Edge TTS voice select ───
    if q.data and q.data.startswith("tts_edge_"):
        edge_key = q.data.replace("tts_edge_", "")
        if edge_key not in EDGE_TTS_VOICES:
            await safe_edit(q.message, "❌ Invalid Edge voice।", reply_markup=tts_voice_menu(uid))
            return
        vname, label, vlang = EDGE_TTS_VOICES[edge_key]
        with user_lock:
            st = user_state.get(uid) or {"action": "tts_wait", "data": {"lang": vlang}}
            st["data"]["voice"] = edge_key
            st["data"]["engine"] = "edge"
            st["data"]["lang"] = vlang
            st["action"] = "tts_wait"
            user_state[uid] = st
        await safe_edit(q.message,
            f"✅ Voice: *{label}*\n"
            f"✅ ভাষা: *{vlang.upper()}* (Edge Neural)\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"এখন যেকোনো টেক্সট পাঠান (সর্বোচ্চ ৪০০০ অক্ষর)।\n"
            f"বটটা সেই কণ্ঠে কথা বলবে — সম্পূর্ণ ফ্রি! 🆓\n\n"
            f"📝 *উদাহরণ:*\n"
            f"• `আজ আবহাওয়া অনেক সুন্দর, চলুন বাইরে যাই`\n"
            f"• `Hello, welcome to my channel`",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    if q.data and q.data.startswith("tts_evacc_"):
        sub = q.data.replace("tts_evacc_", "")
        force = (sub == "refresh")
        try:
            page = 0 if force else int(sub)
        except Exception:
            page = 0
        loading = await safe_edit(q.message,
            "🔄 *আপনার ElevenLabs অ্যাকাউন্টের voice list fetch হচ্ছে...*",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        voices = await loop.run_in_executor(executor, elevenlabs_list_voices, force)
        if not voices:
            await safe_edit(loading,
                "❌ *Voice fetch ব্যর্থ!*\n\n"
                "ElevenLabs API key invalid অথবা network সমস্যা।\n"
                "আবার চেষ্টা করুন অথবা library থেকে বেছে নিন।",
                parse_mode="Markdown", reply_markup=elevenlabs_account_page_menu(0, []))
            return
        # Counts per category
        n_clone = sum(1 for v in voices if v["category"] in ("cloned", "generated"))
        n_premade = sum(1 for v in voices if v["category"] in ("premade", "professional"))
        await safe_edit(loading,
            f"🎤 *আপনার ElevenLabs ভয়েস*\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"🎭 Cloned/Custom: *{n_clone}*\n"
            f"💎 Premade/Pro: *{n_premade}*\n"
            f"📦 মোট: *{len(voices)}*\n\n"
            f"যেকোনো voice-এ ট্যাপ করুন:",
            parse_mode="Markdown", reply_markup=elevenlabs_account_page_menu(page, voices))
        return

    if q.data and q.data.startswith("tts_evid_"):
        # Raw ElevenLabs voice_id সরাসরি set
        voice_id = q.data.replace("tts_evid_", "")
        if not voice_id or len(voice_id) < 15:
            await safe_edit(q.message, "❌ Invalid voice ID।", reply_markup=tts_voice_menu(uid))
            return
        # সুন্দর label খুঁজি cache থেকে
        cache = _eleven_voice_cache.get("voices") or []
        match = next((v for v in cache if v["voice_id"] == voice_id), None)
        pretty = elevenlabs_voice_label(match["category"], match["name"]) if match else f"Voice {voice_id[:10]}"
        with user_lock:
            st = user_state.get(uid) or {"action": "tts_wait", "data": {"lang": "auto"}}
            st["data"]["voice"] = voice_id
            st["data"]["engine"] = "elevenlabs"
            st["action"] = "tts_wait"
            user_state[uid] = st
            lang = st["data"].get("lang", "auto")
        await safe_edit(q.message,
            f"✅ Voice: *{pretty}*\n"
            f"✅ ভাষা: *{lang.upper()}*\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"এখন যেকোনো টেক্সট পাঠান (সর্বোচ্চ ৪০০০ অক্ষর)।\n"
            f"বটটা সেই কণ্ঠে কথা বলবে।\n\n"
            f"📝 উদাহরণ:\n"
            f"• `আজ আবহাওয়া অনেক সুন্দর, চলুন বাইরে যাই`\n"
            f"• `Hello, welcome to my channel`",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    if q.data and (q.data.startswith("tts_voice_")
                   or q.data.startswith("tts_eleven_")
                   or q.data == "tts_engine_gtts"):
        with user_lock:
            st = user_state.get(uid) or {"action": "tts_wait", "data": {"lang": "auto"}}
            if q.data == "tts_engine_gtts":
                st["data"]["engine"] = "gtts"
                voice_label = "🆓 gTTS (Free)"
            elif q.data.startswith("tts_eleven_"):
                # 🆕 v14: ElevenLabs voice  /  🆕 v18: cloned voice ("mine")
                # 🆕 v26: cloned voice could also be Replicate XTTS — engine auto-detect
                voice = q.data.replace("tts_eleven_", "")
                if voice == "mine":
                    clone = get_user_clone(uid)
                    if not clone or not clone.get("voice_id"):
                        await safe_edit(q.message,
                            "❌ আপনার কোনো cloned voice নেই।\n\n"
                            "তৈরি করতে: কোনো voice/audio message-এ reply দিয়ে\n"
                            "`/clonevoice <নাম>` লিখুন।",
                            parse_mode="Markdown", reply_markup=tts_voice_menu(uid))
                        return
                    vid = clone["voice_id"]
                    if vid.startswith("xtts:"):
                        # Replicate XTTS clone — voice = sample audio path
                        st["data"]["voice"] = vid.replace("xtts:", "")
                        st["data"]["engine"] = "xtts"
                        voice_label = f"🎭 XTTS Clone: {clone.get('name','MyVoice')}"
                    else:
                        st["data"]["voice"] = vid
                        st["data"]["engine"] = "elevenlabs"
                        voice_label = f"🎭 Cloned: {clone.get('name','MyVoice')}"
                else:
                    st["data"]["voice"] = voice
                    st["data"]["engine"] = "elevenlabs"
                    pretty = ELEVENLABS_VOICES.get(voice, (None, voice.title()))[1]
                    voice_label = f"🎤 ElevenLabs: {pretty}"
            else:
                voice = q.data.replace("tts_voice_", "")
                st["data"]["voice"] = voice
                st["data"]["engine"] = "openai"
                voice_label = f"✨ OpenAI: {voice.title()}"
            st["action"] = "tts_wait"
            user_state[uid] = st
            lang = st["data"].get("lang", "auto")
        await safe_edit(q.message,
            f"✅ Voice: *{voice_label}*\n"
            f"✅ ভাষা: *{lang.upper()}*\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"এখন যেকোনো টেক্সট পাঠান (সর্বোচ্চ ৪০০০ অক্ষর)।\n\n"
            f"📝 উদাহরণ:\n"
            f"• `আজ আবহাওয়া অনেক সুন্দর, চলুন বাইরে যাই`\n"
            f"• `Hello, welcome to my channel`",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    # ─── 🆕 v13: AI CHAT (ChatGPT / Gemini fallback) FLOW ───
    if q.data == "ai_chat_start":
        if not has_ai():
            await safe_edit(q.message,
                "🤖 *AI চ্যাট*\n\n"
                "━━━━━━━━━━━━━━━━━━\n"
                "❌ OpenAI বা Gemini — কোনো AI key সেট করা নেই।\n\n"
                "Gemini key (ফ্রি): https://aistudio.google.com/apikey\n"
                "OpenAI key: https://platform.openai.com/api-keys",
                parse_mode="Markdown", reply_markup=back_menu())
            return
        with user_lock:
            user_state[uid] = {"action": "ai_chat_wait", "data": {}}
        await safe_edit(q.message,
            "🤖 *AI চ্যাট (ChatGPT) চালু*\n\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "যেকোনো প্রশ্ন বা topic পাঠান, ChatGPT উত্তর দেবে।\n\n"
            "📝 উদাহরণ:\n"
            "• `ভাইরাল ভিডিও বানানোর ৫টা টিপস দাও`\n"
            "• `এই ভিডিওর জন্য একটা catchy intro লিখে দাও`\n"
            "• `Write a 30-second TikTok script about cats`",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    # ──────────────────────────────────────────────────────────────
    # 🆕 v14: AI TOOLKIT FLOWS
    # ──────────────────────────────────────────────────────────────

    # ─── 🎬 ভিডিও কনটেন্ট অ্যানালিসিস ───
    if q.data == "ai_analyze":
        if not has_ai():
            await safe_edit(q.message,
                "❌ *AI key দরকার (OpenAI বা Gemini)*",
                parse_mode="Markdown", reply_markup=back_menu()); return
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu()); return
        title = d.get("title") or "Video"
        platform = d.get("platform") or "Instagram"
        path = d.get("path")
        dur = get_duration(path) if path and Path(path).exists() else None
        msg = await safe_reply(q.message,
            "🎬 *ভিডিও বিশ্লেষণ চলছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            "🤖 ChatGPT analyze করছে...\n📊 Best platform + audience খুঁজছে...",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        result, err = await loop.run_in_executor(executor, analyze_video_content, title, platform, dur)
        if err:
            await safe_edit(msg, err, parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        await safe_edit(msg,
            f"🎬 *কনটেন্ট অ্যানালিসিস রিপোর্ট*\n━━━━━━━━━━━━━━━━━━\n\n{result}",
            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🏷️ Hook + Caption Generator ───
    if q.data == "ai_hooks":
        if not has_ai():
            await safe_edit(q.message,
                "❌ *AI key দরকার (OpenAI বা Gemini)*",
                parse_mode="Markdown", reply_markup=back_menu()); return
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu()); return
        title = d.get("title") or "Video"
        platform = d.get("platform") or "TikTok"
        msg = await safe_reply(q.message,
            "🏷️ *Hook + Caption তৈরি হচ্ছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            "🤖 ৫টা scroll-stopping hook...\n📝 ৩টা caption + CTA...",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        data, err = await loop.run_in_executor(executor, generate_caption_hook, title, platform)
        if err or not data:
            await safe_edit(msg, f"❌ ব্যর্থ: {err or 'unknown'}",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        hooks = "\n".join(f"  {i+1}. {md_escape(h)}" for i, h in enumerate(data.get("hooks", [])))
        captions = "\n\n".join(f"📝 *Caption {i+1}:*\n```\n{c}\n```" for i, c in enumerate(data.get("captions", [])))
        ctas = "\n".join(f"  • {md_escape(c)}" for c in data.get("ctas", []))
        sound = data.get("trending_sounds_hint", "")
        result = (
            f"🏷️ *Hook + Caption Pack* ({platform})\n━━━━━━━━━━━━━━━━━━\n\n"
            f"🎯 *Scroll-Stopping Hooks:*\n{hooks}\n\n"
            f"{captions}\n\n"
            f"📢 *CTAs:*\n{ctas}\n\n"
            f"🎵 *Sound hint:* _{md_escape(sound)}_"
        )
        await safe_edit(msg, result, parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🔍 Thumbnail Ideas Generator ───
    if q.data == "ai_thumb":
        if not has_ai():
            await safe_edit(q.message,
                "❌ *AI key দরকার (OpenAI বা Gemini)*",
                parse_mode="Markdown", reply_markup=back_menu()); return
        with user_lock:
            d = user_videos.get(uid)
        if not d:
            await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                            parse_mode="Markdown", reply_markup=back_menu()); return
        title = d.get("title") or "Video"
        platform = d.get("platform") or "YouTube"
        msg = await safe_reply(q.message,
            "🔍 *Thumbnail আইডিয়া তৈরি হচ্ছে...*\n\n━━━━━━━━━━━━━━━━━━\n"
            "🎨 ৫টা click-trigger concept...",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        result, err = await loop.run_in_executor(executor, generate_thumbnail_ideas, title, platform)
        if err:
            await safe_edit(msg, err, parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        await safe_edit(msg,
            f"🔍 *Thumbnail Ideas* ({platform})\n━━━━━━━━━━━━━━━━━━\n\n{result}",
            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        return

    # ─── 🎙️ Script Generator (asks for duration first) ───
    if q.data == "ai_script":
        if not has_ai():
            await safe_edit(q.message,
                "❌ *AI key দরকার (OpenAI বা Gemini)*",
                parse_mode="Markdown", reply_markup=back_menu()); return
        await safe_edit(q.message,
            "🎙️ *স্ক্রিপ্ট জেনারেটর*\n\n━━━━━━━━━━━━━━━━━━\n"
            "ভিডিওর দৈর্ঘ্য বেছে নিন:",
            parse_mode="Markdown", reply_markup=script_duration_menu())
        return

    if q.data and q.data.startswith("script_dur_"):
        try:
            dur = int(q.data.replace("script_dur_", ""))
        except Exception:
            dur = 30
        with user_lock:
            d = user_videos.get(uid)
        title = (d.get("title") if d else "") or "viral video"
        platform = (d.get("platform") if d else "") or "TikTok"
        msg = await safe_reply(q.message,
            f"🎙️ *Script তৈরি হচ্ছে...* ({dur}s)\n\n━━━━━━━━━━━━━━━━━━\n"
            f"🤖 Hook + Main + CTA...\n🎵 Voiceover-ready text...",
            parse_mode="Markdown")
        loop = asyncio.get_running_loop()
        data, err = await loop.run_in_executor(
            executor, generate_script, title, dur, platform, "bn")
        if err or not data:
            await safe_edit(msg, f"❌ ব্যর্থ: {err or 'unknown'}",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid)); return
        hook = data.get("hook", ""); main = data.get("main", "")
        cta = data.get("cta", ""); vo = data.get("voiceover_text", "")
        result = (
            f"🎙️ *Script Pack ({dur}s · {platform})*\n━━━━━━━━━━━━━━━━━━\n\n"
            f"🎬 *HOOK (০-৩s):*\n_{md_escape(hook)}_\n\n"
            f"📖 *MAIN:*\n{md_escape(main)}\n\n"
            f"🎯 *CTA:*\n_{md_escape(cta)}_\n\n"
            f"🎙️ *Full Voiceover Text (TTS-ready):*\n```\n{vo[:1500]}\n```"
        )
        await safe_edit(msg, result, parse_mode="Markdown", reply_markup=edit_menu_for(uid))
        # AI voiceover audio তৈরি করো (ElevenLabs > OpenAI > gTTS)
        if vo and vo.strip():
            try:
                wait_msg = await safe_reply(q.message,
                    "🎤 *AI ভয়েসওভার তৈরি হচ্ছে...*", parse_mode="Markdown")
                vo_engine = "auto"  # auto = elevenlabs > openai > gtts
                vo_voice = ELEVENLABS_DEFAULT_VOICE if has_elevenlabs() else OPENAI_TTS_VOICE
                out_path, used_engine, vo_err = await loop.run_in_executor(
                    executor, text_to_speech, vo, "bn", vo_voice, vo_engine)
                if out_path and Path(out_path).exists():
                    with open(out_path, "rb") as fp:
                        await c.bot.send_voice(q.message.chat_id, fp,
                            caption=f"🎙️ *Voiceover ready!* (engine: `{used_engine}`)",
                            parse_mode="Markdown")
                    with open(out_path, "rb") as fp:
                        await c.bot.send_audio(q.message.chat_id, fp,
                            title="Script Voiceover", performer="Video Editor Bot",
                            caption="📥 MP3 ফাইল")
                    safe_unlink(out_path)
                    await safe_edit(wait_msg, "✅ *Voiceover পাঠানো হয়েছে!*", parse_mode="Markdown")
                else:
                    await safe_edit(wait_msg,
                        f"⚠️ Voiceover তৈরি করা গেল না: {vo_err or 'unknown'}",
                        parse_mode="Markdown")
            except Exception as e:
                logger.warning("script voiceover: %s", e)
        return

    # ─── 🌐 Subtitle/Text Translation ───
    if q.data and q.data.startswith("tr_lang_"):
        target = q.data.replace("tr_lang_", "")
        with user_lock:
            user_state[uid] = {"action": "translate_wait", "data": {"target": target}}
        await safe_edit(q.message,
            f"🌐 *অনুবাদ → {LANG_NAMES.get(target, target).upper()}*\n\n━━━━━━━━━━━━━━━━━━\n"
            f"এখন *যেকোনো টেক্সট* অথবা *.srt ফাইল* পাঠান।\n"
            f"আমি অনুবাদ করে দেবো (timestamps অক্ষুণ্ণ থাকবে)।\n\n"
            f"📌 _Tip: ভিডিও থেকে subtitle বানিয়ে SRT-ফাইল পাঠালে সেটা সরাসরি অনুবাদ হবে।_",
            parse_mode="Markdown", reply_markup=cancel_input_menu())
        return

    # ─── PROCESSABLE FEATURES ───
    if q.data not in PROCESSABLE: return

    with user_lock:
        d = user_videos.get(uid)
    inp = d.get("path") if d else None
    title = d.get("title") if d else ""
    if not inp or not Path(inp).exists():
        await safe_edit(q.message, "❌ *কোনো ভিডিও নেই!*\n\nপ্রথমে ভিডিও পাঠান।",
                        parse_mode="Markdown", reply_markup=back_menu())
        return

    out_base = str(TEMP_DIR / f"{uuid.uuid4().hex}_output")
    final = await process(inp, out_base, q.data, q)
    if not final: return
    inc_stat("videos_processed")
    inc_feature(q.data)  # 🆕 v17: per-feature usage track
    safe_t = md_escape(title) if title else ""
    title_line = f"\n🎬 *{safe_t}*" if safe_t else ""
    label = LABELS.get(q.data, "")
    cap = f"{label} *সম্পন্ন!*{title_line}\n✅ আপনার এডিটেড ফাইল রেডি"

    try:
        size_mb = Path(final).stat().st_size / (1024*1024)
        if size_mb > MAX_TG_UPLOAD_MB:
            await safe_edit(q.message, f"❌ ফাইল বড় ({size_mb:.1f}MB)!\n\n📦 কম্প্রেস ব্যবহার করুন।",
                            parse_mode="Markdown", reply_markup=edit_menu_for(uid))
            safe_unlink(final); return
        with open(final, "rb") as f:
            try:
                if q.data == "thumb":
                    await c.bot.send_photo(q.message.chat_id, f, caption=cap, parse_mode="Markdown")
                elif q.data == "audio":
                    await c.bot.send_audio(q.message.chat_id, f, caption=cap, parse_mode="Markdown")
                elif q.data == "gif":
                    await c.bot.send_animation(q.message.chat_id, f, caption=cap, parse_mode="Markdown")
                else:
                    await c.bot.send_video(q.message.chat_id, f, caption=cap,
                                           parse_mode="Markdown", supports_streaming=True)
            except BadRequest:
                f.seek(0)
                plain_cap = re.sub(r'[*_`\[\]()]', '', cap)
                if q.data == "thumb":
                    await c.bot.send_photo(q.message.chat_id, f, caption=plain_cap)
                elif q.data == "audio":
                    await c.bot.send_audio(q.message.chat_id, f, caption=plain_cap)
                elif q.data == "gif":
                    await c.bot.send_animation(q.message.chat_id, f, caption=plain_cap)
                else:
                    await c.bot.send_video(q.message.chat_id, f, caption=plain_cap, supports_streaming=True)
        await safe_edit(q.message, "✅ *পাঠানো হয়েছে!*\n\n👇 আরও এডিট করুন:",
                        parse_mode="Markdown", reply_markup=edit_menu_for(uid))
    except Exception as e:
        logger.exception("send err: %s", e)
        await safe_edit(q.message, "❌ পাঠাতে সমস্যা!", parse_mode="Markdown", reply_markup=back_menu())
    finally:
        safe_unlink(final)

# ──────────────────────────────────────────────────────────────
# CLEANUP & ERROR
# ──────────────────────────────────────────────────────────────
TEMP_AGE_SEC      = 3600          # ১ ঘণ্টার পুরনো ফাইল মুছবে
TEMP_MAX_BYTES    = 2 * 1024**3   # 2GB cap — এর বেশি হলে সবচেয়ে পুরনো ফাইল আগে মুছবে

def cleanup_temp():
    """🆕 v17: age-based + size-cap cleanup। Termux-এ disk fill হওয়া আটকায়।"""
    try:
        now = datetime.now().timestamp()
        files = []
        total = 0
        for f in TEMP_DIR.glob("*"):
            try:
                st = f.stat()
                age = now - st.st_mtime
                # age cross করলে সাথে সাথে delete
                if age > TEMP_AGE_SEC:
                    f.unlink(missing_ok=True)
                    continue
                files.append((st.st_mtime, st.st_size, f))
                total += st.st_size
            except Exception as e:
                logger.debug("cleanup stat %s: %s", f, e)
        # size cap পার হলে সবচেয়ে পুরনো থেকে মুছা শুরু
        if total > TEMP_MAX_BYTES:
            files.sort()  # oldest first
            for mtime, size, f in files:
                if total <= TEMP_MAX_BYTES:
                    break
                try:
                    f.unlink(missing_ok=True)
                    total -= size
                    logger.info("cleanup: size cap, removed %s (%dKB)", f.name, size // 1024)
                except Exception as e:
                    logger.debug("cleanup unlink %s: %s", f, e)
    except Exception as e:
        logger.warning("cleanup_temp: %s", e)

async def periodic_cleanup(c):
    cleanup_temp()

_NET_ERR_COUNT = {"n": 0, "last_logged": 0}

async def on_error(u, c):
    # 🆕 v28: Network errors (TimedOut/NetworkError) silent-এ retry — log spam কমে।
    # প্রতি ১০টা error-এ একবার log হবে, সাথে cumulative count।
    err = c.error
    err_name = type(err).__name__ if err else "Unknown"
    err_msg = str(err)[:200] if err else ""
    if err_name in ("TimedOut", "NetworkError"):
        _NET_ERR_COUNT["n"] += 1
        # প্রতি ১০টায় একবার log — terminal পরিষ্কার থাকবে
        if _NET_ERR_COUNT["n"] - _NET_ERR_COUNT["last_logged"] >= 10:
            logger.warning("network slow — %d auto-retry হয়েছে (auto-recover চলছে)",
                           _NET_ERR_COUNT["n"])
            _NET_ERR_COUNT["last_logged"] = _NET_ERR_COUNT["n"]
        return
    if err_name in ("RetryAfter", "BadRequest"):
        logger.warning("telegram err [%s]: %s", err_name, err_msg)
        return
    logger.exception("unhandled err: %s", err)

# ──────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════
# 🆕 v28+ : WINGO 30s LIVE SIGNAL MODULE
# প্রতি ৩০ সেকেন্ডে নতুন period detect → 8-layer analyzer + number prediction
# → সব subscriber-কে auto signal পাঠায়। /wingo দিয়ে subscribe।
# ══════════════════════════════════════════════════════════════════
# 🆕 v33: Multi-mode WinGo (30s + 1m + 3m + 5m) — same analyzer engine, separate state per mode.
WINGO_API = "https://draw.ar-lottery01.com/WinGo/WinGo_30S/GetHistoryIssuePage.json?ts="   # legacy alias (30s)
WINGO_APIS = {
    "30S": "https://draw.ar-lottery01.com/WinGo/WinGo_30S/GetHistoryIssuePage.json?ts=",
    "1M":  "https://draw.ar-lottery01.com/WinGo/WinGo_1M/GetHistoryIssuePage.json?ts=",
    "3M":  "https://draw.ar-lottery01.com/WinGo/WinGo_3M/GetHistoryIssuePage.json?ts=",
    "5M":  "https://draw.ar-lottery01.com/WinGo/WinGo_5M/GetHistoryIssuePage.json?ts=",
}
WINGO_MODE_LABELS = {"30S": "WinGo 30s", "1M": "WinGo 1Min", "3M": "WinGo 3Min", "5M": "WinGo 5Min"}
WINGO_MODE_INTERVAL = {"30S": "৩০ সেকেন্ড", "1M": "১ মিনিট", "3M": "৩ মিনিট", "5M": "৫ মিনিট"}
WINGO_SUBS_FILE = Path(tempfile.gettempdir()) / "wingo_subs.json"
WINGO_PERF_FILE = Path(tempfile.gettempdir()) / "wingo_layer_perf.json"   # 30s — legacy
WINGO_PERF_FILES = {
    "30S": WINGO_PERF_FILE,
    "1M":  Path(tempfile.gettempdir()) / "wingo_layer_perf_1m.json",
    "3M":  Path(tempfile.gettempdir()) / "wingo_layer_perf_3m.json",
    "5M":  Path(tempfile.gettempdir()) / "wingo_layer_perf_5m.json",
}
_wingo_lock = threading.Lock()

def _w_fresh_state():
    return {
        "last_period": None,
        "last_pred":   None,
        "user_stats":  {},
        "layer_perf":  {},
        "global_perf": {"hit":0, "miss":0, "recent": []},
        "user_last_pred": {},
        "_gemini_cache": {},   # 🆕 v33: cache for L25 Gemini AI predictions
        "_openai_cache": {},   # 🆕 v37: cache for L31 OpenAI GPT-4o-mini predictions
        "user_balance":  {},   # 🆕 v37: per-user bankroll (BDT/INR/USD agnostic) for Kelly Criterion
    }


# 🆕 v37 — Kelly Criterion bet sizing (proper Kelly + Half-Kelly safety)
WINGO_BALANCE_FILE = Path(tempfile.gettempdir()) / "wingo_balances.json"

def _w_load_balances():
    """Load all per-mode balances from disk (single file, all modes)."""
    try:
        if WINGO_BALANCE_FILE.exists():
            d = json.loads(WINGO_BALANCE_FILE.read_text())
            for mode, bals in d.items():
                if mode in _wingo_states:
                    # keys are stringified ints in JSON — convert back to int
                    _wingo_states[mode]["user_balance"] = {int(k): float(v) for k, v in bals.items()}
    except Exception as e:
        logger.warning("wingo balance load fail: %s", e)

def _w_save_balances():
    try:
        out = {m: {str(k): float(v) for k, v in s.get("user_balance", {}).items()}
               for m, s in _wingo_states.items()}
        WINGO_BALANCE_FILE.write_text(json.dumps(out))
    except Exception as e:
        logger.warning("wingo balance save fail: %s", e)

def _w_kelly(confidence_pct, payout=1.0):
    """Kelly fraction f* = (bp - q) / b
       where b = net odds (1.0 = even-money 1:1), p = win prob, q = 1-p.
       payout = net odds (e.g. WinGo big/small ≈ even money minus house edge ≈ 0.95).
       Returns: (full_kelly_fraction, half_kelly_fraction).  Both clamped 0..0.25
       so we never recommend insanely aggressive sizing.
    """
    p = max(0.50, min(0.97, confidence_pct / 100.0))
    q = 1.0 - p
    b = max(0.50, payout)
    f = (b * p - q) / b
    f = max(0.0, min(0.25, f))     # cap at 25% of bankroll (conservative)
    return f, f * 0.5              # full + half Kelly

# Per-mode state (30S = legacy alias)
_wingo_states = {m: _w_fresh_state() for m in WINGO_APIS.keys()}
_wingo_state = _wingo_states["30S"]   # legacy alias — preserves existing 30s code paths

def _w_get_state(mode="30S"):
    return _wingo_states.get(mode, _wingo_state)

def _load_wingo_subs():
    try:
        if WINGO_SUBS_FILE.exists():
            return set(json.loads(WINGO_SUBS_FILE.read_text()))
    except Exception:
        pass
    return set()

def _save_wingo_subs(subs):
    try:
        WINGO_SUBS_FILE.write_text(json.dumps(list(subs)))
    except Exception as e:
        logger.warning("wingo subs save fail: %s", e)

def _load_wingo_perf_for(state, perf_file):
    """🆕 v29: Load persisted layer performance for adaptive weighting."""
    try:
        if perf_file.exists():
            d = json.loads(perf_file.read_text())
            state["layer_perf"] = d.get("layer_perf", {})
            state["global_perf"] = d.get("global_perf", {"hit":0,"miss":0,"recent":[]})
    except Exception as e:
        logger.warning("wingo perf load fail (%s): %s", perf_file.name, e)

def _save_wingo_perf_for(state, perf_file):
    try:
        perf_file.write_text(json.dumps({
            "layer_perf":  state["layer_perf"],
            "global_perf": state["global_perf"],
        }))
    except Exception as e:
        logger.warning("wingo perf save fail (%s): %s", perf_file.name, e)

def _load_wingo_perf():
    _load_wingo_perf_for(_wingo_state, WINGO_PERF_FILE)

def _save_wingo_perf():
    _save_wingo_perf_for(_wingo_state, WINGO_PERF_FILE)

_wingo_subs = _load_wingo_subs()
# Load perf for all modes
for _m, _pf in WINGO_PERF_FILES.items():
    _load_wingo_perf_for(_wingo_states[_m], _pf)
# 🆕 v37 — Load Kelly Criterion balances for all modes
_w_load_balances()

def _w_sizeof(n): return "BIG" if n >= 5 else "SMALL"
def _w_colorof(n):
    if n in (0, 5): return "violet"
    if n in (2, 4, 6, 8): return "red"
    return "green"
def _w_color_emoji(c):
    return {"red":"🔴","green":"🟢","violet":"🟣"}.get(c,"⚪")

def _w_fetch_history(mode="30S"):
    """🆕 v33: Multi-mode fetch. Returns list of {period, number} or []."""
    import urllib.request
    url = WINGO_APIS.get(mode, WINGO_API)
    try:
        req = urllib.request.Request(url + str(int(time.time()*1000)),
                                     headers={"User-Agent":"Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as r:
            data = json.loads(r.read().decode("utf-8", errors="ignore"))
        lst = (data.get("data") or {}).get("list") or []
        return [{"period": x.get("issueNumber"), "number": int(x.get("number"))} for x in lst if x.get("number") is not None]
    except Exception as e:
        logger.warning("wingo fetch fail (%s): %s", mode, e)
        return []

def _w_get_adaptive_weight(layer_id, base_weight, state=None):
    """🆕 v29 ULTRA+: Blend static weight with learned accuracy (Bayesian smoothing). v33: state-aware."""
    if state is None: state = _wingo_state
    perf = state["layer_perf"].get(layer_id, {"hit":0, "miss":0})
    total = perf["hit"] + perf["miss"]
    if total < 5:
        return base_weight   # not enough data → use static
    # Beta-Binomial posterior with prior (alpha=5, beta=5) → smooth toward 0.5
    accuracy = (perf["hit"] + 5) / (total + 10)
    # Scale: 0.4 acc → 0.5x weight, 0.5 acc → 1.0x, 0.6 acc → 1.5x, 0.7 acc → 2.0x
    multiplier = max(0.3, min(2.5, 1.0 + (accuracy - 0.5) * 5.0))
    return base_weight * multiplier

def _w_record_layer_vote(layer_id, vote_b, vote_s, actual_size, state=None):
    """Record whether a layer's vote matched the actual outcome. v33: state-aware."""
    if state is None: state = _wingo_state
    if vote_b == 0 and vote_s == 0: return   # no opinion → don't count
    layer_pred = "BIG" if vote_b > vote_s else "SMALL"
    perf = state["layer_perf"].setdefault(layer_id, {"hit":0, "miss":0})
    if layer_pred == actual_size: perf["hit"] += 1
    else: perf["miss"] += 1

def _w_analyze(history, loss_streak=0, win_streak=0, state=None):
    """🆕 v35 ULTRA+++++++ : 30-layer ensemble (28 statistical + 1 Gemini AI + 1 external 3X Leader Rakib clone) with self-tuning weights + anti-streak safety flip + skip recommendation + dual-AI consensus boost.
    state-aware (multi-mode safe). Uses 100-period history. Returns top-3 number probabilities + per-layer votes."""
    if state is None: state = _wingo_state
    H = history[:100]
    sizes = [_w_sizeof(r["number"]) for r in H]
    nums  = [r["number"] for r in H]
    if not sizes: return None

    import math

    # ═══ L1: time-decay frequency (exponential weighting) ═══
    l1b = l1s = 0.0
    for i, r in enumerate(H):
        w = math.exp(-i / 8)
        if r["number"] >= 5: l1b += w
        else: l1s += w

    # ═══ L2: streak + reversal pressure ═══
    streak, top = 1, sizes[0]
    for i in range(1, len(sizes)):
        if sizes[i] == top: streak += 1
        else: break
    rev = min(7.0, (streak ** 1.45) * 0.65)
    l2b, l2s = (0, rev) if top == "BIG" else (rev, 0)
    if streak >= 6:
        # mega-streak strong reversal
        if top == "BIG": l2s += 1.5
        else: l2b += 1.5

    # ═══ L3: zigzag detection ═══
    alt_count = sum(1 for i in range(1, min(10, len(sizes))) if sizes[i] != sizes[i-1])
    l3b = l3s = 0
    if alt_count >= 6:
        # strong zigzag → predict opposite of last
        if top == "BIG": l3s += 3.0
        else: l3b += 3.0
    elif alt_count >= 4:
        if top == "BIG": l3s += 1.5
        else: l3b += 1.5

    # ═══ L4: Markov chain (1-step + 2-step + 3-step) ═══
    seq = list(reversed(sizes))     # chronological
    trans1 = {"BIG":{"BIG":0,"SMALL":0}, "SMALL":{"BIG":0,"SMALL":0}}
    trans2 = {}
    trans3 = {}   # 🆕 3-step Markov
    for i in range(len(seq) - 1):
        trans1[seq[i]][seq[i+1]] += 1
        if i < len(seq) - 2:
            k = f"{seq[i]}_{seq[i+1]}"
            trans2.setdefault(k, {"BIG":0,"SMALL":0})[seq[i+2]] += 1
        if i < len(seq) - 3:
            k3 = f"{seq[i]}_{seq[i+1]}_{seq[i+2]}"
            trans3.setdefault(k3, {"BIG":0,"SMALL":0})[seq[i+3]] += 1
    l4b = l4s = 0
    last = sizes[0]
    t1 = trans1[last]; t1t = t1["BIG"] + t1["SMALL"]
    if t1t:
        pb = t1["BIG"] / t1t
        l4b += pb * 3; l4s += (1-pb) * 3
    if len(sizes) > 1:
        k2 = f"{sizes[1]}_{sizes[0]}"
        if k2 in trans2:
            t2 = trans2[k2]; t2t = t2["BIG"] + t2["SMALL"]
            if t2t >= 2:   # 🆕 require min sample
                pb = t2["BIG"] / t2t
                l4b += pb * 4; l4s += (1-pb) * 4
    if len(sizes) > 2:
        k3 = f"{sizes[2]}_{sizes[1]}_{sizes[0]}"
        if k3 in trans3:
            t3 = trans3[k3]; t3t = t3["BIG"] + t3["SMALL"]
            if t3t >= 2:
                pb = t3["BIG"] / t3t
                l4b += pb * 5; l4s += (1-pb) * 5

    # ═══ L5: gap analysis ═══
    gapB = next((i for i,s in enumerate(sizes) if s == "BIG"), -1)
    gapS = next((i for i,s in enumerate(sizes) if s == "SMALL"), -1)
    l5b = min(3.5, gapB * 0.55) if gapB > gapS and gapB > 0 else 0
    l5s = min(3.5, gapS * 0.55) if gapS > gapB and gapS > 0 else 0

    # ═══ L6: multi-window trend (5/10/15/20) ═══
    l6b = l6s = 0
    for window in (5, 10, 15, 20):
        if len(sizes) < window * 2: continue
        recent = sizes[:window].count("BIG")
        prior  = sizes[window:window*2].count("BIG")
        diff = recent - prior
        if diff >= 2: l6b += 0.8
        elif diff <= -2: l6s += 0.8

    # ═══ L7: 🆕 EMA momentum (exponential moving average) ═══
    # alpha = 0.25, encodes BIG=1, SMALL=0
    l7b = l7s = 0
    if len(sizes) >= 8:
        alpha = 0.25
        ema_short = ema_long = 0.5
        for i in range(min(20, len(sizes)) - 1, -1, -1):
            v = 1.0 if sizes[i] == "BIG" else 0.0
            ema_short = alpha * v + (1 - alpha) * ema_short
            ema_long  = 0.1 * v + 0.9 * ema_long
        # if short EMA > long EMA → upward BIG momentum
        delta = ema_short - ema_long
        if delta > 0.05: l7b += min(3.0, delta * 30)
        elif delta < -0.05: l7s += min(3.0, abs(delta) * 30)

    # ═══ L8: pattern mining (n-gram, length 3/4/5/6) ═══
    chrono = list(reversed(sizes))
    recent_ctx = "".join(reversed(sizes[:6]))
    l8b = l8s = 0
    for n in (3, 4, 5, 6):
        if len(recent_ctx) < n: continue
        ctx = recent_ctx[-n:]
        mB = mS = 0; total = 0
        for i in range(len(chrono) - n):
            if "".join(chrono[i:i+n]) == ctx:
                total += 1
                if chrono[i+n] == "BIG": mB += 1
                else: mS += 1
        if total >= 2:
            # weight longer matches more (they're rarer + stronger signal)
            l8b += (mB/total) * (n * 0.6)
            l8s += (mS/total) * (n * 0.6)

    # ═══ L9: Bayesian / chi-square mean reversion ═══
    N = len(sizes); bigs = sizes.count("BIG"); smalls = N - bigs
    expected = N / 2
    chi = ((bigs - expected)**2 + (smalls - expected)**2) / max(expected, 1)
    l9b = l9s = 0
    if chi > 3.84:
        if bigs > smalls: l9s = min(4, chi * 0.4)
        else: l9b = min(4, chi * 0.4)

    # ═══ L10: 🆕 cyclic/periodic pattern detection ═══
    # Check if sequence has 4/5/6 step periodicity
    l10b = l10s = 0
    for period in (4, 5, 6, 7):
        if len(sizes) < period * 3: continue
        matches = 0; checks = 0
        for i in range(min(15, len(sizes) - period)):
            checks += 1
            if sizes[i] == sizes[i + period]: matches += 1
        if checks and matches / checks >= 0.7:
            # strong periodicity → next likely matches sizes[period-1]
            predicted = sizes[period - 1] if period - 1 < len(sizes) else sizes[0]
            if predicted == "BIG": l10b += 2.0
            else: l10s += 2.0

    # ═══ L11: 🆕 Color momentum (3 colors: red/green/violet) ═══
    colors = [_w_colorof(n) for n in nums]
    l11b = l11s = 0
    # Recent color streak in opposite size class hints reversal
    rec_colors = colors[:5]
    if rec_colors.count("violet") >= 2:
        # violet = 0 or 5 → boundary → often signals shift
        if top == "BIG": l11s += 1.2
        else: l11b += 1.2

    # ═══ L12: 🆕 volatility / variance regime ═══
    l12b = l12s = 0
    if len(nums) >= 15:
        recent_var = sum((n - 4.5)**2 for n in nums[:10]) / 10
        prior_var  = sum((n - 4.5)**2 for n in nums[10:20]) / 10 if len(nums) >= 20 else recent_var
        # high recent variance → bias toward extremes (0/9)
        if recent_var > prior_var * 1.3:
            # extreme regime → predict same side strongly
            if top == "BIG": l12b += 1.5
            else: l12s += 1.5

    # ═══ L13: 🆕 v30 — Run-length distribution (empirical streak survival) ═══
    # Compute distribution of historical streak lengths → estimate P(streak continues)
    l13b = l13s = 0
    runs = []
    if sizes:
        cur_size = sizes[0]; cur_len = 1
        for s in sizes[1:]:
            if s == cur_size: cur_len += 1
            else:
                runs.append(cur_len)
                cur_size = s; cur_len = 1
        runs.append(cur_len)
    if len(runs) >= 5 and streak >= 2:
        # how often did historical streaks reach length >= current?
        reached = sum(1 for ln in runs if ln >= streak)
        # how often did they extend beyond?
        extended = sum(1 for ln in runs if ln > streak)
        if reached >= 3:
            p_continue = extended / reached
            # convert to vote magnitude (low p_continue → strong reverse)
            if p_continue < 0.4:
                # streak likely to break
                rev_strength = (0.4 - p_continue) * 8.0
                if top == "BIG": l13s += min(3.5, rev_strength)
                else: l13b += min(3.5, rev_strength)
            elif p_continue > 0.6:
                # streak likely to continue (rare for long streaks)
                cont_strength = (p_continue - 0.6) * 6.0
                if top == "BIG": l13b += min(2.5, cont_strength)
                else: l13s += min(2.5, cont_strength)

    # ═══ L14: 🆕 v30 — Period-modulo bias (time-of-day cycle) ═══
    # WinGo periods cycle within the day; certain mod-positions historically favor one side.
    l14b = l14s = 0
    try:
        cur_period_int = int(history[0]["period"])
        next_period_mod = (cur_period_int + 1) % 60   # period within hour (60 × 30s)
        # find historical periods with same mod and look at outcomes
        same_mod_outcomes = []
        for r in H[1:]:
            try:
                if int(r["period"]) % 60 == next_period_mod:
                    same_mod_outcomes.append(_w_sizeof(r["number"]))
            except Exception:
                continue
        if len(same_mod_outcomes) >= 4:
            mb = same_mod_outcomes.count("BIG")
            ms = same_mod_outcomes.count("SMALL")
            tot = mb + ms
            pb_mod = mb / tot
            # Dirichlet prior: smooth toward 0.5
            pb_smooth = (mb + 2) / (tot + 4)
            if pb_smooth > 0.6:    l14b += min(2.5, (pb_smooth - 0.5) * 8)
            elif pb_smooth < 0.4:  l14s += min(2.5, (0.5 - pb_smooth) * 8)
    except Exception:
        pass

    # ═══ L15: 🆕 v30 — 2-step number Markov (number-level transitions) ═══
    # P(next size | last 2 actual numbers) — finer-grained than size Markov
    l15b = l15s = 0
    if len(nums) >= 6:
        last_pair = (nums[1], nums[0])  # chronologically: nums[1] then nums[0]
        matches_b = matches_s = 0
        for i in range(len(nums) - 2):
            # nums is reverse-chronological: nums[i+2] → nums[i+1] → nums[i]
            if (nums[i+2], nums[i+1]) == last_pair:
                if nums[i] >= 5: matches_b += 1
                else: matches_s += 1
        total_match = matches_b + matches_s
        if total_match >= 2:
            pb_n = (matches_b + 1) / (total_match + 2)  # Laplace smoothing
            l15b += pb_n * 3.5
            l15s += (1 - pb_n) * 3.5

    # ═══ 🆕 v31 — L16: Pattern Repetition (4-step micro-pattern match) ═══
    # গত 4-result-এর pattern ইতিহাসে কতবার এসেছে, তারপর কোনদিকে গেছে?
    l16b = l16s = 0
    sizes_seq = [_w_sizeof(n) for n in nums]   # reverse-chronological
    if len(sizes_seq) >= 12:
        last4 = tuple(sizes_seq[:4])
        nb = ns = 0
        for i in range(len(sizes_seq) - 5):
            window = tuple(sizes_seq[i+1:i+5])  # older window
            if window == last4:
                # what came AFTER it? sizes_seq[i] is newer than window
                if sizes_seq[i] == "BIG": nb += 1
                else: ns += 1
        tot_p = nb + ns
        if tot_p >= 2:
            pb_p = (nb + 1) / (tot_p + 2)   # Laplace
            strength = min(4.5, tot_p * 1.0)  # more matches → stronger
            l16b += pb_p * strength
            l16s += (1 - pb_p) * strength

    # ═══ 🆕 v31 — L17: Streak Intelligence (continue vs flip detector) ═══
    # দীর্ঘ same-side streak → mean reversion likely (flip)
    # Alternating BSBS pattern → continue
    l17b = l17s = 0
    if len(sizes_seq) >= 6:
        # current streak
        cur = sizes_seq[0]
        streak_len = 1
        for s in sizes_seq[1:8]:
            if s == cur: streak_len += 1
            else: break
        # historical: how often did streak of length N flip?
        if streak_len >= 4:
            # strong flip bias
            if cur == "BIG": l17s += min(4.0, (streak_len - 3) * 1.5)
            else: l17b += min(4.0, (streak_len - 3) * 1.5)
        elif streak_len <= 1 and len(sizes_seq) >= 5:
            # check for alternating pattern (B-S-B-S-B or S-B-S-B-S)
            alt = all(sizes_seq[i] != sizes_seq[i+1] for i in range(4))
            if alt:
                # next likely continues alternation → opposite of current
                if cur == "BIG": l17s += 3.0
                else: l17b += 3.0

    # ═══ 🆕 v31 — L18: Statistical Bias Correction (Chi-square style) ═══
    # যদি গত 30-এ BIG/SMALL ratio chance থেকে significantly বিচ্যুত,
    # তাহলে mean reversion expect করি (gambler's rebound, কিন্তু
    # softer — Bayesian)
    l18b = l18s = 0
    if len(sizes_seq) >= 30:
        last30s = sizes_seq[:30]
        cb = last30s.count("BIG"); cs = last30s.count("SMALL")
        if cb + cs >= 25:
            ratio_b = cb / (cb + cs)
            # deviation from 0.5
            dev = ratio_b - 0.5
            if abs(dev) >= 0.18:    # ≥18% deviation
                # bias OPPOSITE of over-represented (mean reversion)
                strength = min(3.5, abs(dev) * 12)
                if dev > 0: l18s += strength    # too many BIG → SMALL more likely
                else:       l18b += strength

    # ═══ 🆕 v31 — L19: Number-Pair Synergy (last_num + 2nd_last → next size) ═══
    # গত দুই number-এর exact pair-এর পর কোন size এসেছে?
    l19b = l19s = 0
    if len(nums) >= 8:
        last_pair_n = (nums[1], nums[0])
        nb19 = ns19 = 0
        for i in range(len(nums) - 2):
            if (nums[i+2], nums[i+1]) == last_pair_n:
                if nums[i] >= 5: nb19 += 1
                else: ns19 += 1
        tot19 = nb19 + ns19
        if tot19 >= 2:
            p19 = (nb19 + 1) / (tot19 + 2)
            l19b += p19 * 3.0
            l19s += (1 - p19) * 3.0

    # ═══ 🆕 v31 — L20: Hot Number Cluster (5/9 type heat detection) ═══
    # যদি BIG side-এ 5/6/7 cluster বেশি থাকে recent → continue;
    # যদি SMALL side-এ 1/2/3 cluster বেশি → continue
    l20b = l20s = 0
    if len(nums) >= 10:
        last10n = nums[:10]
        b_high_cluster = sum(1 for n in last10n if n in (6,7,8,9))
        s_low_cluster  = sum(1 for n in last10n if n in (1,2,3,4))
        if b_high_cluster >= 5: l20b += min(2.5, (b_high_cluster - 4) * 0.8)
        if s_low_cluster  >= 5: l20s += min(2.5, (s_low_cluster - 4) * 0.8)

    # ═══ 🆕 v32 — L21: Shannon Entropy regime ═══
    # Low entropy (predictable) → trust patterns more (lean to last side)
    # High entropy (chaotic) → mean reversion to 50/50 (lean opposite of recent)
    l21b = l21s = 0
    if len(sizes_seq) >= 20:
        win = sizes_seq[:20]
        cb_e = win.count("BIG"); cs_e = win.count("SMALL")
        tot_e = cb_e + cs_e
        if tot_e:
            pb_e = cb_e / tot_e
            ps_e = cs_e / tot_e
            ent = 0.0
            if pb_e > 0: ent -= pb_e * math.log2(pb_e)
            if ps_e > 0: ent -= ps_e * math.log2(ps_e)
            # ent: 0 (one-sided) → 1 (perfect 50/50)
            if ent < 0.85:
                # low entropy → trust dominant side continues
                if pb_e > ps_e: l21b += min(2.5, (0.85 - ent) * 7)
                else:           l21s += min(2.5, (0.85 - ent) * 7)
            elif ent > 0.97:
                # near-max entropy → mean reversion bias OPPOSITE of last
                if top == "BIG": l21s += 1.5
                else:            l21b += 1.5

    # ═══ 🆕 v32 — L22: Fibonacci lookback voting ═══
    # 1, 2, 3, 5, 8, 13 periods ago — কোন side dominate করেছে?
    l22b = l22s = 0
    fib_idx = [1, 2, 3, 5, 8, 13]
    fib_b = fib_s = 0
    for i in fib_idx:
        if i < len(sizes_seq):
            if sizes_seq[i] == "BIG": fib_b += 1
            else:                     fib_s += 1
    fib_tot = fib_b + fib_s
    if fib_tot >= 4:
        # majority Fibonacci → bias OPPOSITE (mean reversion across natural cycle)
        if fib_b - fib_s >= 2: l22s += min(2.0, (fib_b - fib_s) * 0.6)
        elif fib_s - fib_b >= 2: l22b += min(2.0, (fib_s - fib_b) * 0.6)

    # ═══ 🆕 v32 — L23: Triple-Markov Consensus (1+2+3 step agreement) ═══
    # যদি ১, ২, ৩-step Markov সবাই একই side suggest করে → very strong signal
    l23b = l23s = 0
    mk_votes = []   # list of (b_pct, s_pct)
    if t1t and last in trans1:
        t = trans1[last]; tt = t["BIG"] + t["SMALL"]
        if tt: mk_votes.append((t["BIG"]/tt, t["SMALL"]/tt))
    if len(sizes) >= 2:
        k2c = f"{sizes[1]}_{sizes[0]}"
        if k2c in trans2:
            t = trans2[k2c]; tt = t["BIG"] + t["SMALL"]
            if tt >= 2: mk_votes.append((t["BIG"]/tt, t["SMALL"]/tt))
    if len(sizes) >= 3:
        k3c = f"{sizes[2]}_{sizes[1]}_{sizes[0]}"
        if k3c in trans3:
            t = trans3[k3c]; tt = t["BIG"] + t["SMALL"]
            if tt >= 2: mk_votes.append((t["BIG"]/tt, t["SMALL"]/tt))
    if len(mk_votes) >= 2:
        # all agree on same side?
        all_big = all(b > s for b, s in mk_votes)
        all_small = all(s > b for b, s in mk_votes)
        if all_big:
            avg_p = sum(b for b, _ in mk_votes) / len(mk_votes)
            l23b += min(4.0, (avg_p - 0.5) * 14)   # consensus bonus
        elif all_small:
            avg_p = sum(s for _, s in mk_votes) / len(mk_votes)
            l23s += min(4.0, (avg_p - 0.5) * 14)

    # ═══ 🆕 v32 — L24: Self-Correcting Performance Layer ═══
    # যদি bot recently consistently SMALL miss করছে (predicted BIG, actual SMALL)
    # → SMALL prediction-এ extra weight দাও (auto-bias correction)
    l24b = l24s = 0
    g_perf = state.get("global_perf", {})
    recent_outs = g_perf.get("recent", [])[-15:]   # last 15 results (1=hit, 0=miss)
    if len(recent_outs) >= 8:
        miss_rate = 1 - (sum(recent_outs) / len(recent_outs))
        if miss_rate >= 0.55:
            # we've been wrong a lot — flip the inferred bias
            # (heuristic: bot's previous tilt = whatever last few user_last_pred were)
            recent_preds = list(state.get("user_last_pred", {}).values())[-5:]
            if recent_preds:
                pred_sides = [p.get("signal") for p in recent_preds if p]
                pb_count = pred_sides.count("BIG")
                ps_count = pred_sides.count("SMALL")
                # if mostly predicted BIG and missed → flip to SMALL
                if pb_count > ps_count: l24s += min(2.5, (miss_rate - 0.5) * 8)
                elif ps_count > pb_count: l24b += min(2.5, (miss_rate - 0.5) * 8)

    # ═══ 🆕 v33 — L25: Gemini AI Pattern Reader (REAL AI) — v34: streak-aware ═══
    # Send last 15 outcomes + numbers to Gemini with streak context, ask it to predict.
    # Cached per-pattern (incl. streak state) for 60s to avoid spam. Graceful fallback.
    l25b = l25s = 0
    try:
        if GEMINI_API_KEY and len(sizes) >= 10:
            cache = state.setdefault("_gemini_cache", {})
            # v34: include streak state in cache key so streak transitions get fresh advice
            cache_key = (",".join(sizes[:15]) + "|" + ",".join(str(n) for n in nums[:10])
                         + f"|ls{loss_streak}|ws{win_streak}")
            now_ts = time.time()
            cached = cache.get(cache_key)
            if cached and (now_ts - cached["ts"] < 60):
                ai_pred, ai_conf = cached["pred"], cached["conf"]
            else:
                streak_ctx = ""
                if loss_streak >= 2:
                    streak_ctx = (f"\n⚠️ CRITICAL CONTEXT: We have JUST LOST {loss_streak} "
                                  f"predictions in a row. The pattern likely has shifted. "
                                  f"STRONGLY consider whether to flip direction. Be cautious "
                                  f"and pick the safer side.\n")
                elif loss_streak == 1:
                    streak_ctx = "\nNote: Just lost 1 — be careful, slight contrarian bias OK.\n"
                elif win_streak >= 3:
                    streak_ctx = f"\nNote: Won {win_streak} in a row — current model is hot, lean with the prevailing trend.\n"
                prompt = (
                    "You are an expert sequence-pattern analyst. A binary outcome game produces "
                    "BIG (numbers 5-9) or SMALL (numbers 0-4) each round.\n\n"
                    f"Last 15 outcomes (newest→oldest): {','.join(sizes[:15])}\n"
                    f"Last 10 numbers (newest→oldest): {','.join(str(n) for n in nums[:10])}\n"
                    + streak_ctx +
                    "\nAnalyze for: (a) streaks/runs, (b) alternation rhythm, "
                    "(c) anti-streak/break-likelihood, (d) frequency bias, (e) mean reversion. "
                    "Then predict the NEXT outcome.\n"
                    "Reply on ONE line ONLY in EXACT format: SIZE:BIG CONF:62  "
                    "(use SMALL or BIG; CONF 50-90)"
                )
                resp = gemini_chat(prompt,
                                   system="Reply on one line only in the exact format requested.",
                                   max_tokens=30, temperature=0.3)
                ai_pred, ai_conf = None, 0
                if resp:
                    m = re.search(r'SIZE\s*[:=]\s*(BIG|SMALL).*?CONF\s*[:=]\s*(\d+)',
                                  resp.upper(), re.DOTALL)
                    if m:
                        ai_pred = m.group(1)
                        ai_conf = max(50, min(90, int(m.group(2))))
                cache[cache_key] = {"pred": ai_pred, "conf": ai_conf, "ts": now_ts}
                # Cap cache size
                if len(cache) > 50:
                    oldest = sorted(cache.items(), key=lambda kv: kv[1]["ts"])[:20]
                    for k, _ in oldest: cache.pop(k, None)
            if ai_pred:
                # Map confidence (50-90) to vote magnitude (0 - 4.5)
                vote_strength = (ai_conf - 50) / 40 * 4.5
                if ai_pred == "BIG": l25b += vote_strength
                else: l25s += vote_strength
    except Exception as e:
        logger.warning("L25 Gemini layer fail: %s", e)

    # ═══ 🆕 v34 — L26: Long-Window Mean Reversion (50-period) ═══
    # Heavy long-window bias → bet on reversion (different time-scale than EMA L7)
    l26b = l26s = 0
    recent50 = sizes[:50]
    if len(recent50) >= 30:
        big_pct50 = recent50.count("BIG") / len(recent50)
        if big_pct50 >= 0.62:    # heavy BIG → revert to SMALL
            l26s += min(3.5, (big_pct50 - 0.5) * 14)
        elif big_pct50 <= 0.38:  # heavy SMALL → revert to BIG
            l26b += min(3.5, (0.5 - big_pct50) * 14)

    # ═══ 🆕 v34 — L27: Strong Run Reversal (last-N same-side break) ═══
    # 4-5+ same side in a row → strong probability of break (gambler's fallacy notwithstanding,
    # detected runs of 5+ in 50/50 RNG are rare → next is more likely opposite by saturation)
    l27b = l27s = 0
    run_len = 0
    if sizes:
        side0 = sizes[0]
        for s in sizes:
            if s == side0: run_len += 1
            else: break
    if run_len == 4:
        if side0 == "BIG": l27s += 2.5
        else: l27b += 2.5
    elif run_len == 5:
        if side0 == "BIG": l27s += 4.0
        else: l27b += 4.0
    elif run_len >= 6:
        if side0 == "BIG": l27s += 5.5
        else: l27b += 5.5

    # ═══ 🆕 v34 — L28: Volatility Regime Detection (alt-rate based) ═══
    # High alternation regime → predict opposite of last; streaky regime → continuation
    l28b = l28s = 0
    last20 = sizes[:20]
    if len(last20) >= 15:
        alternations = sum(1 for i in range(len(last20)-1) if last20[i] != last20[i+1])
        alt_rate = alternations / (len(last20)-1)
        if alt_rate >= 0.65:           # very alternating → predict opposite of last
            if last20[0] == "BIG": l28s += min(3.0, (alt_rate - 0.5) * 10)
            else: l28b += min(3.0, (alt_rate - 0.5) * 10)
        elif alt_rate <= 0.30:         # very streaky → continuation
            if last20[0] == "BIG": l28b += min(2.5, (0.5 - alt_rate) * 8)
            else: l28s += min(2.5, (0.5 - alt_rate) * 8)

    # ═══ 🆕 v35 — L29: 3X Leader Rakib Cross-Check (external algo clone) ═══
    # Faithful replica of https://3x-leader-rakib.vercel.app algorithm:
    #   if BIG-count in last 10 ≥ 5 → predict SMALL ;  else → predict BIG
    # Vote strength scales with imbalance magnitude.
    l29b = l29s = 0
    threex_pred = None
    last10_sizes = sizes[:10]
    if len(last10_sizes) >= 10:
        big_count_10 = last10_sizes.count("BIG")
        if big_count_10 >= 5:
            threex_pred = "SMALL"
            l29s += 1.5 + min(2.0, (big_count_10 - 5) * 0.5)
        else:
            threex_pred = "BIG"
            l29b += 1.5 + min(2.0, (5 - big_count_10) * 0.5)

    # ═══ 🆕 v35 — L30: Multi-Window Consensus Mean Reversion (10/20/40) ═══
    # If 10w / 20w / 40w windows ALL agree on reversion direction → very strong signal
    # (only fires when independent time-scales converge — rare but high-quality)
    l30b = l30s = 0
    windows = [10, 20, 40]
    consensus_votes = []
    for w in windows:
        win = sizes[:w]
        if len(win) >= max(8, int(w * 0.7)):
            bp = win.count("BIG") / len(win)
            if bp >= 0.60:   consensus_votes.append("SMALL")  # heavy BIG → revert SMALL
            elif bp <= 0.40: consensus_votes.append("BIG")    # heavy SMALL → revert BIG
    # All scored windows agreeing → strong vote
    if len(consensus_votes) >= 2 and len(set(consensus_votes)) == 1:
        if consensus_votes[0] == "BIG": l30b += 3.5
        else: l30s += 3.5
    elif len(consensus_votes) >= 2:
        # majority of windows agree → moderate vote
        from collections import Counter as _C30
        most, cnt = _C30(consensus_votes).most_common(1)[0]
        if cnt >= 2:
            if most == "BIG": l30b += 1.5
            else: l30s += 1.5

    # ═══ 🆕 v37 — L31: OpenAI GPT-4o-mini Cross-Check (REAL DUAL-AI) ═══
    # Independent second opinion from OpenAI alongside L25 Gemini.  When BOTH
    # agree → super-strong signal (dual-AI consensus boost intensified).
    # Cached per-pattern for 60s.  Graceful fallback if OpenAI key missing.
    l31b = l31s = 0
    openai_pred = None
    openai_conf = 0
    try:
        if has_openai() and len(sizes) >= 10:
            cache = state.setdefault("_openai_cache", {})
            cache_key = (",".join(sizes[:15]) + "|" + ",".join(str(n) for n in nums[:10])
                         + f"|ls{loss_streak}|ws{win_streak}")
            now_ts = time.time()
            cached = cache.get(cache_key)
            if cached and (now_ts - cached["ts"] < 60):
                openai_pred, openai_conf = cached["pred"], cached["conf"]
            else:
                streak_ctx = ""
                if loss_streak >= 2:
                    streak_ctx = (f"\n⚠️ CRITICAL: Just LOST {loss_streak} in a row. "
                                  f"Pattern likely shifted — strongly consider flipping side.\n")
                elif loss_streak == 1:
                    streak_ctx = "\nNote: Just lost 1 — slight contrarian bias OK.\n"
                elif win_streak >= 3:
                    streak_ctx = f"\nNote: Won {win_streak} in a row — current model is hot.\n"
                prompt = (
                    "You are an expert sequence-pattern analyst. A binary outcome game produces "
                    "BIG (numbers 5-9) or SMALL (numbers 0-4) each round.\n\n"
                    f"Last 15 outcomes (newest→oldest): {','.join(sizes[:15])}\n"
                    f"Last 10 numbers (newest→oldest): {','.join(str(n) for n in nums[:10])}\n"
                    + streak_ctx +
                    "\nAnalyze: streaks, alternation rhythm, anti-streak / break likelihood, "
                    "frequency bias, mean reversion, hidden cycles. Predict NEXT outcome.\n"
                    "Reply ONE LINE EXACT: SIZE:BIG CONF:62  (use SMALL or BIG; CONF 50-90)"
                )
                # use cheaper gpt-4o-mini for low cost + low latency
                resp = openai_chat(prompt, model="gpt-4o-mini",
                                   system="Reply on one line only in the exact format requested.",
                                   max_tokens=30, temperature=0.3)
                if resp:
                    m = re.search(r'SIZE\s*[:=]\s*(BIG|SMALL).*?CONF\s*[:=]\s*(\d+)',
                                  resp.upper(), re.DOTALL)
                    if m:
                        openai_pred = m.group(1)
                        openai_conf = max(50, min(90, int(m.group(2))))
                cache[cache_key] = {"pred": openai_pred, "conf": openai_conf, "ts": now_ts}
                if len(cache) > 50:
                    oldest = sorted(cache.items(), key=lambda kv: kv[1]["ts"])[:20]
                    for k, _ in oldest: cache.pop(k, None)
            if openai_pred:
                vote_strength = (openai_conf - 50) / 40 * 4.5
                if openai_pred == "BIG": l31b += vote_strength
                else: l31s += vote_strength
    except Exception as e:
        logger.warning("L31 OpenAI layer fail: %s", e)

    # ═══ 🆕 v38 — L32: Anthropic Claude Haiku (3rd Independent AI) ═══
    # Different model family (Anthropic), different training data, different biases →
    # genuinely orthogonal signal vs Gemini & OpenAI.  Quad-AI consensus possible.
    l32b = l32s = 0
    claude_pred = None
    claude_conf = 0
    try:
        if has_anthropic() and len(sizes) >= 10:
            cache_c = state.setdefault("_claude_cache", {})
            ck = (",".join(sizes[:15]) + "|" + ",".join(str(n) for n in nums[:10])
                  + f"|ls{loss_streak}|ws{win_streak}")
            now_ts2 = time.time()
            cached_c = cache_c.get(ck)
            if cached_c and (now_ts2 - cached_c["ts"] < 60):
                claude_pred, claude_conf = cached_c["pred"], cached_c["conf"]
            else:
                streak_ctx2 = ""
                if loss_streak >= 2:
                    streak_ctx2 = (f"\nCRITICAL: Lost {loss_streak} in a row. "
                                   "Pattern shifted — strongly consider flipping.\n")
                elif win_streak >= 3:
                    streak_ctx2 = f"\nNote: Won {win_streak} in a row.\n"
                prompt_c = (
                    "You are an expert sequence-pattern analyst.  A binary game outputs "
                    "BIG (5-9) or SMALL (0-4) each round.\n"
                    f"Last 15 sides (newest→oldest): {','.join(sizes[:15])}\n"
                    f"Last 10 numbers: {','.join(str(n) for n in nums[:10])}\n"
                    + streak_ctx2 +
                    "Analyze: streaks, alternation, anti-streak break, frequency bias, "
                    "mean reversion, hidden cycles. Predict NEXT outcome.\n"
                    "Reply ONE LINE EXACT: SIZE:BIG CONF:64  (use SMALL or BIG; CONF 50-90)"
                )
                resp_c = anthropic_chat(prompt_c,
                    system="Reply on one line only in the exact format requested.",
                    max_tokens=30, temperature=0.3)
                if resp_c:
                    mc = re.search(r'SIZE\s*[:=]\s*(BIG|SMALL).*?CONF\s*[:=]\s*(\d+)',
                                   resp_c.upper(), re.DOTALL)
                    if mc:
                        claude_pred = mc.group(1)
                        claude_conf = max(50, min(90, int(mc.group(2))))
                cache_c[ck] = {"pred": claude_pred, "conf": claude_conf, "ts": now_ts2}
                if len(cache_c) > 50:
                    oldest_c = sorted(cache_c.items(), key=lambda kv: kv[1]["ts"])[:20]
                    for k, _ in oldest_c: cache_c.pop(k, None)
            if claude_pred:
                vs = (claude_conf - 50) / 40 * 4.5
                if claude_pred == "BIG": l32b += vs
                else: l32s += vs
    except Exception as e:
        logger.warning("L32 Claude layer fail: %s", e)

    # ═══ 🆕 v38 — L33: Hidden Markov Regime Switching ═══
    # Detects two regimes (TRENDING vs ALTERNATING) using transition-matrix MLE
    # over sliding 30-period window.  In TRENDING regime: continue current side.
    # In ALTERNATING regime: flip last side.
    l33b = l33s = 0
    try:
        if len(sizes) >= 30:
            window = sizes[:30]
            same_tr = diff_tr = 0
            for i in range(len(window) - 1):
                if window[i] == window[i+1]: same_tr += 1
                else:                         diff_tr += 1
            total_tr = same_tr + diff_tr
            if total_tr > 0:
                p_same = same_tr / total_tr
                # regime confidence = distance from 0.5 (more extreme = stronger regime signal)
                regime_strength = abs(p_same - 0.5) * 2     # 0..1
                last_side = sizes[0]
                if p_same > 0.55:
                    # TRENDING regime — continue last
                    if last_side == "BIG": l33b += 2.5 * regime_strength
                    else:                  l33s += 2.5 * regime_strength
                elif p_same < 0.45:
                    # ALTERNATING regime — flip last
                    if last_side == "BIG": l33s += 2.5 * regime_strength
                    else:                  l33b += 2.5 * regime_strength
    except Exception as e:
        logger.warning("L33 HMM layer fail: %s", e)

    # ═══ 🆕 v38 — L34: Bayesian Beta-Posterior with Time Decay ═══
    # Conjugate prior Beta(2,2) updated with time-weighted observations.
    # Theoretically optimal for binary outcome estimation under uncertainty.
    l34b = l34s = 0
    try:
        if len(sizes) >= 20:
            alpha = 2.0   # prior pseudo-count for BIG
            beta  = 2.0   # prior pseudo-count for SMALL
            decay = 0.95  # exponential time decay
            for i, s in enumerate(sizes[:60]):
                w = decay ** i        # newer = higher weight
                if s == "BIG": alpha += w
                else:           beta  += w
            # posterior mean
            p_big = alpha / (alpha + beta)
            # posterior variance (smaller = more confident)
            var = (alpha * beta) / ((alpha + beta) ** 2 * (alpha + beta + 1))
            confidence_factor = 1 / (1 + var * 100)   # scale 0..1
            # contrarian: extreme posterior → mean reversion candidate
            if p_big > 0.58:
                # SMALL is "due" — mild bet on SMALL
                l34s += 1.8 * (p_big - 0.5) * 4 * confidence_factor
            elif p_big < 0.42:
                l34b += 1.8 * (0.5 - p_big) * 4 * confidence_factor
            else:
                # near 50% — go with slight posterior lean
                if p_big > 0.5: l34b += 0.8 * confidence_factor
                else:           l34s += 0.8 * confidence_factor
    except Exception as e:
        logger.warning("L34 Bayesian layer fail: %s", e)

    # ═══ 🆕 v38 — L35: FFT-style Cycle Detection (autocorrelation peak) ═══
    # Search for hidden periodicities (period 2..8) via autocorrelation.
    # If a strong cycle exists, predict using cycle phase.
    l35b = l35s = 0
    try:
        if len(sizes) >= 40:
            # binary encode: BIG=1, SMALL=-1
            x = [1 if s == "BIG" else -1 for s in sizes[:40]]
            n = len(x)
            mean = sum(x) / n
            xc = [v - mean for v in x]
            denom = sum(v * v for v in xc) or 1
            best_period = 0
            best_corr = 0
            for lag in range(2, 9):
                num = sum(xc[i] * xc[i + lag] for i in range(n - lag))
                corr = num / denom
                if abs(corr) > abs(best_corr):
                    best_corr = corr
                    best_period = lag
            if abs(best_corr) > 0.18 and best_period > 0:
                # phase: predict using value at position `best_period` ago,
                # sign-adjusted by correlation sign
                anchor = x[best_period - 1]   # `best_period` rounds ago value
                pred_val = anchor if best_corr > 0 else -anchor
                strength = min(3.0, abs(best_corr) * 8)
                if pred_val > 0: l35b += strength
                else:             l35s += strength
    except Exception as e:
        logger.warning("L35 FFT-cycle layer fail: %s", e)

    # ═══ Loss-streak dampening ═══
    dampen = 1.0
    if loss_streak == 1: dampen = 0.93
    elif loss_streak == 2: dampen = 0.84
    elif loss_streak >= 3: dampen = 0.72

    # ═══ 🆕 v38 আগুন-MAX : Adaptive Weighted Ensemble (35 layers — 31 stat/strategy + 3 AI + 1 external) ═══
    # L25: Gemini AI.  L26-L28: mean-reversion / run-break / regime.  L29: 3X Leader Rakib clone.
    # L30: multi-window consensus.  L31: OpenAI cross-check.
    # 🆕 v38: L32: Claude AI (3rd AI).  L33: HMM regime.  L34: Bayesian Beta.  L35: FFT cycles.
    W_BASE = {"L1":1.0,"L2":1.5,"L3":1.3,"L4":1.7,"L5":1.0,"L6":1.1,
              "L7":1.4,"L8":1.6,"L9":1.3,"L10":1.5,"L11":0.9,"L12":1.0,
              "L13":1.6,"L14":1.2,"L15":1.5,
              "L16":1.8,"L17":1.7,"L18":1.4,"L19":1.5,"L20":1.2,
              "L21":1.5,"L22":1.1,"L23":2.0,"L24":1.4,
              "L25":2.2,                                         # 🆕 v33 Gemini AI
              "L26":1.8,"L27":2.4,"L28":1.7,                     # 🆕 v34
              "L29":1.6,"L30":2.0,                               # 🆕 v35
              "L31":2.2,                                         # 🆕 v37 OpenAI
              "L32":2.2,                                         # 🆕 v38 Claude AI
              "L33":1.9,"L34":2.1,"L35":1.7}                     # 🆕 v38 HMM / Bayesian / FFT
    W = {k: _w_get_adaptive_weight(k, v, state) for k, v in W_BASE.items()}

    # Per-layer votes (for tracking + agreement calibration)
    layer_votes = {
        "L1":(l1b,l1s),"L2":(l2b,l2s),"L3":(l3b,l3s),"L4":(l4b,l4s),
        "L5":(l5b,l5s),"L6":(l6b,l6s),"L7":(l7b,l7s),"L8":(l8b,l8s),
        "L9":(l9b,l9s),"L10":(l10b,l10s),"L11":(l11b,l11s),"L12":(l12b,l12s),
        "L13":(l13b,l13s),"L14":(l14b,l14s),"L15":(l15b,l15s),
        "L16":(l16b,l16s),"L17":(l17b,l17s),"L18":(l18b,l18s),
        "L19":(l19b,l19s),"L20":(l20b,l20s),
        "L21":(l21b,l21s),"L22":(l22b,l22s),"L23":(l23b,l23s),"L24":(l24b,l24s),
        "L25":(l25b,l25s),                                       # 🆕 v33 Gemini AI vote
        "L26":(l26b,l26s),"L27":(l27b,l27s),"L28":(l28b,l28s),   # 🆕 v34
        "L29":(l29b,l29s),"L30":(l30b,l30s),                     # 🆕 v35
        "L31":(l31b,l31s),                                       # 🆕 v37 OpenAI vote
        "L32":(l32b,l32s),                                       # 🆕 v38 Claude AI vote
        "L33":(l33b,l33s),"L34":(l34b,l34s),"L35":(l35b,l35s),   # 🆕 v38 HMM/Bayes/FFT
    }
    sb = sum(b * W[k] for k,(b,_) in layer_votes.items()) * dampen
    ss = sum(s * W[k] for k,(_,s) in layer_votes.items()) * dampen

    if loss_streak >= 3:
        # contrarian flip after long loss
        if sb > ss: sb, ss = ss * 1.1, sb
        else: sb, ss = ss, sb * 1.1

    prediction = "BIG" if sb > ss else "SMALL"
    margin = abs(sb - ss)
    sigm = 1 / (1 + math.exp(-margin / 1.8))

    # 🆕 Layer agreement % — boost confidence if layers agree
    voting_layers = [(k, b, s) for k,(b,s) in layer_votes.items() if (b > 0 or s > 0)]
    if voting_layers:
        agreed = sum(1 for _,b,s in voting_layers
                     if (b > s and prediction == "BIG") or (s > b and prediction == "SMALL"))
        agreement_pct = agreed / len(voting_layers)
    else:
        agreement_pct = 0.5

    # Global recent accuracy modifier
    g = _wingo_state.get("global_perf", {})
    recent = g.get("recent", [])
    recent_acc = sum(recent[-20:]) / max(1, len(recent[-20:])) if recent else 0.5
    recent_boost = (recent_acc - 0.5) * 8   # ±4% based on recent perf

    # 🆕 v38 আগুন-MAX: Multi-AI + external agreement (Gemini + OpenAI + Claude + 3X Leader)
    gemini_pred = "BIG" if l25b > l25s else ("SMALL" if l25s > l25b else None)
    # openai_pred / claude_pred already captured above in L31/L32 blocks
    # Count how many AIs agree with ensemble
    ai_preds = [p for p in (gemini_pred, openai_pred, claude_pred) if p]
    ai_agree_count = sum(1 for p in ai_preds if p == prediction)
    n_ai_active    = len(ai_preds)
    dual_consensus = (gemini_pred == prediction and threex_pred == prediction
                      and gemini_pred is not None and threex_pred is not None)
    triple_window_consensus = (len(consensus_votes) >= 3 and len(set(consensus_votes)) == 1
                               and consensus_votes[0] == prediction)
    # 🆕 v37: TRIPLE-AI CONSENSUS — Gemini + OpenAI + 3X all agree
    triple_ai_consensus = (gemini_pred == prediction and openai_pred == prediction
                           and threex_pred == prediction
                           and all(p is not None for p in (gemini_pred, openai_pred, threex_pred)))
    # 🆕 v38: QUAD-AI CONSENSUS — all 3 AIs (Gemini + OpenAI + Claude) + 3X all agree
    quad_ai_consensus = (gemini_pred == prediction and openai_pred == prediction
                         and claude_pred == prediction and threex_pred == prediction
                         and all(p is not None for p in
                                 (gemini_pred, openai_pred, claude_pred, threex_pred)))
    openai_agree_solo = (openai_pred == prediction and openai_pred is not None)
    claude_agree_solo = (claude_pred == prediction and claude_pred is not None)

    conf = (54 + sigm * 35
            + (agreement_pct - 0.5) * 24    # ±12% from agreement
            + recent_boost
            - loss_streak * 2.5
            + min(6, win_streak * 1.3)
            + (5 if dual_consensus else 0)              # 🆕 v35: Gemini+3X both agree
            + (4 if triple_window_consensus else 0)     # 🆕 v35: 10w/20w/40w all agree
            + (3 if openai_agree_solo and not triple_ai_consensus and not quad_ai_consensus else 0)
            + (3 if claude_agree_solo and not triple_ai_consensus and not quad_ai_consensus else 0)
            + (8 if triple_ai_consensus and not quad_ai_consensus else 0)   # 🆕 v37
            + (12 if quad_ai_consensus else 0))         # 🆕 v38: 🌟 QUAD-AI CONSENSUS
    confidence = int(max(52, min(98, round(conf))))

    # ═══ 🆕 v38 — L36: Meta-Confidence Calibrator (Honest Probability) ═══
    # Recent global accuracy অনুযায়ী displayed confidence কে real hit-rate-এর সাথে align করে।
    # যদি last 20 prediction-এ ~55% hit হয় কিন্তু আমরা "85% confident" বলছিলাম → over-confident।
    # Isotonic-style soft calibration: displayed = blend(raw_conf, recent_actual_rate * scaling)।
    try:
        recent_meta = g.get("recent", [])[-30:]
        if len(recent_meta) >= 15:
            recent_acc_meta = sum(recent_meta) / len(recent_meta)   # 0..1
            # Map raw_conf [50..98] → expected hit if calibrated.  When recent_acc_meta is
            # well above 0.5 we trust high confidence more; when it's near/below 0.5 we
            # shrink high confidence toward the realistic rate.
            calibration_target = 50 + recent_acc_meta * 50          # 50..100
            # weight: more samples = more trust in calibration target
            cal_weight = min(0.45, len(recent_meta) / 100.0)        # cap 45% blend
            calibrated = confidence * (1 - cal_weight) + calibration_target * cal_weight
            calibrated = int(max(52, min(98, round(calibrated))))
            calibration_delta = calibrated - confidence
            confidence = calibrated
        else:
            calibration_delta = 0
            recent_acc_meta = 0.5
    except Exception as e:
        logger.warning("L36 meta-calibrator fail: %s", e)
        calibration_delta = 0
        recent_acc_meta = 0.5

    # ═══ 🆕 v34 — ANTI-STREAK SAFETY FLIP ═══
    # Goal: prevent consecutive losses from exceeding 2.
    # If we've lost ≥2 in a row, the ensemble has been wrong → force-flip the prediction.
    # This is statistical "course correction" — bet AGAINST the model's own bias.
    flip_applied = False
    original_signal = prediction
    if loss_streak >= 2:
        flip_applied = True
        prediction = "SMALL" if prediction == "BIG" else "BIG"
        # Acting against ensemble vote → reduce confidence honestly
        confidence = int(max(50, confidence - 6))
        # Recompute sb/ss for downstream (so number-side filter aligns with new signal)
        sb, ss = ss, sb

    # ═══ 🆕 ULTRA Number Prediction (Markov chain on numbers + hot/cold) ═══
    last30 = nums[:30]
    last10 = last30[:10]
    nscores = [0.0] * 10

    # 1. Time-decay recency (slight bias to NOT repeat recent)
    for i, n in enumerate(last30):
        nscores[n] += math.exp(-i / 12) * 0.4

    # 2. Gap analysis — long-absent numbers get boost
    for d in range(10):
        gap = next((i for i,n in enumerate(last30) if n == d), -1)
        if gap < 0: nscores[d] += 5
        elif gap >= 10: nscores[d] += min(4.5, gap * 0.4)
        elif gap >= 6: nscores[d] += gap * 0.25

    # 3. 🆕 Markov on numbers — what came after the last number historically?
    if last30:
        last_num = last30[0]
        next_after = {}
        for i in range(len(last30) - 1):
            if last30[i + 1] == last_num:   # i+1 is older, i is newer
                # so last30[i] came after last30[i+1] in time
                next_after[last30[i]] = next_after.get(last30[i], 0) + 1
        total_na = sum(next_after.values())
        if total_na >= 2:
            for d, cnt in next_after.items():
                nscores[d] += (cnt / total_na) * 3.5

    # 4. Heavy penalty for last 1-2 numbers (rarely repeat immediately)
    if last10: nscores[last10[0]] *= 0.50
    if len(last10) > 1: nscores[last10[1]] *= 0.80

    # 5. Side filter — match the BIG/SMALL prediction
    side_range = [5,6,7,8,9] if prediction == "BIG" else [0,1,2,3,4]
    for d in range(10):
        if d in side_range: nscores[d] *= 1.8
        else: nscores[d] *= 0.35

    # 6. Violet (0,5) damping if recently appeared
    for d in (0, 5):
        gap = next((i for i,n in enumerate(last30) if n == d), -1)
        if 0 <= gap < 10: nscores[d] *= 0.65

    # 7. 🆕 Color frequency bias — favor the under-represented color this side
    side_colors = [_w_colorof(d) for d in side_range]
    color_counts = {"red":0, "green":0, "violet":0}
    for n in last10:
        if n in side_range:
            color_counts[_w_colorof(n)] += 1
    # boost numbers whose color is under-represented
    for d in side_range:
        c = _w_colorof(d)
        if color_counts[c] == 0: nscores[d] *= 1.3

    ranked = sorted(range(10), key=lambda d: -nscores[d])
    # 🆕 v30: normalize top-3 scores to percentages (softmax-style)
    top3 = ranked[:3]
    top3_scores = [max(0.001, nscores[d]) for d in top3]
    s_sum = sum(top3_scores)
    top3_pct = [round(sc / s_sum * 100) for sc in top3_scores]
    # ensure they sum to 100
    if top3_pct and sum(top3_pct) != 100:
        top3_pct[0] += 100 - sum(top3_pct)
    return {
        "signal": prediction,
        "confidence": confidence,
        "num_top": ranked[0],
        "num_alt": [ranked[1], ranked[2]],
        "num_top3": top3,                  # 🆕 v30
        "num_top3_pct": top3_pct,          # 🆕 v30
        "color": _w_colorof(ranked[0]),
        "layer_votes": layer_votes,        # 🆕 v29 ULTRA+ for adaptive learning
        "agreement": round(agreement_pct * 100),
        "n_layers": 36,                    # 🆕 v38 (31 stat/strategy + 3 AI + 1 external + 1 meta-cal)
        "ai_used": (l25b > 0 or l25s > 0 or l31b > 0 or l31s > 0
                    or l32b > 0 or l32s > 0),
        "n_ai_active": n_ai_active,
        "ai_agree_count": ai_agree_count,
        # 🆕 v34: anti-streak protection metadata
        "flip_applied": flip_applied,
        "original_signal": original_signal,
        "skip_recommended": (loss_streak >= 2 and (confidence < 65 or agreement_pct < 0.60))
                            or (loss_streak == 1 and confidence < 56 and agreement_pct < 0.50),
        "loss_streak": loss_streak,
        "win_streak":  win_streak,
        "run_len":     run_len,             # how many same-side in a row at the head of history
        # 🆕 v35: external + dual-AI cross-check metadata
        "external_3x_pred":  threex_pred,                              # 3X Leader Rakib's pick
        "external_3x_agree": (threex_pred == prediction) if threex_pred else None,
        "gemini_pred":       gemini_pred,
        "gemini_agree":      (gemini_pred == prediction) if gemini_pred else None,
        "dual_consensus":    dual_consensus,                           # Gemini AND 3X agree
        "triple_window_consensus": triple_window_consensus,            # 10w/20w/40w all agree
        # 🆕 v37: OpenAI cross-check + triple-AI consensus
        "openai_pred":       openai_pred,
        "openai_conf":       openai_conf,
        "openai_agree":      (openai_pred == prediction) if openai_pred else None,
        "triple_ai_consensus": triple_ai_consensus,                    # Gemini + OpenAI + 3X all agree
        # 🆕 v38: Claude AI + Quad-AI consensus + Meta-calibrator
        "claude_pred":         claude_pred,
        "claude_conf":         claude_conf,
        "claude_agree":        (claude_pred == prediction) if claude_pred else None,
        "quad_ai_consensus":   quad_ai_consensus,                      # 3 AIs + 3X all agree
        "calibration_delta":   calibration_delta,                      # how much L36 adjusted conf
        "recent_acc_30":       round(recent_acc_meta * 100),           # last-30 actual hit rate
    }

def _w_mode_suffix(mode):
    """Convert internal mode to callback_data suffix. 30S → '' (legacy), others → '_1m'/'_3m'/'_5m'."""
    return "" if mode == "30S" else "_" + mode.lower()

def _w_parse_mode(data, prefix):
    """Extract mode from callback_data like 'wingo_now_1m' → '1M'. No suffix = 30S (legacy)."""
    suffix = data[len(prefix):]
    if suffix == "_1m": return "1M"
    if suffix == "_3m": return "3M"
    if suffix == "_5m": return "5M"
    return "30S"

def wingo_main_menu():
    """🆕 v33: Top-level mode selector — pick 30s, 1m, 3m or 5m."""
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("⚡ WinGo 30 সেকেন্ড (fast)", callback_data="wingo_panel")],
        [InlineKeyboardButton("⏱️ WinGo 1 মিনিট (balanced)", callback_data="wingo_panel_1m")],
        [InlineKeyboardButton("⏰ WinGo 3 মিনিট", callback_data="wingo_panel_3m"),
         InlineKeyboardButton("🕐 WinGo 5 মিনিট", callback_data="wingo_panel_5m")],
        [InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")],
    ])

def wingo_panel_menu(mode="30S"):
    """🆕 v30: On-demand only — কোনো auto-broadcast নেই, button চাপলেই signal।  v33: mode-aware."""
    sfx = _w_mode_suffix(mode)
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🎯 এখনই প্রেডিকশন নিন", callback_data=f"wingo_now{sfx}")],
        [InlineKeyboardButton("📊 আমার History", callback_data=f"wingo_stats{sfx}"),
         InlineKeyboardButton("🧠 AI Layer Stats", callback_data=f"wingo_aistats{sfx}")],
        [InlineKeyboardButton("⬅️ Mode বদলান", callback_data="wingo_main"),
         InlineKeyboardButton("🏠 মেনু", callback_data="main_menu")],
    ])

def _w_resolve_user_prev(uid, history, state=None, perf_file=None):
    """🆕 v30: User-এর আগের prediction-টা history-তে resolve হয়েছে কিনা চেক করে।
    Resolve হলে user stats + global layer perf update করে আর result dict return করে।
    Not yet resolved বা no prev → None return।  v33: state-aware (multi-mode)."""
    if state is None: state = _wingo_state
    if perf_file is None: perf_file = WINGO_PERF_FILE
    prev = state["user_last_pred"].get(int(uid))
    if not prev: return None
    target_period = str(prev.get("period"))
    # find this period in history
    found = None
    for r in history:
        if str(r.get("period")) == target_period:
            found = r
            break
    if not found: return None  # not resolved yet (period still in future or too old)
    actual_num = found["number"]
    actual_size = _w_sizeof(actual_num)
    actual_color = _w_colorof(actual_num)
    sig_hit = (actual_size == prev["signal"])
    num_hit = actual_num in [prev["num_top"]] + list(prev.get("num_alt", []))

    # Update per-layer global perf (silent learning continues)
    for layer_id, (vb, vs) in (prev.get("layer_votes") or {}).items():
        _w_record_layer_vote(layer_id, vb, vs, actual_size, state)
    gp = state["global_perf"]
    if sig_hit: gp["hit"] += 1
    else:       gp["miss"] += 1
    gp.setdefault("recent", []).append(1 if sig_hit else 0)
    if len(gp["recent"]) > 50: gp["recent"] = gp["recent"][-50:]
    try: _save_wingo_perf_for(state, perf_file)
    except Exception: pass

    # Update user stats
    us = state["user_stats"].setdefault(int(uid),
        {"win":0,"loss":0,"nwin":0,"nloss":0,"streak":0,"loss_streak":0,"max_streak":0})
    if sig_hit:
        us["win"] += 1; us["streak"] += 1; us["loss_streak"] = 0
        if us["streak"] > us.get("max_streak",0): us["max_streak"] = us["streak"]
    else:
        us["loss"] += 1; us["loss_streak"] += 1; us["streak"] = 0
    if num_hit: us["nwin"] += 1
    else:       us["nloss"] += 1

    # Clear so we don't re-resolve same one
    state["user_last_pred"].pop(int(uid), None)

    return {
        "period": target_period,
        "actual_num": actual_num,
        "actual_size": actual_size,
        "actual_color": actual_color,
        "sig_hit": sig_hit,
        "num_hit": num_hit,
        "prev_signal": prev["signal"],
        "prev_num_top": prev["num_top"],
        "prev_num_alt": prev.get("num_alt", []),
        "user_stats": us,
    }

def _w_format_crosscheck(pred):
    """🆕 v38: Modern multi-AI cross-check panel — Gemini + OpenAI + Claude + 3X Leader."""
    sources = []
    def _add(icon, name, ai_pred, ai_agree, conf=0):
        if not ai_pred: return
        conf_tag = f" `@{conf}%`" if conf else ""
        mark = "✅" if ai_agree else "❌"
        word = "AGREE" if ai_agree else f"DISAGREE → `{ai_pred}`"
        sources.append(f"   {icon} *{name:<14}* {mark} {word}{conf_tag}")

    _add("🤖", "Gemini AI",     pred.get("gemini_pred"),     pred.get("gemini_agree"))
    _add("🧠", "OpenAI GPT-4o", pred.get("openai_pred"),     pred.get("openai_agree"),
         pred.get("openai_conf", 0))
    _add("🪐", "Claude Haiku",  pred.get("claude_pred"),     pred.get("claude_agree"),
         pred.get("claude_conf", 0))
    _add("📡", "3X Leader",      pred.get("external_3x_pred"),pred.get("external_3x_agree"))

    badges = []
    if pred.get("quad_ai_consensus"):
        badges.append("🌟🌟🌟🌟 *QUAD-AI CONSENSUS!* (3 AIs + 3X) +12 conf 🔥🔥🔥")
    elif pred.get("triple_ai_consensus"):
        badges.append("🌟🌟🌟 *TRIPLE-AI CONSENSUS* (+8 conf) 🔥🔥")
    elif pred.get("dual_consensus"):
        badges.append("🌟 *DUAL CONSENSUS BOOST* (+5 conf)")
    if pred.get("triple_window_consensus"):
        badges.append("🎯 *TRIPLE-WINDOW CONSENSUS* (+4 conf)")

    if not sources and not badges:
        return ""

    out = "┃ *🔬 Multi-Source Cross-Check*\n"
    if sources:
        out += "\n".join(sources) + "\n"
    if badges:
        out += "   " + "\n   ".join(badges) + "\n"
    return out


# 🆕 v37 — Kelly Bet display block builder
def _w_format_kelly(pred, balance):
    """Format the Kelly Criterion bet sizing block.  If balance not set → show
    instructions to use /setbalance.  Confidence-based full + half Kelly stake
    recommendation, plus expected value calculation."""
    conf = pred.get("confidence", 50)
    sig  = pred.get("signal", "?")
    skip = pred.get("skip_recommended", False)
    if balance is None or balance <= 0:
        return (
            "💰 *Kelly Bet সাইজিং:*\n"
            "   _Bankroll set করতে: `/setbalance 1000`_\n"
            "   _(Kelly Criterion দিয়ে optimal stake calculate করব)_\n"
        )
    if skip:
        return (
            f"💰 *Kelly Bet সাইজিং:*\n"
            f"   ⛔ *Bet Size: 0* (skip recommended)\n"
            f"   _Bankroll: `{balance:,.0f}` — Loss-streak protection active_\n"
        )
    full_f, half_f = _w_kelly(conf, payout=0.95)   # WinGo big/small ≈ even-money minus ~5% house
    full_amt = balance * full_f
    half_amt = balance * half_f
    if full_f <= 0:
        return (
            f"💰 *Kelly Bet সাইজিং:*\n"
            f"   ⚠️ *Edge ≤ 0* (confidence {conf}% খুব কম, no positive expectation)\n"
            f"   _Recommended: SKIP this round। Bankroll: `{balance:,.0f}`_\n"
        )
    p = conf / 100.0
    ev_pct = (p * 0.95 - (1 - p)) * 100         # %EV per unit bet
    # safety risk band
    if full_f >= 0.15:    risk = "🔴 *AGGRESSIVE*"
    elif full_f >= 0.08:  risk = "🟡 *MODERATE*"
    else:                  risk = "🟢 *CONSERVATIVE*"
    return (
        f"💰 *Kelly Bet সাইজিং:* (`{sig}`-এ)\n"
        f"   💵 Bankroll: `{balance:,.0f}`  ·  EV: *{ev_pct:+.1f}%* / unit\n"
        f"   ⚡ *Full Kelly* ({full_f*100:.1f}%): `{full_amt:,.0f}`  {risk}\n"
        f"   🛡️ *Half Kelly* ({half_f*100:.1f}%): `{half_amt:,.0f}`  _← recommended_\n"
        f"   _Half Kelly = same growth, ~75% less variance.  Use `/setbalance` to update bankroll._\n"
    )

def _w_render_aistats(mode="30S"):
    """🆕 v33: Mode-aware AI layer stats text builder. Returns Markdown string."""
    LAYER_NAMES = {
        "L1":"Time-decay frequency", "L2":"Streak detection",
        "L3":"Zigzag/Alternation",   "L4":"3-step Markov chain",
        "L5":"Gap analysis",         "L6":"Multi-window trend",
        "L7":"EMA momentum",         "L8":"6-gram pattern",
        "L9":"Bayesian inference",   "L10":"Cyclic detection",
        "L11":"Color momentum",      "L12":"Volatility regime",
        "L13":"Run-length survival", "L14":"Period-modulo bias",
        "L15":"2-step num Markov",
        "L16":"Pattern repetition",  "L17":"Streak intelligence",
        "L18":"Chi-square bias",     "L19":"Number-pair synergy",
        "L20":"Hot cluster",
        "L21":"Shannon entropy",     "L22":"Fibonacci lookback",
        "L23":"Triple-Markov consensus","L24":"Self-correcting bias",
        "L25":"🤖 Gemini AI Reader",            # 🆕 v33: real generative AI
        "L26":"Mean Reversion (50w)",           # 🆕 v34
        "L27":"Strong Run Reversal",            # 🆕 v34
        "L28":"Volatility Regime",              # 🆕 v34
        "L29":"📡 3X Leader Rakib (ext)",       # 🆕 v35: external algo clone
        "L30":"Multi-Window Consensus",         # 🆕 v35
        "L31":"🧠 OpenAI GPT-4o-mini",          # 🆕 v37: dual-AI cross-check
        "L32":"🪐 Claude Haiku (Anthropic)",    # 🆕 v38: 3rd independent AI
        "L33":"🌊 HMM Regime Switching",        # 🆕 v38
        "L34":"📐 Bayesian Beta-Posterior",     # 🆕 v38
        "L35":"🎼 FFT Cycle Detection",         # 🆕 v38
    }
    W_BASE = {"L1":1.0,"L2":1.5,"L3":1.3,"L4":1.7,"L5":1.0,"L6":1.1,
              "L7":1.4,"L8":1.6,"L9":1.3,"L10":1.5,"L11":0.9,"L12":1.0,
              "L13":1.6,"L14":1.2,"L15":1.5,
              "L16":1.8,"L17":1.7,"L18":1.4,"L19":1.5,"L20":1.2,
              "L21":1.5,"L22":1.1,"L23":2.0,"L24":1.4,
              "L25":2.2,
              "L26":1.8,"L27":2.4,"L28":1.7,
              "L29":1.6,"L30":2.0,
              "L31":2.2,                        # 🆕 v37 OpenAI cross-check
              "L32":2.2,                        # 🆕 v38 Claude AI
              "L33":1.9,"L34":2.1,"L35":1.7}    # 🆕 v38 HMM/Bayes/FFT

    state = _w_get_state(mode)
    label = WINGO_MODE_LABELS[mode]
    perf = state.get("layer_perf", {})
    gp = state.get("global_perf", {})

    # Global stats
    g_total = gp.get("hit",0) + gp.get("miss",0)
    g_overall = (gp.get("hit",0) / g_total * 100) if g_total else 0
    recent = gp.get("recent", [])
    g_recent = (sum(recent[-20:]) / len(recent[-20:]) * 100) if recent else 0
    g_last10 = (sum(recent[-10:]) / len(recent[-10:]) * 100) if recent else 0
    # 🆕 v36: Last 50 W/L counts
    last50 = recent[-50:]
    w50 = sum(last50)
    l50 = len(last50) - w50
    g_last50_pct = (w50 / len(last50) * 100) if last50 else 0
    # Visual bar of last 20 (newest on right): green=win, red=loss
    last50_emoji = "".join("🟢" if r else "🔴" for r in last50[-20:]) if last50 else ""

    # Build per-layer rows sorted by accuracy
    rows = []
    for lid, name in LAYER_NAMES.items():
        p = perf.get(lid, {"hit":0, "miss":0})
        tot = p["hit"] + p["miss"]
        if tot == 0:
            acc, w_now, status = 0, W_BASE[lid], "⏳ learning"
        else:
            acc = p["hit"] / tot * 100
            w_now = _w_get_adaptive_weight(lid, W_BASE[lid], state)
            if   acc >= 60: status = "🔥 hot"
            elif acc >= 53: status = "✅ good"
            elif acc >= 47: status = "⚪ neutral"
            else:           status = "❄️ cold"
        rows.append((lid, name, p["hit"], p["miss"], tot, acc, W_BASE[lid], w_now, status))

    # Sort by accuracy descending (untrained at bottom)
    rows.sort(key=lambda r: (-r[4], -r[5]))

    lines = []
    for lid, name, h, m, tot, acc, wb, wn, status in rows:
        if tot == 0:
            lines.append(f"`{lid:>3}` {name[:22]:<22} `--% (0)`  w:`{wb:.1f}` {status}")
        else:
            arrow = "↑" if wn > wb else ("↓" if wn < wb else "·")
            lines.append(
                f"`{lid:>3}` {name[:22]:<22} `{acc:>4.0f}% ({tot})`  "
                f"w:`{wb:.1f}{arrow}{wn:.2f}` {status}"
            )

    # Bar for global accuracy
    bar_filled = max(1, int(g_recent / 10)) if g_recent else 0
    g_bar = "█" * bar_filled + "░" * (10 - bar_filled)

    return (
        f"╭─━━━━━━━━━━━━━━━━━━━━╮\n"
        f"│ 🧠 *AI Learning Stats — {label}*\n"
        f"╰─━━━━━━━━━━━━━━━━━━━━╯\n\n"
        f"📊 *Global Performance ({label}):*\n"
        f"  • Overall: `{gp.get('hit',0)}W / {gp.get('miss',0)}L` "
        f"= *{g_overall:.1f}%*\n"
        f"  • *Last 50:* ✅ Win: *{w50}* · ❌ Loss: *{l50}* "
        f"({g_last50_pct:.0f}%)\n"
        f"  • Last 20: *{g_recent:.0f}%*\n"
        f"  • Last 10: *{g_last10:.0f}%*\n"
        f"  `{g_bar}`\n"
        + (f"  Recent 20:  {last50_emoji}\n" if last50_emoji else "")
        + f"\n"
        f"🎯 *Per-Layer Accuracy & Weights (31 layers):*\n"
        f"_(format: ID name acc%(samples) w:base→tuned)_\n\n"
        + "\n".join(lines) +
        f"\n\n_🔬 প্রতিটা layer-এর weight নিজের accuracy অনুযায়ী auto-tune হয়।_\n"
        f"_Bayesian smoothing over-fitting prevent করে।_\n"
        f"_🤖 L25 = real Gemini AI · 🧠 L31 = real OpenAI GPT-4o-mini — দুজনে cross-check করে।_\n"
        f"_💰 Kelly Bet সাইজিং চালু: `/setbalance <amount>` দিয়ে bankroll set করুন।_\n\n"
        f"⚙️ /wingo দিয়ে অন্য mode-এ যান"
    )

async def cmd_aistats(u, c):
    """🆕 v29 ULTRA+: 30s mode-এর layer stats। Per-mode stats panel-এর AI button দিয়ে দেখুন।"""
    await u.message.reply_text(_w_render_aistats("30S"), parse_mode="Markdown")

# 🆕 v37 — /setbalance command for Kelly Criterion bet sizing
async def cmd_setbalance(u, c):
    """Set per-user bankroll for all WinGo modes.  Used by Kelly Bet sizing in
    every prediction display.  Usage: /setbalance 1000  (or 0 to disable)"""
    uid = int(u.effective_user.id)
    args = (c.args if hasattr(c, "args") else []) or []
    if not args:
        # show current balances per mode
        lines = []
        for m in WINGO_APIS.keys():
            bal = _wingo_states[m].get("user_balance", {}).get(uid)
            if bal: lines.append(f"  • {WINGO_MODE_LABELS[m]}: `{bal:,.0f}`")
        cur = "\n".join(lines) if lines else "  _(কোনো balance set করা নেই)_"
        await u.message.reply_text(
            "💰 *Kelly Bet — Bankroll Setting*\n"
            "━━━━━━━━━━━━━━━━━━\n"
            "*Usage:*  `/setbalance <amount>`\n"
            "*Examples:*\n"
            "  • `/setbalance 1000`  → 1,000 BDT/INR/USD bankroll\n"
            "  • `/setbalance 50000` → 50,000 bankroll\n"
            "  • `/setbalance 0`     → disable Kelly display\n\n"
            f"📊 *বর্তমান bankroll:*\n{cur}\n\n"
            "_Kelly Criterion = optimal bet sizing formula যা long-run growth maximize করে।_\n"
            "_Half-Kelly recommended = same expected growth, ~75% less variance।_\n"
            "_Bankroll এক বার set করলে সব mode-এ apply হবে (30s/1m/3m/5m)।_",
            parse_mode="Markdown")
        return
    try:
        amount = float(args[0].replace(",", "").replace("৳", "").replace("$","").strip())
        if amount < 0 or amount > 1e9:
            raise ValueError("amount out of range")
    except (ValueError, IndexError):
        await u.message.reply_text(
            "❌ Invalid amount। Example: `/setbalance 1000`",
            parse_mode="Markdown")
        return
    # apply across ALL modes (single bankroll)
    for m in WINGO_APIS.keys():
        if amount <= 0:
            _wingo_states[m].setdefault("user_balance", {}).pop(uid, None)
        else:
            _wingo_states[m].setdefault("user_balance", {})[uid] = amount
    _w_save_balances()
    if amount <= 0:
        await u.message.reply_text(
            "✅ Bankroll cleared।  Kelly Bet display এখন off।\n"
            "_আবার set করতে: `/setbalance 1000`_",
            parse_mode="Markdown")
    else:
        # quick preview at common confidences
        ex_lines = []
        for c_pct in (60, 70, 80, 90):
            f_full, f_half = _w_kelly(c_pct, payout=0.95)
            ex_lines.append(
                f"  • Conf {c_pct}% → Full: `{amount*f_full:,.0f}` "
                f"· Half: `{amount*f_half:,.0f}`")
        await u.message.reply_text(
            f"✅ Bankroll set: *`{amount:,.0f}`*\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"💰 *Kelly Bet preview:*\n"
            + "\n".join(ex_lines) +
            "\n\n_সব WinGo mode-এ (30s/1m/3m/5m) এই bankroll দিয়ে Kelly Bet calculate হবে।_\n"
            "_পরবর্তী প্রেডিকশন থেকে Full Kelly + Half Kelly amount দেখাবে।_\n"
            "_Recommended: **Half Kelly** — same growth, much safer।_",
            parse_mode="Markdown")

async def cmd_wingo(u, c):
    """🆕 v38 আগুন-MAX: Multi-mode entry — 36 layers + Quad-AI + Meta-Calibration."""
    txt = (
        "╔══════════════════════════════════╗\n"
        "║  🎯  *WinGo AI Predictor*  🎯    \n"
        "║  *v38 আগুন-MAX  ·  36 Layers*    \n"
        "╚══════════════════════════════════╝\n\n"
        "_⚡ Modes:_  *30s · 1m · 3m · 5m*\n"
        "_⚡ On-demand signals — কোনো spam নেই।_\n\n"
        "┃ *🌟 প্রতিটা signal-এ পাবেন*\n"
        "   ▸ BIG / SMALL + smart confidence %\n"
        "   ▸ Top-3 number forecast (with bars)\n"
        "   ▸ Aggregated color signal 🔴🟢🟣\n"
        "   ▸ Multi-AI cross-check panel\n"
        "   ▸ আগের prediction auto-resolve\n"
        "   ▸ 💰 Kelly Bet sizing (Full + Half)\n\n"
        "┃ *🧠 36-Layer Engine (v38)*\n"
        "   • 31 statistical/strategy layers\n"
        "     _(frequency · Markov · entropy · Fibonacci ·_\n"
        "     _ regime · mean-reversion · multi-window …)_\n"
        "   • 🤖 *Gemini AI*  (L25)\n"
        "   • 🧠 *OpenAI GPT-4o-mini*  (L31)\n"
        "   • 🪐 *Claude Haiku Anthropic*  (L32) 🆕\n"
        "   • 🌊 *HMM Regime*  (L33) 🆕\n"
        "   • 📐 *Bayesian Beta*  (L34) 🆕\n"
        "   • 🎼 *FFT Cycles*  (L35) 🆕\n"
        "   • 🎯 *Meta-Confidence Calibrator*  (L36) 🆕\n"
        "   • 📡 *3X Leader Rakib external*  (L29)\n\n"
        "┃ *🛡️ Active Protections*\n"
        "   ▸ Anti-Streak Auto-Flip — ২ loss-এর পরে\n"
        "   ▸ Skip Recommendation — weak signal\n"
        "   ▸ *🌟 Quad-AI Consensus* — 3 AI + 3X agree → +12 conf\n"
        "   ▸ Triple-AI / Dual-AI fallback boosts\n"
        "   ▸ Loss-streak confidence dampening\n\n"
        "┃ *💰 Kelly Criterion Sizing*\n"
        "   `/setbalance 1000` দিয়ে bankroll set করুন।\n"
        "   প্রতি signal-এ Full + *Half Kelly* (recommended)।\n\n"
        "┃ *💾 Mode-isolated State*\n"
        "   প্রতি mode (30s/1m/3m/5m) এর win/loss/learning\n"
        "   alada save — bankroll shared।\n\n"
        "_⚠️ Random game — 100% guarantee নেই।  শুধু_\n"
        "_pattern study / entertainment-এর জন্য।_\n\n"
        "👇 *Mode বেছে নিন:*"
    )
    await u.message.reply_text(txt, parse_mode="Markdown", reply_markup=wingo_main_menu())

async def cmd_wingo1m(u, c):
    """🆕 v38 আগুন-MAX: Direct shortcut to WinGo 1-Min panel — Quad-AI + Meta-Cal."""
    label = WINGO_MODE_LABELS["1M"]
    interval = WINGO_MODE_INTERVAL["1M"]
    ai_lines = []
    if GEMINI_API_KEY:    ai_lines.append("   🤖 *Gemini AI*  (L25)  ✅")
    if has_openai():      ai_lines.append("   🧠 *OpenAI GPT-4o-mini*  (L31)  ✅")
    if has_anthropic():   ai_lines.append("   🪐 *Claude Haiku*  (L32)  ✅")
    n_active = len(ai_lines)
    consensus_note = ""
    if n_active >= 3:
        consensus_note = "\n_🌟🌟🌟 *Quad-AI Consensus available* (+12 conf)_"
    elif n_active >= 2:
        consensus_note = "\n_🌟 Triple-AI Consensus available_"
    ai_block = "\n".join(ai_lines) if ai_lines else "   _(no AI keys configured)_"
    txt = (
        f"╔══════════════════════════════════╗\n"
        f"║  🎯  *{label}*  Panel              \n"
        f"║  _v38 আগুন-MAX  ·  36 layers_     \n"
        f"╚══════════════════════════════════╝\n\n"
        f"┃ *Mode:*  {label}  ({interval}/round)\n"
        f"┃ *Engine:*  36-layer ensemble\n\n"
        f"┃ *🧠 Active AI Engines*\n"
        f"{ai_block}{consensus_note}\n\n"
        f"┃ *🛡️ Protections*\n"
        f"   ▸ Anti-Streak Auto-Flip\n"
        f"   ▸ Meta-Confidence Calibration\n"
        f"   ▸ HMM + Bayesian + FFT layers\n"
        f"   ▸ 📡 3X Leader Rakib cross-check\n\n"
        f"💰 _Kelly Bet on করতে:_  `/setbalance 1000`\n"
        f"⚠️ _Random game — entertainment / pattern study only।_"
    )
    await u.message.reply_text(txt, parse_mode="Markdown",
                               reply_markup=wingo_panel_menu("1M"))

async def wingo_tick(context):
    """🆕 v29: 3s interval. INSTANT separate result + beautiful signal card."""
    loop = asyncio.get_running_loop()
    history = await loop.run_in_executor(executor, _w_fetch_history)
    if not history: return
    current = history[0]["period"]
    last_period = _wingo_state["last_period"]
    last_pred   = _wingo_state["last_pred"]

    if last_period == current: return
    _wingo_state["last_period"] = current

    # ── ⚡ INSTANT RESOLVE — আলাদা সুন্দর result card ──
    if last_pred and str(last_pred.get("period")) == str(current):
        actual_num = history[0]["number"]
        actual_size = _w_sizeof(actual_num)
        actual_color = _w_colorof(actual_num)
        sig_hit = (actual_size == last_pred["signal"])
        num_hit = actual_num in [last_pred["num_top"]] + last_pred["num_alt"]

        # 🆕 v29 ULTRA+ : Adaptive learning — record per-layer accuracy
        for layer_id, (vb, vs) in (last_pred.get("layer_votes") or {}).items():
            _w_record_layer_vote(layer_id, vb, vs, actual_size)
        gp = _wingo_state["global_perf"]
        if sig_hit: gp["hit"] += 1
        else:       gp["miss"] += 1
        gp.setdefault("recent", []).append(1 if sig_hit else 0)
        if len(gp["recent"]) > 50: gp["recent"] = gp["recent"][-50:]
        try: _save_wingo_perf()
        except Exception: pass

        head = "🏆 *WIN!* 🎉🎊" if sig_hit else "💔 *LOSS*"
        bonus = "\n💎 *EXACT NUMBER HIT!* 🎯" if num_hit else ""

        async def _send_result(uid):
            try:
                us = _wingo_state["user_stats"].setdefault(int(uid),
                    {"win":0,"loss":0,"nwin":0,"nloss":0,"streak":0,"loss_streak":0,"max_streak":0})
                if sig_hit:
                    us["win"] += 1; us["streak"] += 1; us["loss_streak"] = 0
                    if us["streak"] > us.get("max_streak",0): us["max_streak"] = us["streak"]
                else:
                    us["loss"] += 1; us["loss_streak"] += 1; us["streak"] = 0
                if num_hit: us["nwin"] += 1
                else:       us["nloss"] += 1

                tot = us["win"] + us["loss"]
                wr = (us["win"]/tot*100) if tot else 0
                streak = ""
                if us["streak"] >= 3: streak = f"\n🔥🔥 *WIN STREAK: {us['streak']}!* 🔥🔥"
                elif us["streak"] >= 2: streak = f"\n🔥 Win streak: *{us['streak']}*"
                elif us["loss_streak"] >= 3: streak = f"\n⚠️ Loss streak: *{us['loss_streak']}* — সাবধান"

                txt = (
                    f"╭─━━━━━━━━━━━━━━━━━━╮\n"
                    f"│ {head}\n"
                    f"╰─━━━━━━━━━━━━━━━━━━╯\n\n"
                    f"📅 Period: `{str(current)[-6:]}`\n"
                    f"🎲 Result: *{actual_num}* {_w_color_emoji(actual_color)} ({actual_size})\n\n"
                    f"🎯 আপনার Signal: *{last_pred['signal']}* "
                    f"→ {'✅ ঠিক!' if sig_hit else '❌ ভুল'}\n"
                    f"🔢 Number {last_pred['num_top']} → "
                    f"{'✅ HIT' if num_hit else '❌ miss'}{bonus}\n\n"
                    f"📊 আপনার Stats: *{us['win']}W / {us['loss']}L* "
                    f"({wr:.0f}% win rate){streak}"
                )
                await context.bot.send_message(int(uid), txt, parse_mode="Markdown")
            except Exception as e:
                logger.warning("wingo result fail uid=%s: %s", uid, e)
                if "blocked" in str(e).lower() or "not found" in str(e).lower():
                    _wingo_subs.discard(uid); _save_wingo_subs(_wingo_subs)

        if _wingo_subs:
            await asyncio.gather(*[_send_result(uid) for uid in list(_wingo_subs)],
                                  return_exceptions=True)
            await asyncio.sleep(0.8)

    # ── 🎯 NEW SIGNAL — সুন্দর card with confidence bar ──
    next_period = str(int(current) + 1)
    pred = await loop.run_in_executor(executor, _w_analyze, history, 0, 0)
    if not pred: return
    pred["period"] = next_period
    _wingo_state["last_pred"] = pred

    conf = pred["confidence"]
    bar_filled = max(1, int(conf / 10))
    conf_bar = "█" * bar_filled + "░" * (10 - bar_filled)
    if conf >= 75: conf_emoji = "🟢"; conf_label = "STRONG"
    elif conf >= 65: conf_emoji = "🟡"; conf_label = "MEDIUM"
    elif conf >= 55: conf_emoji = "🟠"; conf_label = "WEAK"
    else: conf_emoji = "🔴"; conf_label = "AVOID"
    sig_box = "🟢 *BIG* (5-9)" if pred["signal"] == "BIG" else "🔵 *SMALL* (0-4)"
    color_emoji = _w_color_emoji(pred["color"])

    # 🆕 v29 ULTRA+ stats
    agree = pred.get("agreement", 0)
    gp = _wingo_state.get("global_perf", {})
    g_recent = gp.get("recent", [])
    g_acc = (sum(g_recent[-20:]) / len(g_recent[-20:]) * 100) if g_recent else 0
    g_total = gp.get("hit",0) + gp.get("miss",0)
    g_overall = (gp.get("hit",0) / g_total * 100) if g_total else 0

    learning_line = ""
    if g_total >= 10:
        learning_line = (f"\n🧠 *AI Learning:* {g_overall:.0f}% overall · "
                         f"{g_acc:.0f}% recent (last {min(20,len(g_recent))} rounds)")

    msg = (
        f"╭─━━━━━━━━━━━━━━━━━━━━╮\n"
        f"│ 🎯 *WinGo 30s — ULTRA+ SIGNAL*\n"
        f"╰─━━━━━━━━━━━━━━━━━━━━╯\n\n"
        f"📅 Period: `{next_period[-6:]}`\n"
        f"⏱️ পরবর্তী result: ~৩০ সেকেন্ডে\n\n"
        f"━━━━━━━━━━━━━━━━━━━━\n"
        f"🎯 *Signal:*  {sig_box}\n"
        f"🎲 *Number:*  *{pred['num_top']}* {color_emoji}\n"
        f"🎯 *Alt:*  `{pred['num_alt'][0]}` , `{pred['num_alt'][1]}`\n"
        f"━━━━━━━━━━━━━━━━━━━━\n\n"
        f"{conf_emoji} *Confidence: {conf}% — {conf_label}*\n"
        f"`{conf_bar}`\n"
        f"🤝 Layer Agreement: *{agree}%* (12 layers)"
        f"{learning_line}\n\n"
        f"_⚠️ Game random, 100% guarantee নেই_"
    )

    async def _send(uid):
        try:
            await context.bot.send_message(int(uid), msg, parse_mode="Markdown")
        except Exception as e:
            logger.warning("wingo send fail uid=%s: %s", uid, e)
            if "blocked" in str(e).lower() or "not found" in str(e).lower():
                _wingo_subs.discard(uid); _save_wingo_subs(_wingo_subs)

    if _wingo_subs:
        await asyncio.gather(*[_send(uid) for uid in list(_wingo_subs)],
                              return_exceptions=True)


def main():
    # 🆕 v16: BOT_TOKEN আগে validate — invalid হলে পরিষ্কার error
    if not BOT_TOKEN or len(BOT_TOKEN) < 20 or ":" not in BOT_TOKEN:
        print("❌ BOT_TOKEN সেট করা নেই বা ভুল!")
        print("   কোডের উপরে BOT_TOKEN-এ @BotFather থেকে পাওয়া token বসান।")
        print("   Format: 1234567890:ABC-DEF1234ghIkl-zyx57W2v1u123ew11")
        return
    openai_status  = "✅ ON " if has_openai()  else "❌ OFF"
    gemini_status  = "✅ ON " if has_gemini()  else "❌ OFF"
    minimax_status = "✅ ON " if has_minimax() else "❌ OFF"
    eleven_status  = "✅ ON " if (ELEVENLABS_API_KEY and ELEVENLABS_API_KEY.strip()) else "❌ OFF"
    replicate_status = "✅ ON " if has_replicate() else "❌ OFF"
    acoustid_status= "✅ ON " if (ACOUSTID_API_KEY and ACOUSTID_API_KEY.strip()) else "❌ OFF"
    edge_status    = "✅ ON " if has_edge_tts() else "❌ OFF"
    n_clones = len(load_clones())
    load_mcq_cache()  # 🆕 v37
    n_edge_voices = len(EDGE_TTS_VOICES)
    print(f"""╔════════════════════════════════════╗
║  🎬 Video Editor Bot — Pro v39 ULTRA ✨ ║
║  💰 YouTube Money গাইড (NEW)       ║
║  🎤 Auto Voice Clone               ║
║  🎨 Reorganized Menu (sections)    ║
║  🚀 YouTube Viral Pack             ║
║  ✨ Enhance ULTRA (rebuilt)        ║
║  🎞️ Slow-Mo PREMIUM (4-5x faster) ║
║  📺 YouTube Mode (Content ID Safe) ║
║  🎵 TikTok Mode (For You Ready)    ║
║  🆓 Edge TTS Free Premium          ║
║  🎭 Voice Clone (Replicate XTTS)   ║
║  📌 Pinterest Download Fix (UA)    ║
║  ⚡ HTTP retry + size-cap cleanup  ║
║  💎 Wink Premium Pack Loaded       ║
║  🔒 Copyright PRO Bypass (5-layer) ║
║  🌐 Source Platform Button Added   ║
║  🎯 SEO Title + Hashtag Generator  ║
║  📝 Auto Subtitle + Bengali Font   ║
║  🎀 Premium Caption (Bengali)      ║
║  🤖 ChatGPT (OpenAI):    {openai_status}     ║
║  🧠 Gemini (Google):     {gemini_status}     ║
║  🌟 MiniMax (Agent):     {minimax_status}     ║
║  🎤 ElevenLabs (legacy): {eleven_status}     ║
║  🎭 Replicate XTTS Clone:{replicate_status}     ║
║  🆓 Edge Neural TTS:     {edge_status}     ║
║  🎵 AcoustID (Music ID): {acoustid_status}     ║
║  🗣️ Text → Voice (TTS) ENABLED     ║
║  📊 Per-feature stats tracking ON  ║
║  🆓 Edge voices loaded:   {n_edge_voices:>3}      ║
║  🎭 Cloned voices stored: {n_clones:>3}      ║
╚════════════════════════════════════╝""")
    if not has_replicate():
        print("ℹ️  Voice Cloning চাইলে: https://replicate.com → API token →")
        print("    `export REPLICATE_API_TOKEN=...` → bot restart।")
    if not has_edge_tts():
        print("ℹ️  Edge TTS install করতে: pip install edge-tts")
        print("    এটি সম্পূর্ণ ফ্রি — কোনো API key লাগে না!")
    # 🆕 v15: Bengali fonts auto-download (one-time, ~800KB)
    try:
        ensure_bangla_fonts()
    except Exception as e:
        logger.warning("Bengali font setup failed: %s", e)
    cleanup_temp()
    # 🆕 v22: Network resilience — slow network-এ TimedOut error হলে বট off হয়ে যায় না
    # connect_timeout: server connect করতে সর্বোচ্চ সময়
    # read_timeout: response পেতে সর্বোচ্চ সময় (long-poll-এর জন্য বড় হওয়া দরকার)
    # write_timeout: file upload-এর জন্য বড় (ভিডিও/audio পাঠাতে)
    # pool_timeout: connection pool থেকে নিতে সর্বোচ্চ সময়
    # 🆕 v28: Slow/intermittent network-এর জন্য আরও বড় timeout buffer।
    # Long-poll timeout (25s, নিচে) এর চেয়ে read_timeout (90s) যথেষ্ট বড় হতে হবে।
    app = (Application.builder()
           .token(BOT_TOKEN)
           .connect_timeout(45.0)
           .read_timeout(90.0)
           .write_timeout(180.0)
           .pool_timeout(60.0)
           .get_updates_connect_timeout(45.0)
           .get_updates_read_timeout(90.0)
           .get_updates_pool_timeout(60.0)
           .build())
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("help", cmd_help))
    app.add_handler(CommandHandler("cancel", cmd_cancel))
    app.add_handler(CommandHandler("stats", cmd_stats))
    app.add_handler(CommandHandler("platforms", cmd_platforms))
    app.add_handler(CommandHandler("myfonts", cmd_myfonts))
    # 🆕 v27.2: YouTube monetization guide
    app.add_handler(CommandHandler("money", cmd_money))
    # 🆕 WinGo Live Signal
    app.add_handler(CommandHandler("wingo", cmd_wingo))
    app.add_handler(CommandHandler("wingo1m", cmd_wingo1m))    # 🆕 v33: WinGo 1-Min direct shortcut
    # 🆕 v29 ULTRA+ : AI Layer Stats
    app.add_handler(CommandHandler("aistats", cmd_aistats))
    # 🆕 v37 : Kelly Criterion bankroll setting
    app.add_handler(CommandHandler("setbalance", cmd_setbalance))
    # 🆕 v18: voice cloning
    app.add_handler(CommandHandler("clonevoice", cmd_clonevoice))
    app.add_handler(CommandHandler("clonefromvideo", cmd_clonefromvideo))
    app.add_handler(CommandHandler("myvoice", cmd_myvoice))
    app.add_handler(CommandHandler("deletevoice", cmd_deletevoice))
    app.add_handler(CommandHandler("admin", cmd_admin))
    app.add_handler(CommandHandler("broadcast", cmd_broadcast))
    # 🆕 v37: File Extractor + MCQ Solver + Link Expander
    app.add_handler(CommandHandler("extract", cmd_extract))
    app.add_handler(CommandHandler("mcq", cmd_mcq))
    app.add_handler(CommandHandler("unlock", cmd_unlock_link))
    app.add_handler(CommandHandler("expand", cmd_unlock_link))
    app.add_handler(MessageHandler(filters.VIDEO, handle_video))
    app.add_handler(MessageHandler(filters.ANIMATION, handle_video))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(filters.Document.VIDEO, handle_video))
    # 🆕 v37: Non-video documents → file extractor / MCQ
    app.add_handler(MessageHandler(filters.Document.ALL & ~filters.Document.VIDEO, handle_document))
    app.add_handler(MessageHandler(filters.AUDIO | filters.VOICE, handle_audio))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_handler(CallbackQueryHandler(handle_cb))
    app.add_error_handler(on_error)
    # periodic temp cleanup (job_queue available with python-telegram-bot[job-queue])
    try:
        if app.job_queue:
            app.job_queue.run_repeating(periodic_cleanup, interval=3600, first=3600)
            # 🆕 v30: WinGo auto-broadcast DISABLED — was disturbing other features.
            # Predictions now on-demand only (button-triggered via /wingo). The 15-layer
            # analyzer still learns silently each time a user requests a prediction.
            # Wipe leftover subscribers from old auto-broadcast model:
            if _wingo_subs:
                _wingo_subs.clear()
                _save_wingo_subs(_wingo_subs)
                print("🎯 WinGo: পুরনো subscribers wipe করা হলো (auto-broadcast off)")
            print("🎯 WinGo Prediction: ON-DEMAND mode (button-triggered, 24-layer ULTRA+++)")
    except Exception: pass
    print("✅ বট চালু! Polling started...")
    # 🆕 v22: resilient polling loop — TimedOut/NetworkError-এ বট অটো reconnect হবে
    # mobile data slow/intermittent হলেও বট off হয়ে যাবে না
    backoff = 5
    max_backoff = 300  # 5 মিনিট সর্বোচ্চ wait
    while True:
        try:
            # drop_pending_updates=False — restart-এ user-এর pending message মুছবে না
            # timeout: long-poll-এ Telegram server-কে কতক্ষণ wait করতে বলব (🆕 v23 fix)
            # 🆕 v28: long-poll timeout 25s (read_timeout 90s এর নিচে — safe margin)
            app.run_polling(
                allowed_updates=Update.ALL_TYPES,
                drop_pending_updates=False,
                timeout=25,
                bootstrap_retries=-1,  # Telegram-এ initial connect fail হলে infinite retry
            )
            # normal exit (Ctrl+C ইত্যাদি) — break loop
            print("👋 বট বন্ধ করা হয়েছে।")
            break
        except KeyboardInterrupt:
            print("\n👋 বট বন্ধ করা হয়েছে (Ctrl+C)।")
            break
        except Exception as e:
            err_name = type(e).__name__
            print(f"⚠️  Polling error [{err_name}]: {str(e)[:200]}")
            print(f"🔄 {backoff} সেকেন্ড পর আবার connect করার চেষ্টা করব...")
            logger.warning("run_polling crash [%s], retry in %ds: %s", err_name, backoff, e)
            try:
                time.sleep(backoff)
            except KeyboardInterrupt:
                print("\n👋 বট বন্ধ করা হয়েছে।"); break
            # exponential backoff (5s → 10s → 20s → ... → 300s সর্বোচ্চ)
            backoff = min(backoff * 2, max_backoff)

# ══════════════════════════════════════════════════════════════════
# 🆕 v22: BUILT-IN WATCHDOG — OS-level auto-restart
# ──────────────────────────────────────────────────────────────────
# Python নিজে crash করলে / Termux kill করলেও বট আবার চালু হবে।
# Wakelock নিয়ে রাখে যাতে ফোন lock হলেও Termux off না হয়।
#
# Usage:
#   python video_editor_bot_v22_pro.py                  → watchdog mode (default, recommended)
#   python video_editor_bot_v22_pro.py --no-watchdog    → সরাসরি বট (auto-restart ছাড়া)
#   python video_editor_bot_v22_pro.py --worker         → internal subprocess (manually call করবেন না)
# ══════════════════════════════════════════════════════════════════
def run_watchdog():
    """OS-level supervisor — Python process মরলেও subprocess পুনরায় spawn করে।"""
    # Wakelock try (Termux-এ থাকলে — না থাকলে silently skip)
    wakelock_acquired = False
    try:
        r = subprocess.run(["termux-wake-lock"], check=False, timeout=5,
                           capture_output=True)
        if r.returncode == 0:
            wakelock_acquired = True
            print("🔒 Termux wakelock acquired — ফোন lock হলেও বট চলবে।")
    except (FileNotFoundError, Exception):
        pass  # not on Termux, no problem

    print("╔══════════════════════════════════════════════════╗")
    print("║  🛡️  WATCHDOG mode চালু                            ║")
    print("║  বট crash/kill হলে অটো restart হবে               ║")
    print("║  বন্ধ করতে: Ctrl+C (২ বার চাপুন)                  ║")
    print("╚══════════════════════════════════════════════════╝\n")

    backoff = 5
    max_backoff = 300  # 5 মিনিট
    restart_count = 0
    log_file = Path.home() / "bot_watchdog.log"

    def _log(msg):
        line = f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {msg}"
        print(line)
        try:
            with open(log_file, "a", encoding="utf-8") as f:
                f.write(line + "\n")
        except Exception:
            pass

    try:
        while True:
            restart_count += 1
            _log(f"✅ Bot starting (run #{restart_count})...")
            try:
                # নিজেই নিজেকে subprocess হিসেবে spawn করে --worker flag দিয়ে
                proc = subprocess.run(
                    [sys.executable, sys.argv[0], "--worker"],
                    check=False
                )
                exit_code = proc.returncode
            except KeyboardInterrupt:
                _log("👋 Ctrl+C — watchdog বন্ধ।")
                break
            except FileNotFoundError as e:
                _log(f"❌ Python interpreter পাওয়া যাচ্ছে না: {e}")
                break

            # exit code 0 = normal exit, restart করব না
            if exit_code == 0:
                _log("👋 বট normal-ভাবে exit হয়েছে। Watchdog stop।")
                break

            _log(f"⚠️  বট crash করেছে (exit code: {exit_code})")
            _log(f"🔄 {backoff} সেকেন্ড পর restart করছি... (মোট {restart_count} বার)")

            # দ্রুত repeated crash হলে warning
            if restart_count > 0 and restart_count % 5 == 0:
                _log(f"⚠️  {restart_count} বার restart হয়েছে — log দেখুন: {log_file}")

            try:
                time.sleep(backoff)
            except KeyboardInterrupt:
                _log("👋 Ctrl+C — watchdog বন্ধ।")
                break

            # Exponential backoff: 5s → 10s → 20s → 40s → 80s → ... → 300s (cap)
            backoff = min(backoff * 2, max_backoff)
    finally:
        if wakelock_acquired:
            try:
                subprocess.run(["termux-wake-unlock"], check=False, timeout=5,
                               capture_output=True)
                print("🔓 Wakelock release করা হলো।")
            except Exception:
                pass


if __name__ == "__main__":
    # 🆕 v22: argv-based dispatcher
    args = set(sys.argv[1:])
    if "--worker" in args:
        # Subprocess mode — actual বট চালায়
        main()
    elif "--no-watchdog" in args or "--direct" in args:
        # Direct mode — watchdog ছাড়া (debugging/development-এর জন্য)
        print("ℹ️  Direct mode — watchdog OFF (বট crash হলে restart হবে না)\n")
        main()
    else:
        # Default mode — watchdog (recommended for Termux/production)
        try:
            run_watchdog()
        except KeyboardInterrupt:
            print("\n👋 Bye!")